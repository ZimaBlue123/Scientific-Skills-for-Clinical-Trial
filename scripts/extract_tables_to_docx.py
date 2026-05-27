# -*- coding: utf-8 -*-
"""
图片 / 截图 → Word：通用 OCR + 表格流水线（不绑定某一产品说明）。

能力：
1. **仅表格**：img2table + Tesseract → Word 原生表格（含合并单元格）。
2. **整页可编辑**（推荐用于公示页、CSR 截图）：词级 OCR 与表格区域掩膜穿插排版；可选多尺度放大，
   对 img2table 取「单元格数最多」的一份，减轻漏表、正文误吞表内数字。

流程要点：
- 将 tesseract.exe 所在目录加入 PATH；TESSDATA_PREFIX 指向含 .traineddata 的 tessdata。
- img2table 检测表格结构；Tesseract 识别单元格与正文词框。

依赖：python-docx、Pillow、pytesseract、img2table、pandas、beautifulsoup4、
      opencv-contrib-python-headless（需 cv2.ximgproc.niBlackThreshold）。

用法：
  py -3.10 scripts/extract_tables_to_docx.py --image page1.png -o tables.docx
  py -3.10 scripts/extract_tables_to_docx.py --editable --image a.jpg --image b.jpg -o full.docx
  py -3.10 scripts/extract_tables_to_docx.py --editable --portrait --image x.png -o out.docx --tesseract D:\\\\tesseract-ocr\\\\tesseract.exe
"""
from __future__ import annotations

import argparse
import io
import os
import re
import statistics
import sys
from pathlib import Path
from typing import Any, List, Sequence, Tuple

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parent
DEFAULT_TESSERACT = Path(r"D:\tesseract-ocr\tesseract.exe")


def _resolve_tesseract_exe(cli: Path | None) -> Path:
    if cli is not None:
        return cli
    e = os.environ.get("TESSERACT_CMD", "").strip().strip('"')
    return Path(e) if e else DEFAULT_TESSERACT


def _tessdata_dir(tesseract_exe: Path) -> Path:
    prefix = os.environ.get("TESSDATA_PREFIX", "").strip().strip('"')
    if prefix:
        root = Path(prefix)
        if (root / "tessdata").is_dir():
            return root / "tessdata"
        if root.name.lower() == "tessdata" and root.is_dir():
            return root
    return tesseract_exe.parent / "tessdata"


def _prepend_tesseract_to_path(tesseract_exe: Path) -> None:
    d = str(tesseract_exe.parent)
    os.environ["PATH"] = d + os.pathsep + os.environ.get("PATH", "")
    td = _tessdata_dir(tesseract_exe)
    os.environ["TESSDATA_PREFIX"] = str(td)


def _set_cell_run(
    cell: Any,
    text: str,
    *,
    size_pt: float = 7.5,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = align
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        pass
    r = p.add_run(text)
    r.font.size = Pt(size_pt)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        r.bold = True


def _html_table_to_grid(table_tag: Any) -> Tuple[List[List[Any]], int, int]:
    """
    将带 rowspan/colspan 的 HTML 表解析为逻辑网格。
    每个格子为 None（被合并占用）或 dict(text, rs, cs)。
    """
    trs = table_tag.find_all("tr", recursive=False)
    if not trs:
        return [], 0, 0
    n_cols = 0
    for tr in trs:
        tds = tr.find_all(["td", "th"], recursive=False)
        n_cols = max(n_cols, sum(int(td.get("colspan", 1) or 1) for td in tds))

    grid: List[List[Any]] = []
    r = 0
    for tr in trs:
        while len(grid) <= r:
            grid.append([None] * n_cols)
        c = 0
        for td in tr.find_all(["td", "th"], recursive=False):
            while c < n_cols and grid[r][c] is not None:
                c += 1
            if c >= n_cols:
                break
            rs = max(1, int(td.get("rowspan", 1) or 1))
            cs = max(1, int(td.get("colspan", 1) or 1))
            while len(grid) < r + rs:
                grid.append([None] * n_cols)
            txt = td.get_text(separator="\n", strip=True)
            is_th = td.name == "th"
            grid[r][c] = {"text": txt, "rs": rs, "cs": cs, "th": is_th}
            for rr in range(r, r + rs):
                for cc in range(c, c + cs):
                    if rr == r and cc == c:
                        continue
                    grid[rr][cc] = "MERGED"
            c += cs
        r += 1
    return grid, len(grid), n_cols


def _column_char_weights(grid: List[List[Any]], n_cols: int) -> List[float]:
    wts = [10.0] * n_cols
    for ri in range(len(grid)):
        for ci in range(n_cols):
            cell_obj = grid[ri][ci]
            if not isinstance(cell_obj, dict):
                continue
            raw = str(cell_obj.get("text", "")).replace("\n", " ")
            ln = len(raw.strip()) + 4.0
            wts[ci] = max(wts[ci], ln, 8.0)
    return wts


def append_html_table_to_doc(doc: Document, html_fragment: str, *, title: str | None, width_in: float) -> None:
    soup = BeautifulSoup(html_fragment, "html.parser")
    table_tag = soup.find("table")
    if table_tag is None:
        return

    grid, n_rows, n_cols = _html_table_to_grid(table_tag)
    if n_rows == 0 or n_cols == 0:
        return

    if title:
        h = doc.add_paragraph()
        hr = h.add_run(title)
        hr.bold = True
        hr.font.size = Pt(11)
        hr.font.name = "Microsoft YaHei"
        hr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        tbl.style = "Normal Table"

    for ri in range(n_rows):
        for ci in range(n_cols):
            cell_obj = grid[ri][ci]
            if cell_obj in (None, "MERGED"):
                continue
            assert isinstance(cell_obj, dict)
            rs = int(cell_obj["rs"])
            cs = int(cell_obj["cs"])
            c0 = tbl.cell(ri, ci)
            if rs > 1 or cs > 1:
                c1 = tbl.cell(ri + rs - 1, ci + cs - 1)
                c0.merge(c1)
            is_th = bool(cell_obj.get("th"))
            _set_cell_run(
                c0,
                str(cell_obj.get("text", "")),
                size_pt=7.5 if not is_th else 7.8,
                bold=is_th,
                align=WD_ALIGN_PARAGRAPH.CENTER if is_th else WD_ALIGN_PARAGRAPH.LEFT,
            )

    weights = _column_char_weights(grid, n_cols)
    sw = sum(weights) or 1.0
    try:
        for ci in range(n_cols):
            tbl.columns[ci].width = Inches(width_in * float(weights[ci]) / sw)
    except Exception:
        try:
            for ci in range(n_cols):
                tbl.columns[ci].width = Inches(width_in / n_cols)
        except Exception:
            pass

    doc.add_paragraph()


def _raw_html_table(extracted: Any) -> str:
    """不 prettify，便于解析。"""
    from img2table.tables.extraction._utils import (  # type: ignore  # noqa: PLC0415
        create_all_rectangles,
        group_cell_positions,
    )

    cell_span_list = [
        cell_span
        for cells in group_cell_positions(table=extracted)
        for cell_span in create_all_rectangles(cell_positions=cells)
    ]
    cell_span_list = [span for cell_span in cell_span_list for span in cell_span.html_cell_span()]

    rows_html: List[str] = []
    for row_idx in range(len(extracted.content)):
        row_cells = sorted(
            [cell_span for cell_span in cell_span_list if cell_span.top_row == row_idx],
            key=lambda cs: cs.col_left,
        )
        rows_html.append("<tr>" + "".join([cs.html for cs in row_cells]) + "</tr>")
    return "<table>" + "".join(rows_html) + "</table>"


def extract_tables_from_bytes(
    jpeg_bytes: bytes,
    tesseract_exe: Path,
    *,
    lang: str = "chi_sim+eng",
    psm: int = 11,
    borderless: bool = True,
) -> List[Any]:
    _prepend_tesseract_to_path(tesseract_exe)
    from img2table.document import Image as I2Image  # type: ignore  # noqa: PLC0415
    from img2table.ocr import TesseractOCR  # type: ignore  # noqa: PLC0415

    tess_dir = str(_tessdata_dir(tesseract_exe))
    ocr = TesseractOCR(lang=lang, psm=int(psm), tessdata_dir=tess_dir)
    img = I2Image(jpeg_bytes)
    raw = img.extract_tables(
        ocr=ocr,
        implicit_rows=True,
        borderless_tables=borderless,
        min_confidence=20,
    )
    if isinstance(raw, dict):
        out: List[Any] = []
        for k in sorted(raw.keys()):
            out.extend(raw[k])
        return out
    return list(raw)


def _table_cell_count_score(ext: Any) -> int:
    rows = getattr(ext, "content", None) or []
    n = 0
    for row in rows:
        if row is None:
            continue
        try:
            n += len(row)
        except TypeError:
            continue
    return int(n)


def extract_tables_from_bytes_best(
    jpeg_bytes: bytes,
    tesseract_exe: Path,
    *,
    lang: str = "chi_sim+eng",
    psm: int = 11,
    dual_borderless: bool = True,
) -> List[Any]:
    """
    依次尝试 borderless_tables=True/False，按「单元格总数」择优，减少漏检与畸形表。
    """
    if not dual_borderless:
        return extract_tables_from_bytes(
            jpeg_bytes,
            tesseract_exe,
            lang=lang,
            psm=int(psm),
            borderless=True,
        )

    best: List[Any] = []
    best_key: Tuple[int, int, int] = (-1, -1, -1)
    for borderless in (True, False):
        try:
            tabs = extract_tables_from_bytes(
                jpeg_bytes,
                tesseract_exe,
                lang=lang,
                psm=int(psm),
                borderless=borderless,
            )
        except Exception:
            continue
        score = sum(_table_cell_count_score(t) for t in tabs)
        n_tab = len(tabs)
        tie_prio = 1 if borderless else 0
        key = (score, n_tab, tie_prio)
        if key > best_key:
            best_key = key
            best = tabs
    return best


def extract_tables_from_image_path(
    image_path: Path,
    tesseract_exe: Path,
    *,
    lang: str = "chi_sim+eng",
    psm: int = 11,
    borderless: bool = True,
) -> List[Any]:
    return extract_tables_from_bytes(
        Path(image_path).read_bytes(),
        tesseract_exe,
        lang=lang,
        psm=psm,
        borderless=borderless,
    )


def _jpeg_bytes_at_scale(image_bytes: bytes, scale: float) -> bytes:
    """将任意 Pillow 可读图像统一为 JPEG 字节；scale>1 时放大以改善 img2table 线检测。"""
    from PIL import Image  # type: ignore  # noqa: PLC0415

    if scale <= 1.001:
        return image_bytes
    im = Image.open(io.BytesIO(image_bytes))
    im = im.convert("RGB")
    w, h = im.size
    try:
        res = Image.Resampling.LANCZOS
    except AttributeError:
        res = Image.LANCZOS  # type: ignore[attr-defined]
    nw, nh = max(1, int(w * scale)), max(1, int(h * scale))
    im2 = im.resize((nw, nh), res)
    buf = io.BytesIO()
    im2.save(buf, format="JPEG", quality=93, optimize=True)
    return buf.getvalue()


def prepare_jpeg_for_editable_page(
    image_bytes: bytes,
    tesseract_exe: Path,
    *,
    table_psm: int,
    dual_borderless: bool,
    borderless: bool,
    auto_upscale: bool,
    upscale_scales: Sequence[float] = (1.0, 1.5, 2.0),
) -> Tuple[bytes, List[Any]]:
    """
    在多个缩放比例下跑 img2table，选单元格总数最多的一份；返回 (工作图像字节, 已解析表列表)。
    后续词级 OCR 与 bbox 掩膜均基于同一工作图像，避免坐标不一致。
    """
    if not auto_upscale:
        if dual_borderless:
            tabs = extract_tables_from_bytes_best(
                image_bytes, tesseract_exe, psm=int(table_psm)
            )
        else:
            tabs = extract_tables_from_bytes(
                image_bytes,
                tesseract_exe,
                psm=int(table_psm),
                borderless=borderless,
            )
        return image_bytes, tabs

    best_b = image_bytes
    best_t: List[Any] = []
    best_key = (-1, -1)
    for sc in upscale_scales:
        if sc < 1.0:
            continue
        cand = _jpeg_bytes_at_scale(image_bytes, sc) if sc > 1.001 else image_bytes
        if dual_borderless:
            tabs = extract_tables_from_bytes_best(cand, tesseract_exe, psm=int(table_psm))
        else:
            tabs = extract_tables_from_bytes(
                cand,
                tesseract_exe,
                psm=int(table_psm),
                borderless=borderless,
            )
        key = (sum(_table_cell_count_score(t) for t in tabs), len(tabs))
        if key > best_key:
            best_key = key
            best_b, best_t = cand, tabs
    return best_b, best_t


def _pytesseract_words(
    im_rgb: Any,
    tesseract_exe: Path,
    *,
    min_conf: int = 35,
    psm: int = 6,
    lang: str = "chi_sim+eng",
) -> List[dict[str, Any]]:
    import pytesseract  # type: ignore  # noqa: PLC0415
    from pytesseract import Output  # type: ignore  # noqa: PLC0415

    exe = tesseract_exe.expanduser()
    if not exe.is_file():
        raise FileNotFoundError(f"未找到 Tesseract：{exe}")
    _prepend_tesseract_to_path(tesseract_exe)
    pytesseract.pytesseract.tesseract_cmd = str(exe)
    data = pytesseract.image_to_data(
        im_rgb,
        lang=lang,
        output_type=Output.DICT,
        config=f"--psm {int(psm)}",
    )
    words: List[dict[str, Any]] = []
    n = len(data["text"])
    for i in range(n):
        t = (data["text"][i] or "").strip()
        if not t:
            continue
        try:
            conf = int(float(data["conf"][i]))
        except (TypeError, ValueError):
            continue
        if conf < int(min_conf):
            continue
        left = int(data["left"][i])
        top = int(data["top"][i])
        width = int(data["width"][i])
        height = int(data["height"][i])
        words.append(
            {
                "text": t,
                "left": left,
                "top": top,
                "width": width,
                "height": height,
                "right": left + width,
                "bottom": top + height,
            }
        )
    return words


def _bbox_tuple(ext: Any) -> Tuple[float, float, float, float] | None:
    b = getattr(ext, "bbox", None)
    if b is None:
        return None
    return float(b.x1), float(b.y1), float(b.x2), float(b.y2)


def _expand_bbox(
    bb: Tuple[float, float, float, float],
    img_w: int,
    img_h: int,
    *,
    pad_px: float,
    pad_rel: float,
) -> Tuple[float, float, float, float]:
    x1, y1, x2, y2 = bb
    dw = max(x2 - x1, 1.0)
    dh = max(y2 - y1, 1.0)
    extra = float(pad_px) + float(pad_rel) * float(min(dw, dh))
    nx1 = max(0.0, x1 - extra)
    ny1 = max(0.0, y1 - extra)
    nx2 = min(float(img_w), x2 + extra)
    ny2 = min(float(img_h), y2 + extra)
    return nx1, ny1, nx2, ny2


def _table_mask_boxes(
    tables: Sequence[Any],
    img_w: int,
    img_h: int,
    *,
    pad_px: float,
    pad_rel: float,
) -> List[Tuple[float, float, float, float]]:
    out: List[Tuple[float, float, float, float]] = []
    for ext in tables:
        raw = _bbox_tuple(ext)
        if raw is None:
            continue
        out.append(_expand_bbox(raw, img_w, img_h, pad_px=pad_px, pad_rel=pad_rel))
    return out


def _word_overlap_table_fraction(w: dict[str, Any], bb: Tuple[float, float, float, float]) -> float:
    wx1, wy1 = float(w["left"]), float(w["top"])
    wx2, wy2 = float(w["right"]), float(w["bottom"])
    ix1, iy1 = max(wx1, bb[0]), max(wy1, bb[1])
    ix2, iy2 = min(wx2, bb[2]), min(wy2, bb[3])
    if ix2 <= ix1 or iy2 <= iy1:
        return 0.0
    inter = (ix2 - ix1) * (iy2 - iy1)
    wa = max((wx2 - wx1) * (wy2 - wy1), 1.0)
    return inter / wa


def _center_in_box(cx: float, cy: float, bb: Tuple[float, float, float, float]) -> bool:
    return bb[0] <= cx <= bb[2] and bb[1] <= cy <= bb[3]


def _word_in_table_masks(
    w: dict[str, Any],
    boxes: Sequence[Tuple[float, float, float, float]],
    *,
    overlap_frac: float,
) -> bool:
    cx = float(w["left"]) + float(w["width"]) / 2.0
    cy = float(w["top"]) + float(w["height"]) / 2.0
    for bb in boxes:
        if _center_in_box(cx, cy, bb):
            return True
        if _word_overlap_table_fraction(w, bb) >= float(overlap_frac):
            return True
    return False


def _looks_like_leaked_table_prose(s: str) -> bool:
    """正文里误并入的「整表数字墙」启发式过滤。"""
    t = re.sub(r"\s+", "", s)
    if len(t) < 90:
        return False
    digits = sum(c.isdigit() for c in t)
    dr = digits / len(t)
    if "|" in s and dr > 0.18:
        return True
    if dr < 0.24:
        return False
    nparen = s.count("(") + s.count(")")
    if dr > 0.36 and nparen >= 10 and len(t) > 130:
        return True
    if dr > 0.40 and len(t) > 180:
        return True
    return False


def _join_row_words_with_gaps(row: List[dict[str, Any]]) -> str:
    if not row:
        return ""
    sr = sorted(row, key=lambda w: w["left"])
    med_h = statistics.median([float(w["height"]) for w in sr]) if sr else 9.0
    gap_need = max(2.0, med_h * 0.32)
    parts: List[str] = []
    prev_right: float | None = None
    for w in sr:
        if prev_right is not None:
            gap = float(w["left"]) - prev_right
            if gap > gap_need:
                parts.append(" ")
        parts.append(str(w["text"]))
        prev_right = float(w["right"])
    return "".join(parts).strip()


def _cluster_rows_from_words(words: List[dict[str, Any]], y_tol: float) -> List[List[dict[str, Any]]]:
    if not words:
        return []
    items = sorted(words, key=lambda w: w["top"] + w["height"] / 2)
    rows: List[List[dict[str, Any]]] = []
    for w in items:
        cy = w["top"] + w["height"] / 2
        placed = False
        for row in rows:
            ref = statistics.mean(x["top"] + x["height"] / 2 for x in row)
            if abs(cy - ref) <= y_tol:
                row.append(w)
                placed = True
                break
        if not placed:
            rows.append([w])
    for row in rows:
        row.sort(key=lambda x: x["left"])
    rows.sort(key=lambda r: r[0]["top"] + r[0]["height"] / 2)
    return rows


def _paragraphs_from_rows(rows: List[List[dict[str, Any]]]) -> List[Tuple[float, str]]:
    """(段落顶边 y, 多行正文)，行间用换行。"""
    if not rows:
        return []
    heights = [float(w["height"]) for row in rows for w in row]
    med_h = statistics.median(heights) if heights else 12.0
    gap_para = max(med_h * 2.2, 14.0)

    paras: List[Tuple[float, str]] = []
    buf_rows: List[List[dict[str, Any]]] = []
    prev_bottom: float | None = None

    def flush() -> None:
        nonlocal buf_rows
        if not buf_rows:
            return
        ymin = min(float(w["top"]) for row in buf_rows for w in row)
        lines: List[str] = []
        for row in buf_rows:
            row = sorted(row, key=lambda w: w["left"])
            lines.append(_join_row_words_with_gaps(row))
        text = "\n".join(lines).strip()
        if text:
            paras.append((ymin, text))
        buf_rows = []

    for row in rows:
        row_top = min(float(w["top"]) for w in row)
        row_bottom = max(float(w["bottom"]) for w in row)
        if prev_bottom is not None and buf_rows and row_top - prev_bottom > gap_para:
            flush()
        buf_rows.append(row)
        prev_bottom = row_bottom
    flush()
    return paras


def append_editable_page_from_jpeg_bytes(
    doc: Document,
    jpeg_bytes: bytes,
    tesseract_exe: Path,
    *,
    width_in: float,
    page_caption: str | None = None,
    table_title_prefix: str | None = None,
    table_psm: int = 11,
    prose_psm: int = 6,
    bbox_pad: int = 12,
    mask_pad_rel: float = 0.038,
    word_overlap_frac: float = 0.20,
    ocr_min_conf: int = 35,
    borderless: bool = True,
    dual_borderless: bool = True,
    auto_upscale: bool = True,
    upscale_scales: Sequence[float] = (1.0, 1.5, 2.0),
) -> None:
    """
    将单页 JPEG 转为「可编辑段落 + 原生 Word 表格」按纵向顺序写入文档（不嵌入整页图）。
    表格区域由 img2table 得到；正文由 Tesseract 词级 OCR，并剔除落在表格 bbox（加 padding）内的词。
    """
    from PIL import Image  # type: ignore  # noqa: PLC0415

    work_bytes, tables = prepare_jpeg_for_editable_page(
        jpeg_bytes,
        tesseract_exe,
        table_psm=int(table_psm),
        dual_borderless=bool(dual_borderless),
        borderless=bool(borderless),
        auto_upscale=bool(auto_upscale),
        upscale_scales=tuple(float(x) for x in upscale_scales),
    )
    im = Image.open(io.BytesIO(work_bytes)).convert("RGB")
    img_w, img_h = im.size

    if page_caption:
        hp = doc.add_paragraph()
        hr = hp.add_run(page_caption)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.name = "Microsoft YaHei"
        hr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    pad_px = max(float(bbox_pad), 0.007 * float(min(img_w, img_h)))
    boxes = _table_mask_boxes(tables, img_w, img_h, pad_px=pad_px, pad_rel=float(mask_pad_rel))

    words_all = _pytesseract_words(im, tesseract_exe, min_conf=int(ocr_min_conf), psm=int(prose_psm))
    kept: List[dict[str, Any]] = []
    for w in words_all:
        if _word_in_table_masks(w, boxes, overlap_frac=float(word_overlap_frac)):
            continue
        kept.append(w)

    hs = [float(x["height"]) for x in kept]
    med_h = statistics.median(hs) if hs else 12.0
    y_tol = max(med_h * 0.55, 8.0)
    rows = _cluster_rows_from_words(kept, y_tol)
    prose_blocks = _paragraphs_from_rows(rows)

    events: List[Tuple[float, int, str, str | None]] = []
    dropped_leaked_prose = False
    for ymin, txt in prose_blocks:
        if _looks_like_leaked_table_prose(txt):
            dropped_leaked_prose = True
            continue
        events.append((ymin, 0, txt, None))
    for ti, ext in enumerate(tables, start=1):
        b = getattr(ext, "bbox", None)
        if b is None:
            continue
        try:
            html = _raw_html_table(ext)
        except Exception:
            continue
        prefix = table_title_prefix if table_title_prefix is not None else page_caption
        cap = f"{prefix} · 表 {ti}" if prefix else f"表 {ti}"
        events.append((float(b.y1), 1, html, cap))

    events.sort(key=lambda e: (e[0], e[1]))

    if not events:
        note = doc.add_paragraph()
        nr = note.add_run("（本页未识别到可排版内容，请检查图像质量、语言包与 PSM 设置。）")
        nr.font.size = Pt(9)
        nr.font.name = "Microsoft YaHei"
        nr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        return

    for _, kind, payload, cap in events:
        if kind == 0:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(payload)
            r.font.size = Pt(9)
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        else:
            append_html_table_to_doc(doc, payload, title=cap, width_in=float(width_in))

    if dropped_leaked_prose:
        np = doc.add_paragraph()
        nr = np.add_run(
            "（本页已略去部分疑似误并入正文的密集数字块，多为表格掩膜未完全覆盖所致；"
            "若表格仍与图像不一致，可尝试提高截图分辨率或调整 --table-psm / --prose-psm。）"
        )
        nr.font.size = Pt(8)
        nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def editable_docx_append_pages_with_headings(
    doc: Document,
    pages: Sequence[Tuple[str, bytes]],
    tesseract_exe: Path,
    *,
    width_in: float,
    table_psm: int = 11,
    prose_psm: int = 6,
    dual_borderless: bool = True,
    mask_pad_rel: float = 0.038,
    auto_upscale: bool = True,
    upscale_scales: Sequence[float] = (1.0, 1.5, 2.0),
) -> None:
    """多页：每页二级标题 + append_editable_page_from_jpeg_bytes（页间分页）。"""
    for i, (label, image_bytes) in enumerate(pages):
        if i > 0:
            doc.add_page_break()
        idx = i + 1
        h = doc.add_heading(f"{idx}. {label}", level=2)
        h.paragraph_format.keep_with_next = True
        for r in h.runs:
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        append_editable_page_from_jpeg_bytes(
            doc,
            image_bytes,
            tesseract_exe,
            width_in=float(width_in),
            page_caption=None,
            table_title_prefix=label,
            table_psm=int(table_psm),
            prose_psm=int(prose_psm),
            dual_borderless=bool(dual_borderless),
            mask_pad_rel=float(mask_pad_rel),
            auto_upscale=bool(auto_upscale),
            upscale_scales=tuple(float(x) for x in upscale_scales),
        )


def build_editable_docx_from_image_paths(
    image_paths: Sequence[Path],
    out_path: Path,
    tesseract_exe: Path | None,
    *,
    landscape: bool = True,
    table_psm: int = 11,
    prose_psm: int = 6,
    dual_borderless: bool = True,
    mask_pad_rel: float = 0.038,
    auto_upscale: bool = True,
    doc_title: str | None = None,
    doc_subtitle: str | None = None,
) -> Path:
    """
    通用入口：多张图片 → 单个可编辑 Word（段落 + 表格）。返回绝对路径。
    """
    exe = _resolve_tesseract_exe(tesseract_exe)
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for sec in doc.sections:
        if landscape:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = Mm(297), Mm(210)
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width, sec.page_height = Mm(210), Mm(297)
        sec.left_margin = sec.right_margin = Mm(12)
        sec.top_margin = sec.bottom_margin = Mm(10)

    if doc_title:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(doc_title)
        tr.bold = True
        tr.font.size = Pt(15)
        tr.font.name = "Microsoft YaHei"
        tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if doc_subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(3)
        sr = sp.add_run(doc_subtitle)
        sr.font.size = Pt(8.5)
        sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = doc.add_paragraph()
    nr = note.add_run(
        "以下为从图像自动提取的可编辑正文与表格（img2table + Tesseract），未嵌入整页图。"
        "复杂表头、无边框表或扫描质量差时可能漏识或错格，请对照原图核对。"
    )
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    section = doc.sections[0]
    cw_mm = float(section.page_width.mm - section.left_margin.mm - section.right_margin.mm)
    width_in = max(4.0, (cw_mm / 25.4) * 0.94)

    pages: List[Tuple[str, bytes]] = []
    for ip in image_paths:
        pages.append((ip.name, Path(ip).read_bytes()))

    editable_docx_append_pages_with_headings(
        doc,
        pages,
        exe,
        width_in=width_in,
        table_psm=int(table_psm),
        prose_psm=int(prose_psm),
        dual_borderless=bool(dual_borderless),
        mask_pad_rel=float(mask_pad_rel),
        auto_upscale=bool(auto_upscale),
    )

    out_final = Path(out_path).expanduser().resolve()
    out_final.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_final))
    if not out_final.is_file():
        raise RuntimeError(f"Word 保存失败：未找到文件 {out_final}")
    return out_final


def build_docx_from_images(
    image_paths: Sequence[Path],
    out_path: Path,
    tesseract_exe: Path | None,
    *,
    landscape: bool = True,
    borderless: bool = True,
) -> None:
    exe = _resolve_tesseract_exe(tesseract_exe)
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    if landscape:
        for sec in doc.sections:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = Mm(297), Mm(210)
            sec.left_margin = sec.right_margin = Mm(12)
            sec.top_margin = sec.bottom_margin = Mm(10)

    intro = doc.add_paragraph()
    ir = intro.add_run(
        "以下为从图片中自动识别的表格（img2table + Tesseract）。"
        "复杂表头、无边框表或扫描质量差时可能漏表或错格，请对照原图核对。"
    )
    ir.font.size = Pt(9)
    ir.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    content_w_mm = float(doc.sections[0].page_width.mm - doc.sections[0].left_margin.mm - doc.sections[0].right_margin.mm)
    width_in = max(4.0, (content_w_mm / 25.4) * 0.95)

    for ip in image_paths:
        p = doc.add_paragraph()
        pr = p.add_run(f"来源图片：{ip.name}")
        pr.bold = True
        pr.font.size = Pt(10.5)

        try:
            if borderless:
                tables = extract_tables_from_bytes_best(
                    Path(ip).read_bytes(),
                    exe,
                    psm=11,
                )
            else:
                tables = extract_tables_from_image_path(ip, exe, borderless=False)
        except Exception as e:
            doc.add_paragraph().add_run(f"（本图提取失败：{e}）")
            continue

        if not tables:
            doc.add_paragraph().add_run("（未检测到表格区域）")
            continue

        for ti, et in enumerate(tables, start=1):
            try:
                html = _raw_html_table(et)
                append_html_table_to_doc(doc, html, title=f"{ip.name} · 表 {ti}", width_in=width_in)
            except Exception as e:
                doc.add_paragraph().add_run(f"表 {ti} 写入失败：{e}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    p = argparse.ArgumentParser(
        description="图片 → Word：仅表格（默认）或整页可编辑排版（--editable）。通用脚本，不限定某一文档来源。"
    )
    p.add_argument("--image", type=Path, action="append", required=True, help="输入 PNG/JPG，可重复")
    p.add_argument("-o", "--output", type=Path, required=True, help="输出 .docx")
    p.add_argument("--tesseract", type=Path, default=None)
    p.add_argument("--no-borderless", action="store_true", help="仅检测有框线表格（与「仅表格」模式配合）")
    p.add_argument(
        "--editable",
        action="store_true",
        help="整页可编辑（段落+表格掩膜+多尺度放大择优）；默认仅导出检测到的表格",
    )
    p.add_argument(
        "--portrait",
        action="store_true",
        help="与 --editable 同用：纵向 A4（默认横版）",
    )
    p.add_argument(
        "--no-auto-upscale",
        action="store_true",
        help="与 --editable 同用：关闭 1.0×/1.5×/2.0× 放大重试（更快，难表可能更差）",
    )
    p.add_argument("--table-psm", type=int, default=11, help="img2table 内 Tesseract PSM，默认 11")
    p.add_argument("--prose-psm", type=int, default=6, help="与 --editable 同用：正文词级 OCR 的 PSM，默认 6")
    p.add_argument(
        "--no-dual-borderless",
        action="store_true",
        help="与 --editable 同用：不比较 borderless 开/关两种模式",
    )
    p.add_argument(
        "--mask-pad-rel",
        type=float,
        default=0.038,
        help="与 --editable 同用：表格 bbox 相对膨胀比例，默认 0.038",
    )
    args = p.parse_args()
    imgs = [Path(x).expanduser().resolve() for x in (args.image or [])]
    exe0 = _resolve_tesseract_exe(args.tesseract)
    print(f"[extract_tables_to_docx] 使用 Tesseract：{exe0}", file=sys.stderr, flush=True)
    try:
        if args.editable:
            build_editable_docx_from_image_paths(
                imgs,
                args.output,
                args.tesseract,
                landscape=not bool(args.portrait),
                table_psm=int(args.table_psm),
                prose_psm=int(args.prose_psm),
                dual_borderless=not bool(args.no_dual_borderless),
                mask_pad_rel=float(args.mask_pad_rel),
                auto_upscale=not bool(args.no_auto_upscale),
            )
        else:
            build_docx_from_images(
                imgs,
                args.output.resolve(),
                args.tesseract,
                borderless=not args.no_borderless,
            )
    except ImportError as e:
        print(
            "缺少依赖，请执行：pip install img2table opencv-contrib-python-headless pandas beautifulsoup4 pytesseract",
            file=sys.stderr,
        )
        print(e, file=sys.stderr)
        return 1
    except Exception as e:
        print(f"失败：{e}", file=sys.stderr)
        return 1
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
