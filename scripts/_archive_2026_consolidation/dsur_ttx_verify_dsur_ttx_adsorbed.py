# -*- coding: utf-8 -*-
"""自检：核对《吸附破伤风疫苗 DSUR#1》输出文档关键替换点。"""
import docx

OUT = r"review_materials/远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"
doc = docx.Document(OUT)

paras = [p.text for p in doc.paragraphs]
cells = []
for tbl in doc.tables:
    for row in tbl.rows:
        for c in row.cells:
            cells.append(c.text)
full_text = "\n".join(paras) + "\n".join(cells)


def check(label, present, absent=None):
    ok = True
    msgs = []
    for s in present:
        if s not in full_text:
            msgs.append(f"  缺失预期内容: {s}")
            ok = False
    for s in (absent or []):
        if s in full_text:
            msgs.append(f"  仍存在应删除内容: {s}")
            ok = False
    print(("PASS" if ok else "FAIL"), label)
    for m in msgs:
        print(m)
    return ok


results = []
results.append(check("产品名=吸附破伤风疫苗", ["吸附破伤风疫苗"], ["重组破伤风疫苗（大肠埃希菌）", "重组带状疱疹疫苗"]))
results.append(check("报告周期 20250708-20260707", ["20250708", "20260707", "2025年07月08日", "2026年07月07日"]))
results.append(check("报告日期 2026-08-07", ["2026年08月07日"]))
results.append(check("方案编号 YDSWX（TVAX-018-3WT）-001（Ⅰ）", ["YDSWX（TVAX-018-3WT）-001（Ⅰ）"]))
results.append(check("I期60例/III期2200例/删除II期", ["60例", "2200例", "第0、1、12月"], ["480例", "II期临床试验：采用随机"]))
results.append(check("审批号占位=待补充", ["待补充"], ["2024LP01829"]))
results.append(check("CTR暂未登记", ["暂未登记"], ["CTR20250784"]))
results.append(check("IB V0.1 2024-10-28", ["V0.1版，2024年10月28日"]))
results.append(check("留观30分钟", ["留观30分钟"], ["留观60分钟"]))
results.append(check("免疫程序第0天1剂", ["受试者于第0天接种1剂"]))
results.append(check("文献3条+PubMed73篇", ["73篇", "[10]", "[11]", "[12]", "Guan Q", "Xie Z", "Rabadi T"]))
results.append(check("破伤风类毒素效价≥40IU", ["40IU", "破伤风类毒素"]))
results.append(check("无进行中非临床研究（模板原句删除）", [], ["未开展相关的非临床研究"]))

print()
print("=" * 40)
print("总计:", sum(results), "/", len(results), "项通过")
print("=" * 40)
