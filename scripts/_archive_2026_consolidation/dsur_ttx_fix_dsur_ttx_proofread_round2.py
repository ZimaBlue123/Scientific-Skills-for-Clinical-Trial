# -*- coding: utf-8 -*-
"""吸附破伤风疫苗 DSUR#1 校对自检修订（round 2，按用户确认范围执行）。

用户确认范围：
- 问题3（检索周期）→ 采用"报告周期内检索"，改 P122 口径与 P141 一致。
- 问题4（文献编号）→ 用户已自行修改，不动。
- 问题20（附件2 药学/非临床留空）→ 不改。
- 其余按校对建议修改。

用法：
    python fix_dsur_ttx_proofread_round2.py          # dry-run：打印每个替换命中数
    python fix_dsur_ttx_proofread_round2.py --apply  # 正式执行并保存
"""
import sys

from docx import Document

SRC = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"

# 通用子串替换（跨 run，幂等）
REPLACEMENTS = [
    # 1. P9 公司全称补「有限公司」
    ("远大赛威信生命科学（南京）、远大赛威信生命科学（杭州）有限公司",
     "远大赛威信生命科学（南京）有限公司、远大赛威信生命科学（杭州）有限公司"),
    # 2. P195 芽胞→芽孢
    ("芽胞杆菌", "芽孢杆菌"),
    # 3. P10 预充针→预灌封注射器
    ("预充针剂型", "预灌封注射器剂型"),
    # 4. P281 18周岁→18岁
    ("18周岁及以上", "18岁及以上"),
    # 5. P122 检索周期改为报告周期内
    ("检索周期：本次为首次撰写DSUR，避免漏掉以往的安全性文献，拟对DLP前获得的与吸附破伤风疫苗相关且有安全性意义的文献进行总结。",
     "检索周期：2025年07月08日至2026年07月07日（本次DSUR报告周期），对该期间内获得的与吸附破伤风疫苗相关且有安全性意义的文献进行总结。"),
    # 6. T3 可预防性 及其→极其
    ("及其罕见", "极其罕见"),
    # 7. T4 风险机制 Guilian→Guillain、Barre→Barré
    ("Guilian和Barre", "Guillain和Barré"),
    # 8. P254 参考文献 Guillain–Barre´→Guillain-Barré
    ("Guillain–Barre´ syndrome", "Guillain-Barré syndrome"),
    # 9. P256 卷期 6(25)→25(6)
    ("6(25):743-746", "25(6):743-746"),
    # 10. P256 王传林等人→王传林，等
    ("王传林等人", "王传林，等"),
    # 11. 排除标准 糖尿病→伴并发症和/或血糖控制欠佳的糖尿病（T4可预防性、T5总结表）
    ("、糖尿病、", "、伴并发症和/或血糖控制欠佳的糖尿病、"),
    # 12. Down氏综合症→唐氏综合征（3处）
    ("Down氏综合症", "唐氏综合征"),
    # 13. 镰刀细胞贫血→镰状细胞贫血（3处）
    ("镰刀细胞贫血", "镰状细胞贫血"),
    # 14. 格林巴利综合症→格林巴利综合征（2处）
    ("格林巴利综合症", "格林巴利综合征"),
    # 15. 疫苗语境 应慎用药物→应慎用本品（T3风险1临床）
    ("应慎用药物", "应慎用本品"),
    # 16. 疫苗语境 用药过程中和用药后→接种过程中和接种后（T3/T4可预防性，2处）
    ("用药过程中和用药后", "接种过程中和接种后"),
    # 17. 做出对症治疗→给予对症治疗（T3可预防性）
    ("做出对症治疗", "给予对症治疗"),
    # 18. WHO认证→WHO确认（P195）
    ("WHO认证", "WHO确认"),
    # 19. 孕妇及新生儿破伤风→孕产妇及新生儿破伤风（P195）
    ("孕妇及新生儿破伤风", "孕产妇及新生儿破伤风"),
]

# 附件5 说明段（唯一，以"本报告期内"开头）——「严重不良事件」→「严重不良反应」
ATT5_OLD = ("本报告期内，吸附破伤风疫苗无已完成的临床试验，无正在进行的临床试验，"
            "尚未报告严重不良事件，以下表格不适用。")
ATT5_NEW = ("本报告期内，吸附破伤风疫苗无已完成的临床试验，无正在进行的临床试验，"
            "尚未报告严重不良反应，以下表格不适用。")

# 附件6 与 附件R1 说明段文字完全相同，需按出现顺序区分：
#   第1个 = 附件6（SAE，不改）；第2个 = 附件R1（SAR，改）
ATT_SAME = ("自DIBD至本次DLP，吸附破伤风疫苗无已完成的临床试验，无正在进行的临床试验，"
            "尚未报告严重不良事件，以下表格不适用。")

# 附件R2 死亡列表脚注：双星号+「指主要SAR」→ 三星号+「指死亡」
ATT_R2_OLD = "** 指“主要”SAR"
ATT_R2_NEW = "*** 指死亡"


def replace_across_runs(paragraph, old, new):
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return False
    end = idx + len(old)

    pos = 0
    start_run = 0
    for ri, r in enumerate(runs):
        r_end = pos + len(r.text)
        if r_end > idx:
            start_run = ri
            break
        pos = r_end

    pos = 0
    end_run = len(runs) - 1
    for ri, r in enumerate(runs):
        r_start = pos
        r_end = pos + len(r.text)
        if r_start < end:
            end_run = ri
        pos = r_end
        if pos >= end:
            break

    if start_run == end_run:
        runs[start_run].text = runs[start_run].text.replace(old, new, 1)
    else:
        prefix = sum(len(r.text) for r in runs[:start_run])
        start_offset = idx - prefix
        end_prefix = sum(len(r.text) for r in runs[:end_run])
        end_offset = end - end_prefix
        runs[start_run].text = runs[start_run].text[:start_offset] + new
        for r in runs[start_run + 1:end_run]:
            r.text = ""
        runs[end_run].text = runs[end_run].text[end_offset:]
    return True


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    yield p


def main():
    apply = "--apply" in sys.argv
    doc = Document(SRC)

    counts = {old: 0 for old, _ in REPLACEMENTS}
    special = {
        "att5_ae_adr": 0,          # 附件5 严重不良事件→严重不良反应
        "att_r1_ae_adr": 0,        # 附件R1 严重不良事件→严重不良反应（第2个相同句）
        "att_r2_footnote": 0,      # 附件R2 脚注 **指主要SAR→***指死亡
    }

    same_counter = 0

    for p in iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)

        # 通用替换
        for old, new in REPLACEMENTS:
            if old in full and (not apply or replace_across_runs(p, old, new)):
                counts[old] += 1

        # 特殊：附件5（唯一句）
        if ATT5_OLD in full:
            if not apply or replace_across_runs(p, ATT5_OLD, ATT5_NEW):
                special["att5_ae_adr"] += 1

        # 特殊：附件6/附件R1 相同句，第2个改为"严重不良反应"
        if ATT_SAME in full:
            same_counter += 1
            if same_counter == 2:
                if not apply or replace_across_runs(
                    p, "尚未报告严重不良事件", "尚未报告严重不良反应"):
                    special["att_r1_ae_adr"] += 1

        # 特殊：附件R2 脚注（精确整段匹配，避免误伤附件5的三星号脚注）
        if full.strip() == ATT_R2_OLD:
            if not apply or replace_across_runs(p, ATT_R2_OLD, ATT_R2_NEW):
                special["att_r2_footnote"] += 1

    # 输出报告
    print("=" * 72)
    print("通用替换命中数：")
    for old, new in REPLACEMENTS:
        print(f"  [{counts[old]}] {old[:44]!r}  ->  {new[:44]!r}")
    print("-" * 72)
    print("特殊替换命中数：")
    for k, v in special.items():
        print(f"  [{v}] {k}")
    print("  相同句'自DIBD至本次DLP…'出现总次数：", same_counter)
    print("=" * 72)

    if apply:
        doc.save(SRC)
        print("已保存：", SRC)
    else:
        print("DRY-RUN 模式：未修改文件。确认无误后加 --apply 执行。")


if __name__ == "__main__":
    main()
