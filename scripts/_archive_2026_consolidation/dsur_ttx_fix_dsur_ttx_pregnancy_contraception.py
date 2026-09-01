# -*- coding: utf-8 -*-
"""修订「妊娠和哺乳期妇女」排除标准的避孕措辞。

将模板遗留的「未完全禁欲且使用研究者不认可的避孕措施」改为与本产品
I/III 期方案一致的表述「入组前2周内未采取有效的避孕措施」。

用法：
    python fix_dsur_ttx_pregnancy_contraception.py           # dry-run，仅打印命中情况
    python fix_dsur_ttx_pregnancy_contraception.py --apply   # 正式落盘
"""

import sys
from docx import Document

DOCX = "review_materials/远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"

SUB_REPLACEMENTS = [
    ("入组前2周未完全禁欲且使用研究者不认可的避孕措施",
     "入组前2周内未采取有效的避孕措施"),
]


def replace_across_runs(paragraph, old, new):
    """跨 run 替换子串，仅重写 old 覆盖的起始/结束 run，保留两侧 run 格式。"""
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return False
    end = idx + len(old)

    pos = 0
    start_run = 0
    for ri, r in enumerate(runs):
        r_end = pos + len(r.text)
        if r_end > idx:
            start_run = ri
            break
        pos = r_end

    pos = 0
    end_run = len(runs) - 1
    for ri, r in enumerate(runs):
        r_start = pos
        r_end = pos + len(r.text)
        if r_start < end:
            end_run = ri
        pos = r_end
        if pos >= end:
            break

    if start_run == end_run:
        runs[start_run].text = runs[start_run].text.replace(old, new, 1)
    else:
        prefix = sum(len(r.text) for r in runs[:start_run])
        start_offset = idx - prefix
        end_prefix = sum(len(r.text) for r in runs[:end_run])
        end_offset = end - end_prefix
        runs[start_run].text = runs[start_run].text[:start_offset] + new
        for r in runs[start_run + 1:end_run]:
            r.text = ""
        runs[end_run].text = runs[end_run].text[end_offset:]
    return True


def main():
    apply = "--apply" in sys.argv
    doc = Document(DOCX)

    total_hits = 0
    for pi, para in enumerate(doc.paragraphs):
        full = "".join(r.text for r in para.runs)
        for old, new in SUB_REPLACEMENTS:
            if old in full:
                total_hits += 1
                print(f"[命中] 段落[{pi}]")
                print(f"  旧：...{old}...")
                print(f"  新：...{new}...")
                if apply:
                    replace_across_runs(para, old, new)

    if not apply:
        print(f"\n[dry-run] 共命中 {total_hits} 处，未落盘。确认后加 --apply 执行。")
        return

    doc.save(DOCX)
    print(f"\n[已落盘] 共替换 {total_hits} 处，保存至：{DOCX}")

    # 二次验证：残留检查
    doc2 = Document(DOCX)
    for keyword in ["禁欲", "未完全禁欲", "研究者不认可"]:
        count = sum(p.text.count(keyword) for p in doc2.paragraphs)
        print(f"[验证] 「{keyword}」残留计数 = {count}")
    new_count = sum(p.text.count("入组前2周内未采取有效的避孕措施") for p in doc2.paragraphs)
    print(f"[验证] 「入组前2周内未采取有效的避孕措施」计数 = {new_count}")


if __name__ == "__main__":
    main()
