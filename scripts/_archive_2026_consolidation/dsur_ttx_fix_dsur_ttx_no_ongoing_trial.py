# -*- coding: utf-8 -*-
"""修复吸附破伤风疫苗 DSUR 交付稿：报告周期内无正在进行的临床试验（去除套用原模板的"正在开展I期"表述）。

用法：
    python fix_dsur_ttx_no_ongoing_trial.py          # dry-run：打印将执行的替换/删除/清空
    python fix_dsur_ttx_no_ongoing_trial.py --apply  # 正式执行并保存
"""

import sys

from docx import Document

SRC = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"

# 全局子串替换（按最长优先原则，替换后统一为"无正在进行的临床试验"）
SUB_REPLACEMENTS = [
    # 变体1：P172/P178/P213/P251/P256/P263/P284/P293/P228（"开展的I期"）
    ("正在国内开展的I期临床试验尚未有首例入组", "无正在进行的临床试验"),
    # 变体2：P220/P243（"开展1项I期" + 逗号）
    ("正在国内开展1项I期临床试验，尚未有首例受试者入组", "无正在进行的临床试验"),
    # 变体3：P14（"开展1项I期" + 括号方案编号）
    ("正在国内开展1项I期临床试验（方案编号：YDSWX（TVAX-018-3WT）-001（Ⅰ）），尚未有首例受试者入组；", "无正在进行的临床试验；"),
    # 变体4：P77
    ("1项正在国内开展的临床试验为：吸附破伤风疫苗I期临床试验（方案编号：YDSWX（TVAX-018-3WT）-001（Ⅰ））。", "无正在进行的临床试验。"),
    # 变体5：P91
    ("本报告期内，正在进行1项在国内开展的临床试验：", "本报告期内，无正在进行的临床试验。"),
    # 变体6：P111
    ("本报告期内，1项正在国内开展的临床试验：", "本报告期内，无正在进行的临床试验。"),
]

# 需要删除的整段（试验描述段，紧跟在 P91/P111 之后）
DELETE_PARAS = [
    "一项“评价吸附破伤风疫苗在18岁及以上健康人群中接种的安全性及初步免疫原性的随机、盲法、阳性对照的I期临床试验”（方案编号：YDSWX（TVAX-018-3WT）-001（Ⅰ））",
    "吸附破伤风疫苗I期临床试验（方案编号：YDSWX（TVAX-018-3WT）-001（Ⅰ））采用单中心、随机、盲法、阳性对照设计，评价吸附破伤风疫苗在18岁及以上健康受试者中的安全性和免疫原性。截至2026年07月07日，尚未有首例受试者入组，无关于临床安全性的重要发现。",
]


def replace_across_runs(paragraph, old, new):
    """跨 run 子串替换，只改动 old 覆盖的 run，其余 run 保持原样。"""
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return False
    end = idx + len(old)

    pos = 0
    start_run = 0
    for ri, r in enumerate(runs):
        r_end = pos + len(r.text)
        if r_end > idx:
            start_run = ri
            break
        pos = r_end

    pos = 0
    end_run = len(runs) - 1
    for ri, r in enumerate(runs):
        r_start = pos
        r_end = pos + len(r.text)
        if r_start < end:
            end_run = ri
        pos = r_end
        if pos >= end:
            break

    if start_run == end_run:
        runs[start_run].text = runs[start_run].text.replace(old, new, 1)
    else:
        prefix = sum(len(r.text) for r in runs[:start_run])
        start_offset = idx - prefix
        end_prefix = sum(len(r.text) for r in runs[:end_run])
        end_offset = end - end_prefix
        runs[start_run].text = runs[start_run].text[:start_offset] + new
        for r in runs[start_run + 1:end_run]:
            r.text = ""
        runs[end_run].text = runs[end_run].text[end_offset:]
    return True


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    yield p


def clear_cell(cell, text="/"):
    """清空单元格为指定文本（保留第一个 run，清空其余）。"""
    paras = cell.paragraphs
    if not paras:
        return
    first_p = paras[0]
    # 清空其余段落
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # 第一个段落：全部文本写入第一个 run（或新建 run）
    runs = first_p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        first_p.add_run(text)


def main():
    apply = "--apply" in sys.argv
    doc = Document(SRC)

    sub_hits = []   # (old, new, 段落文本片段)
    del_hits = []   # 删除的段落文本
    tbl_hits = []   # 清空的表格行文本

    # 1) 子串替换
    for p in iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        for old, new in SUB_REPLACEMENTS:
            if old in full:
                sub_hits.append((old, new, full))
                if apply:
                    replace_across_runs(p, old, new)

    # 2) 删除整段（仅正文段落，不含表格单元格）
    for p in list(doc.paragraphs):
        t = p.text.strip()
        for target in DELETE_PARAS:
            if t == target.strip():
                del_hits.append(t)
                if apply:
                    p._element.getparent().remove(p._element)

    # 3) 清空"进行中的临床试验"表格数据行（表头含 FVFP 与 计划入组人数）
    for table in doc.tables:
        header = table.rows[0]
        htext = " ".join(c.text for c in header.cells)
        if "FVFP" in htext and "计划入组人数" in htext:
            for row in table.rows[1:]:
                row_text = " | ".join(c.text for c in row.cells)
                tbl_hits.append(row_text)
                if apply:
                    for cell in row.cells:
                        clear_cell(cell, "/")

    # 输出报告
    print("=== 子串替换（%d 处） ===" % len(sub_hits))
    for old, new, full in sub_hits:
        print(f"  {old!r}  ->  {new!r}")
        print(f"      原文片段: {full[:60]}…")

    print("=== 删除整段（%d 段） ===" % len(del_hits))
    for t in del_hits:
        print(f"  DEL: {t[:60]}…")

    print("=== 清空表格数据行（%d 行） ===" % len(tbl_hits))
    for t in tbl_hits:
        print(f"  CLEAR: {t[:80]}…")

    if apply:
        doc.save(SRC)
        print("保存完成:", SRC)
    else:
        print("DRY-RUN 模式：未修改文件。确认无误后加 --apply 正式执行。")


if __name__ == "__main__":
    main()
