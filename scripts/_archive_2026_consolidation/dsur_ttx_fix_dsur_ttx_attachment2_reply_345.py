# -*- coding: utf-8 -*-
"""修订 DSUR#1 附件2 表格第 3、4、5 条临床意见的「状态」列回复。

将第 3（安全性）、4（免疫原性）、5（风险管理计划及知情同意书）条的状态列，
由「复述 CDE 意见」改为「基于方案既有设计 + 将按CDE建议做的调整」。

用法：
    python fix_dsur_ttx_attachment2_reply_345.py           # dry-run
    python fix_dsur_ttx_attachment2_reply_345.py --apply   # 落盘
"""

import copy
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = "review_materials/远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"
TABLE_INDEX = 5

# 第 3、4、5 条对应表格行（0 基）：Row3=第3条、Row4=第4条、Row5=第5条
# 每条为 (row_index, 新的状态列文本)
NEW_STATUS = [
    (
        3,
        "将按CDE建议执行。Ⅲ期方案已设全程免疫人群序贯接种——首批入组60例受试者完成每剂接种后7天安全性数据评估、确认未达暂停标准（≥1例无法排除与疫苗相关的死亡/危及生命SAE、≥15%受试者出现≥3级不良事件等）后方继续后续剂次接种；Ⅰ期方案已设接种前及接种后第4天血常规、血生化、尿常规实验室指标检测。在此基础上，将按CDE建议对全程免疫人群首批60例增加每剂免后4天实验室检测指标检测并纳入试验暂停标准；30天内主动安全性随访尽可能采用面对面方式。",
    ),
    (
        4,
        "将按CDE建议执行。Ⅲ期方案已定义全分析集（FAS，遵循ITT原则）、符合方案集（PPS）及安全性分析集（SS），免疫原性统计分析分别基于FAS和PPS集进行。将按CDE建议以PPS作为免疫原性分析主要分析集、同时对FAS进行分析，并预设两者结果不一致时的原因分析。",
    ),
    (
        5,
        "将按CDE建议执行。根据非临床研究结果，风险管理计划及临床试验方案已关注肌肉注射部位刺激性（疼痛、红肿等）并制定相应风险控制措施；目前Ⅰ期、Ⅲ期已分别制定知情同意书，将按CDE建议进一步针对加强免疫、全程免疫两部分人群分别制定。",
    ),
]


def make_run(text, rPr_template):
    r = OxmlElement('w:r')
    if rPr_template is not None:
        r.append(copy.deepcopy(rPr_template))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def set_cell_text(tc, text):
    """清空单元格并写入单段文本，保留原字体/段落格式。"""
    paras = tc.findall(qn('w:p'))
    if not paras:
        return
    first_p = paras[0]
    p_template = copy.deepcopy(first_p)
    rPr_template = None
    runs = first_p.findall(qn('w:r'))
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            rPr_template = copy.deepcopy(rPr)
    for p in paras:
        tc.remove(p)
    new_p = copy.deepcopy(p_template)
    for r in list(new_p.findall(qn('w:r'))):
        new_p.remove(r)
    new_p.append(make_run(text, rPr_template))
    tc.append(new_p)


def cell_text_exact(tc):
    out = []
    for p in tc.findall(qn('w:p')):
        segs = []
        for r in p.findall(qn('w:r')):
            t_el = r.find(qn('w:t'))
            if t_el is not None:
                segs.append(t_el.text or '')
        out.append(''.join(segs))
    return '\n'.join(out)


def main():
    apply = "--apply" in sys.argv
    doc = Document(DOCX)
    t = doc.tables[TABLE_INDEX]
    rows = list(t.rows)

    for row_idx, new_text in NEW_STATUS:
        if row_idx >= len(rows):
            print(f"[错误] 行 {row_idx} 超出表格行数 {len(rows)}")
            continue
        tcs = rows[row_idx]._tr.findall(qn('w:tc'))
        tc_status = tcs[4]
        old = cell_text_exact(tc_status)
        req = cell_text_exact(tcs[3])
        print(f"--- Row{row_idx}（第{row_idx-2}条）---")
        print(f"  要求列: {req[:40]}...")
        print(f"  旧状态: {old[:50]}...")
        print(f"  新状态: {new_text[:70]}...")
        if apply:
            set_cell_text(tc_status, new_text)

    if not apply:
        print("\n[dry-run 完成] 未落盘。确认后加 --apply 执行。")
        return

    doc.save(DOCX)
    print(f"\n[已落盘] 保存至：{DOCX}")

    # 二次验证
    doc2 = Document(DOCX)
    t2 = doc2.tables[TABLE_INDEX]
    rows2 = list(t2.rows)
    for row_idx, new_text in NEW_STATUS:
        tcs = rows2[row_idx]._tr.findall(qn('w:tc'))
        cur = cell_text_exact(tcs[4])
        ok = (cur == new_text)
        print(f"[验证] Row{row_idx} 状态列 匹配={'OK' if ok else 'FAIL'}")
    # 残留旧句检查
    full = "\n".join(cell_text_exact(r._tr.findall(qn('w:tc'))[4]) for r in rows2[1:])
    for kw in ["以符合方案集（PPS）作为免疫原性分析主要分析集，同时对全分析集（FAS）进行分析", "30天内的主动安全性随访将尽可能采用面对面随访方式"]:
        print(f"[验证] 旧句「{kw[:20]}…」残留 = {full.count(kw)}")


if __name__ == "__main__":
    main()
