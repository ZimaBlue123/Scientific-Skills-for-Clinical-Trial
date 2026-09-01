# -*- coding: utf-8 -*-
"""
优化《冻干b型流感嗜血杆菌结合疫苗（TVAX-018-2）DSUR#2》。

基于已生成的 DSUR#2 clean 版（克隆自 DSUR#1 并已滚动日期/第2次），进一步：
  1) 执行概要 P13 补重要潜在风险列表（已识别：无；潜在：严重过敏反应、热性惊厥）；
  2) 第13章文献：检索结果 4篇 -> 2篇（真实重新检索新增 2 篇，替换旧 4 篇）；
  3) 附件2 状态列 6 条：由套话改为「将按CDE建议执行 + 已做设计->将做调整」；
  4) 附件7 文献列表：替换旧前 4 条为新增 2 条，删除多余 2 条；
  5) 第18章疾病背景参考文献重编号 [5]-[12] -> [3]-[10]；
  6) 「唯一研发者」->「研发者」；「与简称冻干Hib结合疫苗相关」笔误修正。

依赖：python-docx（default venv 已装）。
运行：C:/Users/Administrator/.workbuddy/binaries/python/envs/default/Scripts/python.exe
用法：python optimize_dsur_tvax018_2.py [--apply]
      （不带 --apply 为 dry-run，仅打印命中对照，不落盘）
"""
import re
import sys
import docx

TPL = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#2_20260808-clean.docx"
OUT = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#2_20260808-优化版.docx"

# ---------------------------------------------------------------------------
# 整段替换（正文段落）
# ---------------------------------------------------------------------------
FULL_MAP = {
    # 执行概要 P13 补风险列表
    "报告期内，远大赛威信研发的冻干Hib结合疫苗暂无临床研究数据，根据本品非临床安全性数据、文献报道、同类疫苗的数据进行总结，无重要的潜在风险。":
    "报告期内，远大赛威信研发的冻干Hib结合疫苗暂无临床研究数据，根据本品非临床安全性数据、文献报道、同类疫苗的数据进行总结。（1）重要的已识别风险：无；（2）重要的潜在风险：①严重过敏反应；②热性惊厥。",

    # 第13章 检索结果 4篇 -> 2篇
    "本次检索未发现引用远大赛威信研发的冻干Hib结合疫苗的文献，获得4篇引用同类产品的有意义的安全性相关文献，现对相关安全性的文献总结如下：":
    "本次检索未发现引用远大赛威信研发的冻干Hib结合疫苗的文献，获得2篇引用同类产品的有意义的安全性相关文献，现对相关安全性的文献总结如下：",

    # 附件7 文献列表：旧第1条 -> 新增文献1
    "潘雪娇,周洋,梁辉,等.b型流感嗜血杆菌疾病和疫苗应用综述[J].中国疫苗和免疫,2025,31(03):376-384.DOI:10.19914/j.CJVI.2025060.":
    "Zhou Y, Tang R, Wei M, et al. Post-marketing safety surveillance of Haemophilus influenzae Type b conjugate vaccine in children aged 3 months to 5 years in Jiangsu, China[J]. Hum Vaccin Immunother. 2026. doi: 10.1080/21645515.2026.2677876.",

    # 附件7 文献列表：旧第2条 -> 新增文献2
    "J. Guo, B. Zheng, T. Chen, S. Guo, M. Liu, and S. Dong, “Prioritizing vaccine introduction to China’s National Immunization Program: a multi-criteria decision analysis.,” BMC Public Health, vol. 24, no. 1, p. 3458, Dec. 2024, doi: 10.1186/s12889-024-21013-7.":
    "Zhang C, Liu X, Han L, et al. Reporting profile of adverse events following DTaP-IPV/Hib pentavalent vaccine: A 13-year real-world surveillance data analysis[J]. Hum Vaccin Immunother. 2025. doi: 10.1080/21645515.2025.2597061.",
}

# 附件7 待删除的旧文献条目（旧第3、4条，已不再被正文引用）
DELETE_PARAS = [
    "M. Ren et al., “Post-Marketing Surveillance of Adverse Events Following Immunization with Haemophilus Influenzae Type b Conjugate Vaccine - China, 2010-2021.,” China CDC Wkly, vol. 6, no. 33, pp. 834–840, Aug. 2024, doi: 10.46234/ccdcw2024.180.",
    "X. Pan, Y. Chen, H. Liang, L. Shen, and X. Qi, “Adverse Events Following Haemophilus influenzae Type b (Hib) Monovalent Vaccines in Zhejiang Province, China, from 2017 to 2023.,” Vaccines (Basel), vol. 13, no. 4, Mar. 2025, doi: 10.3390/vaccines13040349.",
]

# 全局子串替换
SUB = [
    ("唯一研发者", "研发者"),
    ("与简称冻干Hib结合疫苗相关", "与冻干Hib结合疫苗相关"),
]

# 第18章疾病背景参考文献重编号：[5]-[12] -> [3]-[10]
RENUM = {12: 10, 11: 9, 10: 8, 9: 7, 8: 6, 7: 5, 6: 4, 5: 3}

# ---------------------------------------------------------------------------
# 表3（第13章文献表）新增 2 行内容（5 列：编号/文献类型/名称/摘要/安全性信息）
# ---------------------------------------------------------------------------
NEW_LIT_ROWS = [
    [
        "1",
        "上市后安全性监测",
        "Post-marketing safety surveillance of Haemophilus influenzae Type b conjugate vaccine in children aged 3 months to 5 years in Jiangsu, China[1]",
        "Zhou Y等开展了一项前瞻性、观察性、多中心上市后安全性监测研究，评价某b型流感嗜血杆菌（Hib）结合疫苗在江苏3月龄～5岁儿童中的安全性（Hum Vaccin Immunother, 2026）。研究于2020年6月至2021年12月在江苏省20个区县开展，主动监测组入组2714例、被动监测组6347例。主动监测显示总体不良事件发生率为6.7%（181/2714），其中不良反应6.3%（170/2714），以轻中度为主（1～2级6.0%，3级0.3%）；被动监测显示不良反应发生率为0.33%（21/6347），以发热（0.25%）最常见。主动监测与被动监测均未报告疫苗相关严重不良反应。结论：该Hib结合疫苗在3月龄～5岁儿童中安全性良好，未识别到安全性相关风险。",
        "局部不良反应：注射部位红肿、硬结；全身不良反应：发热等；未报告疫苗相关严重不良反应。",
    ],
    [
        "2",
        "上市后监测研究",
        "Reporting profile of adverse events following DTaP-IPV/Hib pentavalent vaccine: A 13-year real-world surveillance data analysis[2]",
        "Zhang C等对湖州市2012～2024年DTaP-IPV/Hib五联疫苗13年疑似预防接种异常反应（AEFI）监测数据进行了回顾性描述性分析（Hum Vaccin Immunother, 2025）。总体AEFI报告发生率为19.67/10万，呈显著上升趋势（P<0.001），2018年达峰值47.01/10万；非严重AEFI占99.4%（19.56/10万），严重AEFI罕见（0.6%，0.11/10万）。最常见不良反应为注射部位红肿（13.48/10万）、硬结（8.06/10万）和发热（7.64/10万），过敏性皮疹发生率为0.53/10万。共报告3例严重AEFI，包括1例过敏性紫癜及2例偶合症（血小板减少性紫癜、脓毒症）。结论：绝大多数反应轻微、自限，严重事件罕见，为该疫苗的安全性提供初步证据。",
        "局部不良反应：红肿、硬结；全身不良反应：发热；其他十分罕见不良反应：过敏性皮疹（0.53/10万）、过敏性紫癜等。",
    ],
]

# ---------------------------------------------------------------------------
# 附件2 状态列：要求细则前缀 -> 新状态文本（6 条）
# ---------------------------------------------------------------------------
ATT2_STATUS = {
    "1.试验设计：":
    "将按CDE建议执行。Ⅰ/Ⅱ期临床试验方案已规定采用一次揭盲法，受试者完成基础免疫接种后6个月安全性观察、经盲态审核并认定可靠无误后锁定数据库进行揭盲；Ⅲ期临床试验方案已规定采用一次揭盲法，受试者完成基础免疫接种后6个月安全性观察、经盲态审核后锁定数据库揭盲。将按CDE建议，Ⅰ/Ⅱ期在完成基础免疫6个月随访后由非盲团队进行阶段性分析（对研究者和受试者保持盲态），至少完成加强免疫30天随访后再揭盲；明确Ⅲ期临床试验启动时机，至少获得Ⅰ/Ⅱ期临床试验所有受试者阶段性分析结果后再启动Ⅲ期入组；Ⅲ期在获得加强免疫后30天的安全性和免疫原性数据后揭盲、统计分析，并继续按照随访计划收集全程免疫后6个月内的SAE。",
    "2.研究人群：":
    "将按CDE建议执行。Ⅰ/Ⅱ期、Ⅲ期临床试验方案已设置相应的入选/排除标准。将按CDE建议排除有Hib感染史者；对于小月龄婴幼儿，将排除早产儿、低出生体重儿等特殊人群。",
    "3.安全性方面：":
    "将按CDE建议执行。Ⅰ/Ⅱ期临床试验方案已规定第一阶段2岁及以上受试者于接种前和接种后第4天采集血样、尿样进行血生化、血常规和尿常规检测。将按CDE建议尽可能在1岁以上婴幼儿中开展实验室指标检测，并增加凝血指标检测；增加征集期内主动随访频次；参考同类品种完善征集性AE（如增加嗜睡等条目）；核实Ⅲ期临床试验的暂停、终止标准相关表述。",
    "4.免疫原性方面：":
    "将按CDE建议执行。Ⅲ期临床试验方案已定义主要研究假设（2-5月龄免前阴性人群基础免疫后30天抗体阳转率非劣于对照组、2-5月龄全人群基础免疫后30天抗体长期保护率非劣于对照组，非劣效界值均为-10%）。将按CDE建议自早期临床试验对加强免疫后的免疫原性进行评价，并适当收紧免疫原性采血时间窗；将Ⅲ期临床试验主要研究假设修订为不同年龄层试验疫苗组基础免疫后血清抗PRP抗体浓度>1.0 μg/mL、>0.15 μg/mL的受试者百分比非劣效于阳性对照组（非劣效界值为-5%），组间基础免疫后抗体GMC非劣效比较以及加强免疫后抗体长期保护率及GMC非劣效比较作为次要假设；根据修订后的研究假设分别计算不同年龄层的样本量，同时兼顾安全性评价要求；增加检测抗TTc抗体，并分析其对安全有效性的影响。",
    "6.风险管理计划和知情同意书：":
    "将按CDE建议执行。风险管理计划已参照同类产品梳理重要的潜在风险（严重过敏反应、热性惊厥）并制定相应风险控制措施。将按CDE建议参照同类产品进一步完善重要的潜在风险并制定相应措施；按照《药物警戒质量管理规范》及时更新临床试验期间的风险控制计划；知情同意书列出可能影响受试者安全或本品安全有效性评价的关键排除标准。",
}


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def set_para_text(p, text):
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def set_cell_text(cell, text):
    paras = cell.paragraphs
    first = paras[0]
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    set_para_text(first, text)


def renumber(text):
    """[5]-[12] -> [3]-[10]，用占位符避免连锁替换。"""
    for n in sorted(RENUM.keys(), reverse=True):
        text = text.replace("[%d]" % n, "\x00%d\x00" % n)
    for n, m in RENUM.items():
        text = text.replace("\x00%d\x00" % n, "[%d]" % m)
    return text


def apply_sub(text):
    for old, new in SUB:
        text = text.replace(old, new)
    return text


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main(apply=False):
    doc = docx.Document(TPL)
    hits = 0

    # 1. 正文段落：FULL_MAP / 子串替换 / 重编号
    for p in doc.paragraphs:
        full = p.text
        if full in FULL_MAP:
            print("[FULL] %r  =>  %r" % (full[:28], FULL_MAP[full][:28]))
            hits += 1
            if apply:
                set_para_text(p, FULL_MAP[full])
            continue
        new = full
        new = apply_sub(new)
        new = renumber(new)
        if new != full:
            print("[EDIT] %r  =>  %r" % (full[:40], new[:40]))
            hits += 1
            if apply:
                set_para_text(p, new)

    # 2. 附件7 删除旧文献条目
    for p in doc.paragraphs:
        if p.text in DELETE_PARAS:
            print("[DEL ] %r" % p.text[:40])
            hits += 1
            if apply:
                p._element.getparent().remove(p._element)

    # 3. 表3（第13章文献表）：旧 4 行 -> 新 2 行
    lit_tbl = None
    for tbl in doc.tables:
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if hdr and hdr[0] == "编号" and "文献类型" in hdr:
            lit_tbl = tbl
            break
    if lit_tbl is not None:
        rows = lit_tbl.rows
        print("[TBL ] 文献表 rows=%d（头+%d 数据行）" % (len(rows), len(rows) - 1))
        for ri, cells in enumerate(NEW_LIT_ROWS):
            if ri + 1 < len(rows):
                tr = rows[ri + 1]
                for ci, val in enumerate(cells):
                    if ci < len(tr.cells):
                        print("  [TBL] 行%d 列%d  =>  %r" % (ri + 1, ci, val[:30]))
                        hits += 1
                        if apply:
                            set_cell_text(tr.cells[ci], val)
        # 删除多余数据行（旧第3、4行）
        extra = len(rows) - 1 - len(NEW_LIT_ROWS)
        for _ in range(extra):
            if len(lit_tbl.rows) > 1 + len(NEW_LIT_ROWS):
                dead = lit_tbl.rows[-1]
                print("  [TBL] 删除多余行  =>  %r" % dead.cells[0].text[:20])
                hits += 1
                if apply:
                    dead._element.getparent().remove(dead._element)
    else:
        print("[WARN] 未找到文献表（编号/文献类型）")

    # 4. 附件2（T5）状态列：按“要求细则”前缀匹配
    att2 = None
    for tbl in doc.tables:
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if "日期" in hdr and "状态" in hdr:
            att2 = tbl
            break
    if att2 is not None:
        print("[TBL ] 附件2 rows=%d" % len(att2.rows))
        for row in att2.rows[1:]:
            cells = row.cells
            if len(cells) < 5:
                continue
            detail = cells[3].text.strip()
            for prefix, newstatus in ATT2_STATUS.items():
                if detail.startswith(prefix):
                    print("  [TBL] 附件2 %r  =>  %r" % (detail[:20], newstatus[:40]))
                    hits += 1
                    if apply:
                        set_cell_text(cells[4], newstatus)
                    break
    else:
        print("[WARN] 未找到附件2表（日期/状态）")

    if apply:
        doc.save(OUT)
        print("SAVED ->", OUT)
    else:
        print("\n[dry-run] 命中 %d 处；未落盘，如需写入请追加 --apply" % hits)


if __name__ == "__main__":
    main(apply="--apply" in sys.argv)
