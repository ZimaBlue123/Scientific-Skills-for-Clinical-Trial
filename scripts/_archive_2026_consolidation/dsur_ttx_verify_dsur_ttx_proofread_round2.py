# -*- coding: utf-8 -*-
"""验证 DSUR 校对修订 round2 的结果。"""
from docx import Document

SRC = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    yield p


def main():
    doc = Document(SRC)
    full_texts = ["".join(r.text for r in p.runs) for p in iter_paragraphs(doc)]
    blob = "\n".join(full_texts)

    print("=" * 60)
    print("【应清零项】")
    should_be_zero = [
        "及其罕见", "Guilian", "Down氏综合症", "镰刀细胞贫血",
        "格林巴利综合症", "综合症", "芽胞", "预充针", "18周岁",
        "王传林等人", "6(25)", "Guillain–Barre´", "、糖尿病、",
    ]
    for kw in should_be_zero:
        cnt = blob.count(kw)
        mark = "OK" if cnt == 0 else "!! 残留"
        print(f"  [{cnt}] {mark} {kw!r}")

    print("-" * 60)
    print("【AE/ADR 三处】")
    for kw in ["尚未报告严重不良事件", "尚未报告严重不良反应"]:
        print(f"  [{blob.count(kw)}] {kw!r}")

    print("-" * 60)
    print("【应生效的新表述】")
    should_present = [
        "唐氏综合征", "镰状细胞贫血", "伴并发症和/或血糖控制欠佳的糖尿病",
        "Guillain和Barré", "极其罕见", "预灌封注射器", "25(6)", "王传林，等",
        "*** 指死亡", "检索周期：2025年07月08日至2026年07月07日",
        "远大赛威信生命科学（南京）有限公司、远大赛威信生命科学（杭州）有限公司",
    ]
    for kw in should_present:
        cnt = blob.count(kw)
        mark = "OK" if cnt > 0 else "!! 缺失"
        print(f"  [{cnt}] {mark} {kw!r}")

    print("=" * 60)


if __name__ == "__main__":
    main()
