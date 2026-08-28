import os
import re

import docx
import openpyxl

os.chdir("review_materials")
wb = openpyxl.load_workbook("清洁残留计算所需信息.xlsx", data_only=True)
ws = wb.active
print("Excel Header:", [cell.value for cell in ws[1]])


# Read text from docx to find "剂量"
def extract_text(filename):
    doc = docx.Document(filename)
    text = []
    for p in doc.paragraphs:
        if (
            "剂量" in p.text
            or "用量" in p.text
            or "mg" in p.text
            or "治疗" in p.text
            or "用法" in p.text
        ):
            text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if (
                    "剂量" in cell.text
                    or "用量" in cell.text
                    or "mg" in cell.text
                    or "治疗" in cell.text
                    or "用法" in cell.text
                ):
                    text.append(cell.text)
    return text


for file in os.listdir("."):
    if file.endswith(".docx") and any(proj in file for proj in ["009", "009B", "006"]):
        print(f"\n========== {file} ==========")
        lines = extract_text(file)
        snippets = set()
        for l in lines:
            l = l.replace("\n", " ").strip()
            # Try to match patterns like "每次xxx mg", "每日xxx mg"
            if (
                (
                    re.search(r"\d+\s*(mg|g|ml)", l, re.IGNORECASE)
                    or "最大" in l
                    or "最低" in l
                    or "用法用量" in l
                )
                and len(l) > 5
                and len(l) < 500
            ):
                snippets.add(l)

        for s in list(snippets)[:10]:
            print(" -", s)
