# -*- coding: utf-8 -*-
"""
《冻干b型流感嗜血杆菌结合疫苗（TVAX-018-2）DSUR#1》修订版后续优化：

1) 19.2 本周期重要风险 段落后新增「重要潜在风险控制措施」表（3 列 × 3 行，
   头=风险分类/风险概述/措施概要；2 行=严重过敏反应、热性惊厥），
   「措施概要」列含「排除标准」「医学干预措施」两个子节；
2) 附件2 重要监管要求汇总表「状态」列详细重写（5 条 CDE 临床意见），
   详细到方案/RMP/ICF 实际条款；第4行同时覆盖「4.免疫原性方面」「5.其他」两条意见。

依据：
  - 1.8.3 风险管理计划 V1.0（2024年03月05日）
  - 1.3.4.1 1-2期临床试验方案 与 1.3.4.1 3期临床试验方案（入选/排除/急救条款）
  - 1.3.4.2 知情同意书样稿（关键排除标准条目）
  - 通知书-临床意见 2024年07月11日 CDE 临床意见 6 条

运行：C:/Users/Administrator/.workbuddy/binaries/python/envs/default/Scripts/python.exe
用法：python update_dsur_tvax018_1_risk_controls.py [--apply]
"""
import sys
import docx
from copy import deepcopy

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

TPL = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#1_20250808-修订版.docx"
OUT = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#1_20250808-修订版.docx"

# ---------------------------------------------------------------------------
# 19.2 表格内容（风险分类 / 风险概述 / 措施概要）
# ---------------------------------------------------------------------------
# 措施概要 含两个子节：「排除标准」「医学干预措施」。
# 在 docx 表格单元格里用换行符 \n 表示子节小标题（加粗）与列表项。
TABLE_19_2_HEADERS = ["风险分类", "风险概述", "措施概要"]

TABLE_19_2_ROWS = [
    {
        "风险分类": "重要潜在风险",
        "风险概述": "严重过敏反应（MedDRA PT：10002198）",
        "措施概要": [
            ("排除标准：", True),
            ("对试验疫苗的任何成分、任何疫苗或药物有严重过敏史，如过敏性休克、过敏性喉头水肿、过敏性紫癜、局部过敏坏死反应（Arthus反应）；", False),
            ("已被诊断为患有可能干扰研究进行或完成的严重的疾病或先天畸形（包括但不限于：遗传性过敏体质、格林巴利综合症等）；", False),
            ("接种前3天内，曾服用退热镇痛或抗过敏药物。", False),
            ("医学干预措施：", True),
            ("接种现场应具备适当的急救治疗措施及配备肾上腺素等药物以备接种后偶有发生严重过敏反应时使用（依据Ⅰ-Ⅱ期/Ⅲ期临床试验方案「提醒和防范」），接种疫苗后应留下观察30分钟；", False),
            ("现场医务人员需掌握严重过敏反应及心肺复苏处置流程，准备基础急救设备及药物（肾上腺素、糖皮质激素、抗组胺药H1受体拮抗剂、心电监护、除颤仪、简易呼吸器、氧气）；", False),
            ("一旦出现过敏性休克等危及受试者生命的情况，现场有资质的医师应立即按照抢救流程组织原地抢救，第一时间给予肾上腺素和监护、吸氧等措施，同时寻求急诊医疗中心支援，并呼叫120转往急救医疗中心继续观察治疗。", False),
        ],
    },
    {
        "风险分类": "重要潜在风险",
        "风险概述": "热性惊厥（MedDRA PT：10010908）",
        "措施概要": [
            ("排除标准：", True),
            ("本人或亲生父母有惊厥、癫痫和精神病史；", False),
            ("＜12月龄者：婴儿异常产程出生（难产、器械助产）或有窒息、神经系统损害史，现患病理性黄疸、肛周脓肿、严重湿疹；", False),
            ("后续剂次：接种后48小时内出现高热（腋窝温度≥39.5°C），伴或不伴惊厥；接种试验用疫苗后7天之内出现脑病、癫痫发作。", False),
            ("医学干预措施：", True),
            ("接种前详细询问受试者及家族成员的惊厥、癫痫及精神疾病史，接种前详细查体并告知受试者及家属发生热性惊厥的可能；", False),
            ("接种后按规定留观，注意监测生命体征，必要时可延长留观时间；", False),
            ("告知受试者及家属若出现发热及惊厥表现时需注意防护，避免因惊厥发作而受伤；热性惊厥患者的神志多于惊厥后很快恢复，无需过多治疗，但若惊厥呈持续状态或频繁发生则需尽快入院进一步诊治，避免造成延误或危及受试者的生命安全。", False),
        ],
    },
]


# ---------------------------------------------------------------------------
# 附件2 状态列详细文本（5 条 CDE 临床意见，按"要求"列前缀匹配）
# ---------------------------------------------------------------------------
ATT2_STATUS = {
    "1.试验设计：":
    "将按CDE建议执行。Ⅰ/Ⅱ期临床试验方案已规定采用一次揭盲法，受试者完成基础免疫接种后6个月安全性观察、经盲态审核并认定可靠无误后锁定数据库进行揭盲；Ⅲ期临床试验方案已规定采用一次揭盲法，受试者完成基础免疫接种后6个月安全性观察、经盲态审核后锁定数据库揭盲。将按CDE建议，Ⅰ/Ⅱ期在完成基础免疫6个月随访后由非盲团队进行阶段性分析（对研究者和受试者保持盲态），至少完成加强免疫30天随访后再揭盲；明确Ⅲ期临床试验启动时机，至少获得Ⅰ/Ⅱ期临床试验所有受试者阶段性分析结果后再启动Ⅲ期入组；Ⅲ期在获得加强免疫后30天的安全性和免疫原性数据后揭盲、统计分析，并继续按照随访计划收集全程免疫后6个月内的SAE。",

    "2.研究人群：":
    "将按CDE建议执行。Ⅰ/Ⅱ期、Ⅲ期临床试验方案已设置相应的入选/排除标准（参见方案「入选标准」「首剂排除标准」「后续剂次推迟标准」「后续剂次排除标准」章节）。将按CDE建议排除有Hib感染史者；对于小月龄婴幼儿，将排除早产儿、低出生体重儿等特殊人群，相关排除标准将在Ⅰ/Ⅱ期/Ⅲ期临床试验方案修订时补充明确。",

    "3.安全性方面：":
    "将按CDE建议执行。Ⅰ/Ⅱ期临床试验方案已规定第一阶段2岁及以上受试者于接种前和接种后第4天采集血样、尿样进行血生化、血常规和尿常规检测；并已规定对1-5岁、6-11月龄、2-5月龄受试者每个年龄层均按至少2天入组、第一天入组人数不超过5例的分批入组策略，在上一年龄层完成基础免疫接种后7天安全性观察且评估未触发暂停/终止标准后，方可入组下一年龄层受试者；对所有受试者每剂接种后8-30天内再进行1次主动随访，随访形式可为电话/微信，并做好相应记录。将按CDE建议尽可能在1岁以上婴幼儿中开展实验室指标检测，并增加凝血指标检测；增加征集期内主动随访频次；参考同类品种完善征集性AE（如增加嗜睡等条目）；核实Ⅲ期临床试验的暂停、终止标准相关表述。",

    "4.免疫原性方面：":  # 该行同时含「5.其他：增加检测抗TTc抗体」的意见
    "将按CDE建议执行。Ⅲ期临床试验方案已定义主要研究假设（2-5月龄免前阴性人群基础免疫后30天抗体阳转率非劣于对照组、2-5月龄全人群基础免疫后30天抗体长期保护率非劣于对照组，非劣效界值均为-10%）。将按CDE建议自早期临床试验对加强免疫后的免疫原性进行评价，并适当收紧免疫原性采血时间窗；将Ⅲ期临床试验主要研究假设修订为不同年龄层试验疫苗组基础免疫后血清抗PRP抗体浓度>1.0 μg/mL、>0.15 μg/mL的受试者百分比非劣效于阳性对照组（非劣效界值为-5%），组间基础免后抗体GMC非劣效比较以及加强免后抗体长期保护率及GMC非劣效比较作为次要假设；根据修订后的研究假设分别计算不同年龄层的样本量，同时兼顾安全性评价要求。同时，将按CDE建议增加检测抗TTc抗体，并分析其对安全有效性的影响。",

    "6.风险管理计划和知情同意书：":
    "将按CDE建议执行。风险管理计划已参照同类产品梳理重要的潜在风险（①严重过敏反应；②热性惊厥）并制定相应风险控制措施，详见《冻干b型流感嗜血杆菌结合疫苗风险管理计划》（版本号V1.0，2024年03月05日）。将按CDE建议参照同类产品进一步完善重要的潜在风险并制定相应措施；按照《药物警戒质量管理规范》（GVP）及时更新临床试验期间的风险控制计划；Ⅰ/Ⅱ期/Ⅲ期临床试验知情同意书已列出可能影响受试者安全或本品安全有效性评价的关键排除标准，包括对试验疫苗的任何成分或任何疫苗/药物有严重过敏史、本人或亲生父母有惊厥/癫痫/精神病史、严重疾病或先天畸形、免疫功能异常相关用药史等，并将在获得任何新的安全性信息时按方案要求更新知情同意书并由受试者/法定监护人重新签署。",
}


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def make_cell_paragraphs(cell, segments):
    """
    在 cell 中写入多个段落（清空原有段落）。
    segments = [(text, is_bold), ...]
    """
    paras = cell.paragraphs
    first = paras[0]
    # 移除多余段落
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # 重置 first 段落
    for r in list(first.runs):
        r._element.getparent().remove(r._element)

    cur_p = first
    for i, (text, is_bold) in enumerate(segments):
        if i > 0:
            # 后续段落需新建
            from docx.oxml import OxmlElement
            new_p = OxmlElement("w:p")
            cur_p._element.addnext(new_p)
            # cur_p 切换到新建段落（通过 docx 包装）
            # 简化：从 cell 中找最后一个段落
            cur_p = cell.paragraphs[-1]
        run = cur_p.add_run(text)
        if is_bold:
            run.bold = True


def set_simple_cell(cell, text):
    """把 cell 内容替换为单段文本。"""
    paras = cell.paragraphs
    first = paras[0]
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    runs = first.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        first.add_run(text)


def make_new_table_after(target_p, headers, rows):
    """
    在 target_p 段落后插入新表（与现有附件2表同款样式 af5）。
    headers = ["风险分类", "风险概述", "措施概要"]
    rows = [{"风险分类": ..., "风险概述": ..., "措施概要": [(text,is_bold),...]}, ...]
    """
    from docx.oxml import OxmlElement

    n_cols = len(headers)
    n_rows = 1 + len(rows)  # 头 + 数据

    # 直接构建 XML，更可控
    tbl_xml = OxmlElement("w:tbl")

    # tblPr（参照现有 af5 表格）
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(W + "val", "af5")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(W + "w", "0")
    tblW.set(W + "type", "auto")
    tblPr.append(tblW)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(W + "val", "04A0")
    tblPr.append(tblLook)
    tbl_xml.append(tblPr)

    # tblGrid
    tblGrid = OxmlElement("w:tblGrid")
    # 列宽：第一列 1500、第二列 2200、第三列 5000（约 6:8:18）
    col_widths = [1500, 2200, 5400]
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(W + "w", str(w))
        tblGrid.append(gc)
    tbl_xml.append(tblGrid)

    # 构造行的辅助函数
    def make_row(cells_content, header=False):
        tr = OxmlElement("w:tr")
        # 行高自适应
        trPr = OxmlElement("w:trPr")
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)
        tr.append(trPr)
        for i, content in enumerate(cells_content):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(W + "w", str(col_widths[i]))
            tcW.set(W + "type", "dxa")
            tcPr.append(tcW)
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(W + "val", "center")
            tcPr.append(vAlign)
            tc.append(tcPr)

            if header:
                # 头单元格：单段文本
                p = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                jc = OxmlElement("w:jc")
                jc.set(W + "val", "center")
                pPr.append(jc)
                p.append(pPr)
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rPr.append(b)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.text = content
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r.append(t)
                p.append(r)
                tc.append(p)
            else:
                # 数据单元格：每段一条（文本, 是否加粗）
                # 兼容两种 content：list-of-tuples（措施概要多段） 或 str（单段）
                segs = content if isinstance(content, list) else [(content, False)]
                for j, (text, is_bold) in enumerate(segs):
                    p = OxmlElement("w:p")
                    pPr = OxmlElement("w:pPr")
                    p.append(pPr)
                    r = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    if is_bold:
                        b = OxmlElement("w:b")
                        rPr.append(b)
                    r.append(rPr)
                    t = OxmlElement("w:t")
                    t.text = text
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    r.append(t)
                    p.append(r)
                    tc.append(p)
            tr.append(tc)
        return tr

    # 头行
    tbl_xml.append(make_row(headers, header=True))
    # 数据行
    for row in rows:
        tbl_xml.append(make_row([
            row["风险分类"],
            row["风险概述"],
            row["措施概要"],
        ], header=False))

    # 在 target_p 段落后插入
    target_p._element.addnext(tbl_xml)
    return tbl_xml


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main(apply=False):
    doc = docx.Document(TPL)
    body = doc.element.body
    hits = 0

    # ---------- 1) 19.2 段落后插入风险控制措施表 ----------
    # 19.2 content paragraph = doc.paragraphs[208]（首次 DSUR 段落实文）
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "本期报告为本品首次DSUR" in p.text:
            target_idx = i
            break
    if target_idx is None:
        print("[ERR] 未找到 19.2 content paragraph")
        return
    target_p = doc.paragraphs[target_idx]
    print("[19.2] 目标段落 idx=%d: %s" % (target_idx, target_p.text[:60]))

    if apply:
        make_new_table_after(target_p, TABLE_19_2_HEADERS, TABLE_19_2_ROWS)
        print("[19.2] 已插入 %dx%d 表（头+2 行数据）" % (len(TABLE_19_2_HEADERS), len(TABLE_19_2_ROWS)))
        hits += 1

    # ---------- 2) 附件2 状态列详细重写 ----------
    att2 = None
    for tbl in doc.tables:
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        # 附件2表头 = 日期 | 机构/国家 | 要求 | 要求 | 状态
        if "日期" in hdr and "状态" in hdr and any("机构" in c for c in hdr):
            att2 = tbl
            break
    if att2 is None:
        print("[ERR] 未找到附件2表")
        return
    print("[TBL ] 附件2 rows=%d" % len(att2.rows))

    for row in att2.rows[1:]:
        cells = row.cells
        if len(cells) < 5:
            continue
        detail = cells[3].text.strip()  # col3 = 要求详细列
        for prefix, newstatus in ATT2_STATUS.items():
            if detail.startswith(prefix):
                print("  [ATT2] %r  =>  %r..." % (prefix, newstatus[:40]))
                hits += 1
                if apply:
                    set_simple_cell(cells[4], newstatus)
                break

    if apply:
        doc.save(OUT)
        print("SAVED ->", OUT)
    else:
        print("\n[dry-run] 命中 %d 处；未落盘，如需写入请追加 --apply" % hits)


if __name__ == "__main__":
    main(apply="--apply" in sys.argv)