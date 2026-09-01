# -*- coding: utf-8 -*-
"""
生成《冻干b型流感嗜血杆菌结合疫苗（TVAX-018-2）DSUR#2》DOCX。

策略：克隆 DSUR#1 DOCX，保留版式/样式/表格/页眉页脚，仅做：
  1) 去除全部 8 条批注（commentRangeStart/End、commentReference 及 comments 部件）；
  2) 整段替换（FULL_MAP）——修正 DSUR#1 批注指出的笔误，更新"第2次DSUR"实质内容；
  3) 全局子串替换（SUB）——日期滚动、第1次→第2次。

依赖：python-docx（default venv 已装）。
运行：C:/Users/Administrator/.workbuddy/binaries/python/envs/default/Scripts/python.exe
"""
import docx
from docx.oxml.ns import qn

TPL = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#1_20250808-clean（有笔误）.docx"
OUT = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#2_20260808-clean.docx"

# ---------------------------------------------------------------------------
# 1. 整段替换（old 全文 -> new 全文）
# ---------------------------------------------------------------------------
FULL_MAP = {
    # 标题页报告日期：2025-08-08 -> 2026-08-08
    "报告日期：2025年08月08日":
    "报告日期：2026年08月08日",

    # 批注1（笔误）：无重要风险 -> 无重要的潜在风险
    "报告期内，远大赛威信研发的冻干Hib结合疫苗暂无临床研究数据，根据本品非临床安全性数据、文献报道、同类疫苗的数据进行总结，无重要风险。":
    "报告期内，远大赛威信研发的冻干Hib结合疫苗暂无临床研究数据，根据本品非临床安全性数据、文献报道、同类疫苗的数据进行总结，无重要的潜在风险。",

    # 第13章文献 检索周期：首次撰写 -> 第2次撰写
    "检索周期：本次为首次撰写DSUR，避免漏掉以往的安全性文献，拟对DLP前获得的与冻干Hib结合疫苗相关且有安全性意义的文献进行总结。":
    "检索周期：本次为第2次撰写DSUR，拟对本报告期内（2025年07月11日至2026年07月10日）获得的与冻干Hib结合疫苗相关且有安全性意义的文献进行总结。",

    # 第14章 其他DSUR：补记第1次DSUR
    "本报告期内，未撰写或递交过其它与冻干Hib结合疫苗相关的DSUR。":
    "已撰写并递交冻干Hib结合疫苗第1次DSUR（报告日期：2025年08月08日）。除第1次DSUR外，本报告期内未撰写或递交过其它与冻干Hib结合疫苗相关的DSUR。",

    # 批注3（笔误）：第17章日期 08月09日 -> 08月08日
    "本DSUR自数据锁定点（2025年07月10日）后至2025年08月09日，无新的潜在重要安全性发现。":
    "本DSUR自数据锁定点（2026年07月10日）后至2026年08月08日，无新的潜在重要安全性发现。",

    # 批注4（笔误）：18.1.2 潜在的风险 重新表述
    "根据本品非临床安全性数据、文献报道和同类疫苗临床研究数据等安全性信息，本品暂无重要已识别和重要的潜在风险。":
    "根据本品非临床安全性数据、文献报道和同类疫苗临床研究数据等安全性信息，整理了现阶段重要的潜在风险，对风险进行评估，并将在临床试验中进行密切监测。本品现阶段重要的潜在风险包括：①严重过敏反应；②热性惊厥。",

    # 19.1 上周期重要风险
    "不适用，本次DSUR为第1次撰写。":
    "第1次DSUR（报告日期：2025年08月08日）识别的重要潜在风险为：①严重过敏反应；②热性惊厥。",

    # 19.2 本周期重要风险
    "本期报告为本品首次DSUR，本报告期内，远大赛威信研发的冻干Hib结合疫苗尚未在任何国家/地区获批上市，未开展相关的非临床研究。无正在进行的临床试验，自DIBD起截至DLP，无严重不良事件发生，无可疑且非预期严重不良反应（SUSAR）发生，未提示新的安全性风险。从目前本品非临床安全性数据、文献报道和同类疫苗临床研究数据的安全性信息可知，预计接种本品可能会出现常见的疫苗接种后一般反应，无影响风险获益的重要风险。":
    "本期报告为本品第2次DSUR，本报告期内，远大赛威信研发的冻干Hib结合疫苗尚未在任何国家/地区获批上市，未开展相关的非临床研究。无正在进行的临床试验，自DIBD起截至DLP，无严重不良事件发生，无可疑且非预期严重不良反应（SUSAR）发生，未提示新的安全性风险。本报告期内未发现新的重要风险，本品现阶段重要的潜在风险维持为：①严重过敏反应；②热性惊厥。",
}

# ---------------------------------------------------------------------------
# 2. 全局子串替换（按顺序应用；仅作用于未命中 FULL_MAP 的段落/单元格）
# ---------------------------------------------------------------------------
SUB = [
    # 报告周期：先整体替换报告周期字符串，再替换残留的 DLP 日期
    ("2024年07月11日至2025年07月10日", "2025年07月11日至2026年07月10日"),
    # 第1次DSUR -> 第2次DSUR（标题页 + 执行概要；"首次获得临床批准"不受影响）
    ("第1次DSUR", "第2次DSUR"),
    # 残留 DLP 日期（截至/数据锁定点/表格表头）-> 2026-07-10
    ("2025年07月10日", "2026年07月10日"),
]


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def apply_sub_to_text(text):
    for old, new in SUB:
        text = text.replace(old, new)
    return text


def set_para_text(p, text):
    """将段落文本替换为 text，保留首 run 格式。"""
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def set_cell_text(cell, text):
    """将单元格文本替换为 text，保留首段首 run 格式。"""
    paras = cell.paragraphs
    first = paras[0]
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    set_para_text(first, text)


def strip_comments_from_root(root):
    """删除 root 中的批注引用元素（commentRangeStart/End 及含 commentReference 的 run）。"""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
        for el in root.findall(".//{%s}%s" % (W, tag.split(":")[1])):
            el.getparent().remove(el)
    for ref in root.findall(".//{%s}commentReference" % W):
        run = ref.getparent()
        if run is not None:
            run.getparent().remove(run)


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
doc = docx.Document(TPL)

# 处理正文段落
for p in doc.paragraphs:
    full = p.text
    if full in FULL_MAP:
        set_para_text(p, FULL_MAP[full])
        continue
    new = apply_sub_to_text(full)
    if new != full:
        set_para_text(p, new)

# 处理表格单元格
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            ctext = cell.text
            if ctext in FULL_MAP:
                set_cell_text(cell, FULL_MAP[ctext])
                continue
            new = apply_sub_to_text(ctext)
            if new != ctext:
                set_cell_text(cell, new)

# 去除批注：正文 + 表格 + 页眉页脚
strip_comments_from_root(doc.element.body)
for tbl in doc.tables:
    strip_comments_from_root(tbl._tbl)
for sec in doc.sections:
    for hf in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
        strip_comments_from_root(hf._element)

# 删除 comments 相关部件关系（comments/commentsExtended/commentsIds/commentsExtensible），使批注彻底消失
part = doc.part
removed = []
for rid in list(part.rels):
    rel = part.rels[rid]
    if "comment" in rel.reltype.lower():
        removed.append(rid)
for rid in removed:
    try:
        part.drop_rel(rid)
    except Exception:
        pass

doc.save(OUT)
print("SAVED ->", OUT)
print("removed comment rels:", removed)
