# -*- coding: utf-8 -*-
"""DSUR#2 第13章文献安全性汇总段落优化 + 表3单位修正。

背景：
- DSUR#2 第13章"根据上述文献资料将同类产品提示的安全性信息汇总如下"段
  （局部/全身/其他十分罕见三清单）原为照搬 DSUR#1 旧内容，与本次新增的两篇
  文献（Zhou 2026、Zhang 2025）不符。现依据两篇新文献实际数据重写。
- 表3（doc.tables[2]）文献2（Zhang 2025）行发生率单位"/10万"有误，
  源文献分母为 10,000 剂，应为"/万剂"。

修改点：
1. 替换"局部不良反应"清单段落（疼痛、红斑、肿胀、硬结。）
2. 替换"全身不良反应"清单段落（发热、哭闹、疲劳……嗜睡。）
3. 替换"其他十分罕见不良反应"清单段落（过敏性皮疹、血管性水肿……流涕。）
4. 表3 文献2行 C3/C4 单元格内 "10万" -> "万剂"。
"""
from docx import Document

TPL = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#2_20260808-优化版.docx"

NEW_LOCAL = ("注射部位红斑、肿胀、硬结、疼痛、皮疹、瘙痒，偶见蜂窝织炎；"
             "注射部位红肿和硬结为最常见的局部反应。")

NEW_SYSTEMIC = "发热、腹泻、食欲下降、呕吐；发热为最常见的全身反应。"

NEW_RARE = ("过敏性皮疹、过敏性紫癜、荨麻疹、斑丘疹、麻疹猩红热样皮疹、无菌性脓肿等；"
            "其中1例过敏性紫癜为疫苗相关性严重不良反应，另报告的血小板减少性紫癜和脓毒症各1例均为偶合症。"
            "两篇文献均未识别到新的疫苗相关严重不良反应。")

OLD_LOCAL = "疼痛、红斑、肿胀、硬结。"
OLD_SYSTEMIC = "发热、哭闹、疲劳、肌肉痛、皮疹、呕吐、腹泻、瘙痒、食欲不振、虚弱、嗜睡。"
OLD_RARE = ("过敏性皮疹、血管性水肿、过敏性休克、喉头水肿、热性惊厥、非热性惊厥、吉兰-巴雷综合征、"
            "急性播散性脑脊髓炎、血小板减少性紫癜、过敏性紫癜、腹痛、咳嗽、喉部红肿、恶心、出汗、流涕。")


def set_paragraph_text(p, text):
    """将段落文本整体替换为新文本，保留首段 run 的格式，删除其余 run。"""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def main():
    doc = Document(TPL)

    # 1) 替换三个清单段落
    targets = {OLD_LOCAL: NEW_LOCAL, OLD_SYSTEMIC: NEW_SYSTEMIC, OLD_RARE: NEW_RARE}
    hit = {k: False for k in targets}
    for p in doc.paragraphs:
        t = p.text
        for old, new in targets.items():
            if t == old:
                set_paragraph_text(p, new)
                hit[old] = True

    # 2) 修正表3（tables[2]）文献2行 C3/C4 单位
    tbl = doc.tables[2]
    fixed_cells = 0
    for ci in (3, 4):
        cell = tbl.rows[2].cells[ci]
        for cp in cell.paragraphs:
            for r in cp.runs:
                if "10万" in r.text:
                    r.text = r.text.replace("10万", "万剂")
                    fixed_cells += 1

    doc.save(TPL)

    # 3) 校验
    print("替换命中：")
    for old, ok in hit.items():
        print("  [%s] %s" % ("OK" if ok else "MISS", old[:24]))
    print("单位修正单元格 run 数：", fixed_cells)

    doc2 = Document(TPL)
    print("\n=== 校验第13章汇总 ===")
    for p in doc2.paragraphs:
        if p.text.strip() in (NEW_LOCAL, NEW_SYSTEMIC, NEW_RARE):
            print("  ", p.style.name, "|", p.text)
    print("\n=== 校验表3文献2行 ===")
    t2 = doc2.tables[2]
    print("  C3:", t2.rows[2].cells[3].text)
    print("  C4:", t2.rows[2].cells[4].text)


if __name__ == "__main__":
    main()
