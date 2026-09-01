# -*- coding: utf-8 -*-
"""修复吸附破伤风疫苗 DSUR 交付稿的文本问题（run 级精确替换，保留脚注上标格式）。

目前已执行/支持：
1. T3R3 错别字「格利巴利」→「格林巴利」
2. T5R3「风险管理计划V1.0（2024/08/08）」→「风险管理计划V0.1（2024年11月21日）」
3. 前言「唯一研发者」→「研发者」（仿制苗，不宜称唯一研发者）
4. 全文「临床试验批准通知书编号：待补充」→「临床试验批准通知书编号：2025LP01760」

用法：
    python fix_dsur_ttx_final_typos.py          # dry-run：打印 run 结构与将执行的替换
    python fix_dsur_ttx_final_typos.py --apply  # 正式执行替换并保存
"""

import sys

from docx import Document

SRC = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"

REPLACEMENTS = [
    ("格利巴利", "格林巴利"),
    ("风险管理计划V1.0（2024/08/08）", "风险管理计划V0.1（2024年11月21日）"),
    ("唯一研发者", "研发者"),
    ("编号：待补充", "编号：2025LP01760"),
]


def replace_across_runs(paragraph, old, new):
    """跨 run 子串替换，只改动 old 覆盖的 run，其余 run（含上标）保持原样。"""
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return False
    end = idx + len(old)

    # 定位 old 起始所在 run（第一个 r_end > idx 的 run）
    pos = 0
    start_run = 0
    for ri, r in enumerate(runs):
        r_end = pos + len(r.text)
        if r_end > idx:
            start_run = ri
            break
        pos = r_end

    # 定位 old 结束所在 run（最后一个 r_start < end 的 run）
    pos = 0
    end_run = len(runs) - 1
    for ri, r in enumerate(runs):
        r_start = pos
        r_end = pos + len(r.text)
        if r_start < end:
            end_run = ri
        pos = r_end
        if pos >= end:
            break

    if start_run == end_run:
        runs[start_run].text = runs[start_run].text.replace(old, new, 1)
    else:
        prefix = sum(len(r.text) for r in runs[:start_run])
        start_offset = idx - prefix
        end_prefix = sum(len(r.text) for r in runs[:end_run])
        end_offset = end - end_prefix
        runs[start_run].text = runs[start_run].text[:start_offset] + new
        for r in runs[start_run + 1:end_run]:
            r.text = ""
        runs[end_run].text = runs[end_run].text[end_offset:]
    return True


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    yield p


def main():
    apply = "--apply" in sys.argv
    doc = Document(SRC)

    found = []  # (old, run结构描述)
    counts = {old: 0 for old, _ in REPLACEMENTS}

    for p in iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        for old, new in REPLACEMENTS:
            if old in full:
                # 记录 run 结构
                run_desc = [
                    {"text": r.text, "sup": bool(r.font.superscript), "bold": bool(r.bold)}
                    for r in p.runs
                ]
                found.append((old, new, run_desc))
                if apply:
                    if replace_across_runs(p, old, new):
                        counts[old] += 1

    if not found:
        print("未找到任何目标文本，请检查。")
        return

    for old, new, run_desc in found:
        print("=" * 70)
        print(f"目标子串: {old!r}  ->  {new!r}")
        print(f"所在段落 run 数: {len(run_desc)}")
        for i, rd in enumerate(run_desc):
            t = rd["text"]
            shown = t if len(t) <= 40 else t[:20] + "…" + t[-20:]
            flag = ""
            if old in rd["text"]:
                flag = "  <== 含目标子串"
            print(f"  run[{i}] sup={rd['sup']} bold={rd['bold']} len={len(t)}: {shown!r}{flag}")

    if apply:
        doc.save(SRC)
        print("=" * 70)
        for old, cnt in counts.items():
            print(f"已替换 {old!r} -> {cnt} 处")
        print("保存完成:", SRC)
    else:
        print("=" * 70)
        print("DRY-RUN 模式：未修改文件。确认无误后加 --apply 正式执行。")


if __name__ == "__main__":
    main()
