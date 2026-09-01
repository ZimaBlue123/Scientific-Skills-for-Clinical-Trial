# -*- coding: utf-8 -*-
"""
修复《吸附破伤风疫苗 III 期临床试验方案》中的笔误（run 级精确替换，保留格式）。

修复点（均位于 III 期方案，经与 I 期方案正确写法对照确认）：
  1. 正文"吸附破伤风疫苗（大肠埃希菌）" → "吸附破伤风疫苗"
     —— "（大肠埃希菌）"为重组疫苗表达系统残留，吸附破伤风疫苗为类毒素疫苗，与之无关；
  2. "Ⅱ期临床试验研究" → "Ⅲ期临床试验研究"（III 期方案误写为 II 期）；
  3. 正文误嵌引用著录"《疫苗临床试验技术指导原则（征求意见稿）》，NMPA，2024年9月3日》[3]"
     → "《疫苗临床试验技术指导原则（征求意见稿）》[3]"（删除误嵌入的 NMPA/日期及多余右书名号；
     参考文献条目中正确的"…，NMPA，2024年9月3日."结尾句号者不动）；
  4. 表格"制品名称"单元格"吸附破伤风疫苗（大肠埃希菌）" → "吸附破伤风疫苗"。
"""
import docx

F = r"review_materials/远大吸附破伤风疫苗-III期临床试验方案（V0.3，2025年06月10日）clean-updated.docx"

doc = docx.Document(F)

TTX = "（大肠埃希菌）"

# 期望被清空的引用著录 run 序列（仅当紧跟正文误嵌标题 run 之后精确匹配才清空）
REF_TAIL = ["NMPA", "，", "2024", "年", "9", "月", "3", "日", "》"]


def fix_para(p):
    """对段落做 run 级修复，返回是否做了修改。"""
    changed = False
    runs = p.runs
    n = len(runs)
    i = 0
    while i < n:
        r = runs[i]
        t = r.text
        # 1) Ⅱ → Ⅲ（独立 run）
        if t == "Ⅱ":
            r.text = "Ⅲ"
            changed = True
        # 2) 删除"（大肠埃希菌）"前缀 / 独立 run
        elif t == TTX:
            r.text = ""
            changed = True
        elif t.startswith(TTX):
            r.text = t[len(TTX):]
            changed = True
        i += 1
    # 3) 正文误嵌引用著录（仅正文句，结尾为 "》[3]"；参考文献条目结尾为 "." 不处理）
    #    仅在段尾存在 "日》" 的序列时才删除误嵌的 "，NMPA，2024年9月3日》"。
    if "日》" not in p.text:
        return changed
    i = 0
    while i < n:
        r = runs[i]
        if r.text == "疫苗临床试验技术指导原则（征求意见稿）》，":
            r.text = "疫苗临床试验技术指导原则（征求意见稿）》"
            changed = True
            j = i + 1
            for k, exp in enumerate(REF_TAIL):
                idx = j + k
                if idx < n and runs[idx].text == exp:
                    runs[idx].text = ""
                    changed = True
                else:
                    break
            i = j + len(REF_TAIL)
        else:
            i += 1
    return changed


# 处理正文段落
para_hits = 0
for p in doc.paragraphs:
    if TTX in p.text or "Ⅱ期临床试验研究" in p.text or "日》" in p.text:
        if fix_para(p):
            para_hits += 1

# 处理表格单元格
cell_hits = 0
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            if TTX in cell.text:
                for pp in cell.paragraphs:
                    for r in pp.runs:
                        if r.text == TTX:
                            r.text = ""
                            cell_hits += 1
                        elif r.text.startswith(TTX):
                            r.text = r.text[len(TTX):]
                            cell_hits += 1

doc.save(F)
print("段落命中段落数:", para_hits)
print("单元格修复 run 数:", cell_hits)
print("SAVED ->", F)
