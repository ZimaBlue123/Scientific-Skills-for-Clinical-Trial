# -*- coding: utf-8 -*-
"""
TVAX-009 糖尿病亚组免疫原性结果 —— Word 文档生成器（仅 PPS 版本）。

背景：在完整版（PPS+FAS 6 表）基础上，按用户要求仅保留 PPS 分析集，
去掉 FAS 绝对/相对表及相对表中的 FAS 行与「分析集」列，重新编号为 4 张表。

数据源：review_materials/diabetes_immuno_result.json（与完整版一致）。
产出 4 张表：
  1. 抗-HBs 阳转率（PPS）绝对时间点
  2. 抗-HBs 阳转率（PPS）相对时间点
  3. 抗-HBs GMC（PPS）绝对时间点
  4. 抗-HBs GMC（PPS）相对时间点

页面：横向 A4；表格 8pt，表头浅蓝填充，程序/组别列垂直合并。
"""
import json
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = 'E:/Cursor Project/2-Scientific-Skills-for-Clinical_Trial/review_materials'
OUT = BASE + '/TVAX-009-糖尿病亚组免疫原性结果表-PPS-20260903.docx'

with open(BASE + '/diabetes_immuno_result.json', encoding='utf-8') as f:
    DATA = json.load(f)

PPS = DATA['pps']

GROUPS = [
    ('A1', '低剂量组(A1)', '0,1月程序', None),
    ('A2', '高剂量组(A2)', '0,1月程序', None),
    ('B1', '低剂量组(B1)', '0,2月程序', None),
    ('B2', '高剂量组(B2)', '0,2月程序', None),
    ('C1', '阳性对照组(C1)', '0,1,6月程序', '18-59岁'),
    ('C3', '0,1,6月高剂量组(C3)', '0,1,6月程序', '≥60岁'),
    ('C2', '阳性对照组(C2)', '0,1,6月程序', '≥60岁'),
]
PROGRAMS = [
    ('0,1月程序', ['A1', 'A2']),
    ('0,2月程序', ['B1', 'B2']),
    ('0,1,6月程序', ['C1', 'C3', 'C2']),
]
PROG_OF = {code: prog for code, _, prog, _ in GROUPS}
LABEL = {code: lab for code, lab, _, _ in GROUPS}
ANNO = {code: a for code, _, _, a in GROUPS}

PROG_TIMEPOINTS = {
    '0,1月程序': ['M1', 'M2', 'M3', 'M6', 'M7', 'M8'],
    '0,2月程序': ['M1', 'M2', 'M3', 'M4', 'M7', 'M8'],
    '0,1,6月程序': ['M1', 'M2', 'M6', 'M7', 'M8'],
}
ALL_TP = ['M1', 'M2', 'M3', 'M4', 'M6', 'M7', 'M8']
REL_MAP = {
    '全免后1个月': {'0,1月程序': 'M2', '0,2月程序': 'M3', '0,1,6月程序': 'M7'},
    '全免后2个月': {'0,1月程序': 'M3', '0,2月程序': 'M4', '0,1,6月程序': 'M8'},
}

HDR_FILL = 'D9E2F3'
TAG_FILL = 'F2F2F2'


def set_run(run, size, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Arial'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, size, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run(run, size, bold, color)
    return p


def set_cell_fill(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def write_cell(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ''
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run(run, size, bold)
    if fill:
        set_cell_fill(cell, fill)


def set_grid_widths(table, widths_mm):
    tbl = table._tbl
    old = tbl.find(qn('w:tblGrid'))
    if old is not None:
        tbl.remove(old)
    grid = OxmlElement('w:tblGrid')
    for w in widths_mm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(round(w * 56.6929))))
        grid.append(gc)
    tbl.insert(0, grid)


def vmerge(table, col, start_row, end_row):
    head = table.cell(start_row, col)
    tcPr = head._tc.get_or_add_tcPr()
    vm = OxmlElement('w:vMerge')
    vm.set(qn('w:val'), 'restart')
    tcPr.append(vm)
    for r in range(start_row + 1, end_row + 1):
        c = table.cell(r, col)
        c.text = ''
        tcPr2 = c._tc.get_or_add_tcPr()
        tcPr2.append(OxmlElement('w:vMerge'))


def sero_val(data, code, tag, tp):
    rec = data[code][tag][tp]
    if rec['N_阳转'] == 0 and rec['pct'] is None:
        return '—'
    return '%d/%d\n(%.2f%%)' % (rec['n_阳转'], rec['N_阳转'], rec['pct'])


def gmc_val(data, code, tag, tp):
    rec = data[code][tag][tp]
    if rec['N_gmc'] == 0 and rec['gmc'] is None:
        return '—'
    return '%.2f\n(N=%d)' % (rec['gmc'], rec['N_gmc'])


def group_label(code):
    lab = LABEL[code]
    if ANNO[code]:
        lab += '\n(%s)' % ANNO[code]
    return lab


def build_absolute_table(doc, mode, title, note):
    add_para(doc, title, size=11, bold=True, space_before=6, space_after=2)
    n_rows = 1 + 14
    table = doc.add_table(rows=n_rows, cols=10)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_grid_widths(table, [22, 34, 13, 28, 28, 28, 28, 28, 28, 28])

    headers = ['程序', '组别', '糖尿病病史'] + ALL_TP
    for j, h in enumerate(headers):
        write_cell(table.cell(0, j), h, size=8, bold=True, fill=HDR_FILL)

    row = 1
    prog_blocks = []
    for prog, codes in PROGRAMS:
        prog_start = row
        for code in codes:
            for tag in ('有', '无'):
                if row == prog_start:
                    write_cell(table.cell(row, 0), prog, size=8)
                else:
                    table.cell(row, 0).text = ''
                if tag == '有':
                    write_cell(table.cell(row, 1), group_label(code), size=8)
                else:
                    table.cell(row, 1).text = ''
                write_cell(table.cell(row, 2), tag, size=8, fill=TAG_FILL)
                for j, tp in enumerate(ALL_TP):
                    if tp in PROG_TIMEPOINTS[PROG_OF[code]]:
                        val = sero_val(PPS, code, tag, tp) if mode == 'sero' else gmc_val(PPS, code, tag, tp)
                    else:
                        val = '—'
                    write_cell(table.cell(row, 3 + j), val, size=8)
                row += 1
        prog_blocks.append((prog_start, row - 1, prog))

    for start, end, _ in prog_blocks:
        vmerge(table, 0, start, end)
    row = 1
    for _, codes in PROGRAMS:
        for code in codes:
            vmerge(table, 1, row, row + 1)
            row += 2

    add_para(doc, note, size=9, color='333333', space_before=2, space_after=6)


def build_relative_table(doc, mode, title, note):
    add_para(doc, title, size=11, bold=True, space_before=6, space_after=2)
    n_rows = 1 + 14
    table = doc.add_table(rows=n_rows, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_grid_widths(table, [22, 34, 13, 98, 98])

    headers = ['程序', '组别', '糖尿病病史', '全免后1个月', '全免后2个月']
    for j, h in enumerate(headers):
        write_cell(table.cell(0, j), h, size=8, bold=True, fill=HDR_FILL)

    row = 1
    prog_blocks = []
    for prog, codes in PROGRAMS:
        prog_start = row
        for code in codes:
            for tag in ('有', '无'):
                if row == prog_start:
                    write_cell(table.cell(row, 0), prog, size=8)
                else:
                    table.cell(row, 0).text = ''
                if tag == '有':
                    write_cell(table.cell(row, 1), group_label(code), size=8)
                else:
                    table.cell(row, 1).text = ''
                write_cell(table.cell(row, 2), tag, size=8, fill=TAG_FILL)
                for col_idx, tp_label in enumerate(('全免后1个月', '全免后2个月')):
                    m = REL_MAP[tp_label][PROG_OF[code]]
                    val = sero_val(PPS, code, tag, m) if mode == 'sero' else gmc_val(PPS, code, tag, m)
                    write_cell(table.cell(row, 3 + col_idx), val, size=8)
                row += 1
        prog_blocks.append((prog_start, row - 1, prog))

    for start, end, _ in prog_blocks:
        vmerge(table, 0, start, end)
    row = 1
    for _, codes in PROGRAMS:
        for code in codes:
            vmerge(table, 1, row, row + 1)
            row += 2

    add_para(doc, note, size=9, color='333333', space_before=2, space_after=6)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Mm(297)
    sec.page_height = Mm(210)
    sec.left_margin = sec.right_margin = Mm(15)
    sec.top_margin = sec.bottom_margin = Mm(15)

    add_para(doc, '重组乙型肝炎疫苗（汉逊酵母，CpG和铝佐剂）Ⅲ期临床试验', size=14, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, '合并糖尿病病史亚组免疫原性结果表（PPS）', size=14, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_para(doc,
             '本分析纳入基线合并糖尿病病史的受试者共 30 例，其中 18-59 岁 18 例'
             '（0,1月低剂量组 5 例、0,1月高剂量组 3 例、0,2月低剂量组 2 例、'
             '0,2月高剂量组 3 例、阳性对照组 5 例）、≥60 岁 12 例'
             '（0,1,6月高剂量组 8 例、阳性对照组 4 例）。',
             size=10.5, space_after=4)
    add_para(doc,
             '注：本版本仅呈现 PPS（符合方案集）分析结果。阳转率表中“n/N”为阳转例数/该时间点可评价例数，'
             '“（百分比）”为阳转率；GMC 表中括号内为该时间点计入计算的例数；抗-HBs 浓度低于检测下限'
             '（<2.00 mIU/mL）的样本按 LLOQ/2（1.00 mIU/mL）计入 GMC 计算；“—”表示该时间点无数据或未采集。',
             size=9, color='555555', space_after=8)

    rel_note = ('注：全免后 1 个月/2 个月在各程序的对应时间点——'
                '0,1月程序为 M2/M3，0,2月程序为 M3/M4，0,1,6月程序为 M7/M8。')

    build_absolute_table(doc, 'sero',
                         '表 1  有、无糖尿病病史人群接种后抗-HBs阳转率（PPS）——绝对时间点（首剂接种后）',
                         '结果显示：有糖尿病病史人群完成全程 2 剂/3 剂接种后的阳转率总体达到较高水平，'
                         '0,2月程序两组在 M3 均达 100.00%，0,1,6月高剂量组及两个阳性对照组在 M7 均达 100.00%。')

    build_relative_table(doc, 'sero',
                         '表 2  有、无糖尿病病史人群接种后抗-HBs阳转率（PPS）——相对时间点（全免后）',
                         '全免后 1 个月、2 个月，各组有、无糖尿病病史人群的阳转率均维持在较高水平。' + rel_note)

    build_absolute_table(doc, 'gmc',
                         '表 3  有、无糖尿病病史人群接种后抗-HBs GMC（PPS）——绝对时间点（首剂接种后）',
                         '各组有、无糖尿病病史人群完成全程 2 剂/3 剂接种后各时间点的抗体 GMC 均呈现上升趋势。')

    build_relative_table(doc, 'gmc',
                         '表 4  有、无糖尿病病史人群接种后抗-HBs GMC（PPS）——相对时间点（全免后）',
                         '全免后 1 个月、2 个月，各组有、无糖尿病病史人群的抗体 GMC 均维持在一定水平。' + rel_note)

    doc.save(OUT)
    print('saved:', OUT)
    print('tables:', len(doc.tables))


if __name__ == '__main__':
    main()
