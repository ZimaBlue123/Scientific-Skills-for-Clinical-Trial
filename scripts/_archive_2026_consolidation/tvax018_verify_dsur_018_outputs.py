# -*- coding: utf-8 -*-
import docx

def load(p):
    return docx.Document(p)

def all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(c.text)
    return "\n".join(parts)

D1 = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#1_20250808-修订版.docx"
D2 = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#2_20260808-优化版.docx"

print("="*70)
print("DSUR#1 修订版 校验")
d1 = load(D1); t1 = all_text(d1)
checks1 = [
    ("无重要风险", "应已修复为风险列表", "出现"),
    ("唯一研发者", "应改为研发者", "出现"),
    ("08月09日", "应为08月08日", "出现"),
    ("暂无重要已识别和重要的潜在风险", "应改为风险列表", "出现"),
    ("无影响风险获益的重要风险", "应改为风险列表", "出现"),
    ("重要的已识别风险：无", "应存在", "缺失"),
    ("①严重过敏反应；②热性惊厥", "应存在（至少一处）", "缺失"),
]
for s, desc, bad in checks1:
    cnt = t1.count(s)
    status = "OK" if cnt == 0 else ("FAIL" if bad == "出现" else "OK")
    print("  [%s] %-40s 出现 %d 次  %s" % (status, s, cnt, desc))

# 批注残留检查
comment_rels = [r.reltype for rid, r in d1.part.rels.items() if "comment" in r.reltype.lower()]
print("  批注关系部件残留:", comment_rels if comment_rels else "无（已清除）")

print("="*70)
print("DSUR#2 优化版 校验")
d2 = load(D2); t2 = all_text(d2)
checks2 = [
    ("无重要的潜在风险", "执行概要应补列表", "出现"),
    ("唯一研发者", "应改为研发者", "出现"),
    ("获得4篇", "应为获得2篇", "出现"),
    ("简称冻干Hib结合疫苗", "笔误应修正", "出现"),
    ("潘雪娇,周洋", "旧文献应替换", "出现"),
    ("J. Guo, B. Zheng", "旧文献应替换", "出现"),
    ("M. Ren et al", "旧文献应删除", "出现"),
    ("X. Pan, Y. Chen", "旧文献应删除", "出现"),
    ("将按CDE建议执行", "附件2状态列", "缺失"),
    ("Zhou Y, Tang R", "新增文献1", "缺失"),
    ("Zhang C, Liu X", "新增文献2", "缺失"),
    ("获得2篇", "检索结果2篇", "缺失"),
    ("重要的已识别风险：无", "执行概要", "缺失"),
]
for s, desc, bad in checks2:
    cnt = t2.count(s)
    ok = (cnt == 0) if bad == "出现" else (cnt > 0)
    print("  [%s] %-40s 出现 %d 次  %s" % ("OK" if ok else "FAIL", s, cnt, desc))

# 重编号检查：疾病背景不应再有 [5]-[12]，而应有 [3]-[10]
import re
cits = re.findall(r"\[(\d+)\]", t2)
from collections import Counter
cc = Counter(int(x) for x in cits)
print("  正文引用编号分布:", dict(sorted(cc.items())))
bad_cit = [n for n in cc if n >= 5 and n <= 12]
print("  残留 [5]-[12] 引用:", bad_cit if bad_cit else "无（重编号完成）")
print("  [3]-[10] 是否存在:", all(n in cc for n in range(3, 11)))
