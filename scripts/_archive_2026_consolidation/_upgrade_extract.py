"""Insert Markdown rendering into scripts/extract_docx_full.py (priority 1)."""

from pathlib import Path

p = Path(r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\scripts\extract_docx_full.py")
src = p.read_text(encoding="utf-8")

OLD_DOC = """Extract text from .docx / .doc files.

Supports a single file, a list of files, or every supported file in a
folder. Modern .docx is handled via python-docx; legacy .doc is handled
via the Windows-only Word COM interface.

Usage
-----
    # Single file
    py -3 scripts/extract_docx_full.py document.docx [output.txt]

    # Folder batch (all docx/doc files combined into one output)
    py -3 scripts/extract_docx_full.py folder/ [output.txt]
"""

NEW_DOC = """Extract text from .docx / .doc files.

Supports a single file, a list of files, or every supported file in a
folder. Modern .docx is handled via python-docx; legacy .doc is handled
via the Windows-only Word COM interface.

Output formats
--------------
* --format text  (default) plain paragraphs + tables-as-pipes
* --format md    Markdown-flavoured rendering. Replaces the now-archived
                 scripts/extract_docx_to_md.py.

Usage
-----
    py -3 scripts/extract_docx_full.py document.docx [output.txt]
    py -3 scripts/extract_docx_full.py document.docx --format md output.md
    py -3 scripts/extract_docx_full.py folder/ [output.txt]
"""

assert OLD_DOC in src, "OLD_DOC not found"
src = src.replace(OLD_DOC, NEW_DOC)

# Insert new function after `extract_docx`
OLD_FUNC = '''def extract_docx(filepath: Path) -> str:
    """Extract all text from a .docx using python-docx."""
    from docx import Document  # local import keeps module importable when missing

    doc = Document(str(filepath))
    full_text: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\\n\\n".join(full_text)'''

NEW_FUNC = (
    OLD_FUNC
    + '''



def extract_docx_to_markdown(filepath: Path) -> str:
    """Render a .docx to Markdown.

    Paragraph styles named "Heading 1..3" / "Title" map to ``#``-style
    headings; tables render as GFM pipe tables. Replaces the archived
    ``extract_docx_to_md.py``.
    """
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(filepath))
    out: list[str] = []

    def _heading_level(para: Paragraph) -> int | None:
        name = (para.style.name or "").strip().lower()
        if name == "title":
            return 0
        if name.startswith("heading "):
            try:
                return int(name.split()[1])
            except (IndexError, ValueError):
                return None
        return None

    def _esc(s: str) -> str:
        return s.replace("|", "\\\\|").replace("\\n", " ")

    def _render_table(tbl: Table) -> list[str]:
        rows = tbl.rows
        if not rows:
            return []
        md: list[str] = []
        header = [_esc(c.text.strip()) for c in rows[0].cells]
        md.append("| " + " | ".join(header) + " |")
        md.append("|" + "|".join(["---"] * len(header)) + "|")
        for row in rows[1:]:
            cells = [_esc(c.text.strip()) for c in row.cells]
            md.append("| " + " | ".join(cells) + " |")
        return md

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.rstrip()
            if not text:
                continue
            level = _heading_level(para)
            if level == 0:
                out.append(f"# {text}")
            elif level is not None and 1 <= level <= 6:
                out.append("#" * (level + 1) + " " + text)
            else:
                out.append(text)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            out.extend(_render_table(tbl))
    return "\\n\\n".join(out) + "\\n"'''
)

assert OLD_FUNC in src, "OLD_FUNC not found"
src = src.replace(OLD_FUNC, NEW_FUNC)

# Add --format to parser
OLD_PARSER = """    parser.add_argument(
        "--log-level", default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging verbosity (default: %(default)s).",
    )
    return parser"""

NEW_PARSER = """    parser.add_argument(
        "--log-level", default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging verbosity (default: %(default)s).",
    )
    parser.add_argument(
        "--format", default="text",
        choices=("text", "md"),
        dest="fmt",
        help="Output format (default: %(default)s).",
    )
    return parser"""

assert OLD_PARSER in src, "OLD_PARSER not found"
src = src.replace(OLD_PARSER, NEW_PARSER)

# Inject --format dispatch in main()
OLD_MAIN_DISPATCH = """    if input_p.is_dir():
        extract_folder(input_p, output_p)
        return 0

    try:
        text = extract_file(input_p)
    except Exception as exc:  # noqa: BLE001
        logger.exception("failed to extract %s: %s", input_p, exc)
        return 1"""

NEW_MAIN_DISPATCH = """    if input_p.is_dir():
        extract_folder(input_p, output_p)
        return 0

    if args.fmt == "md":
        try:
            text = extract_docx_to_markdown(input_p)
        except Exception as exc:  # noqa: BLE001
            logger.exception("failed to extract md from %s: %s", input_p, exc)
            return 1
    else:
        try:
            text = extract_file(input_p)
        except Exception as exc:  # noqa: BLE001
            logger.exception("failed to extract %s: %s", input_p, exc)
            return 1"""

assert OLD_MAIN_DISPATCH in src, "OLD_MAIN_DISPATCH not found"
src = src.replace(OLD_MAIN_DISPATCH, NEW_MAIN_DISPATCH)

p.write_text(src, encoding="utf-8")
print(f"Updated {p}")
print(f"New size: {p.stat().st_size} bytes")
