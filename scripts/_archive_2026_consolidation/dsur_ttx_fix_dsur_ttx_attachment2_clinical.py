# -*- coding: utf-8 -*-
"""重写 DSUR#1 附件2「重要监管要求汇总」表格的临床部分（一）。

依据临床试验批准通知书（编号 2025LP01760）中的 5 条临床意见，将
「（一）临床试验开展前」的临床部分由 3 行扩展为 5 行，更新节标题，
并逐条填写 CDE 意见原文（要求列）与申办者回复（状态列）。

仅修改临床部分（一），「（二）药学」「（三）非临床」及文档其余内容不动。

用法：
    python fix_dsur_ttx_attachment2_clinical.py           # dry-run，仅打印
    python fix_dsur_ttx_attachment2_clinical.py --apply   # 正式落盘
"""

import copy
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = "review_materials/远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"
TABLE_INDEX = 5

# 节标题（第2列，vMerge=restart 的单元格）
SECTION_HEADER = "（一）临床试验开展前，进一步完善临床试验方案、风险管理计划、知情同意书等资料："

# 5 条临床意见：(要求列多行, 状态列回复)
CLINICAL_ROWS = [
    (
        [
            "1. 研究人群方面：",
            "（1）请结合我国非新生儿破伤风病例的年龄分布特征，合理设置受试人群的年龄构成，关注老年人群占比；",
            "（2）全程免疫人群既往接种史较为复杂，建议根据既往接种剂次分层分析。",
        ],
        "将按CDE建议执行。Ⅲ期临床试验纳入18岁及以上健康人群（年龄不设上限），可覆盖老年人群，并结合我国非新生儿破伤风病例以老年人群为主（平均发病年龄约69岁）的流行病学特征设置年龄构成；受试人群按既往含破伤风类毒素疫苗接种史分为加强免疫人群（接种≥3剂）与全程免疫人群（接种<3剂或免疫史不详），全程免疫人群按免疫史不详、接种<3剂等情形进一步细分并按既往接种剂次分层分析。",
    ),
    (
        [
            "2. 试验设计方面：提醒关注试验用疫苗包装规格可能存在的差异，并应制定相应的盲态维持措施。",
        ],
        "将按CDE建议执行。试验用疫苗采用预灌封注射器包装（0.5mL/支，1支/盒），与对照疫苗包装规格存在差异；临床试验采用随机、双盲、阳性对照设计，通过盲态数据审核、非盲团队管理等相应措施维持盲态。",
    ),
    (
        [
            "3. 安全性方面：",
            "（1）鼓励30天内的主动安全性随访尽可能采用面对面随访方式；",
            "（2）建议全程免疫人群首批入组的60例受试者增加每剂免后4天的实验室检测指标检测，并将其纳入试验暂停标准。",
        ],
        "将按CDE建议执行。30天内的主动安全性随访将尽可能采用面对面随访方式；全程免疫人群首批入组60例受试者增加每剂免后4天的实验室检测指标检测，并纳入试验暂停标准。",
    ),
    (
        [
            "4. 免疫原性方面：建议以PPS作为免疫原性分析主要分析集，同时对FAS进行分析。两者结果不一致时进一步分析原因。",
        ],
        "将按CDE建议执行。Ⅲ期临床试验以符合方案集（PPS）作为免疫原性分析主要分析集，同时对全分析集（FAS）进行分析，两者结果不一致时进一步分析原因。",
    ),
    (
        [
            "5. 风险管理计划及知情同意书：根据非临床研究结果，临床试验期间关注肌肉注射刺激性并制定相应风险控制措施。",
            "鉴于加强免疫、全程免疫试验操作流程差异较大，建议针对两部分人群分别制定知情同意书。",
        ],
        "将按CDE建议执行。根据非临床研究结果，临床试验期间关注肌肉注射部位刺激性（疼痛、红肿等）并制定相应风险控制措施；针对加强免疫与全程免疫人群操作流程差异，分别制定知情同意书。",
    ),
]


def make_run(text, rPr_template):
    r = OxmlElement('w:r')
    if rPr_template is not None:
        r.append(copy.deepcopy(rPr_template))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def set_cell_text(tc, lines):
    """清空单元格并写入多行文本，保留原字体/段落格式。"""
    if isinstance(lines, str):
        lines = [lines]
    paras = tc.findall(qn('w:p'))
    if not paras:
        return
    first_p = paras[0]
    p_template = copy.deepcopy(first_p)
    rPr_template = None
    runs = first_p.findall(qn('w:r'))
    if runs:
        rPr = runs[0].find(qn('w:rPr'))
        if rPr is not None:
            rPr_template = copy.deepcopy(rPr)
    for p in paras:
        tc.remove(p)
    for line in lines:
        new_p = copy.deepcopy(p_template)
        for r in list(new_p.findall(qn('w:r'))):
            new_p.remove(r)
        new_p.append(make_run(line, rPr_template))
        tc.append(new_p)


def get_tcs(tr):
    return tr.findall(qn('w:tc'))


def main():
    apply = "--apply" in sys.argv
    doc = Document(DOCX)
    t = doc.tables[TABLE_INDEX]

    # 定位现有行：Row0 表头，Row1-3 临床，Row4 药学起始
    trs = t._tr_lst if hasattr(t, '_tr_lst') else t.rows  # fallback
    rows = list(t.rows)

    if len(rows) < 4:
        print("[错误] 表格行数不足，无法处理。")
        return

    row1 = rows[1]._tr  # 临床第1行
    row2 = rows[2]._tr  # 临床第2行
    row3 = rows[3]._tr  # 临床第3行

    print(f"[dry-run] 表格共 {len(rows)} 行，临床部分现有 3 行（Row1-3），将扩展为 5 行。")

    # 打印节标题变更
    tc2 = get_tcs(row1)[2]
    old_header = "".join(tc2.itertext())
    print(f"[节标题] 旧：{old_header!r}")
    print(f"[节标题] 新：{SECTION_HEADER!r}")

    # 打印每行内容
    print("\n[要求/状态] 5 条内容如下：")
    for i, (req_lines, status) in enumerate(CLINICAL_ROWS, 1):
        print(f"--- 第{i}条 ---")
        for rl in req_lines:
            print(f"  要求: {rl}")
        print(f"  状态: {status[:60]}{'...' if len(status) > 60 else ''}")

    print(f"\n[插入] 将在 Row3 之后插入 2 行（要求4、要求5），药学部分自动下移。")

    if not apply:
        print("\n[dry-run 完成] 未落盘。确认后加 --apply 执行。")
        return

    # === 正式执行 ===
    # 1. 更新节标题
    set_cell_text(get_tcs(row1)[2], SECTION_HEADER)

    # 2. 填充 Row1-3
    for idx, tr in enumerate([row1, row2, row3]):
        req_lines, status = CLINICAL_ROWS[idx]
        tcs = get_tcs(tr)
        set_cell_text(tcs[3], req_lines)
        set_cell_text(tcs[4], status)

    # 3. 插入 2 行（复制 Row3 结构：日期/机构/节标题均为 vMerge=continue）
    ref_tr = row3
    for idx in [3, 4]:
        new_tr = copy.deepcopy(ref_tr)
        ref_tr.addnext(new_tr)
        ref_tr = new_tr
        req_lines, status = CLINICAL_ROWS[idx]
        tcs = get_tcs(new_tr)
        set_cell_text(tcs[3], req_lines)
        set_cell_text(tcs[4], status)

    doc.save(DOCX)
    print(f"\n[已落盘] 保存至：{DOCX}")

    # 4. 二次验证
    doc2 = Document(DOCX)
    t2 = doc2.tables[TABLE_INDEX]
    rows2 = list(t2.rows)
    print(f"[验证] 表格现共 {len(rows2)} 行（原12行 + 2 = 14行）。")
    for ri in range(1, 6):
        tcs = get_tcs(rows2[ri]._tr)
        req = "".join(tcs[3].itertext())
        sta = "".join(tcs[4].itertext())
        print(f"  Row{ri} 要求[{req[:30]}...] 状态[{sta[:30]}...]")
    # 残留检查
    full = "\n".join("".join(c.text for c in row.cells) for row in rows2)
    for kw in ["已落实", "已分别制定全程免疫人群", "风险管理计划.0", "（一）临床试验开展前，进一步完善临床试验方案与风险管理计划"]:
        print(f"[验证] 「{kw}」残留 = {full.count(kw)}")


if __name__ == "__main__":
    main()
