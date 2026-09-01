# -*- coding: utf-8 -*-
"""
生成《冻干b型流感嗜血杆菌结合疫苗（TVAX-018-2）DSUR#1》笔误修订版。

策略：克隆 DSUR#1（有笔误）DOCX，保留版式/样式/表格/页眉页脚，仅做：
  1) 选择性删除批注——删除已解决的 4 条问题批注（id 1/174/242/258 = ①–④），
     保留 4 条占位批注（id 130/281/361/362 = ⑤–⑧）；
  2) 整段替换（FULL_MAP）——修正 4 处笔误/补列表（批注[1][174][242][258]）；
  3) 全局子串替换（SUB）——「唯一研发者」→「研发者」；
  4) 保持 4 处占位（第12章、附件2药学、附件R5非临床、附件R5药学）。

批注删除需同步清理四个部件：
  word/comments.xml（按 w:id）、word/commentsExtended.xml（按位置）、
  word/commentsIds.xml（按位置，并收集 durableId）、
  word/commentsExtensible.xml（按 durableId）。

依赖：python-docx + lxml（default venv 已装）。
运行：C:/Users/Administrator/.workbuddy/binaries/python/envs/default/Scripts/python.exe
用法：python fix_dsur_tvax018_1_typos.py [--apply]
      （不带 --apply 为 dry-run，仅打印命中对照，不落盘）
"""
import sys
import docx
from lxml import etree

TPL = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#1_20250808-clean（有笔误）.docx"
OUT = r"review_materials/018-2 DSUR/远大赛威信冻干HIB结合疫苗_DSUR#1_20250808-修订版.docx"

# 整段替换（old 全文 -> new 全文）
FULL_MAP = {
    # 批注[1] 执行概要 P13：无重要风险 -> 补风险列表
    "报告期内，远大赛威信研发的冻干Hib结合疫苗暂无临床研究数据，根据本品非临床安全性数据、文献报道、同类疫苗的数据进行总结，无重要风险。":
    "报告期内，远大赛威信研发的冻干Hib结合疫苗暂无临床研究数据，根据本品非临床安全性数据、文献报道、同类疫苗的数据进行总结。（1）重要的已识别风险：无；（2）重要的潜在风险：①严重过敏反应；②热性惊厥。",

    # 批注[174] 第17章：08月09日 -> 08月08日
    "本DSUR自数据锁定点（2025年07月10日）后至2025年08月09日，无新的潜在重要安全性发现。":
    "本DSUR自数据锁定点（2025年07月10日）后至2025年08月08日，无新的潜在重要安全性发现。",

    # 批注[242] 18.1.2 潜在的风险：暂无重要风险 -> 补风险列表
    "根据本品非临床安全性数据、文献报道和同类疫苗临床研究数据等安全性信息，本品暂无重要已识别和重要的潜在风险。":
    "根据本品非临床安全性数据、文献报道和同类疫苗临床研究数据等安全性信息，整理了现阶段重要的潜在风险，对风险进行评估，并将在临床试验中进行密切监测。本品现阶段重要的潜在风险包括：①严重过敏反应；②热性惊厥。",

    # 批注[258] 第19章 19.2 本周期重要风险：无影响风险获益的重要风险 -> 补列表
    "本期报告为本品首次DSUR，本报告期内，远大赛威信研发的冻干Hib结合疫苗尚未在任何国家/地区获批上市，未开展相关的非临床研究。无正在进行的临床试验，自DIBD起截至DLP，无严重不良事件发生，无可疑且非预期严重不良反应（SUSAR）发生，未提示新的安全性风险。从目前本品非临床安全性数据、文献报道和同类疫苗临床研究数据的安全性信息可知，预计接种本品可能会出现常见的疫苗接种后一般反应，无影响风险获益的重要风险。":
    "本期报告为本品首次DSUR，本报告期内，远大赛威信研发的冻干Hib结合疫苗尚未在任何国家/地区获批上市，未开展相关的非临床研究。无正在进行的临床试验，自DIBD起截至DLP，无严重不良事件发生，无可疑且非预期严重不良反应（SUSAR）发生，未提示新的安全性风险。本品现阶段重要的潜在风险为：①严重过敏反应；②热性惊厥。",
}

# 全局子串替换
SUB = [
    ("唯一研发者", "研发者"),
]

# ---------------------------------------------------------------------------
# 批注选择性删除相关常量
# ---------------------------------------------------------------------------
# 要删除的问题批注 id（①–④）；占位批注（⑤–⑧）保留
DELETE_IDS = {"1", "174", "242", "258"}

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
W16CID = "http://schemas.microsoft.com/office/word/2016/wordml/cid"
W16CEX = "http://schemas.microsoft.com/office/word/2018/wordml/cex"

REL_COMMENTS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
REL_CEX = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
REL_CIDS = "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
REL_CEXT = "http://schemas.microsoft.com/office/2018/08/relationships/commentsExtensible"


def set_para_text(p, text):
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def apply_sub(text):
    for old, new in SUB:
        text = text.replace(old, new)
    return text


def strip_anchors_selective(root):
    """删除 DELETE_IDS 对应批注的锚点（commentRangeStart/End/Reference），保留其余批注。"""
    for tag in ("w:commentRangeStart", "w:commentRangeEnd"):
        local = tag.split(":")[1]
        for el in root.findall(".//{%s}%s" % (W, local)):
            if el.get("{%s}id" % W) in DELETE_IDS:
                el.getparent().remove(el)
    for ref in root.findall(".//{%s}commentReference" % W):
        if ref.get("{%s}id" % W) in DELETE_IDS:
            run = ref.getparent()
            if run is not None:
                run.getparent().remove(run)


def clean_comment_parts(doc):
    """删除 comments.xml 及三个辅助部件中 DELETE_IDS 对应的条目。"""
    comments_part = cex_part = cids_part = cext_part = None
    for rel in doc.part.rels.values():
        rt = rel.reltype
        if rt == REL_COMMENTS:
            comments_part = rel.target_part
        elif rt == REL_CEX:
            cex_part = rel.target_part
        elif rt == REL_CIDS:
            cids_part = rel.target_part
        elif rt == REL_CEXT:
            cext_part = rel.target_part

    # 1. comments.xml（CommentsPart，含 _element）
    croot = comments_part._element
    comment_els = croot.findall("{%s}comment" % W)
    ordered_ids = [el.get("{%s}id" % W) for el in comment_els]
    del_positions = [i for i, cid in enumerate(ordered_ids) if cid in DELETE_IDS]
    removed = []
    for el in comment_els:
        if el.get("{%s}id" % W) in DELETE_IDS:
            removed.append(el.get("{%s}id" % W))
            croot.remove(el)
    print("[comments.xml] 删除 w:comment id =", removed, "（保留", [i for i in ordered_ids if i not in DELETE_IDS], "）")

    # 2. commentsExtended.xml（按位置对齐 comments.xml 顺序）
    if cex_part is not None:
        root = etree.fromstring(cex_part.blob)
        els = root.findall("{%s}commentEx" % W15)
        if len(els) == len(ordered_ids):
            for i in sorted(del_positions, reverse=True):
                root.remove(els[i])
            cex_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            print("[commentsExtended.xml] 删除", len(del_positions), "条，剩余", len(root.findall('{%s}commentEx' % W15)), "条")
        else:
            print("[WARN] commentsExtended 条数(%d) != comments 条数(%d)，跳过辅助部件清理" % (len(els), len(ordered_ids)))

    # 3. commentsIds.xml（按位置，收集 durableId）
    del_durable = set()
    if cids_part is not None:
        root = etree.fromstring(cids_part.blob)
        els = root.findall("{%s}commentId" % W16CID)
        if len(els) == len(ordered_ids):
            for i in del_positions:
                if i < len(els):
                    d = els[i].get("{%s}durableId" % W16CID)
                    if d:
                        del_durable.add(d)
            for i in sorted(del_positions, reverse=True):
                root.remove(els[i])
            cids_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            print("[commentsIds.xml] 删除", len(del_positions), "条，durableId =", del_durable)
        else:
            print("[WARN] commentsIds 条数(%d) != comments 条数(%d)，跳过" % (len(els), len(ordered_ids)))

    # 4. commentsExtensible.xml（按 durableId）
    if cext_part is not None and del_durable:
        root = etree.fromstring(cext_part.blob)
        kept = []
        for el in root.findall("{%s}commentExtensible" % W16CEX):
            if el.get("{%s}durableId" % W16CEX) in del_durable:
                root.remove(el)
            else:
                kept.append(el.get("{%s}durableId" % W16CEX))
        cext_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        print("[commentsExtensible.xml] 删除 durableId in", del_durable, "，剩余", kept)


def main(apply=False):
    doc = docx.Document(TPL)

    # 1. 整段替换 + 子串替换（正文段落）
    for p in doc.paragraphs:
        full = p.text
        if full in FULL_MAP:
            print("[FULL] %r  =>  %r" % (full[:30], FULL_MAP[full][:30]))
            if apply:
                set_para_text(p, FULL_MAP[full])
            continue
        new = apply_sub(full)
        if new != full:
            print("[SUB ] %r  =>  %r" % (full[:40], new[:40]))
            if apply:
                set_para_text(p, new)

    # 2. 选择性删除批注锚点 + 清理批注部件
    if apply:
        strip_anchors_selective(doc.element.body)
        for sec in doc.sections:
            for hf in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
                if hf is not None:
                    strip_anchors_selective(hf._element)
        clean_comment_parts(doc)

        doc.save(OUT)
        print("SAVED ->", OUT)
    else:
        print("\n[dry-run] 未落盘；如需写入请追加 --apply")


if __name__ == "__main__":
    main(apply="--apply" in sys.argv)
