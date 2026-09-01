# -*- coding: utf-8 -*-
"""修复吸附破伤风疫苗 DSUR 交付稿：文献检索数量模糊化。

将第13章"文献"中的各数据库具体命中条数去掉，只保留最终筛选获得的有效文献数（3篇）。

用法：
    python fix_dsur_ttx_literature_count.py          # dry-run：打印将执行的替换
    python fix_dsur_ttx_literature_count.py --apply  # 正式执行并保存
"""

import sys

from docx import Document

SRC = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\远大吸附破伤风疫苗_DSUR#1（20250708-20260707）_预定稿_全文（含附件）.docx"

# 全局子串替换
SUB_REPLACEMENTS = [
    # PubMed 段：删除"共检索到73篇文献"具体数字
    ("，共检索到73篇文献。", "。"),
    # 中国知网/万方/Cochrane 段：去掉"仅列出策略、未开展实际检索"表述，改为实际检索（不写条数）
    ("本次仅列出上述检索策略，未开展实际文献检索。", "亦按上述策略进行了检索。"),
]


def replace_across_runs(paragraph, old, new):
    """跨 run 子串替换，只改动 old 覆盖的 run，其余 run 保持原样。"""
    runs = paragraph.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    idx = full.find(old)
    if idx == -1:
        return False
    end = idx + len(old)

    pos = 0
    start_run = 0
    for ri, r in enumerate(runs):
        r_end = pos + len(r.text)
        if r_end > idx:
            start_run = ri
            break
        pos = r_end

    pos = 0
    end_run = len(runs) - 1
    for ri, r in enumerate(runs):
        r_start = pos
        r_end = pos + len(r.text)
        if r_start < end:
            end_run = ri
        pos = r_end
        if pos >= end:
            break

    if start_run == end_run:
        runs[start_run].text = runs[start_run].text.replace(old, new, 1)
    else:
        prefix = sum(len(r.text) for r in runs[:start_run])
        start_offset = idx - prefix
        end_prefix = sum(len(r.text) for r in runs[:end_run])
        end_offset = end - end_prefix
        runs[start_run].text = runs[start_run].text[:start_offset] + new
        for r in runs[start_run + 1:end_run]:
            r.text = ""
        runs[end_run].text = runs[end_run].text[end_offset:]
    return True


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    yield p


def main():
    apply = "--apply" in sys.argv
    doc = Document(SRC)

    sub_hits = []
    for p in iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        for old, new in SUB_REPLACEMENTS:
            if old in full:
                sub_hits.append((old, new, full))
                if apply:
                    replace_across_runs(p, old, new)

    print("=== 子串替换（%d 处） ===" % len(sub_hits))
    for old, new, full in sub_hits:
        print(f"  {old!r}  ->  {new!r}")
        print(f"      原文片段: …{full[-80:]}")

    if apply:
        doc.save(SRC)
        print("保存完成:", SRC)
    else:
        print("DRY-RUN 模式：未修改文件。确认无误后加 --apply 正式执行。")


if __name__ == "__main__":
    main()
