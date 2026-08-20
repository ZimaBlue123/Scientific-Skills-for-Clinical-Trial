#!/usr/bin/env python3
"""
generator_base.py — shared building blocks for the ``scripts/generate_*.py``
family.

This module abstracts cross-cutting concerns that every audit/CSR/MMR report
generator re-implements:

* **Logging**  : every script starts logging the same way. Use
  :func:`setup_logging` to get a project-style ``logging.Logger``.
* **CLI**      : :func:`make_argparser` returns a parser with sane defaults
  (``--log-level``, ``--output``, ``--config``). Sub-classes can extend it.
* **Loading**  : :func:`load_template` finds a template markdown file by name
  and returns its UTF-8 contents.
* **Document** : :func:`build_document` returns a ``docx.Document()`` whose
  default styles already have :func:`apply_cn_en_fonts` applied, so each
  generator only needs to inject its own content.
* **Saving**   : :func:`save_document` writes a ``Document`` to disk and
  emits a single ``logger.info`` line that mirrors the convention used by
  generate_csr_docx.py et al.

This module previously lived as duplicated code across 10 generate_*.py
files (~195 KB total). The cleanup is part of the 2026-07 scripts/
consolidation plan (priority 3).
"""

from __future__ import annotations

import argparse
import logging
import sys
from collections.abc import Sequence
from pathlib import Path

from docx import Document

# Re-use the established helper to keep cn/en font behaviour identical.
from .docx_utils import apply_cn_en_fonts

LOG_FORMAT = "%(asctime)s [%(levelname)s] %(name)s: %(message)s"
logger = logging.getLogger("generator_base")


def setup_logging(level: str = "INFO") -> None:
    """Configure the root logger with a single console handler.

    Idempotent: safe to call from multiple scripts in the same process.
    """
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.INFO),
        format=LOG_FORMAT,
        stream=sys.stderr,
        force=True,
    )


def make_argparser(
    prog: str,
    description: str,
    *,
    default_output: str | None = None,
    config_help: str = "Path to a JSON/YAML config (optional).",
) -> argparse.ArgumentParser:
    """Return a project-style ArgumentParser with common flags."""
    p = argparse.ArgumentParser(prog=prog, description=description)
    p.add_argument(
        "--output",
        default=default_output,
        help=(
            "Output file path; default %(default)s"
            if default_output
            else "Output file path."
        ),
    )
    p.add_argument(
        "--log-level",
        default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging verbosity (default %(default)s).",
    )
    p.add_argument("--config", default=None, help=config_help)
    return p


def load_template(name: str, templates_dir: Path | None = None) -> str:
    """Read ``templates_dir/name`` and return its UTF-8 text.

    Parameters
    ----------
    name : str
        Filename under ``templates_dir``.
    templates_dir : Path | None
        Folder containing templates. Defaults to
        ``scripts/common_templates/`` (kept in the repo).

    Raises
    ------
    FileNotFoundError
        If the template file does not exist.
    OSError
        If the template file cannot be read.
    """
    if templates_dir is None:
        templates_dir = Path(__file__).resolve().parent.parent / "common_templates"
    path = templates_dir / name
    if not path.exists():
        raise FileNotFoundError(f"template not found: {path}")
    try:
        return path.read_text(encoding="utf-8")
    except OSError:
        logger.exception("failed to read template: %s", path)
        raise


def build_document(template_name: str | None = None) -> Document:
    """Create a Document and apply CN/EN font defaults.

    If ``template_name`` is provided, ``load_template`` is called but its
    text is *not* injected — generators are free to insert content as
    needed. ``template_name`` is currently informational; reserved for a
    future Jinja2-style templating step.
    """
    if template_name is not None:
        _ = load_template(template_name)
    doc = Document()
    apply_cn_en_fonts(doc)
    logger.debug("fresh Document() created; default styles cn/en-fonted")
    return doc


def save_document(doc: Document, output: Path | str) -> Path:
    """Persist ``doc`` to ``output`` (mkdir parents first).

    Raises
    ------
    OSError
        If the target directory cannot be created or the file cannot be written.
    """
    out = Path(output)
    try:
        out.parent.mkdir(parents=True, exist_ok=True)
    except OSError:
        logger.exception("cannot create output directory: %s", out.parent)
        raise
    try:
        doc.save(str(out))
    except OSError:
        logger.exception("failed to write document: %s", out)
        raise
    logger.info("wrote %s", out)
    return out


__all__ = [
    "setup_logging",
    "make_argparser",
    "load_template",
    "build_document",
    "save_document",
]


def main(argv: Sequence[str] | None = None) -> int:
    """Tiny self-test: prints the version & exits."""
    setup_logging("INFO")
    logger.info("generator_base loaded; helpers: %s", ", ".join(__all__))
    return 0


if __name__ == "__main__":
    sys.exit(main())
