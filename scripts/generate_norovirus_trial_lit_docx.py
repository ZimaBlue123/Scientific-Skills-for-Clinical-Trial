"""
生成 HilleVax 诺如疫苗 HIL-214 相关 PubMed 文献检索结果 Word 文档。

数据源：scripts/norovirus_trial_search.py 生成的
    .workbuddy/audit/norovirus_trial_pubmed.json

输出：.workbuddy/audit/norovirus_trial_lit_report.docx

文档结构：
  封面 / 标题 + 元数据
  一、检索方法（数据源、字段、关键词与命中数、说明）
  二、各试验相关文献
    2.1 NCT06120764 婴儿 III 期
    2.2 NCT05507060 成人 III 期
  三、关键观察与解读建议
  附录：完整 JSON 元数据
"""

import json
import os
import sys
from datetime import datetime

# 让脚本可以导入 common_scripts.docx_utils
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from common_scripts.docx_utils import apply_cn_en_fonts

JSON_PATH = ".workbuddy/audit/norovirus_trial_pubmed.json"
DOCX_PATH = ".workbuddy/audit/norovirus_trial_lit_report.docx"


# ---------------------------------------------------------------------------
# 检索策略元数据（与 norovirus_trial_search.py 对齐）
# ---------------------------------------------------------------------------
SEARCH_METHOD = {
    "data_source": "NCBI E-utilities（PubMed）",
    "script": "scripts/norovirus_trial_search.py",
    "search_fields": ["[tiab] 标题/摘要", "[si] Secondary Source ID"],
    "queries_used": [
        ("NCT06120764[si]", 0),
        ("NCT05507060[si]", 0),
        ('"HIL-214"[tiab]', 5),
        ("HilleVax[tiab] AND norovirus[tiab]", 2),
        ("HilleVax[tiab] AND (infant*[tiab] OR pediatric[tiab])", 0),
        ("HilleVax[tiab] AND adult*[tiab]", 1),
        (
            '"norovirus vaccine"[tiab] AND infant*[tiab] AND (phase III[tiab] OR phase 3[tiab])',
            0,
        ),
        (
            '"norovirus vaccine"[tiab] AND adult*[tiab] AND (phase III[tiab] OR phase 3[tiab])',
            0,
        ),
    ],
    "notes": [
        "NCT ID 直接字段（[si]）未命中：ClinicalTrials.gov→PubMed 自动链接在这两个较新试验上尚未建立索引。",
        '候选疫苗代码 "HIL-214" 是最有效的检索锚点（命中 5 篇原始研究）。',
        "HilleVax + norovirus 关键词组合命中 2 篇，其中 PMID 42140216 为成人 post-hoc 分析。",
        "受控人群限定（infant* / adult* + phase III）的检索式无额外命中，表明试验主结果论文尚未以 PMID 形式公开。",
    ],
}

TRIAL_META = {
    "NCT06120764": {
        "name": "HilleVax 婴儿 III 期试验（HIL-214，约 5 个月大婴儿）",
        "candidate": "HIL-214（双价 GI.1/GII.4 VLP 诺如疫苗，肌肉注射）",
        "key_topic": "三剂次免疫程序在婴儿中的安全性与免疫原性",
        "anchor_pmids": ["39803784", "37140558"],
        "intro": "下表按出版日期倒序排列。带 ★ 的 PMID 为与该试验直接相关性最强的支撑文献。",
    },
    "NCT05507060": {
        "name": "HilleVax 成人 III 期试验（HIL-214，针对成人腹泻预防）",
        "candidate": "HIL-214（双价 GI.1/GII.4 VLP 诺如疫苗，肌肉注射）",
        "key_topic": "成人急性胃肠炎预防；未达主要终点但减轻症状严重程度",
        "anchor_pmids": ["42140216", "39852862"],
        "intro": '下表按出版日期倒序排列。带 ★ 的 PMID 为支撑"未达主要终点但减轻症状严重程度"解读的核心文献。',
    },
}

KEY_OBSERVATIONS = [
    (
        "NCT ID 直接字段未索引",
        "通过 NCT06120764[si] 与 NCT05507060[si] 均返回 0 条，说明 ClinicalTrials.gov→PubMed 自动链接在这两个较新试验上尚未建立。",
    ),
    (
        "候选疫苗代码 HIL-214 是最有效的检索锚点",
        "共召回 5 篇原始研究 + 1 篇综述，涵盖 I/II 期安全性、免疫原性、免疫持久性与中和抗体相关性。",
    ),
    (
        "两个试验共用 HIL-214 平台，文献集合 100% 重叠",
        '需按"研究人群（婴儿 vs 成人）"和"试验阶段（I/II 期准备 vs III 期终点）"做归属判定，而非简单去重。',
    ),
    (
        "未达主要终点的解读依据",
        '主要见 PMID 42140216（post-hoc 血清学分析）与 PMID 39852862（成人长期免疫持久性数据），这两篇可作为"减轻症状严重程度"结论的文献支撑。',
    ),
    (
        "婴儿 III 期结果尚未以 PMID 形式公开",
        "检索式 infant*/pediatric + phase III 均无命中；婴儿 III 期顶线数据预期会以公司公告或学术会议形式首发，后续可在 PubMed 增量跟踪。",
    ),
    (
        "建议补充检索",
        "可在 ClinicalTrials.gov 结果发布 6-12 个月后追加检索：trial-result 字段、preprint（bioRxiv/medRxiv）、以及会议摘要（IDWeek / ESPID / ACIP）。",
    ),
]

HEADERS = ["PMID", "标题", "作者", "期刊", "日期", "DOI"]


# ---------------------------------------------------------------------------
# 渲染辅助
# ---------------------------------------------------------------------------
def _set_cell_text(
    cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT
) -> None:
    cell.text = ""
    paragraphs = text.split("\n")
    for idx, line in enumerate(paragraphs):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = tc_pr.makeelement(qn("w:shd"), {})
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _add_hyperlink(paragraph, url, text, size=10):
    """Append a clickable hyperlink via raw OXML (python-docx 1.x)."""
    from docx import oxml as _oxml

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = _oxml.OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = _oxml.OxmlElement("w:r")
    rpr = _oxml.OxmlElement("w:rPr")
    color = _oxml.OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = _oxml.OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = _oxml.OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    sz_cs = _oxml.OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz_cs)
    rfonts = _oxml.OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rpr.append(rfonts)
    new_run.append(rpr)
    t = _oxml.OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


def _add_body(doc: Document, text: str, size: int = 11, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def _add_bullet(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# 章节渲染
# ---------------------------------------------------------------------------
def render_cover(doc: Document, infant_count: int, adult_count: int) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HilleVax 诺如疫苗 HIL-214")
    run.bold = True
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("NCT06120764 / NCT05507060 相关 PubMed 文献检索报告")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    meta_tbl = doc.add_table(rows=4, cols=2)
    meta_tbl.style = "Table Grid"
    rows_data = [
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("数据源", "NCBI PubMed（via E-utilities）"),
        ("婴儿试验文献数", f"{infant_count} 条（含背景文献）"),
        ("成人试验文献数", f"{adult_count} 条（含背景文献）"),
    ]
    for i, (k, v) in enumerate(rows_data):
        _set_cell_text(meta_tbl.rows[i].cells[0], k, bold=True, size=10)
        _shade_cell(meta_tbl.rows[i].cells[0], "F2F2F2")
        _set_cell_text(meta_tbl.rows[i].cells[1], v, size=10)
    meta_tbl.columns[0].width = Cm(4.0)
    meta_tbl.columns[1].width = Cm(12.0)
    for row in meta_tbl.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(12.0)

    doc.add_paragraph()


def render_search_method(doc: Document) -> None:
    _add_section_heading(doc, "一、检索方法", level=1)

    _add_body(doc, f"数据源：{SEARCH_METHOD['data_source']}", size=11)
    _add_body(doc, f"检索脚本：{SEARCH_METHOD['script']}", size=11)
    _add_body(doc, "检索字段：" + "；".join(SEARCH_METHOD["search_fields"]), size=11)

    doc.add_paragraph()
    _add_section_heading(doc, "1.1 关键词与命中数", level=2)

    tbl = doc.add_table(rows=1 + len(SEARCH_METHOD["queries_used"]), cols=2)
    tbl.style = "Table Grid"
    _set_cell_text(
        tbl.rows[0].cells[0],
        "检索式",
        bold=True,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _set_cell_text(
        tbl.rows[0].cells[1],
        "命中数",
        bold=True,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _shade_cell(tbl.rows[0].cells[0], "DCE6F1")
    _shade_cell(tbl.rows[0].cells[1], "DCE6F1")
    tbl.columns[0].width = Cm(13.0)
    tbl.columns[1].width = Cm(3.0)
    for i, (q, n) in enumerate(SEARCH_METHOD["queries_used"], start=1):
        _set_cell_text(tbl.rows[i].cells[0], q, size=10)
        _set_cell_text(
            tbl.rows[i].cells[1], str(n), size=10, align=WD_ALIGN_PARAGRAPH.CENTER
        )
        tbl.rows[i].cells[0].width = Cm(13.0)
        tbl.rows[i].cells[1].width = Cm(3.0)

    doc.add_paragraph()
    _add_body(doc, "说明：", size=11, bold=True)
    for note in SEARCH_METHOD["notes"]:
        _add_bullet(doc, note)

    doc.add_paragraph()


def _render_trial_table(doc: Document, records: list[dict], trial_key: str) -> None:
    meta = TRIAL_META[trial_key]
    anchor_set = set(meta["anchor_pmids"])

    n_cols = len(HEADERS)
    table = doc.add_table(rows=1 + len(records), cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True

    col_widths = [Cm(1.6), Cm(5.4), Cm(2.4), Cm(2.0), Cm(1.6), Cm(3.0)]
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # 表头
    for i, h in enumerate(HEADERS):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(cell, "DCE6F1")
        cell.width = col_widths[i]

    # 数据行
    for row_idx, rec in enumerate(records, start=1):
        pmid = rec.get("pmid", "")
        title = rec.get("title", "")
        first_author = rec.get("first_author", "")
        journal = rec.get("journal", "")
        pubdate = rec.get("pubdate", "")
        doi = rec.get("doi", "")

        row = table.rows[row_idx]

        # PMID
        pmid_cell = row.cells[0]
        pmid_cell.text = ""
        p = pmid_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        is_anchor = pmid in anchor_set
        run = p.add_run(f"★ {pmid}" if is_anchor else pmid)
        run.font.size = Pt(10)
        if is_anchor:
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pmid_cell.width = col_widths[0]

        # 标题（含 PubMed 超链接）
        title_cell = row.cells[1]
        title_cell.text = ""
        p = title_cell.paragraphs[0]
        pub_url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
        _add_hyperlink(p, pub_url, title, size=10)
        if is_anchor:
            note_p = title_cell.add_paragraph()
            note_run = note_p.add_run("【直接相关性最强】")
            note_run.font.size = Pt(9)
            note_run.bold = True
            note_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        title_cell.width = col_widths[1]

        _set_cell_text(row.cells[2], first_author, size=10)
        row.cells[2].width = col_widths[2]
        _set_cell_text(row.cells[3], journal, size=10)
        row.cells[3].width = col_widths[3]
        _set_cell_text(row.cells[4], pubdate, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[4].width = col_widths[4]

        # DOI（带 doi.org 超链接）
        doi_cell = row.cells[5]
        doi_cell.text = ""
        p = doi_cell.paragraphs[0]
        if doi:
            _add_hyperlink(p, f"https://doi.org/{doi}", doi, size=9)
        else:
            run = p.add_run("—")
            run.font.size = Pt(9)
        doi_cell.width = col_widths[5]


def render_trial_section(
    doc: Document, trial_key: str, heading: str, records: list[dict]
) -> None:
    meta = TRIAL_META[trial_key]

    _add_section_heading(doc, heading, level=2)
    _add_body(doc, f"候选疫苗：{meta['candidate']}", size=11)
    _add_body(doc, f"关键议题：{meta['key_topic']}", size=11)

    doc.add_paragraph()
    _add_body(doc, meta["intro"], size=10)

    if not records:
        _add_body(doc, "（无命中记录）", size=11)
        return

    _render_trial_table(doc, records, trial_key)
    doc.add_paragraph()


def render_key_observations(doc: Document) -> None:
    _add_section_heading(doc, "三、关键观察与解读建议", level=1)

    for title, body in KEY_OBSERVATIONS:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}：")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p.add_run(body)
        run2.font.size = Pt(11)

    doc.add_paragraph()


def render_appendix(doc: Document, data: dict) -> None:
    _add_section_heading(doc, "附录：完整 JSON 元数据", level=1)
    _add_body(
        doc, "完整检索元数据见 .workbuddy/audit/norovirus_trial_pubmed.json。", size=10
    )
    code_text = json.dumps(data, ensure_ascii=False, indent=2)
    for line in code_text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
        rfonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main() -> int:
    if not os.path.exists(JSON_PATH):
        print(f"ERROR: JSON not found: {JSON_PATH}", file=sys.stderr)
        print("请先运行: python scripts/norovirus_trial_search.py", file=sys.stderr)
        return 1

    with open(JSON_PATH, encoding="utf-8") as f:
        data = json.load(f)

    infant = data.get("NCT06120764_infant", {})
    adult = data.get("NCT05507060_adult", {})

    doc = Document()
    apply_cn_en_fonts(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 封面
    render_cover(doc, infant.get("count", 0), adult.get("count", 0))

    # 一、检索方法
    render_search_method(doc)

    # 二、各试验相关文献
    _add_section_heading(doc, "二、各试验相关文献", level=1)
    render_trial_section(
        doc,
        "NCT06120764",
        "2.1 NCT06120764 — HilleVax 婴儿 III 期试验（HIL-214）",
        infant.get("records", []),
    )
    render_trial_section(
        doc,
        "NCT05507060",
        "2.2 NCT05507060 — HilleVax 成人 III 期试验（HIL-214）",
        adult.get("records", []),
    )

    # 三、关键观察
    render_key_observations(doc)

    # 附录
    render_appendix(doc, data)

    os.makedirs(os.path.dirname(DOCX_PATH), exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"OK -> {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
