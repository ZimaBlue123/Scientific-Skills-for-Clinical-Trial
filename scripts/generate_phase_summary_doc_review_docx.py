#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Word review report for phase CSR interim summary document."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _apply_cn_en_fonts(doc: Document) -> None:
    def set_style(style_name: str) -> None:
        if style_name not in doc.styles:
            return
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:cs"), "Times New Roman")

    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Table Grid"):
        set_style(name)


def _set_run_font(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run)
    run.bold = bold


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                _set_run_font(r)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    _set_run_font(r)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    _apply_cn_en_fonts(doc)

    today = date.today().strftime("%Y年%m月%d日")
    src_doc = (
        "YDSWX（TVAX-006）-002（II）阶段性小结（安全性&体液免疫原性）"
        " V1.0，2026年05月25日"
    )

    title = doc.add_heading("临床试验阶段性小结文档审核报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_para(doc, f"审核日期：{today}", bold=False)
    _add_para(doc, f"受审文件：{src_doc}", bold=False)
    _add_para(doc, "审核范围：全文逐节核对（摘要、正文、小结、讨论与结论及主要表格叙述）", bold=False)
    _add_para(doc, "审核类型：数据与表述一致性、分析结论完整性、语言文字与术语规范", bold=False)
    doc.add_paragraph()

    doc.add_heading("审核说明", level=1)
    _add_para(
        doc,
        "本报告基于对受审 Word 文档全文的通读与交叉核对，重点对照摘要、"
        "各章节正文、免疫原性/安全性小结及讨论与结论之间的逻辑与数据一致性。"
        "所列问题按严重程度与修改优先级排序，供撰写方修订参考。",
    )

    doc.add_heading("一、前后矛盾或易误导的数据与表述", level=1)

    doc.add_heading("1. 摘要：抗原指标名称写错（高优先级）", level=2)
    _add_para(doc, "位置：摘要 → 40~49岁 → 抗VZV抗原免疫应答 → 基于PPS-h2、第2剂接种后30天。")
    _add_bullets(
        doc,
        [
            "问题：将指标误写为“试验组 gE 抗原血清抗体 SCR…（100.00% vs. 5.71%）”，"
            "与同段 gE 结果重复，且与正文表3-6及§3.3中VZV结果不一致。",
            "建议：改为“试验组 VZV 抗原血清抗体 SCR/GMC/GMI”，并与表3-6逐项核对。",
        ],
    )

    doc.add_heading("2. 讨论与结论：与摘要/正文/统计结果不一致（高优先级）", level=2)
    _add_para(doc, "位置：§5 讨论与结论，关于50~59岁、≥60岁、≥50岁年龄层的概括段。")
    _add_bullets(
        doc,
        [
            "问题：先写“第2剂接种前和接种后30天，两组gE、VZV的SCR、校正后GMC及GMI均无显著差异（P>0.05）”，"
            "与摘要、§3.3及表3-4、表3-6矛盾——≥60岁、PPS-h2时试验组gE GMC（P=0.011）"
            "及VZV GMC（P=0.030）显著低于阳性对照组2。",
            "建议：按年龄层、访视时点、终点分别表述；≥60岁PPS-h2应写明“SCR/GMI相当，校正GMC低于阳性对照组2”。",
        ],
    )

    doc.add_heading("3. §3.4 免疫原性小结：章节内表述不统一", level=2)
    _add_bullets(
        doc,
        [
            "第2剂前（h1）段对≥60岁、≥50岁写“各参数组间无显著差异”，对SCR略低（如79.71% vs 89.86%）的表述偏弱。",
            "第2剂后（h2）段对≥60岁已正确写明校正GMC显著低于对照，与§5讨论首段不一致，需全文统一。",
        ],
    )

    doc.add_heading("4. 安全性：摘要与正文细节不完全一致（中优先级）", level=2)
    _add_table(
        doc,
        ["议题", "摘要", "正文（如§4.1）"],
        [
            [
                "40~49岁 0~14天 AE",
                "写“均与试验用疫苗相关”",
                "存在非征集性AE，并非全部与疫苗相关",
            ],
            [
                "≥50岁 摘要安全性段",
                "写“阳性对照组1例受试者…”",
                "该亚组应对照阳性对照组2，属组别笔误",
            ],
        ],
    )

    doc.add_heading("5. 统计表：率差置信区间笔误（中优先级）", level=2)
    _add_para(
        doc,
        "位置：§3.3，≥50岁、gE SCR率差。正文写“09.74%”，应为“-0.74%”"
        "（与SCR 99.26% vs 100.00%及CI一致）。",
    )

    doc.add_heading("6. 研究人群表述不统一（低~中优先级）", level=2)
    _add_bullets(
        doc,
        [
            "方案/标题为“40岁及以上”，试验目的写“健康成人”，范围偏宽。",
            "方案要求“≥60岁中≥70岁不少于该层20%”，正文未见是否满足的汇总说明。",
        ],
    )

    doc.add_heading("7. 对照设计差异（非数据矛盾，结论需限定）", level=2)
    _add_bullets(
        doc,
        [
            "40~49岁：试验组 vs 安慰剂→感维®（阳性对照组1）。",
            "≥50岁：试验组 vs 欣安立适®（阳性对照组2）。",
            "不宜将“与阳性对照相当”笼统用于全年龄，应分年龄、分对照说明。",
        ],
    )

    doc.add_heading("8. 报告范围说明（非矛盾，结论需限定）", level=2)
    _add_para(
        doc,
        "SAP载明细胞免疫结果待获得后补充；本次仅体液免疫+第2剂后30天安全性。"
        "结论不宜暗示细胞免疫或12/24/36个月持久性已完成。",
    )

    doc.add_heading("二、建议补充的分析与结论", level=1)

    doc.add_heading("1. 建议在结论/局限性中明确写明", level=2)
    _add_bullets(
        doc,
        [
            "细胞免疫亚组（90例）本次未纳入，预计补充时间及是否影响III期决策。",
            "长期免疫原性（12/24/36个月）为方案次要终点，本次仅至第2剂后30天。",
            "≥60岁GMC低于欣安立适而SCR/GMI相当：补充与方案预设标准的关系及临床意义讨论。",
            "≥70岁占≥60岁20%的入组要求是否在140例中达标。",
            "AESI/SAE/妊娠观察至全程免后12个月，本次仅覆盖数据冻结日前30天窗口。",
        ],
    )

    doc.add_heading("2. 建议加强的分析性结论", level=2)
    _add_table(
        doc,
        ["主题", "建议"],
        [
            [
                "40~49岁",
                "明确相对第0天安慰剂的免疫原性；解读第60天接种减毒活疫苗后对照SCR仍低的原因",
            ],
            ["≥50岁", "分述与欣安立适SCR/GMI相当；≥60岁GMC差异的临床意义"],
            ["安全性", "40~49岁反应高于对照1与亚单位疫苗预期一致；3级发热短暂，可增benefit-risk表述"],
            ["分析集", "可补充因方案偏离剔除PPS的例数及对主要终点的定性影响"],
            ["合并用药", "试验组合并用药率更高，如与AE无关建议一句说明，避免读者误解为混杂"],
        ],
    )

    doc.add_heading("3. III期支持性结论（宜加限定语）", level=2)
    _add_para(
        doc,
        "建议将结论调整为：在“第2剂后30天”体液免疫原性及安全性方面，结果支持开展/推进III期；"
        "细胞免疫及长期随访结果待后续报告补充。",
    )

    doc.add_heading("三、语句、术语与错别字问题", level=1)

    doc.add_heading("明确错误（建议修改）", level=2)
    _add_table(
        doc,
        ["位置", "问题", "建议"],
        [
            ["摘要 VZV（PPS-h2，40~49岁）", "gE误作VZV指标名", "改为VZV"],
            ["§5讨论首段", "笼统写全部“无显著差异”", "按年龄层、访视、终点拆分"],
            ["§3.4小结 h1", "“GMT”", "应为“GMI”"],
            ["§3.4小结", "“阳性对照徐2”", "阳性对照组2"],
            ["§3.4小结", "“阳性对照组1gE”", "加空格：阳性对照组1 gE"],
            ["≥50岁 SCR率差", "“09.74%”", "“-0.74%”"],
            ["§4摘要 ≥50岁", "“阳性对照组1例”", "阳性对照组2"],
            ["声明", "“福塔雷萨2013版”", "福尔塔莱萨或Fortaleza 2013"],
            ["统计分析", "MedDRA英文全称", "…Regulatory Activities"],
            ["方案/终点", "“CD42+T细胞”", "核对方案是否为CD4+等"],
            ["约第319行", "句末缺句号", "补“。”"],
        ],
    )

    doc.add_heading("用语与格式建议统一", level=2)
    _add_bullets(
        doc,
        [
            "“试验疫苗”/“试验用疫苗”混用，建议统一为“试验用疫苗”。",
            "表号“表2-1”与目录“表21”并存，建议统一编号体系。",
            "目录“目录1”“表目录1”等似为页码占位未更新。",
            "方案V1.3与报告V1.0，建议在概述说明阶段性分析所依据方案版本。",
            "“校正后GMC”/“LS GMC”混用，首次出现宜定义一致。",
        ],
    )

    doc.add_heading("表述可优化（非硬性错误）", level=2)
    _add_bullets(
        doc,
        [
            "“抗体水平均显著升高”宜注明为接种后较基线或明确组间比较对象。",
            "安全性小结“征集性AE发生率分别为54.76%…”缺主语，宜改为“总体征集性AE发生率为…”。",
            "40~49岁“相对于安慰剂”可写清“阳性对照组1（含首剂安慰剂设计）”。",
        ],
    )

    doc.add_heading("四、总体评价", level=1)
    _add_table(
        doc,
        ["维度", "评价"],
        [
            ["核心数据一致性", "摘要、§2、§3主体与TFL趋势基本一致"],
            ["主要风险点", "摘要VZV笔误；讨论对≥60岁GMC过度概括；3.4错别字；率差09.74%"],
            ["结论力度", "支持III期方向合理，但应限定在本次数据范围并单列待报内容"],
        ],
    )

    doc.add_heading("五、优先修改清单（建议顺序）", level=1)
    _add_bullets(
        doc,
        [
            "修正摘要40~49岁VZV段 gE→VZV。",
            "重写§5讨论中50~59/≥60/≥50岁免疫原性概括（与§3.4、摘要对齐）。",
            "通读修正§3.4小结（GMT、阳性对照徐2、空格、h2≥60岁表述）。",
            "改正率差09.74%及≥50岁安全性“阳性对照组2”。",
            "结论/局限性补充：细胞免疫待报、长期随访、≥70岁比例、≥60岁GMC临床解读。",
            "统一术语（试验用疫苗、MedDRA、Fortaleza、CD4相关指标名）。",
        ],
    )

    doc.add_paragraph()
    _add_para(
        doc,
        "本报告由文档审核流程自动生成，具体数值修订请以受审文件最新版TFL及SAP为准。",
    )

    return doc


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "review_materials"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"YDSWX_phase_summary_document_review_{date.today().isoformat()}.docx"

    doc = build_document()
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
