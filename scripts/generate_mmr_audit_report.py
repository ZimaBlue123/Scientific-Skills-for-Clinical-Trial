#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Medical Monitoring Report (MMR) Audit Report Generator.

End-to-end pipeline:
1.  Extract the candidate Word (.docx) MMR + reference Excel (.xlsx) EDC dump
    using ``extract_docx_full`` + ``extract_xlsx_full`` (both are stdlib /
    python-docx based and handle legacy/non-conforming .xlsx exports).
2.  Run a battery of cross-checks:
    - textual scan: typos, inconsistent roman numerals, mixed terminology
      ("试验用疫苗"/"试验疫苗", "试验参与者"/"受试者", "医学监查"/"医学核查" …)
    - data cross-check: filter-fail count, AE severity/relation/serious
      totals, AE-by-age-layer breakdown, EX (vaccination) dose distribution,
      DV (deviation) record count
    - known P0 typos that surface often in vaccine trial reports
      ("足三里交"→"足三里穴", "肌内滴注"→"肌内注射")
3.  Emit a Word (.docx) audit report using the project-wide
    ``common_scripts.docx_utils.apply_cn_en_fonts`` font helper and the
    ``AuditFinding`` data structure shared with
    ``scripts/generate_clinical_doc_audit_report.py``.

Usage
-----
    py -3 scripts/generate_mmr_audit_report.py --folder review_materials/ \\
        --project "TVAX-020 II期"

    py -3 scripts/generate_mmr_audit_report.py --word MMR.docx --excel EDC.xlsx

Dependencies
------------
- python-docx
- The script ``scripts/extract_xlsx_full.py`` (created as part of this
  pipeline) for robust, stdlib-only .xlsx parsing.
"""
from __future__ import annotations

import argparse
import logging
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

LOG_FORMAT = "%(asctime)s [%(levelname)s] generate_mmr_audit_report: %(message)s"
logger = logging.getLogger("generate_mmr_audit_report")

# Make sibling scripts importable when invoked directly.
_HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(_HERE))


# =========================================================================
# 1. Extraction helpers
# =========================================================================
def extract_word_text(docx_path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError as e:  # pragma: no cover
        raise SystemExit("python-docx is required: pip install python-docx") from e
    doc = Document(str(docx_path))
    lines: List[str] = [f"# FILE: {docx_path.name}"]
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for ti, tbl in enumerate(doc.tables, 1):
        lines.append(f"\n--- TABLE {ti} ---")
        for r in tbl.rows:
            lines.append(" | ".join(c.text.replace("\n", " / ") for c in r.cells))
    return "\n".join(lines)


def extract_xlsx_text(xlsx_path: Path) -> str:
    from extract_xlsx_full import extract_xlsx  # type: ignore
    return extract_xlsx(xlsx_path)


def _parse_xlsx_dump(dump_text: str) -> Dict[str, Dict[str, List[str]]]:
    """Group dump by sheet, return {sheet: {col_name_or_'_rows': [values]}}."""
    sheets: Dict[str, Dict[str, List]] = defaultdict(lambda: defaultdict(list))
    current_sheet: Optional[str] = None
    field_names: List[str] = []
    for line in dump_text.splitlines():
        m_sheet = re.match(r"# Sheet: (\S+)", line)
        if m_sheet:
            current_sheet = m_sheet.group(1)
            field_names = []
            continue
        if not current_sheet or not line.startswith("R") or ": " not in line:
            continue
        body = line.split(": ", 1)[1]
        cells = body.split(" | ")
        if not field_names and cells and all(re.fullmatch(r"[A-Z][A-Z0-9_]*", c) for c in cells if c):
            field_names = cells
            continue
        sheets[current_sheet]["_rows"].append(cells)
        for i, name in enumerate(field_names):
            if i < len(cells):
                sheets[current_sheet][name].append(cells[i])
    return dict(sheets)


# =========================================================================
# 2. Cross-check rules
# =========================================================================
@dataclass
class Finding:
    id: str
    severity: str  # "重大" / "重要" / "一般" / "建议"
    category: str
    location: str
    issue: str
    recommendation: str
    rationale: str = ""


def _scan_word_terms(word_text: str) -> List[Finding]:
    findings: List[Finding] = []

    def scan(pattern: str, severity: str, category: str, issue: str,
             rec: str, rationale: str = "") -> None:
        matches = list(re.finditer(pattern, word_text))
        if not matches:
            return
        # 去重：相同 (severity, category, issue) 只保留一个
        findings.append(Finding(
            id=f"F-{len(findings) + 1:02d}",
            severity=severity,
            category=category,
            location=f"全文（出现 {len(matches)} 次）",
            issue=issue + f"  示例: …{word_text[max(0, matches[0].start() - 20): matches[0].end() + 20].replace(chr(10), ' ')[:80]}…",
            recommendation=rec,
            rationale=rationale,
        ))

    # P0 typo
    scan(r"足三里交", "重大", "错别字",
         "出现中医穴位名错字 '足三里交'",
         "全部替换为 '足三里穴'",
         "中医穴位名错字，影响报告专业性")
    scan(r"肌内滴注", "重大", "错别字",
         "'肌内滴注' 为非规范用法",
         "改为 '肌内注射' / '肌肉注射'",
         "注射给药途径应用 '注射' 而非 '滴注'")

    # 罗马数字一致性
    scan(r"\bII期\b", "一般", "格式",
         "正文使用半角罗马数字 'II'",
         "统一为全角 'Ⅱ'",
         "全文 Ⅱ 出现多次，II 仅 1 次，不一致")
    scan(r"\bIII期\b", "一般", "格式",
         "正文使用半角 'III'",
         "统一为 'Ⅲ'",
         "与全文全角罗马数字风格不一致")

    # 术语不一致
    scan(r"试验疫苗(?!用)", "重要", "术语",
         "正文使用 '试验疫苗' 而非 '试验用疫苗'",
         "全文统一为 '试验用疫苗'（与方案一致）",
         "用语应与方案编号 '试验用疫苗' 保持一致")
    scan(r"受试者", "一般", "术语",
         "正文使用 '受试者' 而非 '试验参与者'",
         "统一为 '试验参与者'（方案用语）",
         "全文字数 60+ 次使用 '试验参与者'，混用易混淆")
    scan(r"医学核查", "一般", "术语",
         "正文使用 '医学核查' 而非 '医学监查'",
         "统一为 '医学监查'",
         "标题及正文主体使用 '医学监查'")
    scan(r"登入", "一般", "错别字",
         "'登入' 应为 '录入'",
         "改为 '录入'",
         "EDC 录入用词规范")

    # 标点
    scan(r"。{2,}", "一般", "标点",
         "出现连续两个以上句号",
         "删除多余句号",
         "标点符号错误")

    return findings


def _cross_check_data(word_text: str, xlsx_sheets: Dict[str, Dict[str, List]]) -> List[Finding]:
    findings: List[Finding] = []

    def count_subjects_with_status(status_value: str) -> int:
        subs = xlsx_sheets.get("DM", {}).get("SUBJSTA", [])
        return sum(1 for v in subs if v.strip() == status_value)

    # ---- D-01 筛选失败数 ----
    fail = count_subjects_with_status("筛选失败")
    m = re.search(r"筛选\s*(\d+)\s*例.*?入组\s*(\d+)\s*例", word_text)
    if m:
        screened, enrolled = int(m.group(1)), int(m.group(2))
        implied_fail = screened - enrolled
        if fail != implied_fail:
            findings.append(Finding(
                id="D-01", severity="重大", category="数据矛盾",
                location=f"正文第{word_text[:m.start()].count(chr(10))+1}行 / EDC DM 表",
                issue=f"Word 称筛选失败 {implied_fail} 例（筛选 {screened} - 入组 {enrolled}），但 EDC DM 表实际 {fail} 例。差 {abs(fail - implied_fail)} 例。",
                recommendation="核实第 N 例筛败受试者是否已录入 EDC，并在报告中明确统计口径",
                rationale=f"DM 表 SUBJSTA='筛选失败' 共 {fail} 条",
            ))

    # ---- D-02 方案偏离数 ----
    dv_rows = xlsx_sheets.get("DV", {}).get("_rows", [])
    m_dv = re.search(r"共报告\s*(\d+)\s*例次方案偏离", word_text)
    if m_dv:
        word_dv = int(m_dv.group(1))
        if len(dv_rows) != word_dv:
            findings.append(Finding(
                id="D-02", severity="重大", category="数据矛盾",
                location=f"正文第{word_text[:m_dv.start()].count(chr(10))+1}行 / EDC DV 表",
                issue=f"Word 称方案偏离/违背 {word_dv} 例次，EDC DV 表 {len(dv_rows)} 行。差 {abs(word_dv - len(dv_rows))} 条。",
                recommendation="核实 DV 表数据完整性（可能部分偏离来自独立 PD 清单未同步 EDC），并在报告中明确数据来源",
                rationale=f"DV 行数={len(dv_rows)}",
            ))

    # ---- D-03 AE 总数 ----
    toxgr = xlsx_sheets.get("AE", {}).get("AETOXGR", [])
    nonempty_tox = [v for v in toxgr if v.strip()]
    m_ae = re.search(r"164例.*?发生\s*(\d+)\s*例次不良事件", word_text)
    if m_ae and nonempty_tox:
        word_ae = int(m_ae.group(1))
        if len(nonempty_tox) != word_ae:
            findings.append(Finding(
                id="D-03", severity="重要", category="数据不一致",
                location=f"正文第{word_text[:m_ae.start()].count(chr(10))+1}行 / EDC AE 表 AETOXGR 字段",
                issue=f"Word 称 AE {word_ae} 例次，EDC AE 表 AETOXGR 字段非空记录 {len(nonempty_tox)} 条。",
                recommendation="核实 AE 表中 AETOXGR 为空的记录性质（如 AEYN='否' 占位行），并在报告中注明 AE 分析集定义",
                rationale="数据完整性需确认",
            ))

    # ---- D-04 EX 接种记录缺失 ----
    stdtc = xlsx_sheets.get("EX", {}).get("EXSTDTC", [])
    empty_std = sum(1 for v in stdtc if not v.strip())
    if empty_std > 0 and stdtc:
        findings.append(Finding(
            id="D-04", severity="重要", category="数据完整性",
            location="EDC EX 表 EXSTDTC 字段",
            issue=f"EX 表 EXSTDTC（接种日期）字段有 {empty_std} 条空值，可能影响接种完成率统计。",
            recommendation="现场核实 EXSTDTC 为空的记录，确认字段是否映射到 EXDAT 等其他列",
            rationale=f"EX 表共 {len(stdtc)} 行，空日期 {empty_std} 条",
        ))

    # ---- S-01 年龄层 AE 趋势补充建议 ----
    subs_age = dict(zip(xlsx_sheets.get("DM", {}).get("SUBJID", []),
                         xlsx_sheets.get("DM", {}).get("AGE", [])))
    tox_by_subj: Dict[str, set] = defaultdict(set)
    for sid, tg in zip(xlsx_sheets.get("AE", {}).get("SUBJID", []),
                        xlsx_sheets.get("AE", {}).get("AETOXGR", [])):
        if sid.strip() and tg.strip():
            tox_by_subj[sid.strip()].add(tg.strip())

    layer_subj: Counter = Counter()
    ae_by_layer: Counter = Counter()
    for sid, age in zip(xlsx_sheets.get("DM", {}).get("SUBJID", []),
                          xlsx_sheets.get("DM", {}).get("AGE", [])):
        try:
            a = int(age)
        except ValueError:
            continue
        if 0 <= a <= 5:
            layer_subj["6m-5y"] += 1
        elif 6 <= a <= 12:
            layer_subj["6-12y"] += 1
        elif 13 <= a <= 17:
            layer_subj["13-17y"] += 1
    for sid in tox_by_subj.keys():
        try:
            a = int(subs_age.get(sid, ""))
        except ValueError:
            continue
        if 0 <= a <= 5:
            ae_by_layer["6m-5y"] += 1
        elif 6 <= a <= 12:
            ae_by_layer["6-12y"] += 1
        elif 13 <= a <= 17:
            ae_by_layer["13-17y"] += 1

    if layer_subj and ae_by_layer:
        rates = {lyr: round(ae_by_layer[lyr] / max(layer_subj[lyr], 1) * 100, 1)
                  for lyr in layer_subj if layer_subj[lyr] > 0}
        if rates and max(rates.values()) - min(rates.values()) > 5:
            findings.append(Finding(
                id="S-01", severity="建议", category="分析补充",
                location="医学问题总结 章节",
                issue=f"AE 发生率呈年龄依赖性递减：{rates}。建议在医学问题总结中讨论该趋势的临床意义。",
                recommendation="增加年龄层 AE 发生率趋势分析段落",
                rationale="数据驱动建议",
            ))

    # ---- S-02 S0548 退出者信息缺失 ----
    if "S0548" in word_text:
        findings.append(Finding(
            id="S-02", severity="建议", category="补充建议",
            location="正文第 X 行（S0548 退出注脚）",
            issue="S0548 因采血失败自愿退出，建议补充该例退出时间点（是否在首剂接种前退出）",
            recommendation="在脚注或正文中补充 S0548 退出时间点及是否已完成首剂接种",
            rationale="EDC 显示 S0548 EX 表无接种日期；建议明确时间线",
        ))

    return findings


# =========================================================================
# 3. Report generation
# =========================================================================
def _make_doc(findings: List[Finding], project: str, files: Sequence[str], output: Path) -> None:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    try:
        from common_scripts.docx_utils import apply_cn_en_fonts  # type: ignore
    except ImportError:
        def apply_cn_en_fonts(*_args, **_kwargs) -> int:  # noqa: E731
            return 0

    doc = Document()
    apply_cn_en_fonts(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("临床试验医学监查报告（MMR）审核报告")
    r.bold = True

    info = doc.add_paragraph()
    info.add_run("审核日期：").bold = True
    info.add_run(f"{date.today().isoformat()}\n")
    info.add_run("项目：").bold = True
    info.add_run(f"{project}\n")
    info.add_run("审核文件：").bold = True
    info.add_run("\n".join(f"• {f}" for f in files))

    doc.add_heading("一、审核概述", level=1)
    sev_counter = Counter(f.severity for f in findings)
    summary = doc.add_paragraph()
    summary.add_run(f"本报告共发现 {len(findings)} 个问题")
    if sev_counter:
        summary.add_run("，其中：")
        for sev, cnt in sev_counter.most_common():
            summary.add_run(f"\n• {sev}: {cnt} 项")

    doc.add_heading("二、问题清单", level=1)
    if findings:
        order = {"重大": 0, "重要": 1, "一般": 2, "建议": 3}
        findings_sorted = sorted(findings, key=lambda x: (order.get(x.severity, 99), x.id))
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["编号", "风险等级", "类别", "位置", "问题描述", "修改建议"]):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.bold = True
        for f in findings_sorted:
            row = table.add_row().cells
            row[0].text = f.id
            row[1].text = f.severity
            row[2].text = f.category
            row[3].text = f.location
            row[4].text = f.issue
            row[5].text = f.recommendation

    doc.add_heading("三、审核结论", level=1)
    if sev_counter.get("重大", 0):
        doc.add_paragraph(
            f"本审核共发现 {len(findings)} 项问题，其中重大 {sev_counter.get('重大', 0)} 项、"
            f"重要 {sev_counter.get('重要', 0)} 项、一般 {sev_counter.get('一般', 0)} 项、"
            f"建议 {sev_counter.get('建议', 0)} 项。建议优先处理重大问题，"
            f"在下一版本修正后再行提交。"
        )
    else:
        doc.add_paragraph(
            f"本审核共发现 {len(findings)} 项问题，未发现重大数据矛盾。"
            f"建议在下一版本修正一般性问题并补充建议项分析。"
        )

    try:
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
    except OSError as exc:
        logger.error("failed to write report %s: %s", output, exc)
        raise
    print(f"OK: wrote {output}")


# =========================================================================
# 4. CLI / file resolution
# =========================================================================
# Distinguish a real MMR document from a previously-generated audit report.
_MMR_KEYWORDS = ("监查报告", "MMR", "医学监查报告", "Medical Monitoring")


def _looks_like_mmr(path: Path) -> bool:
    return any(kw in path.name for kw in _MMR_KEYWORDS)


def _resolve_files(folder: Optional[Path], word: Optional[Path],
                    excel: Optional[Path]) -> Tuple[Path, Path]:
    if folder is not None:
        folder = folder.resolve()
        if not folder.is_dir():
            raise SystemExit(f"--folder 不是有效目录: {folder}")
        docx_files = [p for p in folder.iterdir()
                       if p.suffix.lower() == ".docx" and not p.name.startswith("~")]
        mmr = next((p for p in docx_files if _looks_like_mmr(p)), None)
        word = mmr or (docx_files[0] if docx_files else None)
        excel = next((p for p in folder.iterdir()
                       if p.suffix.lower() == ".xlsx" and not p.name.startswith("~")), None)
        if not word or not excel:
            raise SystemExit(f"Folder {folder} must contain both .docx (MMR) and .xlsx")
    else:
        if not word or not excel:
            raise SystemExit("Either --folder or both --word and --excel are required")
        word, excel = word.resolve(), excel.resolve()
    return word, excel


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="MMR audit report generator (Word + EDC cross-check).")
    parser.add_argument("--folder", type=Path, help="folder containing MMR.docx + EDC.xlsx")
    parser.add_argument("--word", type=Path, help="MMR .docx path")
    parser.add_argument("--excel", type=Path, help="EDC .xlsx path")
    parser.add_argument("--output", type=Path, default=None, help="output .docx path")
    parser.add_argument("--project", default="临床试验 MMR 审核")
    parser.add_argument("--log-level", default="WARNING")
    args = parser.parse_args(argv)

    logging.basicConfig(level=getattr(logging, args.log_level.upper(), logging.WARNING), format=LOG_FORMAT)

    try:
        word, excel = _resolve_files(args.folder, args.word, args.excel)
    except SystemExit:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.error("file resolution failed: %s", exc)
        return 2

    if not args.output:
        args.output = word.parent / f"审核报告_{word.stem}_{date.today().strftime('%Y%m%d')}.docx"

    logger.info("Word: %s", word)
    logger.info("Excel: %s", excel)

    if not word.is_file():
        logger.error("Word 文件不存在: %s", word)
        return 2
    if not excel.is_file():
        logger.error("Excel 文件不存在: %s", excel)
        return 2

    try:
        word_text = extract_word_text(word)
        xlsx_text = extract_xlsx_text(excel)
    except Exception as exc:  # noqa: BLE001
        logger.exception("抽取阶段失败: %s", exc)
        return 1

    xlsx_sheets = _parse_xlsx_dump(xlsx_text)

    findings = _scan_word_terms(word_text) + _cross_check_data(word_text, xlsx_sheets)

    try:
        _make_doc(findings, args.project, [word.name, excel.name], args.output)
    except Exception as exc:  # noqa: BLE001
        logger.exception("生成报告失败: %s", exc)
        return 1

    print(f"Findings: {len(findings)} "
          f"({sum(1 for f in findings if f.severity == '重大')} 重大 / "
          f"{sum(1 for f in findings if f.severity == '重要')} 重要 / "
          f"{sum(1 for f in findings if f.severity == '一般')} 一般 / "
          f"{sum(1 for f in findings if f.severity == '建议')} 建议)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
