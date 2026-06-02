#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Word review report for CTD 2.5 clinical overview document."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _apply_cn_en_fonts(doc: Document) -> None:
    def set_style(style_name: str) -> None:
        if style_name not in doc.styles:
            return
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:cs"), "Times New Roman")

    for name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Table Grid"):
        set_style(name)


def _set_run_font(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "Times New Roman")


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run)
    run.bold = bold


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                _set_run_font(r)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font(r)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    _apply_cn_en_fonts(doc)

    today = date.today().strftime("%Y年%m月%d日")
    src_doc = (
        "《重组带状疱疹疫苗（CHO细胞）（TVAX-006）2.5 临床综述》"
        " V1.0（版本日期 2026年05月22日；清洁版文件名日期 2026年05月25日）"
    )

    title = doc.add_heading("2.5 临床综述文档审核报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_para(doc, f"审核日期：{today}")
    _add_para(doc, f"受审文件：{src_doc}")
    _add_para(
        doc,
        "审核范围：全文通读（封面与版本说明、2.5.1~2.5.7、有效性/安全性数据叙述、"
        "有效性评价、获益-风险结论及嵌入表格脚注）。",
    )
    _add_para(
        doc,
        "审核方法：将 markitdown 提取的全文与 II 期主要结果表（表4~7、表60~67等）、"
        "药物暴露叙述、获益-风险章节进行交叉核对。",
    )
    doc.add_paragraph()

    doc.add_heading("审核说明", level=1)
    _add_para(
        doc,
        "本报告针对注册申报用 CTD 2.5 临床综述，从数据与表述一致性、应补充的分析性结论、"
        "以及语言文字与格式规范三方面提出审阅意见。所列位置以 Word 章节标题或表格编号定位，"
        "修订时请以最新版 Word 与 TFL 为准。",
    )

    # --- Section 1 ---
    doc.add_heading("一、前后矛盾或易误导的数据与表述", level=1)

    doc.add_heading("1. 有效性评价：中老年层免疫原性概括与统计结果矛盾（高优先级）", level=2)
    _add_para(doc, "位置：2.5.4.6 有效性评价，关于 50~59 岁、≥60 岁与阳性对照组 2 的概括段。")
    _add_bullets(
        doc,
        [
            "问题：文中写“第2剂接种前和接种后30天，两组 gE、VZV 的 SCR、校正后 GMC 及 GMI "
            "均无显著差异（P>0.05）”。但同文档 2.5.4.5 主要有效性结果中，≥60 岁、PPS-h2、"
            "第2剂免后30天：试验组 gE LS GMC 显著低于阳性对照组 2（校正比值 0.80，P=0.011）；"
            "VZV LS GMC 显著低于阳性对照组 2（比值 0.83，P=0.030）。该句与后文“GMC 虽显著低于"
            "阳性对照组 2”亦存在同段自相矛盾。",
            "建议：按年龄层、访视时点（h1/h2）、终点（SCR/GMC/GMI）分别表述；≥60 岁 h2 建议写为"
            "“SCR 与 GMI 组间相当，校正 GMC 低于欣安立适®（阳性对照组 2），差异有统计学意义”，"
            "并说明与预设非劣效/优效标准及临床意义的关系（若方案或 SAP 有规定）。",
        ],
    )

    doc.add_heading("2. II 期药物暴露：入组例数与组分描述错误（高优先级）", level=2)
    _add_para(doc, "位置：2.5.5.1 药物暴露程度，YDSWX（TVAX-006）-002（Ⅱ）段落。")
    _add_table(
        doc,
        ["项目", "正文现写", "与表3/设计应一致", "风险"],
        [
            [
                "阳性对照组 1 例数",
                "阳性对照组 170 例",
                "40~49 岁阳性对照组 1 应为 70 例",
                "合计变为 520 例，与“入组 420 例”矛盾",
            ],
            [
                "试验组疫苗组成",
                "仅写 QS-21、CpG-QCX1",
                "应含 gE 50 μg + TVA01 佐剂（见表 3 注释）",
                "暴露描述不完整，影响审评对试验药理解",
            ],
            [
                "阳性对照组 1 接种剂次",
                "共 1 剂",
                "第 0 天安慰剂 + 第 60 天感维®，应为 2 次接种",
                "与方案设计不一致",
            ],
        ],
    )

    doc.add_heading("3. 试验状态与数据范围表述需统一（中~高优先级）", level=2)
    _add_bullets(
        doc,
        [
            "多处标注“试验进行中”，但已写入完整 II 期 PPS-h2 免疫原性及 0~30 天安全性汇总，"
            "建议在 2.5.4/2.5.5 开篇明确：数据截止日、分析集、是否盲态/阶段性锁库，以及与"
            "长期随访（12/24/36 个月体液免疫、细胞免疫亚组）的关系。",
            "方案要求 ≥60 岁层 ≥70 岁者不少于该层 20%，正文未见 II 期是否达标的汇总；"
            "若已入组完成，建议在研究人群或局限性中补充。",
            "II 期设计含细胞免疫亚组 90 例及 12/24/36 个月访视，本次有效性仅至第 2 剂后 30 天"
            "体液免疫；获益-风险（2.5.6.4）写“激发体液免疫和细胞免疫……各年龄层保护效力相当”"
            "易被理解为细胞免疫与长期保护已完成验证，与当前数据范围不符。",
        ],
    )

    doc.add_heading("4. 对照设计差异未在结论中限定（中优先级）", level=2)
    _add_bullets(
        doc,
        [
            "40~49 岁：试验组 vs 阳性对照组 1（首剂安慰剂 + 第 60 天感维®），免疫原性显著优于对照。",
            "50~59 岁、≥60 岁：试验组 vs 阳性对照组 2（欣安立适®），不宜笼统写“与阳性对照相当”"
            "或“满足免疫保护需求”，应区分 SCR/GMI 相当与 ≥60 岁 GMC 低于对照。",
            "2.5.4.3 非劣效界值写“不适用”，但若注册路径需与 Shingrix 比较，建议在综述中说明"
            "免疫原性比较属探索性/描述性，或引用方案/SAP 中的预设标准。",
        ],
    )

    doc.add_heading("5. 结构/编号与章节嵌入问题（中优先级）", level=2)
    _add_bullets(
        doc,
        [
            "2.5.4.1 研究人群中，II 期排除标准以“6.”嵌套在 I 期主要排除标准列表内，"
            "易造成 I/II 期标准混淆，建议拆分为独立小节。",
            "2.5.5.1 药物暴露程度下，I 期与 II 期均使用“1.”编号，目录层级重复。",
            "封面“历次版本变更摘要”为 V1.0/2026-05-22，文件名含 20260525，建议在版本说明中"
            "注明清洁版与版本表关系。",
        ],
    )

    doc.add_heading("6. 表格脚注重复粘贴（低~中优先级）", level=2)
    _add_para(
        doc,
        "位置：表 4、表 6、表 60、表 64 等脚注。",
    )
    _add_bullets(
        doc,
        [
            "“Clopper-PearsonClopper-Pearson”“精确法精确法”等为重复录入，影响专业呈现。",
            "建议统一为“Clopper-Pearson 方法”“精确法”并全文检索同类脚注。",
        ],
    )

    # --- Section 2 ---
    doc.add_heading("二、建议补充的分析结论", level=1)

    doc.add_heading("1. 建议在 2.5.4.6 或 2.5.6 中明确写明的局限性", level=2)
    _add_bullets(
        doc,
        [
            "II 期细胞免疫亚组（90 例）结果尚未纳入本综述；预计补充时间及是否影响 III 期决策。",
            "长期免疫原性（12/24/36 个月）为方案计划访视，本次仅呈现至第 2 剂后 30 天。",
            "SAE/AESI/妊娠观察至全程接种后 12 个月，安全性汇总若以 0~30 天为主，应说明"
            "更长随访期数据状态。",
            "≥60 岁第 2 剂后 GMC 低于欣安立适而 SCR/GMI 相当：补充与方案预设标准的关系、"
            "及文献/说明书语境下的临床解读（避免仅写“免疫功能衰退”推测）。",
            "≥70 岁占 ≥60 岁层 20% 的入组要求是否满足（建议给出例数或比例）。",
        ],
    )

    doc.add_heading("2. 建议加强的分析性结论", level=2)
    _add_table(
        doc,
        ["主题", "建议补充内容"],
        [
            [
                "40~49 岁",
                "明确相对首剂安慰剂设计下，第 2 剂前/后免疫应答；说明感维® 在第 60 天接种后"
                "对照 SCR 仍低的原因及对解读的影响",
            ],
            [
                "≥50 岁",
                "分述与欣安立适 SCR/GMI 相当；≥60 岁 GMC 差异的统计学与临床意义",
            ],
            [
                "安全性",
                "40~49 岁试验组 AE 高于对照 1 与佐剂/重组疫苗预期一致；3 级发热短暂，"
                "可在获益-风险中单列 benefit-risk 表述",
            ],
            [
                "I 期",
                "剂量探索结论（低/高剂量相当）与 II 期剂量选择的衔接一句",
            ],
            [
                "注册定位",
                "2.5.6 获益不宜仅引用上市疫苗人群研究，应扣回本品 I/II 期实际数据范围",
            ],
        ],
    )

    doc.add_heading("3. 2.5.6 获益-风险评估（宜加限定语）", level=2)
    _add_para(
        doc,
        "建议将综合结论调整为：基于现有 I/II 期至第 2 剂后 30 天体液免疫原性及短期安全性，"
        "支持继续临床开发；细胞免疫、长期随访及保护效力（如 III 期）待后续数据补充。"
        "避免在未呈现细胞免疫结果时写“各年龄层保护效力相当”。",
    )

    # --- Section 3 ---
    doc.add_heading("三、语句、术语与错别字问题", level=1)

    doc.add_heading("明确错误（建议修改）", level=2)
    _add_table(
        doc,
        ["位置", "问题", "建议"],
        [
            ["2.5.5.1 II 期暴露", "阳性对照组 170 例", "阳性对照组 1，70 例"],
            ["2.5.5.1 II 期暴露", "试验组仅写佐剂成分", "补充 gE 50 μg + TVA01 佐剂"],
            ["2.5.5.1 II 期暴露", "阳性对照组 1“共 1 剂”", "第 0 天安慰剂 + 第 60 天感维®，共 2 次接种"],
            ["2.5.4.6 有效性评价", "中老年“均无显著差异”", "按终点拆分；≥60 岁 h2 写明 GMC P=0.011/0.030"],
            ["2.5.6.1.2 / 2.5.6.4", "Zostava", "Zostavax 或统一商品名"],
            ["2.5.6.1.2", "HZ 的效率", "效力"],
            ["感维说明书转录", "失眼", "失眠"],
            ["感维说明书转录", "上连不良反应", "上述"],
            ["2.5.6.4", "由于这些事件由是", "由于这些事件是由于"],
            ["2.5.5.1 I 期", "gE蛋白）50μg", "gE 蛋白 50 μg（括号配对）"],
            ["表脚注多处", "Clopper-Pearson 重复", "删除重复词"],
        ],
    )

    doc.add_heading("用语与格式建议统一", level=2)
    _add_bullets(
        doc,
        [
            "“试验参与者”与“受试者”混用（如入选标准写“受试者”），申报资料宜统一为“受试者”或“试验参与者”。",
            "I 期叙述中 GMT/LS GMT 与 II 期 GMC/LS GMC 并存，建议在首次出现处说明检测指标与命名。",
            "目录链接保留“目 录 3”等占位页码，定稿前需更新域代码。",
            "2.5.6 开篇保留“（盖章）”字样，正式排版时移至签章页。",
            "文件名“重组带疱疫苗”应为“重组带状疱疹疫苗”（与正文一致）。",
        ],
    )

    doc.add_heading("表述可优化（非硬性错误）", level=2)
    _add_bullets(
        doc,
        [
            "“抗体水平均显著升高”宜注明相对基线或相对对照的比较对象。",
            "“能满足该人群的免疫保护需求”属效力外推，建议改为“体液免疫应答与对照相当/更优”等可核对表述。",
            "I 期安全性亚层分析标题为“40~59岁”，内容为五组设计，与 II 期年龄层分析并列时建议加小标题区分研究编号。",
        ],
    )

    # --- Section 4 & 5 ---
    doc.add_heading("四、总体评价", level=1)
    _add_table(
        doc,
        ["维度", "评价"],
        [
            [
                "数据呈现完整性",
                "I/II 期免疫原性表与叙述主体较完整，II 期 PPS-h1/h2 与表 4~7、60~67 可对应",
            ],
            [
                "主要风险点",
                "有效性评价段笼统“无显著差异”；II 期暴露例数 170 与疫苗描述错误；获益-风险超出当前数据范围",
            ],
            [
                "注册可用性",
                "修订上述高优先级项并补充局限性后，更适合作为 2.5 临床综述定稿基础",
            ],
        ],
    )

    doc.add_heading("五、优先修改清单（建议顺序）", level=1)
    _add_bullets(
        doc,
        [
            "更正 2.5.5.1 II 期入组例数（70+140+210=420）及试验组/阳性对照组 1 接种方案描述。",
            "重写 2.5.4.6 中 50~59/≥60 岁免疫原性概括，与 2.5.4.5 及 P 值一致。",
            "在 2.5.4/2.5.6 增加数据截止、细胞免疫与长期随访待报说明。",
            "收紧 2.5.6.4 获益-风险表述，避免未验证的“保护效力相当”。",
            "通读修正错别字（Zostava、失眼、上连、由是、效率等）及表脚注重复。",
            "理顺 2.5.4.1  I/II 期入选排除标准结构与 2.5.5.1 编号。",
        ],
    )

    doc.add_paragraph()
    _add_para(
        doc,
        "本报告由文档审核流程生成；数值与统计结论修订请以受审文件最新版 Word、"
        "统计分析报告及 TFL 为准。",
    )

    return doc


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "review_materials"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"TVAX-006_2.5临床综述_文档审核报告_{date.today().isoformat()}.docx"

    doc = build_document()
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
