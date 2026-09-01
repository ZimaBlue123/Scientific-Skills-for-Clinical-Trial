"""
edit_office_utils.py
Unified editing utilities for DOCX files to prevent code duplication in future scripts.
"""

from docx.text.paragraph import Paragraph
from docx.table import _Cell
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.document import Document
from docx.table import Table

def make_run(paragraph, text, bold=False, italic=False, underline=False, font_name=None, font_size=None):
    """Create a new run in a paragraph with specified formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    return run

def set_cell_text(cell: _Cell, text: str):
    """Set the text of a cell, clearing any existing paragraphs."""
    cell.text = ""  # Clears all paragraphs except the first empty one
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.add_paragraph(text)

def iter_paragraphs(parent):
    """Yield all paragraphs in a document, including those in tables and nested tables."""
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)

def replace_across_runs(paragraph, search_text, replace_text):
    """
    Naively replaces search_text with replace_text within a paragraph's text.
    Warning: This destroys individual run formatting and merges all text into a single run.
    For precise preservation, more advanced logic is needed.
    """
    if search_text in paragraph.text:
        new_text = paragraph.text.replace(search_text, replace_text)
        paragraph.clear()
        paragraph.add_run(new_text)
        return True
    return False

def clear_cell(cell: _Cell):
    """Clears all content from a cell."""
    cell.text = ""
