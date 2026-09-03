# -*- coding: utf-8 -*-
"""
TVAX-009 糖尿病病史亚组免疫原性计算
- 从 清单1(第8册) 提取 30 例糖尿病受试者
- 从 清单2(第9册) 免疫原性清单(FAS/PPS) 提取每例各访视抗-HBs浓度与阳转标志
- 计算 7 组 x 有/无糖尿病 的 阳转率(N/n/%) 与 GMC(例数/GMC)
"""
import json, math, os
from collections import defaultdict
from docx import Document

BASE = "E:/Cursor Project/2-Scientific-Skills-for-Clinical_Trial/review_materials"
F1 = os.path.join(BASE, "YDSWX(TVAX-009)-002(Ⅱ)-基础阶段-统计分析报告-第8册(清单1).docx")
F2 = os.path.join(BASE, "YDSWX(TVAX-009)-002(Ⅱ)-基础阶段-统计分析报告-第9册(清单2).docx")

# ---------- 1. 提取糖尿病受试者 ----------
d1 = Document(F1)
t_disease = d1.tables[10]  # 16.2.4.6 既往疾病/现病史
diabetes = set()
diabetes_detail = []
for r in t_disease.rows[1:]:
    cells = [c.text.strip() for c in r.cells]
    if len(cells) < 7:
        continue
    if "糖尿病" in cells[6]:
        age = cells[0]; grp = cells[1]; sid = cells[2]
        key = (grp, sid)
        diabetes.add(key)
        diabetes_detail.append({"年龄组": age, "组别": grp, "研究编号": sid,
                                "疾病名称": cells[6], "开始日期": cells[7]})

print("=== 糖尿病受试者(去重后) ===")
print("count:", len(diabetes))
from collections import Counter
gc = Counter(k[0] for k in diabetes)
for g in sorted(gc):
    print(f"  {g}: {gc[g]}")

# ---------- 2. 提取免疫原性清单 ----------
VISIT_MAP = {
    "首剂接种前": "M0",
    "首剂接种后1个月": "M1",
    "首剂接种后2个月": "M2",
    "首剂接种后3个月": "M3",
    "首剂接种后4个月": "M4",
    "首剂接种后5个月": "M5",
    "首剂接种后6个月": "M6",
    "首剂接种后7个月": "M7",
    "首剂接种后8个月": "M8",
}

def parse_conc(s):
    s = s.strip()
    if not s:
        return None
    if s.startswith("<"):
        return 1.00  # LLOQ=2.00 -> LLOQ/2 = 1.00
    try:
        return float(s)
    except:
        return None

def load_listing(doc, tindex):
    """返回 dict[(组别,研究编号)] -> dict[visit_Mx] -> {conc, yangzhuan, yangxing}"""
    t = doc.tables[tindex]
    data = defaultdict(dict)
    for r in t.rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if len(cells) < 11:
            continue
        grp = cells[1]; sid = cells[2]; visit_raw = cells[5]
        if not sid or not grp or grp.startswith("("):
            # 跳过脚注行/空行
            continue
        m = VISIT_MAP.get(visit_raw)
        if not m:
            continue
        conc = cells[7]
        yangxing = cells[8]
        yangzhuan = cells[9]
        key = (grp, sid)
        data[key][m] = {
            "conc": parse_conc(conc),
            "conc_raw": conc,
            "yangxing": yangxing,
            "yangzhuan": yangzhuan,
        }
    return data

d2 = Document(F2)
print("tables in 第9册:", len(d2.tables))
fas = load_listing(d2, 8)   # 表16.2.6.1 免疫原性清单(FAS)
pps = load_listing(d2, 9)   # 表16.2.6.2 免疫原性清单(PPS)

print("FAS 唯一受试者:", len(fas), " PPS 唯一受试者:", len(pps))

# ---------- 组别定义 ----------
GROUPS = [
    ("A1", "0,1月低剂量组(A1)", "0,1月程序"),
    ("A2", "0,1月高剂量组(A2)", "0,1月程序"),
    ("B1", "0,2月低剂量组(B1)", "0,2月程序"),
    ("B2", "0,2月高剂量组(B2)", "0,2月程序"),
    ("C1", "阳性对照组(C1)", "0,1,6月程序"),
    ("C3", "0,1,6月高剂量组(C3)", "0,1,6月程序"),
    ("C2", "阳性对照组(C2)", "0,1,6月程序"),
]
# 各程序适用的绝对时间点
PROG_TIMEPOINTS = {
    "0,1月程序": ["M1","M2","M3","M6","M7","M8"],
    "0,2月程序": ["M1","M2","M3","M4","M7","M8"],
    "0,1,6月程序": ["M1","M2","M6","M7","M8"],
}
ALL_TP = ["M1","M2","M3","M4","M6","M7","M8"]

def gmean(vals):
    if not vals:
        return None
    vals = [v for v in vals if v is not None and v > 0]
    if not vals:
        return None
    return math.exp(sum(math.log(v) for v in vals) / len(vals))

def compute(data):
    """按组 x 有/无糖尿病 x 时间点 计算"""
    result = {}
    for code, grp, prog in GROUPS:
        # 该组所有受试者
        subs = {k for k in data if k[0] == grp}
        di = {k for k in subs if k in diabetes}
        nondi = subs - di
        result[code] = {"group": grp, "prog": prog,
                        "有糖尿病_N": len(di), "无糖尿病_N": len(nondi),
                        "有": {}, "无": {}}
        for tag, sset in (("有", di), ("无", nondi)):
            for m in ALL_TP:
                # 阳转率
                N = 0; n = 0
                gmc_vals = []
                for k in sset:
                    rec = data[k].get(m)
                    if rec:
                        yz = rec["yangzhuan"]
                        if yz in ("是", "否"):
                            N += 1
                            if yz == "是":
                                n += 1
                        c = rec["conc"]
                        if c is not None:
                            gmc_vals.append(c)
                pct = round(n / N * 100, 2) if N else None
                g = gmean(gmc_vals)
                g = round(g, 2) if g is not None else None
                result[code][tag][m] = {
                    "N_阳转": N, "n_阳转": n, "pct": pct,
                    "N_gmc": len(gmc_vals), "gmc": g,
                }
    return result

fas_res = compute(fas)
pps_res = compute(pps)

# ---------- 打印 ----------
def show(res, label):
    print(f"\n########## {label} ##########")
    for code, _, _ in GROUPS:
        r = res[code]
        print(f"\n== {code} {r['group']}  有糖尿病={r['有糖尿病_N']} 无糖尿病={r['无糖尿病_N']}")
        for tag in ("有", "无"):
            print(f"  [{tag}糖尿病] ", end="")
            for m in ALL_TP:
                d = r[tag][m]
                print(f"{m}:N={d['N_阳转']}/n={d['n_阳转']}/{d['pct']}% GMCn={d['N_gmc']}/GMC={d['gmc']} | ", end="")
            print()

show(fas_res, "FAS")
show(pps_res, "PPS")

# ---------- 保存 ----------
out = {"diabetes": sorted([list(k) for k in diabetes]),
       "diabetes_detail": diabetes_detail,
       "fas": {g: fas_res[g] for g in [x[0] for x in GROUPS]},
       "pps": {g: pps_res[g] for g in [x[0] for x in GROUPS]},
       "groups": [list(x) for x in GROUPS],
       "timepoints": ALL_TP,
       "prog_timepoints": PROG_TIMEPOINTS}
with open(os.path.join(BASE, "diabetes_immuno_result.json"), "w", encoding="utf-8") as f:
    json.dump(out, f, ensure_ascii=False, indent=1)
print("\nJSON saved to", os.path.join(BASE, "diabetes_immuno_result.json"))
