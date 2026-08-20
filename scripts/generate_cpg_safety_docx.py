"""
国内外 CpG 佐剂预防性疫苗安全性数据汇总 - Word 文档生成脚本
输出: review_materials/国内外CpG佐剂预防性疫苗安全性数据汇总_20260820.docx
版式: 无封面/无目录/无章节编号; 横向 A4; 三模块网格表, 按地域分块
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_PATH = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\国内外CpG佐剂预防性疫苗安全性数据汇总_20260820.docx"

# ---------------------------------------------------------------- data ---
# 模块一: (标识, 疫苗名称, 适应症, 地域, 分期, 设计, 样本量, 分组, 研究终点)
TRIALS = [
    # ---- 国外(已上市/已发表) ----
    (
        "NCT00435812 / NCT01005407 / NCT02117934",
        "HEPLISAV-B (HBsAg-1018)",
        "乙型肝炎 (HBV)",
        "国际",
        "上市(FDA 2017); III期×3",
        "随机、观察者盲、活性对照(Engerix-B)、多中心",
        "14,238 (10,038疫苗/4,200对照, 三试验合并)",
        "2剂(0,4周)+24周安慰剂 vs Engerix-B 3剂(0,4,24周)",
        "主: 28周血清保护率(anti-HBs≥10 mIU/mL); 次: 安全性",
    ),
    (
        "NCT01770743",
        "AV7909 (AVA+CpG 7909, NuThrax)",
        "炭疽(暴露后预防)",
        "国际",
        "II期(已上市 FDA 2023)",
        "随机、双盲、BioThrax活性对照、多中心",
        "168",
        "AV7909 0/14d, 0/28d, 0/14/28d, 半剂量0/14/28d 各若干例; BioThrax 0/14/28d",
        "主: D63 TNA应答率(免疫原性); 次: AE/SAE/反应原性发生率",
    ),
    (
        "NCT04672395",
        "SCB-2019 (S-Trimer + CpG 1018 + Alum)",
        "COVID-19 (SARS-CoV-2)",
        "国际(含中国)",
        "II/III期",
        "多中心、双盲、随机、安慰剂对照(5国)",
        "30,174入组; 30,128接种 (15,064/15,064)",
        "SCB-2019 2剂(0,21天) vs 安慰剂",
        "主: 任何严重程度COVID-19保护效力; 次: 中重度/重症保护效力、安全性",
    ),
    (
        None,
        "BK-SE36/CpG (CpG-ODN K3 + SE36 + Alum)",
        "疟疾(恶性疟原虫)",
        "国际",
        "Ib期",
        "单中心、双盲、随机、对照、年龄降阶梯",
        "135 (91疫苗/44对照)",
        "BK-SE36/CpG 3剂(0,28,112天) vs 狂犬疫苗对照; 队列: 21-45岁/5-10岁/12-24月龄",
        "主: 安全性与反应原性; 次: 免疫原性",
    ),
    (
        "NCT03172975 (CHHI)",
        "Na-GST-1/Alhydrogel ± CpG 10104",
        "钩虫(美洲钩虫)",
        "国际",
        "I期(Phase 1)",
        "随机、观察者盲、对照、剂量递增、单中心",
        "24 (CpG组16/无CpG组8)",
        "30μg+500μg CpG / 100μg / 100μg+500μg CpG 各8例; 3剂(0,56,112天)",
        "主: 安全性/耐受性; 次: 抗原特异性IgG",
    ),
    (
        "NCT05506969",
        "rF1V + CpG 1018",
        "鼠疫(肺鼠疫)",
        "国际",
        "II期(已完成)",
        "随机、活性对照、观察者盲、多中心",
        "200",
        "rF1V-1018 2剂(0,4周) vs rF1V单疫苗 3剂",
        "主: 反应原性与安全性(D1-W56); 数据未公开",
    ),
    (
        None,
        "H5N1 流感疫苗 + CpG 1018",
        "流感(H5N1大流行)",
        "国际",
        "I/II期(Part 1完成)",
        "随机、活性对照",
        "101 (18-49岁)",
        "单剂/双剂多种CpG 1018配方",
        "主: 安全性与免疫原性; Part 2数据预计2026",
    ),
    (
        None,
        "Z-1018 带状疱疹疫苗 + CpG 1018",
        "带状疱疹(VZV)",
        "国际",
        "I/II期(Part 2进行中)",
        "头对头 vs Shingrix (70岁+)",
        "待披露",
        "Z-1018 vs AS01B佐剂Shingrix",
        "主: 安全性与免疫原性; 数据未公开",
    ),
    # ---- 国内(已发表) ----
    (
        None,
        "重组乙型肝炎疫苗(汉逊酵母, CpG ODN)",
        "乙型肝炎 (HBV)",
        "国内",
        "I期(完成, 2020发表)",
        "随机、双盲、对照、单中心(浙江开化)",
        "48 (24/24)",
        "CpG ODN 250μg试验组 vs 市售铝佐剂对照(大连汉信); 0-1-6月",
        "主: 安全性; 次: anti-HBs阳转率/GMC",
    ),
    (
        "NCT04954131",
        "SCB-2019 中国II期 (CpG 1018 + Alum)",
        "COVID-19 (SARS-CoV-2)",
        "国内",
        "II期",
        "双盲、随机、安慰剂对照(山西太原/四川成都CDC)",
        "766",
        "SCB-2019(30μg抗原+1.5mg CpG 1018+0.75mg Alum) 2剂 vs 安慰剂",
        "主: 免疫原性; 次: 安全性(SAE/AESI与安慰剂无差异)",
    ),
    (
        "NCT04982068",
        "ZR202-CoV (CpG 7909 + Alum)",
        "COVID-19 (SARS-CoV-2)",
        "国内",
        "I期(完成, 2023发表)",
        "随机、双盲、安慰剂对照、剂量递增(河南许昌)",
        "72 (18/18/18/18)",
        "A:50μg+250μg CpG / B:25μg+500μg CpG / C:50μg+500μg CpG / 安慰剂; 2剂(0,28天)",
        "主: 安全性/耐受性; 次: 免疫原性",
    ),
    # ---- 国内(进行中, 数据未公开) ----
    (
        "NCT06851832",
        "重组带状疱疹疫苗 (gE + CpG 1018 + MF59)",
        "带状疱疹(VZV)",
        "国内",
        "I/II期(进行中)",
        "随机、双盲、活性/安慰剂对照、多中心",
        "960",
        "低剂量gE50μg+MF59+CpG1018 50μg / 标准gE50μg+MF59+CpG1018 100μg / 减毒活疫苗 / Shingrix配方 / 佐剂对照 / 安慰剂",
        "主: 各剂0-14天征集性AE、0-30天非征集性AE、SAE/AESI; 未公开",
    ),
    (
        "CXSL2500371",
        "华普 HP2001 带状疱疹疫苗 (HP007)",
        "带状疱疹(VZV)",
        "国内",
        "I期(2026-03启动)",
        "双盲、2剂、40岁+",
        "待披露",
        "HP007佐剂带状疱疹疫苗 2剂",
        "安全性/免疫原性; 未公开",
    ),
    (
        "CTR20261430",
        "华普 HP2002 乙肝疫苗 (CpG+铝)",
        "乙型肝炎 (HBV)",
        "国内",
        "I期(2026-04启动)",
        "随机、双盲、平行对照、单中心",
        "待披露",
        "不同CpG配比浓度, 0-1-6月3剂, 18岁+",
        "安全性/耐受性(主); 免疫原性(探索); 未公开",
    ),
    (
        "CTR20233405",
        "重组乙型肝炎疫苗(汉逊酵母, CpG和铝佐剂)",
        "乙型肝炎 (HBV)",
        "国内",
        "I期(2023-11公示, 招募完成)",
        "单中心、随机、双盲、对照、平行分组",
        "120 (目标)",
        "不同CpG-QCX1佐剂配比浓度, 0-1-6月3剂, 18岁+",
        "主: 安全性和耐受性; 次: 免疫原性; 未公开",
    ),
    (
        "ChiCTR2500108408",
        "远大乙肝疫苗 (CpG+铝)",
        "乙型肝炎 (HBV)",
        "国内",
        "II期(进行中)",
        "随机、双盲、对照",
        "待披露",
        "CpG+铝佐剂乙肝疫苗",
        "安全性/免疫原性; 未公开",
    ),
    (
        "CTR20233253",
        "怡道/中慧元通 带状疱疹疫苗 (CpG)",
        "带状疱疹(VZV)",
        "国内",
        "III期(2023-11启动, 已报NDA)",
        "随机、双盲、对照(40岁+)",
        "待披露",
        "CpG佐剂带状疱疹疫苗 vs 对照",
        "保护效力(主); 安全性; 未公开",
    ),
    (
        "CXSL2600443",
        "明瑞佳 MRJ103 带状疱疹疫苗 (CpG1018+QS-21+脂质体)",
        "带状疱疹(VZV)",
        "国内",
        "临床批件(IND获批2026)",
        "待启动",
        "待披露",
        "MRJ103复合佐剂带状疱疹疫苗",
        "安全性/免疫原性; 未公开",
    ),
]

# 模块二: (标识简写, 局部AE/ADR, 全身AE/ADR, 数据性质, SAE/SADR, AESI)
SAFETY = [
    # 国外
    (
        "HEPLISAV-B",
        "未征集AE(28d): 42.0% vs 41.3%(T1); 35.4% vs 36.2%(T2); 20.1% vs 20.1%(T3)[1]",
        "轻中度AE: 45.6% vs 45.7%; 心血管事件: 0.27% vs 0.14%[1]",
        "AE",
        "SAE: T1 1.5% vs 2.1%; T2 3.9% vs 4.8%; T3 6.2% vs 5.3%[1]",
        "免疫介导AE(IMAE): T1 0.2% vs 0.7%; T2 0.3% vs 0.0%; T3 0.1% vs 0%[1]; 汇总0.32% vs 0.38%[9]",
    ),
    (
        "AV7909",
        "总体AE: 79% (AV7909) vs 65% (BioThrax); 92% vs 87%为1-2级[2]",
        "反应原性各组无显著模式差异[2]",
        "AE",
        "SAE: 2/44(4.55%, 0/14d组), 0/34, 0/23, 0/44, 1/23(4.35%, BioThrax)[CTR结果页]; 无疫苗相关SAE[2]",
        "无自身免疫病因AE报告[2]",
    ),
    (
        "SCB-2019 SPECTRA",
        "征集性局部反应: 36% vs 安慰剂10%; 以轻中度注射部位疼痛为主, 第2剂后降低[3]",
        "征集性全身反应: 与安慰剂组相似, 无显著差异[3]",
        "AE",
        "SAE/重度AE: 罕见, 与安慰剂组无显著差异[3]",
        "AESI: 无异常信号(与安慰剂均衡)[3]",
    ),
    (
        "BK-SE36/CpG 疟疾",
        "局部AE: 46% vs 18%; 疼痛/活动受限(成人17-33%, 儿童40-57%, 婴幼儿0-19%)[4]",
        "发热(婴幼儿第1剂29%); 头痛(成人第3剂20%); AEFI 56% vs 19%[4]",
        "AE",
        "SAE 5例(全部重症疟疾, 均判定与疫苗无关); SUSAR 0; 死亡0[4]",
        "Grade3事件6例(4人)均不相关; 自身免疫标志物(ANA/anti-dsDNA/ANCA)无临床异常[4]",
    ),
    (
        "Na-GST-1/CpG 钩虫",
        "注射部位疼痛: 30μg+CpG 7/8, 100μg无CpG 7/8, 100μg+CpG 9例次; 压痛最常见; 重度仅2例(压痛1例CpG组, 肿胀1例无CpG组)[5]",
        "头痛最常见(25-50%); 肌痛62.5%(100μg+CpG组, 其中中度4例); 发热0-25%[5]",
        "AE",
        "SAE: 0例; 无疫苗相关SAE[5]",
        "AESI: 0例; 无自身免疫症状; 1例ANA升高停药(揭盲后为无CpG组)[5]",
    ),
    (
        "rF1V/CpG 1018 鼠疫",
        "未公开(试验已完成, 结果未发布)[6]",
        "未公开[6]",
        "—",
        "未公开[6]",
        "未公开[6]",
    ),
    (
        "H5N1/CpG 1018 流感",
        "未公开(Part 1完成, Part 2数据预计2026)[7]",
        "未公开[7]",
        "—",
        "未公开[7]",
        "未公开[7]",
    ),
    (
        "Z-1018 带状疱疹",
        "未公开(Part 2 vs Shingrix进行中)[8]",
        "未公开[8]",
        "—",
        "未公开[8]",
        "未公开[8]",
    ),
    # 国内
    (
        "重组乙肝(CpG ODN) I期",
        "总体AE: 66.67%(16/24) vs 54.17%(13/24), P=0.556; 全部1-2级, 无≥3级[10]",
        "同上(原文未细分局部/全身报道)[10]",
        "AE",
        "SAE: 未报告[10]",
        "无免疫介导AE报道[10]",
    ),
    (
        "SCB-2019 中国II期",
        "见SPECTRA全球数据(含中国受试者): 局部36% vs 10%[3]",
        "全身反应与安慰剂无显著差异[3]",
        "AE",
        "SAE/重度AE/AESI罕见, 与安慰剂无差异[3]",
        "AESI与安慰剂无差异[3]",
    ),
    (
        "ZR202-CoV I期",
        "任一局部症状: A 5.6%(1/18), B 22.2%(4/18), C 38.9%(7/18), 安慰剂5.6%(1/18); 疼痛5.6%/22.2%/27.8%/5.6%; 肿胀C组16.7%(3/18)[11]",
        "任一全身症状: A 11.1%, B 11.1%, C 5.6%, 安慰剂5.6%; 发热(腋温≥37.3℃)同比例; 恶心A组5.6%(1/18)[11]",
        "AE(均为1-2级, 无≥3级征集性AE)",
        "SAE 1例(A组, 输尿管结石, 判定与疫苗无关, 5.6%)[11]",
        "无AESI; 非征集性AE≥3级: C组1例(尿潜血, 月经相关), 安慰剂2例; 纤维蛋白原升高22.2%/22.2%/17.3%/5.6%[11]",
    ),
    # 国内进行中
    (
        "安百胜带状疱疹 I/II",
        "未公开(进行中)[12]",
        "未公开[12]",
        "—",
        "未公开[12]",
        "未公开[12]",
    ),
    (
        "华普 HP2001",
        "未公开(进行中)[13]",
        "未公开[13]",
        "—",
        "未公开[13]",
        "未公开[13]",
    ),
    (
        "华普 HP2002",
        "未公开(进行中)[14]",
        "未公开[14]",
        "—",
        "未公开[14]",
        "未公开[14]",
    ),
    (
        "CTR20233405 乙肝I期",
        "未公开(招募完成)[15]",
        "未公开[15]",
        "—",
        "未公开[15]",
        "未公开[15]",
    ),
    (
        "远大乙肝 II期",
        "未公开(进行中)[16]",
        "未公开[16]",
        "—",
        "未公开[16]",
        "未公开[16]",
    ),
    (
        "怡道带状疱疹 III期",
        "未公开(已报NDA)[17]",
        "未公开[17]",
        "—",
        "未公开[17]",
        "未公开[17]",
    ),
    (
        "明瑞佳 MRJ103",
        "未公开(IND获批)[18]",
        "未公开[18]",
        "—",
        "未公开[18]",
        "未公开[18]",
    ),
]

# 模块三: (标识, 文献, DOI, PMID, 链接)
REFS = [
    (
        "HEPLISAV-B / 国内乙肝",
        "Hyer R, et al. Safety of a two-dose investigational hepatitis B vaccine, HBsAg-1018, using a toll-like receptor 9 agonist adjuvant in adults (3项III期合并分析). Vaccine. 2018;36(19):2604-2611",
        "10.1016/j.vaccine.2018.03.067",
        "29628151",
        "https://pubmed.ncbi.nlm.nih.gov/29628151/",
    ),
    (
        "AV7909",
        "Hopkins RJ, et al. Randomized, double-blind, active-controlled study evaluating the safety and immunogenicity of three vaccination schedules and two dose levels of AV7909 vaccine for anthrax post-exposure prophylaxis in healthy adults. Vaccine. 2016;34(18):2096-2105",
        "10.1016/j.vaccine.2016.03.006",
        "26979136",
        "https://clinicaltrials.gov/study/NCT01770743",
    ),
    (
        "SCB-2019 SPECTRA",
        "Bravo L, et al. Efficacy of the adjuvanted subunit protein COVID-19 vaccine, SCB-2019: a phase 2 and 3 multicentre, double-blind, randomised, placebo-controlled trial. Lancet. 2022;399:461-472",
        "10.1016/S0140-6736(22)00055-1",
        "35065005",
        "https://clinicaltrials.gov/study/NCT04672395",
    ),
    (
        "BK-SE36/CpG 疟疾",
        "Safety and immunogenicity of BK-SE36/CpG malaria vaccine in healthy Burkinabe adults and children: a phase 1b randomised, controlled, double-blinded, age de-escalation trial. Front Immunol. 2023;14:1267372",
        "10.3389/fimmu.2023.1267372",
        "37638048",
        "https://www.frontiersin.org/journals/immunology/articles/10.3389/fimmu.2023.1267372/full",
    ),
    (
        "Na-GST-1/CpG 钩虫",
        "Diemert DJ, et al. Randomized, observer-blind, controlled Phase 1 study of the Na-GST-1/Alhydrogel hookworm vaccine with or without CpG ODN adjuvant. PLoS Negl Trop Dis. 2024;18(12):e0012788",
        "10.1371/journal.pntd.0012788",
        "39775205",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC11717351",
    ),
    (
        "rF1V/CpG 1018 鼠疫",
        "Dynavax/DoD rF1V+CpG 1018 Phase 2 (NCT05506969), 注册信息",
        "—",
        "—",
        "https://clinicaltrials.gov/study/NCT05506969",
    ),
    (
        "H5N1/CpG 1018 流感",
        "Dynavax Pandemic Influenza Adjuvant Program (Phase 1/2, Part 1 完成; 公司公告2025)",
        "—",
        "—",
        "https://www.nasdaq.com/press-release/dynavax-reports-third-quarter-2025-financial-results-and-announces-new-100-million",
    ),
    (
        "Z-1018 带状疱疹",
        "Dynavax Z-1018 (CpG 1018+gE) Phase 1/2 管线公告(2025)",
        "—",
        "—",
        "https://synapse.patsnap.com/drug/9e98f53e66bf47af84a078b8ed92d770",
    ),
    (
        "免疫介导AE综述",
        "Campbell JD. A narrative review of immune-mediated adverse events in clinical trials of CpG ODN TLR9 agonists. Vaccine. 2026",
        "10.1016/j.vaccine.2026(S0264410X26002458)",
        "—",
        "https://doi.org/10.1016/j.vaccine.2026",
    ),
    (
        "重组乙肝(CpG ODN) I期",
        "梁贞贞, 邵彦, 王慎玉, 等. 成人接种重组(汉逊酵母)乙型肝炎疫苗(CpG ODN佐剂)的安全性及免疫原性: I期临床试验初步结果. 中华预防医学杂志. 2020;54(8):854-860",
        "10.3760/cma.j.cn112150-20200401-00490",
        "32842315",
        "https://pubmed.ncbi.nlm.nih.gov/32842315/",
    ),
    (
        "ZR202-CoV",
        "Feng GW, et al. Safety, tolerability, and immunogenicity of a CpG/Alum adjuvanted SARS-CoV-2 recombinant protein vaccine (ZR202-CoV): Phase 1, randomized, double-blind, placebo-controlled, dose-escalation trial. Hum Vaccin Immunother. 2023;19(3):2262635",
        "10.1080/21645515.2023.2262635",
        "37881130",
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC10644802/",
    ),
    (
        "安百胜带状疱疹",
        "重组带状疱疹疫苗(gE+CpG1018+MF59) I/II期注册信息",
        "—",
        "—",
        "https://clinicaltrials.gov/study/NCT06851832",
    ),
    (
        "华普 HP2001",
        "华普生物 HP2001 带状疱疹疫苗I期受理信息(CXSL2500371)",
        "—",
        "—",
        "https://parrbio.com/newsinfo/11232956.html",
    ),
    (
        "华普 HP2002",
        "华普 HP2002 乙肝疫苗I期登记信息(CTR20261430)",
        "—",
        "—",
        "https://www.chinadrugtrials.org.cn",
    ),
    (
        "CTR20233405",
        "重组乙型肝炎疫苗(汉逊酵母,CpG和铝佐剂)I期临床研究登记信息(CTR20233405)",
        "—",
        "—",
        "https://www.chinadrugtrials.org.cn",
    ),
    (
        "远大乙肝 II期",
        "远大乙肝疫苗II期注册信息(ChiCTR2500108408)",
        "—",
        "—",
        "https://www.chictr.org.cn",
    ),
    (
        "怡道带状疱疹 III期",
        "怡道/中慧元通带状疱疹疫苗III期登记信息(CTR20233253)",
        "—",
        "—",
        "https://www.chinadrugtrials.org.cn",
    ),
    (
        "明瑞佳 MRJ103",
        "明瑞佳带状疱疹疫苗临床批件(CXSL2600443)",
        "—",
        "—",
        "https://www.chinadrugtrials.org.cn",
    ),
]

# 数据缺口(不列入表格, 文末带过)
GAPS = (
    "HPV疫苗: 仅见CpG佐剂治疗性(肿瘤)HPV疫苗临床前/临床研究, 无CpG佐剂预防性HPV疫苗人体试验, 按规则不列入; "
    "吉诺卫、昆明简达(带状疱疹/流感)、远大带状疱疹(TVA01)、长春卓谊(CpG狂犬疫苗, 已获临床批件、I期待启动): "
    "临床试验信息未公开或试验未启动, 按规则不列入; "
    "昆明简达CpG复合佐剂流感疫苗已于2026-07获FDA临床许可(临床试验信息尚未公示)。"
)


# ------------------------------------------------------------ helpers ----
def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_font(run, size=9, bold=False, name_ea="宋体", name_ascii="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_ascii
    r = run._element.rPr
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_ea)


def add_para(
    doc,
    text,
    size=10.5,
    bold=False,
    align=None,
    space_after=6,
    ea="宋体",
    ascii_f="Times New Roman",
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, name_ea=ea, name_ascii=ascii_f)
    return p


def make_table(doc, headers, rows, widths_cm, header_fill="D9E2F3", body_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # header
    hdr = t.rows[0]
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        c.width = Cm(widths_cm[j])
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9, bold=True, name_ea="黑体", name_ascii="Arial")
        set_cell_shading(c, header_fill)
    # body
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            c = cells[j]
            c.width = Cm(widths_cm[j])
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            if j == 0:
                run = p.add_run(v)
                set_font(run, size=body_size, bold=True)
            else:
                run = p.add_run(v)
                set_font(run, size=body_size, bold=False)
    return t


def add_block_row(table, text, ncols, fill="BDD7EE"):
    """在表格内插入分块标题行(合并单元格)"""
    row = table.add_row()
    row.cells[0].merge(row.cells[ncols - 1])
    c = row.cells[0]
    p = c.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=9.5, bold=True, name_ea="黑体", name_ascii="Arial")
    set_cell_shading(c, fill)


# ---------------------------------------------------------------- build ---
doc = Document()
sec = doc.sections[0]
# 横向 A4
sec.orientation = 1  # WD_ORIENT.LANDSCAPE
sec.page_width = Cm(29.7)
sec.page_height = Cm(21.0)
sec.left_margin = Cm(1.5)
sec.right_margin = Cm(1.5)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)

# 标题
add_para(
    doc,
    "国内外CpG佐剂预防性疫苗安全性数据汇总",
    size=16,
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
add_para(
    doc,
    "检索日期: 2026-08-20 | 数据截止: 2026-08-20 | 数据源: ClinicalTrials.gov / 中国药物临床试验登记与信息公示平台(Chinadrugtrials.org.cn) / PubMed / 期刊原文",
    size=9,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=8,
    ea="宋体",
)

# 绪论
add_para(
    doc,
    "本报告汇总国内外使用CpG类佐剂(含CpG 1018、CpG 7909、CpG ODN K3、CpG 10104、HP007等)的预防性疫苗临床试验安全性数据，"
    "覆盖乙型肝炎、炭疽、COVID-19、疟疾、钩虫、鼠疫、流感、带状疱疹等适应症；治疗性肿瘤疫苗不在范围内。"
    "安全性数据均来自文献原文的仔细阅读与提取，含具体发生率数值；数据未公开的试验(进行中)在表中明确标注。"
    "已上市产品数据以关键注册试验/汇总分析为准，国内试验数据以已发表文献为准。",
    size=10.5,
    space_after=10,
)

# ============================ 模块一 ============================
add_para(
    doc,
    "模块一: 临床试验基本信息",
    size=13,
    bold=True,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
H1 = [
    "试验标识(NCT/CTR)",
    "疫苗名称",
    "适应症",
    "地域",
    "试验分期",
    "整体设计",
    "样本量",
    "试验分组",
    "研究终点",
]
W1 = [2.6, 2.9, 2.2, 1.2, 2.4, 3.0, 2.6, 4.4, 3.4]
t1 = make_table(doc, H1, [], W1, body_size=8)
add_block_row(t1, "◆ 国外部分 (已上市 / 已发表)", 9)
for row in TRIALS[:8]:
    t1.add_row().cells  # placeholder to keep loop simple
    cells = t1.rows[-1].cells
    vals = row
    for j, v in enumerate(vals):
        cells[j].width = Cm(W1[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(v) if v else "—")
        set_font(run, size=8, bold=(j == 0))
add_block_row(t1, "◆ 国内部分 (已发表安全性数据)", 9)
for row in TRIALS[8:11]:
    cells = t1.add_row().cells
    for j, v in enumerate(row):
        cells[j].width = Cm(W1[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(v) if v else "—")
        set_font(run, size=8, bold=(j == 0))
add_block_row(t1, "◆ 国内部分 (临床试验信息可查, 安全性数据未公开)", 9)
for row in TRIALS[11:]:
    cells = t1.add_row().cells
    for j, v in enumerate(row):
        cells[j].width = Cm(W1[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(v) if v else "—")
        set_font(run, size=8, bold=(j == 0))
doc.add_paragraph()

# ============================ 模块二 ============================
add_para(
    doc,
    "模块二: 核心安全性数据 (量化)",
    size=13,
    bold=True,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
H2 = [
    "试验",
    "局部 AE/ADR (名称+发生率)",
    "全身 AE/ADR (名称+发生率)",
    "数据性质",
    "SAE / SADR",
    "AESI (特殊关注不良事件)",
]
W2 = [2.8, 5.2, 4.6, 1.8, 4.6, 4.8]
t2 = make_table(doc, H2, [], W2, body_size=8)
add_block_row(t2, "◆ 国外部分", 6)
for row in SAFETY[:8]:
    cells = t2.add_row().cells
    for j, v in enumerate(row):
        cells[j].width = Cm(W2[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(v)
        set_font(run, size=8, bold=(j == 0))
add_block_row(t2, "◆ 国内部分 (已发表)", 6)
for row in SAFETY[8:11]:
    cells = t2.add_row().cells
    for j, v in enumerate(row):
        cells[j].width = Cm(W2[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(v)
        set_font(run, size=8, bold=(j == 0))
add_block_row(t2, "◆ 国内部分 (安全性数据未公开)", 6)
for row in SAFETY[11:]:
    cells = t2.add_row().cells
    for j, v in enumerate(row):
        cells[j].width = Cm(W2[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(v)
        set_font(run, size=8, bold=(j == 0))
add_para(
    doc,
    "注: 数据性质标注为AE(不良事件); 各试验原文以AE报告为主, 未见以ADR(与疫苗因果相关不良反应)单独报道的安全性汇总。",
    size=8.5,
    space_after=8,
)
doc.add_paragraph()

# ============================ 模块三 ============================
add_para(
    doc,
    "模块三: 详细参考文献",
    size=13,
    bold=True,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
H3 = ["对应试验", "文献标题 (作者/期刊/年份)", "DOI", "PMID", "原文链接"]
W3 = [3.2, 7.6, 3.4, 1.8, 4.8]
t3 = make_table(doc, H3, [], W3, body_size=8)
for row in REFS:
    cells = t3.add_row().cells
    for j, v in enumerate(row):
        cells[j].width = Cm(W3[j])
        p = cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(v)
        set_font(run, size=8, bold=(j == 0))
doc.add_paragraph()

# ============================ 数据缺口 ============================
add_para(
    doc,
    "数据缺口与排除说明",
    size=13,
    bold=True,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
add_para(doc, GAPS, size=9.5, space_after=6)

# 关键结论
add_para(
    doc, "关键安全性结论", size=13, bold=True, space_after=4, ea="黑体", ascii_f="Arial"
)
for s in [
    "1. 已上市产品(Heplisav-B、AV7909)及大型III期试验(SPECTRA, n=30,128)显示CpG佐剂疫苗总体耐受性良好, 严重不良事件与对照无显著差异;",
    "2. 国内已发表数据(乙肝I期48例、ZR202-CoV I期72例)中, 不良事件以1-2级为主, 无疫苗相关≥3级AE或SAE;",
    "3. 免疫介导不良事件(AESI)未见异常信号: Heplisav-B汇总IMAE 0.32% vs 铝佐剂0.38%, 疟疾/钩虫疫苗均无自身免疫事件;",
    "4. 局部反应(注射部位疼痛)为最常见不良事件类型, 发生率略高于无佐剂/其他佐剂对照, 但多为轻度、一过性;",
    "5. 国内带状疱疹/流感/狂犬等CpG佐剂疫苗管线尚在临床推进中, 安全性数据未公开, 待后续补充。",
]:
    add_para(doc, s, size=10, space_after=2)

doc.save(OUT_PATH)
print("SAVED:", OUT_PATH)
