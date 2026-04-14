from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path


@dataclass(frozen=True)
class Finding:
    id: str
    severity: str  # 重大/重要/一般/建议
    category: str
    location: str
    issue: str
    recommendation: str
    rationale: str
    cross_ref: str


def _apply_cn_en_fonts(doc) -> None:
    """
    Enforce document-wide fonts:
    - Chinese (East Asia): 宋体
    - English (ASCII/HAnsi): Times New Roman
    """
    from docx.oxml.ns import qn

    def set_style(style_name: str) -> None:
        if style_name not in doc.styles:
            return
        style = doc.styles[style_name]
        font = style.font
        font.name = "Times New Roman"
        # East Asia font mapping (Chinese)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:cs"), "Times New Roman")

    # Core styles that typically cover body/headings/tables.
    for name in [
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Table Grid",
    ]:
        set_style(name)


def _add_heading(doc, text: str, level: int = 1) -> None:
    # python-docx: level 0=Title, 1..9=Heading n
    doc.add_heading(text, level=level)


def main() -> int:
    import argparse

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    parser = argparse.ArgumentParser(description="Generate strict audit report docx.")
    parser.add_argument(
        "--root",
        type=str,
        default=str(Path(__file__).resolve().parents[1]),
        help="Project root path (default: inferred from this script location)",
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        default="review_materials",
        help="Output directory relative to root unless absolute path is provided",
    )
    args = parser.parse_args()

    today = date.today().isoformat()

    project_root = Path(args.root)
    output_dir = Path(args.output_dir)
    if not output_dir.is_absolute():
        output_dir = project_root / output_dir

    main_doc = "YDSWX(TVAX-006)-002(II)_数据审核报告_V0.7_20260312.docx"
    ref_xlsx = "YDSWX(TVAX-006)-002(II)_数据审核报告清单_20260312.xlsx"
    ref_protocol = "远大重组带状疱疹疫苗Ⅱ期-方案（版本号：1.3，版本日期：2026年01月05日）-清洁版-定稿.docx"

    # Use ASCII filename to avoid Windows console encoding issues.
    out_name = f"audit_report_data_review_v0.7_strict_audit_{today}.docx"
    out_path = output_dir / out_name

    findings: list[Finding] = [
        Finding(
            id="F-01",
            severity="重要",
            category="合规/可追溯性（疫苗有效期与批次信息）",
            location="主审报告 1.2 疫苗有效期（表1.2）",
            issue=(
                "表1.2存在关键字段格式/单位不一致：如“0.5m1”疑为“0.5mL”；"
                "“50ug”建议统一为“50 μg”；有效期字段存在“20280112/2026年8月/2027-5”等混用，"
                "且部分缺少“日”或不满足可直接核对的日期格式。"
            ),
            recommendation=(
                "统一单位与符号（mL、μg），统一有效期日期格式（建议YYYY-MM-DD），"
                "对仅到“月”的有效期补充具体到日的来源（标签/检验报告/批签发/放行文件）；"
                "在表后补充“给药日期与有效期核对逻辑”（逐批次/逐受试者或逐接种日）与核对结论。"
            ),
            rationale=(
                "疫苗有效期与批次是GCP/疫苗试验质量管理的关键可追溯要素；"
                "日期粒度不足会削弱“在有效期内使用”的可核验性，也不利于外部稽查复核。"
            ),
            cross_ref=f"交叉核对：方案V1.3中明确“有效期以实际使用批次为准，详见标签或检验报告等文件”。",
        ),
        Finding(
            id="F-02",
            severity="重要",
            category="数据范围/截止（cut-off）与统计口径一致性",
            location="主审报告 1.3.2 数据审核范围",
            issue=(
                "报告定义通用访视界面cut-off为“末次接种+30天内（包含30天）”，并对"
                "合并用药/用疫苗/非药物治疗/不良事件在“日缺失且开始月=末次接种同月或+1月”时"
                "纳入30天内分析。该规则属于时间信息缺失时的“包含性判定/推断”，但报告未说明"
                "其与统计分析计划（SAP）或数据管理缺失值处理规则的一致性，也未说明可能导致的偏倚方向。"
            ),
            recommendation=(
                "将该规则明确标注为“缺失日期的处理/纳入判定规则”，补充其依据（SAP/DM计划/预设算法）；"
                "建议增加敏感性说明：对“月+1且日缺失”是否可能超出30天窗口的情形，"
                "给出保守处理（如按最早日/最晚日两种极端）或列出受影响记录数量。"
            ),
            rationale=(
                "cut-off与缺失日期处理直接影响AE/用药/用疫苗等的纳入，属于可能改变安全性结论的关键方法学要素；"
                "需要可追溯、可复现且与SAP一致。"
            ),
            cross_ref="交叉核对：方案V1.3安全性/合并用药等收集窗口与阶段性分析范围描述。",
        ),
        Finding(
            id="F-03",
            severity="重要",
            category="数据一致性（叙述 vs 汇总 vs 注释）",
            location="主审报告 7 免疫血样采集情况（7.1.1/7.1.3）",
            issue=(
                "免疫血样采集叙述存在口径混杂：报告称“第二剂免后第30天体液免疫采血417例”，"
                "又说明“其中2例未完成第2剂接种…即纳入…415例”。同时在“采集缺失情况”章节中"
                "汇总缺失均为0，但脚注又提到1例受试者在第二剂免后第30天未采血且退出（细胞免疫亚组）。"
            ),
            recommendation=(
                "明确并拆分指标：区分“实际完成采血人数”“满足访视定义可纳入分析人数（如需V4接种）”"
                "与“缺失定义（因退出/访视未进行/未在窗口内等）”；"
                "将脚注事件在主表中以一致的缺失分类体现（例如：退出导致的未采血是否计入缺失，需写清标准）。"
            ),
            rationale=(
                "阶段性分析通常要求对样本量、缺失与纳入标准给出一致且可复核的口径；"
                "当前“正文=0缺失、脚注=存在未采血”会在稽查与统计复核时被判为不一致。"
            ),
            cross_ref=f"交叉核对：参考清单XLSX包含个体层面采血记录；方案V1.3规定采血时点与随访流程。",
        ),
        Finding(
            id="F-04",
            severity="一般",
            category="数据一致性（合计/分项/重叠说明）",
            location="主审报告 10.2 血样采集超窗情况汇总（表10.2与文字）",
            issue=(
                "文字描述称“免疫血样采集超窗共计5例次，其中体液5例次，细胞1例次”，"
                "若按“例次”计数则5+1=6；表10.2脚注说明细胞超窗包含在体液超窗受试者内。"
                "当前未统一“合计”为“例次”还是“受试者数/去重后例次”。"
            ),
            recommendation=(
                "在章节开头明确统计口径：合计为“去重受试者数”还是“事件/例次”；"
                "如存在重叠，建议同时给出：去重后受试者数、未去重例次数，并在表注中保持一致术语。"
            ),
            rationale=(
                "超窗属于方案偏离的重要来源；口径不清会影响偏离数量与严重程度判断。"
            ),
            cross_ref="交叉核对：参考清单XLSX中有超窗明细，可支持去重与例次统计。",
        ),
        Finding(
            id="F-05",
            severity="重要",
            category="安全性范围与方案一致性（阶段性 vs 全程）",
            location="主审报告 1.3.2 数据审核范围 & 17 严重不良事件核查",
            issue=(
                "方案V1.3规定SAE/AESI/妊娠事件监测至“全程接种后12个月”，"
                "而本报告的数据截止定义为“末次接种后30天安全性随访”。报告虽包含SAE汇总（2例次），"
                "但未在SAE章节显式说明：本次仅覆盖截止点前SAE（如存在更长随访期事件，尚未纳入）。"
            ),
            recommendation=(
                "在SAE/AESI/妊娠相关章节增加“分析期/截止点说明”，并明确“长期随访期（31天~12个月）”"
                "数据是否已锁定/是否纳入本次报告；如未纳入，写明将于最终分析补充。"
            ),
            rationale=(
                "不明确分析期可能被解读为已完成12个月SAE监测并汇总为0/2例次，存在合规与解释风险。"
            ),
            cross_ref="交叉核对：方案V1.3“安全性观察”与“阶段性分析揭盲”条款。",
        ),
        Finding(
            id="F-06",
            severity="一般",
            category="随机化/盲法证据充分性",
            location="主审报告 5 随机化核查",
            issue=(
                "仅给出结论“线上随机…未发现问题”，未呈现可复核证据（如分层/分组例数、"
                "1:1分配实现情况、IWRS审计追踪、随机化偏差检查）。"
            ),
            recommendation=(
                "补充关键证据摘要：按年龄层（40~49/50~59/≥60）及组别的随机化例数；"
                "说明IWRS/随机系统版本、权限控制与审计追踪；对是否存在重复随机/跳号/分配异常给出核查方法。"
            ),
            rationale="随机化与盲法完整性是监管与稽查关注重点，建议在阶段性数据审核报告中留痕。",
            cross_ref="交叉核对：方案V1.3“随机化与盲法”章节对分层与比例的描述。",
        ),
        Finding(
            id="F-07",
            severity="一般",
            category="方案偏离/违背分类透明度",
            location="主审报告 8 方案偏离/违背情况汇总",
            issue=(
                "报告给出偏离10条、违背4条及分类，但未在正文解释“偏离 vs 违背”的判定准则，"
                "仅在违背分类中用括号描述“第2剂缺失（符合第2剂接种排除标准者除外）”。"
            ),
            recommendation=(
                "在8章增加“判定标准”小节，引用方案V1.3对偏离/违背的定义及示例；"
                "对“第2剂缺失但因排除标准终止接种”的处理给出可复核规则与受试者例数。"
            ),
            rationale="有助于确保偏离/违背统计与SAP分析集（如PPS纳入）一致，并利于审计追溯。",
            cross_ref="交叉核对：方案V1.3第9.6节（方案偏离/违背）与分析集定义（PPS条款）。",
        ),
        Finding(
            id="F-08",
            severity="建议",
            category="数据管理质疑（queries）呈现完整性",
            location="主审报告 1.4 数据管理质疑情况",
            issue=(
                "报告给出质疑总量742条且全部关闭，但未提供关键质量维度（如平均关闭时长、"
                "主要问题类型Top N、是否存在反复开启、与关键终点相关的高风险质疑）。"
            ),
            recommendation=(
                "建议补充一页质量摘要：按模块（入排/接种/AE/免疫原性等）或严重度的质疑分布；"
                "列示与关键终点相关的关键质疑是否全部解决及其证据。"
            ),
            rationale="阶段性锁库/揭盲前的质量证明材料越充分，越利于内部决策与外部检查。",
            cross_ref="交叉核对：EDC导出/质疑日志（如可获取）。",
        ),
        Finding(
            id="F-09",
            severity="一般",
            category="一致性/术语规范",
            location="主审报告 全文（单位、符号、日期格式）",
            issue=(
                "存在“ml/m1”“ug/μg”“YYYYMMDD/YYYY-MM-DD/中文年月”等混用风险，"
                "以及部分英文字母大小写/符号（如cut-off、cutoff）不统一。"
            ),
            recommendation=(
                "建立并应用格式规范：单位用法（mL、μg）、日期格式（YYYY-MM-DD）、"
                "术语表（cut-off等）统一；对表格字段设定强制格式。"
            ),
            rationale="提升可读性与可核验性，降低外部提交/归档被退回的概率。",
            cross_ref="交叉核对：申办方文档模板/统计报告格式规范（如有）。",
        ),
    ]

    doc = Document()
    _apply_cn_en_fonts(doc)

    title = doc.add_paragraph(
        "YDSWX(TVAX-006)-002(Ⅱ)\n数据审核报告（V0.7/2026-03-12）\n严格审计式交叉审核报告"
    )
    # Keep it centered; font inherits from Normal/Title settings.
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成日期：{today}")

    _add_heading(doc, "一、审核对象与参考文件", level=1)
    doc.add_paragraph(f"主审文件：{main_doc}")
    doc.add_paragraph(f"参考文件1（清单/明细）：{ref_xlsx}")
    doc.add_paragraph(f"参考文件2（方案）：{ref_protocol}")
    doc.add_paragraph("审核目标：从合规性、一致性与数据逻辑三条主线识别问题点，并提出可执行修改建议。")

    _add_heading(doc, "二、总体结论（摘要）", level=1)
    doc.add_paragraph(
        "总体上，主审报告的关键里程碑数字（筛选/入组/完成/退出）在内部叙述中呈现自洽，"
        "且与方案的阶段性分析目标（全程接种后30天）方向一致。"
    )
    doc.add_paragraph(
        "但在“免疫血样采集缺失/口径”“血样超窗合计口径”“疫苗有效期表格字段规范化”"
        "及“SAE监测期在阶段性报告中的范围说明”等方面存在需要修订的审计点。"
    )

    _add_heading(doc, "三、问题清单（分条：问题/修改建议/考量原因）", level=1)
    table = doc.add_table(rows=1, cols=8)
    hdr = table.rows[0].cells
    hdr[0].text = "编号"
    hdr[1].text = "风险等级"
    hdr[2].text = "类别"
    hdr[3].text = "位置"
    hdr[4].text = "问题"
    hdr[5].text = "修改建议"
    hdr[6].text = "考量原因/依据"
    hdr[7].text = "交叉核对点"

    for f in findings:
        row = table.add_row().cells
        row[0].text = f.id
        row[1].text = f.severity
        row[2].text = f.category
        row[3].text = f.location
        row[4].text = f.issue
        row[5].text = f.recommendation
        row[6].text = f.rationale
        row[7].text = f.cross_ref

    _add_heading(doc, "四、参考的公开资料链接（联网检索）", level=1)
    refs = [
        ("ICH E3（临床研究报告结构与内容，PDF）", "https://database.ich.org/sites/default/files/E3_Guideline.pdf"),
        ("EMA：ICH E3 科学指南页面", "https://www.ema.europa.eu/en/ich-e3-structure-content-clinical-study-reports-scientific-guideline"),
        ("FDA：E3 指南下载页", "https://www.fda.gov/media/84857/download"),
        ("中国政府网：疫苗临床试验质量管理指导原则（试行）", "https://www.gov.cn/gongbao/content/2014/content_2580994.htm"),
        ("药审中心（CDE）官网入口（指导原则检索）", "https://www.cde.org.cn/"),
        ("CIO在线：关于发布《疫苗临床试验技术指导原则》的通告（2025年第16号）（便于快速检索）", "https://m.ciopharma.com/supervise/44545"),
    ]
    for name, url in refs:
        p = doc.add_paragraph()
        p.add_run(f"- {name}：{url}")

    doc.add_paragraph("\n注：上述链接用于方法学与合规性口径参考；本次审核的事实依据以项目文件与EDC导出为准。")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    print(f"OK: wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

