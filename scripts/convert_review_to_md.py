#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os, sys
from pathlib import Path
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
INPUT_DIR = Path(__file__).resolve().parent.parent / "review_materials"
OUTPUT_DIR = INPUT_DIR / "converted"
def safe_md_name(src_name):
    base = Path(src_name).stem
    for c in '\\/:*?"<>|':
        base = base.replace(c, "_")
    return base + ".md"
def convert_with_markitdown(src):
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        result = md.convert(str(src))
        return result.text_content if hasattr(result, "text_content") else str(result)
    except Exception as e:
        print("  [markitdown]", e, file=sys.stderr)
        return None
def convert_docx_fallback(src):
    try:
        from docx import Document
        doc = Document(str(src))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n\n".join(parts) if parts else None
    except Exception as e:
        print("  [python-docx]", e, file=sys.stderr)
        return None
def convert_pdf_fallback(src):
    try:
        import pypdf
        reader = pypdf.PdfReader(str(src))
        return "\n\n".join(p.extract_text() or "" for p in reader.pages).strip() or None
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(str(src)) as pdf:
                return "\n\n".join((p.extract_text() or "") for p in pdf.pages).strip() or None
        except Exception as e:
            print("  [pdfplumber]", e, file=sys.stderr)
            return None
    except Exception as e:
        print("  [pypdf]", e, file=sys.stderr)
        return None
def convert_file(src):
    content = convert_with_markitdown(src)
    if not content or not content.strip():
        if src.suffix.lower() == ".docx":
            content = convert_docx_fallback(src)
        elif src.suffix.lower() == ".pdf":
            content = convert_pdf_fallback(src)
    if not content or not content.strip():
        return None
    out_name = safe_md_name(src.name)
    out_path = OUTPUT_DIR / out_name
    out_path.write_text(content, encoding="utf-8")
    return out_path
def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = []
    for f in sorted(INPUT_DIR.iterdir()):
        if not f.is_file() or f.suffix.lower() not in (".docx", ".pdf"):
            continue
        print("Converting:", f.name)
        out = convert_file(f)
        if out:
            created.append(str(out))
            print("  ->", out.name)
    print("\n--- Created Markdown files ---")
    for p in created:
        print(p)
    return created
if __name__ == "__main__":
    main()
