from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from common_scripts.docx_utils import apply_cn_en_fonts


@dataclass
class Finding:
    id: str
    severity: str
    category: str
    location: str
    issue: str
    recommendation: str


def _add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def main() -> int:
    today = date.today().isoformat()
    output_dir = Path("review_materials")
    out_name = "远大赛威信_重组带状疱疹疫苗_DSUR#2_审核报告.docx"
    out_path = output_dir / out_name

    findings = [
        Finding(
            id="F-01",
            severity="重要",
            category="数据遗漏/占位符未更新",
            location="执行概要，第4个条目",
            issue="存在未填写的占位符：“其中xx例参与者发生了xx例次AE，共xx例参与者发生了xx例次与试验疫苗有关的AE”。",
            recommendation="根据上一版DSUR数据，此处应补充完整。参考DSUR#1中的表述（94例报告AE，91例发生与疫苗接种相关的AE），并结合具体例次进行修正。",
        ),
        Finding(
            id="F-02",
            severity="重要",
            category="内容更新遗漏（方案与研究者手册更新）",
            location="附录R4及R5",
            issue="正文第4节提到本报告期内研究者手册更新到V1.3，但在附录R4“报告周期内发生的药物临床试验方案变更...总结表”中，各项均为“无”或“NA”。同时，DSUR资料包中显示I期方案更新至V1.4，II期方案更新至V1.3，未在R4体现。附录R5仅保留了模板提示语，未填写实质内容。",
            recommendation="将研究者手册（V1.3）的更新及各期临床试验方案的更新（若发生在报告期内）如实填入附录R4表格中。并在附录R5补充下一周期总体研究计划概要。",
        ),
        Finding(
            id="F-03",
            severity="建议",
            category="安全性风险分析补充",
            location="13. 文献 及 19. 重要风险总结",
            issue="文献检索部分识别了前庭神经炎、吉兰-巴雷综合征（GBS）、巨噬细胞活化综合征、痛风等散发安全性信号，结论指出暂不形成新的风险信号。但在第19节“重要风险总结”及最终结论中，未进一步强调对这些新识别潜在信号的监控建议。",
            recommendation="建议在第19节及结论部分增加补充说明：针对文献检索提示的散发神经系统及免疫介导的安全性信号（如吉兰-巴雷综合征），将在后续的临床试验实施及药物警戒活动中予以持续定向的主动监测，以进一步确保受试者安全并为疫苗的获益-风险评价提供长期数据支持。",
        ),
        Finding(
            id="F-04",
            severity="一般",
            category="语句表述语病及格式错别字",
            location="执行概要、正文相关段落",
            issue="1. “TVA01佐剂系统由QS-21...构成的脂质体混悬液”句式杂糅；\n2. “细胞毒性：T淋巴细胞（CTL）”与“炎症小体：并促进”冒号使用错误；\n3. “Toll-likereceptor9”缺少空格；\n4. “报告日期：2026年09月xx日”待更新具体日期。",
            recommendation="1. 改为“TVA01佐剂系统是...构成的”或删除“构成的”后的“的”；\n2. 删除多余的冒号；\n3. 修改为“Toll-like receptor 9”；\n4. 定稿前补充完整报告日期。",
        ),
        Finding(
            id="F-05",
            severity="一般",
            category="逻辑与前后一致性",
            location="8.1 已完成的临床试验 / 执行概要",
            issue="执行概要提及累计3项试验（含1项已完成），8.1节仅描述“本报告期内，无已完成的临床试验”。此表述在逻辑上是准确的（TVAX-006-01在上个报告期已完成），但可能引起审阅者误解。",
            recommendation="建议在8.1节补充一句话：“（注：TVAX-006-01试验已于上一报告周期完成，详见前次DSUR，本报告周期内无新增已完成试验）”，以增强上下文衔接性。",
        ),
    ]

    doc = Document()
    apply_cn_en_fonts(doc)

    title = doc.add_paragraph("重组带状疱疹疫苗（CHO细胞）DSUR#2 初稿\n审核报告")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成日期：{today}")

    _add_heading(doc, "一、审核说明与参考文件", level=1)
    doc.add_paragraph(
        "目标审核文件：远大赛威信_重组带状疱疹疫苗_DSUR#2（20250731-20260730）-初稿-SDM.docx"
    )
    doc.add_paragraph(
        "参考资料1：远大赛威信_重组带状疱疹疫苗_DSUR#1（20230731-20250730）_预定稿_全文（含附件）.docx"
    )
    doc.add_paragraph(
        "参考资料2：DSUR资料（1期、2期临床试验方案及研究者手册相关PDF文件）"
    )
    doc.add_paragraph(
        "审核目标：针对前后矛盾的数据与表述、语句表述及错别字、补充的分析结论进行全面排查。"
    )

    _add_heading(doc, "二、总体结论摘要", level=1)
    doc.add_paragraph(
        "经审核，该DSUR初稿在临床试验安全性数据的汇总及严重不良事件（SAE）统计上逻辑严密，"
        "SAE例数及参与者数量（21例发生25例次，与各分项1+5+15例完全一致）数据一致性较好。"
    )
    doc.add_paragraph(
        "需重点关注并修改的问题包括：1) 执行概要部分存在的关键数据占位符遗漏；"
        "2) 附录R4及R5对本报告周期内研究者手册（更新至V1.3）及相关方案更新未作登记；"
        "3) 个别文句的格式、标点错用与语病；"
        "4) 建议针对文献新挖掘的安全信号增加后续监控建议。"
    )

    _add_heading(doc, "三、详细问题清单与修改建议", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "编号"
    hdr[1].text = "风险等级"
    hdr[2].text = "类别"
    hdr[3].text = "位置"
    hdr[4].text = "问题描述"
    hdr[5].text = "修改建议"

    for f in findings:
        row = table.add_row().cells
        row[0].text = f.id
        row[1].text = f.severity
        row[2].text = f.category
        row[3].text = f.location
        row[4].text = f.issue
        row[5].text = f.recommendation

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"OK: wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
