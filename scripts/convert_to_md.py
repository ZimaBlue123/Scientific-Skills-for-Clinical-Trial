#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Unified document converter: Convert docx/pdf/rtf/doc to markdown.

Supports batch folder conversion or single-file conversion.

Usage
-----
    # Single file (standard markdown)
    py -3 scripts/convert_to_md.py input.docx -o output.md

    # Single file with numbered paragraphs/tables (##P1, ##T1 markers)
    py -3 scripts/convert_to_md.py input.docx -o output.md --mode numbered

    # Batch folder
    py -3 scripts/convert_to_md.py --folder review_materials -o review_materials/converted
"""
from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Set, Tuple

# Optional / heavy dependencies are imported lazily inside their respective
# helper functions to keep the module importable for tooling and to give
# clearer error messages when a backend is missing.

try:
    from docx import Document  # type: ignore
    from docx.oxml.table import CT_Tbl  # type: ignore
    from docx.oxml.text.paragraph import CT_P  # type: ignore
    from docx.table import Table as DocxTable  # type: ignore
    from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore
except ModuleNotFoundError:  # pragma: no cover - import guard
    Document = None  # type: ignore[assignment]
    CT_Tbl = None  # type: ignore[assignment]
    CT_P = None  # type: ignore[assignment]
    DocxTable = None  # type: ignore[assignment]
    DocxParagraph = None  # type: ignore[assignment]

LOG_FORMAT = "%(asctime)s [%(levelname)s] convert_to_md: %(message)s"
logger = logging.getLogger("convert_to_md")

DEFAULT_EXTENSIONS: Set[str] = {".docx", ".pdf", ".rtf"}
MARKER_TERMINATOR = "\n"

# markitdown handle (lazy). False means we tried and the import failed.
_markitdown = None  # type: ignore[var-annotated]


def _get_markitdown():
    """Return a cached markitdown instance, or False if unavailable."""
    global _markitdown
    if _markitdown is None:
        try:
            from markitdown import MarkItDown  # type: ignore

            _markitdown = MarkItDown()
        except ImportError:
            _markitdown = False
    return _markitdown


# ---------------------------------------------------------------------------
# Basic / numbered backends
# ---------------------------------------------------------------------------


def _convert_docx_basic(filepath: Path) -> Optional[str]:
    """Fallback: read .docx via python-docx and emit plain text."""
    if Document is None:  # pragma: no cover
        logger.error("python-docx is not installed; cannot read %s", filepath)
        return None
    try:
        doc = Document(str(filepath))  # type: ignore[misc]
        parts: List[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        joined = "\n\n".join(parts).strip()
        return joined or None
    except Exception as exc:  # noqa: BLE001
        logger.error("[python-docx] %s: %s", filepath.name, exc)
        return None


def _convert_pdf_basic(filepath: Path) -> Optional[str]:
    """Fallback: read .pdf via pypdf (preferred) or pdfplumber."""
    try:
        import pypdf  # type: ignore

        reader = pypdf.PdfReader(str(filepath))
        return "\n\n".join((p.extract_text() or "") for p in reader.pages).strip() or None
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        logger.error("[pypdf] %s: %s", filepath.name, exc)
        return None

    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(filepath)) as pdf:
            return "\n\n".join((p.extract_text() or "") for p in pdf.pages).strip() or None
    except ImportError:
        logger.error("Install pypdf or pdfplumber to read PDF files.")
    except Exception as exc:  # noqa: BLE001
        logger.error("[pdfplumber] %s: %s", filepath.name, exc)
    return None


def _clean_cell_text(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def table_to_md(table: "DocxTable", max_cols: Optional[int] = None) -> str:  # type: ignore[name-defined]
    """Render a docx Table as a markdown table."""
    rows: List[List[str]] = [[_clean_cell_text(cell.text) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    ncol = max(len(r) for r in rows)
    if max_cols is not None:
        ncol = min(ncol, max_cols)

    norm: List[List[str]] = [(r + [""] * (ncol - len(r)))[:ncol] for r in rows]
    header, body = norm[0], norm[1:] if len(norm) > 1 else []

    md_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * ncol) + " |",
    ]
    md_lines.extend("| " + " | ".join(r) + " |" for r in body)
    return "\n".join(md_lines)


def iter_block_items(doc: "Document") -> Iterable[Tuple[str, object]]:  # type: ignore[name-defined]
    """Yield ``("p", Paragraph)`` and ``("tbl", Table)`` in document order."""
    if CT_P is None or CT_Tbl is None or DocxParagraph is None or DocxTable is None:  # pragma: no cover
        raise RuntimeError("python-docx is not installed; cannot iterate docx blocks.")
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield ("p", DocxParagraph(child, doc))
        elif isinstance(child, CT_Tbl):
            yield ("tbl", DocxTable(child, doc))


def docx_to_md_numbered(src_path: Path, max_cols: Optional[int] = None) -> str:
    """Convert docx to markdown with ``##P{n}`` and ``##T{n}`` markers."""
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx is not installed; cannot read docx.")
    doc = Document(str(src_path))  # type: ignore[misc]

    out: List[str] = [f"# Source: {src_path.name}", ""]
    para_idx = 0
    table_idx = 0

    for kind, item in iter_block_items(doc):
        if kind == "p":
            text = " ".join(item.text.split())  # type: ignore[union-attr]
            if not text:
                continue
            para_idx += 1
            out.extend([f"## P{para_idx}", text, ""])
        else:
            md = table_to_md(item, max_cols=max_cols)  # type: ignore[arg-type]
            if not md.strip():
                continue
            table_idx += 1
            out.extend([f"## T{table_idx}", md, ""])

    return ("\n".join(out).rstrip() + MARKER_TERMINATOR).rstrip("\n") + MARKER_TERMINATOR


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def _read_rtf(src: Path) -> Optional[str]:
    try:
        from striprtf.striprtf import rtf_to_text  # type: ignore
    except ImportError:
        logger.error("Install 'striprtf' to read .rtf files.")
        return None
    try:
        return rtf_to_text(src.read_text(encoding="utf-8", errors="ignore")).strip() or None
    except Exception as exc:  # noqa: BLE001
        logger.error("[striprtf] %s: %s", src.name, exc)
        return None


def convert_file(
    src: Path,
    output_path: Optional[Path] = None,
    *,
    mode: str = "standard",
    max_cols: Optional[int] = None,
) -> Optional[Path]:
    """Convert a single file to markdown and return the output path (or None)."""
    if not src.exists() or not src.is_file():
        logger.error("file not found: %s", src)
        return None

    suffix = src.suffix.lower()
    if suffix not in {".docx", ".pdf", ".rtf", ".doc"}:
        logger.error("unsupported format: %s", suffix)
        return None

    content: Optional[str] = None

    if mode == "numbered" and suffix == ".docx":
        try:
            content = docx_to_md_numbered(src, max_cols=max_cols)
        except Exception as exc:  # noqa: BLE001
            logger.error("[numbered mode] %s: %s", src.name, exc)
            content = None  # fall through to standard backends

    if content is None:
        md = _get_markitdown()
        if md and suffix in {".docx", ".pdf"}:
            try:
                result = md.convert(str(src))
                content = (getattr(result, "text_content", "") or "").strip() or None
            except Exception as exc:  # noqa: BLE001
                logger.error("[markitdown] %s: %s", src.name, exc)

        if content is None:
            if suffix == ".docx":
                content = _convert_docx_basic(src)
            elif suffix == ".pdf":
                content = _convert_pdf_basic(src)
            elif suffix == ".rtf":
                content = _read_rtf(src)
            elif suffix == ".doc":
                logger.error(
                    ".doc format requires Word COM (Windows) — skipping %s", src.name
                )

    if not content or not content.strip():
        logger.error("[FAILED] could not extract content: %s", src.name)
        return None

    out_path = output_path or src.with_suffix(".md")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(content, encoding="utf-8")
    print(f"OK: {src.name} -> {out_path.name}")
    return out_path


def convert_folder(
    input_dir: Path,
    output_dir: Optional[Path] = None,
    extensions: Optional[Set[str]] = None,
    *,
    mode: str = "standard",
    max_cols: Optional[int] = None,
) -> List[Path]:
    """Convert every supported file in ``input_dir`` to ``output_dir``."""
    if extensions is None:
        extensions = set(DEFAULT_EXTENSIONS)
    extensions = {e.lower() if e.startswith(".") else f".{e.lower()}" for e in extensions}

    if output_dir is None:
        output_dir = input_dir / "converted"
    output_dir.mkdir(parents=True, exist_ok=True)

    created: List[Path] = []
    skipped: List[str] = []

    for f in sorted(input_dir.iterdir()):
        if not f.is_file() or f.suffix.lower() not in extensions:
            continue
        out_name = f.stem
        for c in '\\/:*?"<>|':
            out_name = out_name.replace(c, "_")
        out_path = output_dir / (out_name + ".md")

        result = convert_file(f, out_path, mode=mode, max_cols=max_cols)
        if result is not None:
            created.append(result)
        else:
            skipped.append(f.name)

    print("\n--- Summary ---")
    print(f"Converted: {len(created)}")
    print(f"Skipped/Failed: {len(skipped)}")
    if skipped:
        print(f"Skipped files: {', '.join(skipped)}")
    return created


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _parse_extensions(raw: str) -> Set[str]:
    parts = {p.strip().lower() for p in raw.split(",") if p.strip()}
    return {p if p.startswith(".") else f".{p}" for p in parts}


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="convert_to_md",
        description="Convert documents (docx/pdf/rtf/doc) to markdown.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  py -3 scripts/convert_to_md.py document.docx -o output.md
  py -3 scripts/convert_to_md.py document.docx -o output.md --mode numbered
  py -3 scripts/convert_to_md.py --folder review_materials
  py -3 scripts/convert_to_md.py --folder review_materials -o markdown_output/
        """,
    )
    parser.add_argument("input", nargs="?", help="Input file or folder")
    parser.add_argument(
        "--folder", "-f", help="Input folder (alternative to positional)"
    )
    parser.add_argument("-o", "--output", help="Output file or folder")
    parser.add_argument(
        "--extensions",
        default=",".join(sorted(DEFAULT_EXTENSIONS)),
        help="Comma-separated extensions to process (default: .docx,.pdf,.rtf).",
    )
    parser.add_argument(
        "--mode", "-m", default="standard",
        choices=["standard", "numbered"],
        help="Output mode: standard (default) or numbered (##P1/##T1 markers).",
    )
    parser.add_argument(
        "--max-cols", type=int, default=None,
        help="Max columns when rendering tables (numbered mode only).",
    )
    parser.add_argument(
        "--log-level", default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging verbosity (default: %(default)s).",
    )
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    logging.basicConfig(level=getattr(logging, args.log_level), format=LOG_FORMAT)

    if args.input:
        input_path = Path(args.input)
    elif args.folder:
        input_path = Path(args.folder)
    else:
        input_path = Path(__file__).resolve().parents[1] / "review_materials"

    if not input_path.exists():
        logger.error("input path does not exist: %s", input_path)
        return 1

    extensions = _parse_extensions(args.extensions)
    output_path = Path(args.output) if args.output else None

    if input_path.is_file():
        result = convert_file(
            input_path, output_path, mode=args.mode, max_cols=args.max_cols
        )
        return 0 if result else 1

    results = convert_folder(
        input_path,
        output_path,
        extensions,
        mode=args.mode,
        max_cols=args.max_cols,
    )
    return 0 if results else 1


if __name__ == "__main__":
    sys.exit(main())
