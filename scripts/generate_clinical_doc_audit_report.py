# -*- coding: utf-8 -*-
"""
Clinical Document Audit Report Generator

Generates a Word audit report for clinical trial documents with findings
organized by severity, category, and provides recommendations.

Usage:
    py -3 generate_clinical_doc_audit_report.py --folder <folder_path> [--output <output_path>]
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from common_scripts.docx_utils import apply_cn_en_fonts


@dataclass
class AuditFinding:
    """Represents a single audit finding."""
    id: str
    severity: str  # 重大/重要/一般/建议
    category: str
    location: str
    issue: str
    recommendation: str
    rationale: str = ""
    cross_ref: str = ""


@dataclass
class AuditReportConfig:
    """Configuration for audit report generation."""
    title: str = "临床试验文档审核报告"
    project_name: str = ""
    files_reviewed: list[str] = field(default_factory=list)
    findings: list[AuditFinding] = field(default_factory=list)


def create_audit_report(config: AuditReportConfig, output_path: Path) -> None:
    """Generate audit report Word document from config."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    apply_cn_en_fonts(doc)
    today = date.today().isoformat()

    # Title
    title = doc.add_paragraph(config.title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info section
    info = doc.add_paragraph()
    info.add_run('审核日期：').bold = True
    info.add_run(f'{today}\n')
    if config.project_name:
        info.add_run('项目名称：').bold = True
        info.add_run(f'{config.project_name}\n')
    if config.files_reviewed:
        info.add_run('审核文件：').bold = True
        info.add_run('\n'.join(f'• {f}' for f in config.files_reviewed))

    # Executive Summary
    doc.add_heading('审核概述', level=1)
    total = len(config.findings)
    severity_counts = {}
    for f in config.findings:
        severity_counts[f.severity] = severity_counts.get(f.severity, 0) + 1

    summary = doc.add_paragraph()
    summary.add_run(f'本报告共发现 {total} 个问题')
    if severity_counts:
        summary.add_run('，其中：')
        for sev, count in severity_counts.items():
            summary.add_run(f'\n• {sev}: {count}项')
    else:
        summary.add_run('。')

    # Findings by category
    if config.findings:
        doc.add_heading('问题清单', level=1)
        table = doc.add_table(rows=1, cols=7)
        hdr = table.rows[0].cells
        headers = ['编号', '风险等级', '类别', '位置', '问题', '修改建议', '依据/参考']
        for i, h in enumerate(headers):
            hdr[i].text = h

        for f in config.findings:
            row = table.add_row().cells
            row[0].text = f.id
            row[1].text = f.severity
            row[2].text = f.category
            row[3].text = f.location
            row[4].text = f.issue
            row[5].text = f.recommendation
            row[6].text = f.rationale if f.rationale else f.cross_ref

    # Conclusion
    doc.add_heading('审核结论', level=1)
    conclusion = doc.add_paragraph()
    if total == 0:
        conclusion.add_run('未发现明显问题。')
    else:
        conclusion.add_run(f'文档整体结构完整，内容详实。发现 {total} 处问题，')
        if severity_counts.get('重大', 0) > 0:
            conclusion.add_run('其中包含重大问题，建议优先处理。')
        else:
            conclusion.add_run('建议在正式使用前修正上述问题。')

    doc.save(str(output_path))
    print(f'Audit report generated: {output_path}')


def main() -> int:
    import argparse
    from pathlib import Path

    parser = argparse.ArgumentParser(description="Generate clinical document audit report.")
    parser.add_argument("--folder", type=str, required=True, help="Folder containing documents to audit")
    parser.add_argument("--output", type=str, default=None, help="Output path for audit report")
    parser.add_argument("--title", type=str, default="临床试验文档审核报告", help="Report title")
    parser.add_argument("--project", type=str, default="", help="Project name")
    args = parser.parse_args()

    folder = Path(args.folder)
    if not folder.is_absolute():
        folder = Path(__file__).resolve().parents[1] / folder

    # List files
    docx_files = sorted([f.name for f in folder.glob("*.docx")])
    
    if not docx_files:
        print(f'No docx files found in {folder}')
        return 1

    # Create config with findings for this specific project
    config = AuditReportConfig(
        title=args.title,
        project_name=args.project or folder.name,
        files_reviewed=docx_files,
        findings=[
            AuditFinding(
                id="SP-01", severity="已确认", category="申办公司信息",
                location="风险管理计划签名页、III期方案第168行等",
                issue="✓ 已确认申办公司包含：远大赛威信生命科学（杭州）有限公司、远大赛威信生命科学（南京）有限公司",
                recommendation="无需修改，信息完整",
                rationale="两家公司地址分别在杭州和南京，完整覆盖两个分公司"
            ),
            AuditFinding(
                id="S-01", severity="一般", category="语句重复",
                location="III期临床试验方案 第230行",
                issue="语句重复：'在本试验条件下，在本试验条件下，TVAX-018-1分别...'",
                recommendation="删除重复的'在本试验条件下'",
                rationale="编辑错误，影响文档专业性"
            ),
            AuditFinding(
                id="S-02", severity="一般", category="语句重复",
                location="I期统计分析计划 第596行",
                issue="语句重复：'安全性观察后安全性观察后'",
                recommendation="删除重复的'安全性观察后'",
                rationale="编辑错误，影响文档专业性"
            ),
            AuditFinding(
                id="S-03", severity="重要", category="语句不完整",
                location="III期临床试验方案 第603行",
                issue="'调整后把握度为'后缺少具体数值",
                recommendation="补充完整数值或删除不完整语句",
                rationale="关键统计参数不完整"
            ),
            AuditFinding(
                id="T-01", severity="一般", category="标点错误",
                location="I期统计分析计划 第831行",
                issue="双句号错误：'收集严重不良事件。。'",
                recommendation="删除多余句号",
                rationale="标点符号错误"
            ),
            AuditFinding(
                id="F-01", severity="建议", category="格式不一致",
                location="III期/I期临床试验方案",
                issue="试验编号格式不一致：I期用R001（三位），III期用R0001（四位）",
                recommendation="统一使用R+四位数字格式",
                rationale="两阶段方案格式应统一"
            ),
        ]
    )

    # Output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = folder / f"临床试验文档审核报告_{date.today().isoformat()}.docx"

    create_audit_report(config, output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
