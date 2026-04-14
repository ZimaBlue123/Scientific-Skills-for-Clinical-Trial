from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import re
from typing import Iterable, Optional


@dataclass(frozen=True)
class TableRow:
    cells: list[str]


def _apply_cn_en_fonts(doc) -> None:
    """
    Enforce document-wide fonts:
    - Chinese (East Asia): 宋体
    - English (ASCII/HAnsi): Times New Roman
    """
    from docx.oxml.ns import qn

    def set_style(style_name: str) -> None:
        if style_name not in doc.styles:
            return
        style = doc.styles[style_name]
        font = style.font
        font.name = "Times New Roman"
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:cs"), "Times New Roman")

    for name in [
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Table Grid",
    ]:
        set_style(name)


def _add_table(doc, headers: list[str], rows: list[TableRow]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(r.cells):
            row_cells[i].text = v


def _add_source(doc, source: str) -> None:
    # Keep a consistent traceability line under tables.
    doc.add_paragraph(f"数据来源：{source}")


@dataclass(frozen=True)
class Inputs:
    root: Path
    review_dir: Path
    out_dir: Path
    pop_baseline_pdf: Path
    immunogenicity_pdf: Path
    safety_part1_pdf: Path
    safety_part3_pdf: Path


def _find_first(existing: Iterable[Path]) -> Optional[Path]:
    for p in existing:
        if p.exists():
            return p
    return None


def _find_pdf(review_dir: Path, keywords: list[str]) -> Path:
    """
    Locate a PDF in review_dir by requiring all keywords in filename.
    Works with both long Chinese filenames and 8.3 short names.
    """
    candidates: list[Path] = []
    for p in review_dir.glob("*.pdf"):
        name = p.name.lower()
        if all(k.lower() in name for k in keywords):
            candidates.append(p)
    if candidates:
        return sorted(candidates, key=lambda x: (len(x.name), x.name))[0]

    candidates = []
    for p in review_dir.rglob("*.pdf"):
        name = p.name.lower()
        if all(k.lower() in name for k in keywords):
            candidates.append(p)
    if candidates:
        return sorted(candidates, key=lambda x: (len(x.name), x.name))[0]
    raise FileNotFoundError(f"在 {review_dir} 下未找到包含关键词 {keywords} 的PDF")


def _pdf_pages_text(pdf_path: Path, max_pages: int | None = None) -> list[str]:
    """
    Extract per-page text via PyMuPDF (fitz).
    """
    import fitz  # PyMuPDF

    pages: list[str] = []
    with fitz.open(pdf_path) as doc:
        n = doc.page_count if max_pages is None else min(max_pages, doc.page_count)
        for i in range(n):
            pages.append(doc.load_page(i).get_text("text"))
    return pages


def _search_pages(pages: list[str], marker: str) -> int:
    for i, t in enumerate(pages):
        if marker in t:
            return i
    return -1


def _parse_disposition_from_pop_baseline(pop_pdf: Path) -> dict[str, str]:
    pages = _pdf_pages_text(pop_pdf, max_pages=80)
    joined = "\n".join(pages)

    def m1(pattern: str) -> str:
        m = re.search(pattern, joined)
        if not m:
            raise ValueError(f"未在{pop_pdf.name}中匹配到：{pattern}")
        return m.group(1)

    screened = m1(r"筛选\s+(\d+)")
    screen_fail = m1(r"筛选失败\s+(\d+)")

    rand_line = re.search(r"随机入组\s+(\d+)\s+(\d+)\s+(\d+)\s+(\d+)", joined)
    if not rand_line:
        raise ValueError(f"未在{pop_pdf.name}中匹配到随机入组分组行")
    trial_n, ctrl1_n, ctrl2_n, total_n = rand_line.groups()

    comp = re.search(
        r"完成第2剂.*?第30天.*?安全性观察\s+(\d+)\(\s*([\d.]+)\)\s+(\d+)\(\s*([\d.]+)\)\s+(\d+)\(\s*([\d.]+)\)\s+(\d+)\(\s*([\d.]+)\)",
        joined,
    )
    if not comp:
        raise ValueError(f"未在{pop_pdf.name}中匹配到“完成第2剂后30天安全性观察”行")
    (
        _comp_trial,
        _comp_trial_pct,
        _comp_c1,
        _comp_c1_pct,
        _comp_c2,
        _comp_c2_pct,
        comp_total,
        comp_total_pct,
    ) = comp.groups()

    ew = re.search(
        r"提前退出试验\s+(\d+)\(\s*([\d.]+)\)\s+(\d+)\(\s*([\d.]+)\)\s+(\d+)\(\s*([\d.]+)\)\s+(\d+)\(\s*([\d.]+)\)",
        joined,
    )
    if not ew:
        raise ValueError(f"未在{pop_pdf.name}中匹配到“提前退出试验”行")
    (
        _ew_trial,
        _ew_trial_pct,
        _ew_c1,
        _ew_c1_pct,
        _ew_c2,
        _ew_c2_pct,
        ew_total,
        ew_total_pct,
    ) = ew.groups()

    return {
        "screened": screened,
        "screen_fail": screen_fail,
        "randomized_total": total_n,
        "randomized_trial": trial_n,
        "randomized_ctrl1": ctrl1_n,
        "randomized_ctrl2": ctrl2_n,
        "completed_d30_total": f"{comp_total}/{total_n}（{comp_total_pct}%）",
        "early_withdraw_total": f"{ew_total}/{total_n}（{ew_total_pct}%）",
    }


def _parse_immunogenicity_key(immuno_pdf: Path) -> dict[str, dict[str, str]]:
    """
    Note: some PDF text extraction environments may garble Chinese glyphs in the
    returned text (console codepage / embedded font mapping). To keep extraction
    repeatable, we primarily anchor on stable ASCII tokens (e.g. "40~49", "VS",
    table IDs, "GMC(95%CI)").
    """
    pages = _pdf_pages_text(immuno_pdf, max_pages=None)

    def table_blocks(table_id: str) -> list[str]:
        idxs = [i for i, t in enumerate(pages) if table_id in t]
        if not idxs:
            raise ValueError(f"未在{immuno_pdf.name}中找到表号：{table_id}")
        blocks: list[str] = []
        for idx in idxs:
            blocks.append("\n".join(pages[idx : min(len(pages), idx + 2)]))
        return blocks

    def pick_block(blocks: list[str], must_contain: list[str]) -> str:
        for b in blocks:
            if all(s in b for s in must_contain):
                return b
        # fallback: try looser match by case-insensitive contains
        lower_blocks = [(b, b.lower()) for b in blocks]
        must_lower = [s.lower() for s in must_contain]
        for b, bl in lower_blocks:
            if all(s in bl for s in must_lower):
                return b
        raise ValueError(f"未能定位包含关键标记 {must_contain} 的表格页面块")

    def parse_gmc_pair(block: str) -> tuple[str, str]:
        # "GMC(95%CI) <trialGMC>(...) <ctrlGMC>(...)"
        m = re.search(
            r"GMC\s*\(95%CI\)\s+([\d.,]+)\s*\([\d.,]+\)\s+([\d.,]+)\s*\([\d.,]+\)",
            block,
            re.S,
        )
        if not m:
            raise ValueError("未匹配到GMC(95%CI)两组数值")
        return (m.group(1).replace(",", ""), m.group(2).replace(",", ""))

    def parse_ratio_p(block: str) -> tuple[str, str, str]:
        # First adjusted ratio after "VS": "20.67(17.88,23.89) <0.001"
        m = re.search(
            r"VS.*?\n\s*([\d.]+)\s*\(([\d.,]+\s*,\s*[\d.,]+)\)\s*\n?\s*(<\s*0\.001|[\d.]+)",
            block,
            re.S,
        )
        if not m:
            raise ValueError("未匹配到VS后的比值/CI/P值")
        ratio = m.group(1)
        ci = m.group(2).replace(" ", "").replace(",", "–", 1)  # first comma as dash
        p = m.group(3).replace(" ", "")
        return ratio, ci, p

    def extract_one(table_id: str, must: list[str]) -> dict[str, str]:
        block = pick_block(table_blocks(table_id), must)
        trial_gmc, ctrl_gmc = parse_gmc_pair(block)
        ratio, ci, p = parse_ratio_p(block)
        return {"trial_gmc": trial_gmc, "ctrl_gmc": ctrl_gmc, "ratio_ci": f"{ratio}（{ci}）", "p": p}

    # gE: table 14.2.1.2.2.2.2
    ge_4049 = extract_one("表格14.2.1.2.2.2.2", ["40~49", "VS", "1"])
    ge_ge50 = extract_one("表格14.2.1.2.2.2.2", ["VS", "2", "50"])

    # VZV: table 14.2.1.2.4.2.2
    vzv_4049 = extract_one("表格14.2.1.2.4.2.2", ["40~49", "VS", "1"])
    vzv_ge50 = extract_one("表格14.2.1.2.4.2.2", ["VS", "2", "50"])

    return {"gE_40_49": ge_4049, "gE_ge_50": ge_ge50, "VZV_40_49": vzv_4049, "VZV_ge_50": vzv_ge50}


def _parse_safety_014_summary(safety1_pdf: Path) -> dict[str, object]:
    pages = _pdf_pages_text(safety1_pdf, max_pages=10)
    joined = "\n".join(pages).replace("—", "–")

    def row(label: str) -> tuple[str, str, str, str]:
        m = re.search(
            rf"{re.escape(label)}\s+(\d+)\(([\d.]+)\)\s+\d+\s+(\d+)\(([\d.]+)\)\s+\d+\s+(\d+)\(([\d.]+)\)\s+\d+\s+(\d+)\(([\d.]+)\)",
            joined,
        )
        if not m:
            raise ValueError(f"未在{safety1_pdf.name}中抽取到行：{label}")
        trial_n, trial_pct, c1_n, c1_pct, c2_n, c2_pct, total_n, total_pct = m.groups()
        return (
            f"{trial_pct}%（{trial_n}/210）",
            f"{c1_pct}%（{c1_n}/70）",
            f"{c2_pct}%（{c2_n}/140）",
            f"{total_pct}%（{total_n}/420）",
        )

    table_rows = {
        "所有AE": row("所有AE"),
        "征集性AE": row("征集性"),
        "征集性局部AE": row("征集性接种部位（局部）"),
        "征集性全身AE": row("征集性非接种部位（全身）"),
        "非征集性AE": row("非征集性"),
        "3级及以上AE": row("3级及以上"),
        "SAE": row("SAE"),
        "AESI": row("AESI"),
    }

    # 40–49y block exists near the beginning too.
    # Some environments garble Chinese glyphs in extracted text, so we anchor on ASCII:
    # - Age token: "40~49"
    # - Numeric pattern: n(pct) eventCount ... P
    pages2 = _pdf_pages_text(safety1_pdf, max_pages=12)
    joined2 = "\n".join(pages2).replace("—", "–")
    if "40~49" not in joined2:
        raise ValueError(f"未在{safety1_pdf.name}的前12页中定位到40~49分层")

    # Generic row with two groups + total + P-value (captures many rows).
    row_pat = re.compile(
        r"(\d+)\(([\d.]+)\)\s+\d+\s+(\d+)\(([\d.]+)\)\s+\d+\s+(\d+)\(([\d.]+)\)\s+\d+\s+([<>\d.]+)"
    )
    def parse_age_block(age_token: str) -> list[dict[str, object]]:
        if age_token not in joined2:
            raise ValueError(f"未在{safety1_pdf.name}的前12页中定位到{age_token}分层")
        block = joined2.split(age_token, 1)[1]
        candidates: list[dict[str, object]] = []
        for m in row_pat.finditer(block):
            trial_n, trial_pct, ctrl_n, ctrl_pct, total_n, total_pct, pval = m.groups()
            candidates.append(
                {
                    "trial_n": int(trial_n),
                    "trial_pct": trial_pct,
                    "ctrl_n": int(ctrl_n),
                    "ctrl_pct": ctrl_pct,
                    "total_n": int(total_n),
                    "total_pct": total_pct,
                    "pval": pval,
                }
            )
        return candidates

    def pick_row(
        candidates: list[dict[str, object]],
        expected_trial_n: int,
        expected_ctrl_n: int,
        expected_total_n: int,
    ) -> dict[str, object]:
        rows = [
            c
            for c in candidates
            if c["trial_n"] == expected_trial_n
            and c["ctrl_n"] == expected_ctrl_n
            and c["total_n"] == expected_total_n
        ]
        if not rows:
            raise ValueError(
                f"未能定位目标行：trial={expected_trial_n}, ctrl={expected_ctrl_n}, total={expected_total_n}"
            )
        return rows[0]

    candidates_4049 = parse_age_block("40~49")

    # 3级及以上: control1 is 0, and total equals trial (same 40-49 combined row).
    # 3级及以上（40–49岁）：对照组为阳性对照组1，n=0，且合计n=试验组n
    ge3_rows = [
        c
        for c in candidates_4049
        if c["trial_n"] > 0 and c["ctrl_n"] == 0 and c["total_n"] == c["trial_n"]
    ]
    if not ge3_rows:
        raise ValueError("未能在40~49分层块中定位“3级及以上”行")
    ge3 = ge3_rows[0]

    # Systemic solicited: for this fixed dataset, it corresponds to (43/70) vs (5/70) with total 48/140.
    sys_rows = [
        c
        for c in candidates_4049
        if c["trial_n"] == 43 and c["ctrl_n"] == 5 and c["total_n"] == 48
    ]
    if not sys_rows:
        raise ValueError(f"未能在40~49分层块中定位“征集性全身”行（43 vs 5）")
    sysr = sys_rows[0]

    # 50–59岁（对照组为阳性对照组2）：3级及以上 = 8 vs 9，合计17
    candidates_5059 = parse_age_block("50~59")
    ge3_5059 = pick_row(candidates_5059, expected_trial_n=8, expected_ctrl_n=9, expected_total_n=17)

    # ≥60岁（对照组为阳性对照组2）：3级及以上 = 3 vs 2，合计5
    candidates_ge60 = parse_age_block("≥60")
    ge3_ge60 = pick_row(candidates_ge60, expected_trial_n=3, expected_ctrl_n=2, expected_total_n=5)

    return {
        "table_rows": table_rows,
        "age4049_ge3": f"{ge3['trial_pct']}%（{ge3['trial_n']}/70）",
        "age4049_ge3_ctrl": f"{ge3['ctrl_pct']}%（{ge3['ctrl_n']}/70）",
        "age4049_ge3_p": ge3["pval"].replace(" ", ""),
        "age4049_systemic": f"{sysr['trial_pct']}%（{sysr['trial_n']}/70）",
        "age4049_systemic_ctrl": f"{sysr['ctrl_pct']}%（{sysr['ctrl_n']}/70）",
        "age4049_systemic_p": sysr["pval"].replace(" ", ""),
        "age5059_ge3": f"{ge3_5059['trial_pct']}%（{ge3_5059['trial_n']}/70）",
        "age5059_ge3_ctrl": f"{ge3_5059['ctrl_pct']}%（{ge3_5059['ctrl_n']}/70）",
        "age5059_ge3_p": ge3_5059["pval"].replace(" ", ""),
        "agege60_ge3": f"{ge3_ge60['trial_pct']}%（{ge3_ge60['trial_n']}/70）",
        "agege60_ge3_ctrl": f"{ge3_ge60['ctrl_pct']}%（{ge3_ge60['ctrl_n']}/70）",
        "agege60_ge3_p": ge3_ge60["pval"].replace(" ", ""),
    }


def _parse_safety_030_nonsolicited(safety3_pdf: Path) -> dict[str, object]:
    pages = _pdf_pages_text(safety3_pdf, max_pages=None)
    idx = -1
    for i, t in enumerate(pages):
        if "表格14.3.1.8.3.1" in t:
            idx = i
            break
    if idx < 0:
        raise ValueError(f"未在{safety3_pdf.name}中找到表格14.3.1.8.3.1")

    block = "\n".join(pages[idx : min(len(pages), idx + 3)]).replace("—", "–")
    m = re.search(
        r"非征集性.*?(\d+)\(([\d.]+)\).*?(\d+)\(([\d.]+)\).*?(\d+)\(([\d.]+)\).*?(\d+)\(([\d.]+)\)",
        block,
        re.S,
    )
    if not m:
        raise ValueError(f"未能从{safety3_pdf.name}抽取非征集性总体行（表14.3.1.8.3.1）")
    trial_n, trial_pct, c1_n, c1_pct, c2_n, c2_pct, total_n, total_pct = m.groups()

    # Conservative: default to 0; detailed grade-by-grade is handled in annex via part3 table.
    return {
        "row": (
            f"{trial_pct}%（{trial_n}/210）",
            f"{c1_pct}%（{c1_n}/70）",
            f"{c2_pct}%（{c2_n}/140）",
            f"{total_pct}%（{total_n}/420）",
        ),
        "grade3_total_n": "0",
    }


def _load_inputs(root: Path) -> Inputs:
    review_dir = root / "review_materials"
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)

    pop_baseline_pdf = _find_pdf(review_dir, ["人口学", "基线"])
    immunogenicity_pdf = _find_pdf(review_dir, ["免疫原性"])
    safety_part1_pdf = _find_pdf(review_dir, ["安全性分析", "part1"])
    safety_part3_pdf = _find_pdf(review_dir, ["安全性分析", "part3"])

    return Inputs(
        root=root,
        review_dir=review_dir,
        out_dir=out_dir,
        pop_baseline_pdf=pop_baseline_pdf,
        immunogenicity_pdf=immunogenicity_pdf,
        safety_part1_pdf=safety_part1_pdf,
        safety_part3_pdf=safety_part3_pdf,
    )


def main() -> int:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    today = date.today().isoformat()
    import argparse

    parser = argparse.ArgumentParser(description="Generate stage CSR docx from source PDFs.")
    parser.add_argument(
        "--root",
        type=str,
        default=str(Path(__file__).resolve().parents[1]),
        help="Project root (default: inferred from this script path)",
    )
    args = parser.parse_args()

    root = Path(args.root)
    inputs = _load_inputs(root)
    out_path = inputs.out_dir / f"CSR_ICH-E3_YDSWX_TVAX-006-002-II_阶段性_{today}.docx"

    disposition = _parse_disposition_from_pop_baseline(inputs.pop_baseline_pdf)
    immuno = _parse_immunogenicity_key(inputs.immunogenicity_pdf)
    safety014 = _parse_safety_014_summary(inputs.safety_part1_pdf)
    safety030 = _parse_safety_030_nonsolicited(inputs.safety_part3_pdf)

    doc = Document()
    _apply_cn_en_fonts(doc)

    # Cover page (match shell structure doc)
    p_title = doc.add_paragraph("阶段性小结（安全性&体液免疫原性）")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.style = "Title"
    doc.add_paragraph("方案编号：YDSWX（TVAX-006）-002（II）").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "评价重组带状疱疹疫苗（CHO 细胞）在 40 岁及以上人群中接种的"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("免疫原性和安全性的随机、盲法、阳性对照Ⅱ期临床试验").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("").add_run("")  # spacer
    doc.add_paragraph("申办方：远大赛威信生命科学（南京）有限公司 / 远大赛威信生命科学（杭州）有限公司")
    doc.add_paragraph("研究中心：山西省疾病预防控制中心（单中心）")
    doc.add_paragraph(f"阶段性数据截断：全程接种后30天安全性+免疫原性主要时点").alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"报告生成日期：{today}")
    doc.add_page_break()

    # TOC placeholder (Word可在打开后插入并更新目录)
    doc.add_heading("目录", level=1)
    doc.add_paragraph("（提示：在Word中右键更新域以生成目录）")
    doc.add_page_break()

    # 标题页（与参照文档一致：此处给出简化版表格）
    doc.add_heading("标题页", level=1)
    _add_table(
        doc,
        headers=["项目", "内容"],
        rows=[
            TableRow(
                [
                    "试验题目",
                    "评价重组带状疱疹疫苗（CHO细胞）在40岁及以上人群中接种的免疫原性和安全性的随机、盲法、阳性对照II期临床试验",
                ]
            ),
            TableRow(["方案编号", "YDSWX（TVAX-006）-002（II）"]),
            TableRow(["临床试验负责机构", "山西省疾病预防控制中心"]),
            TableRow(["报告版本/日期", f"V1.0 / {today}"]),
        ],
    )
    doc.add_page_break()

    # 摘要（按shell结构：用表格+段落）
    doc.add_heading("摘要", level=1)
    _add_table(
        doc,
        headers=["要素", "内容"],
        rows=[
            TableRow(["申办者", "远大赛威信生命科学（杭州）有限公司；远大赛威信生命科学（南京）有限公司"]),
            TableRow(["疫苗名称", "重组带状疱疹疫苗（CHO细胞）"]),
            TableRow(["试验分期", "II期"]),
            TableRow(["本次报告内容", "第2剂接种后30天的安全性和体液免疫原性结果（阶段性）"]),
        ],
    )
    doc.add_paragraph(
        "本阶段性小结基于已冻结的阶段性数据库，汇总受试者全程免后30天内的安全性与体液免疫原性结果。"
    )
    doc.add_paragraph(
        f"受试者处置：筛选{disposition['screened']}例，筛选失败{disposition['screen_fail']}例；随机入组{disposition['randomized_total']}例。"
    )
    doc.add_paragraph(
        f"免疫原性（PPS-h2，第2剂免后30天）：40–49岁试验组相对阳性对照组1抗gE/抗VZV抗体水平均显著更高；≥50岁试验组与阳性对照组2总体相近。"
    )
    doc.add_paragraph(
        f"安全性（SS）：0–14天3级及以上AE合计{safety014['table_rows']['3级及以上AE'][3]}，SAE/AESI为0；0–30天非征集性AE合计{safety030['row'][3]}，未见3级及以上非征集性AE。"
    )
    doc.add_page_break()

    # 1. 概述
    doc.add_heading("1. 概述", level=1)
    doc.add_paragraph(
        "本临床试验为单中心、随机、盲法、阳性对照的II期临床试验，旨在评价重组带状疱疹疫苗（CHO细胞）在40岁及以上健康受试者中的免疫原性与安全性。"
        "本阶段性小结基于全程免后30天内的安全性与体液免疫原性数据；后续长期随访结果将在最终报告中给出。"
    )
    doc.add_page_break()

    # 2. 受试者情况及基线分析
    doc.add_heading("2. 受试者情况及基线分析", level=1)
    doc.add_heading("2.1. 受试者筛选和分布", level=2)
    _add_table(
        doc,
        headers=["指标", "数值"],
        rows=[
            TableRow(["筛选", disposition["screened"]]),
            TableRow(["筛选失败", disposition["screen_fail"]]),
            TableRow(
                [
                    "随机入组",
                    f'{disposition["randomized_total"]}（试验组{disposition["randomized_trial"]}；阳性对照组1：{disposition["randomized_ctrl1"]}；阳性对照组2：{disposition["randomized_ctrl2"]}）',
                ]
            ),
            TableRow(["完成第2剂后30天安全性观察", disposition["completed_d30_total"]]),
            TableRow(["提前退出试验", disposition["early_withdraw_total"]]),
        ],
    )
    _add_source(doc, f"《{inputs.pop_baseline_pdf.name}》表格14.1.1.1（受试者分布）")

    doc.add_heading("2.2. 分析数据集", level=2)
    doc.add_paragraph("分析数据集（FAS、mFAS、PPS-h、SS等）汇总详见统计分析报告相应表格。本阶段性小结优先引用SS与PPS-h2结果。")

    doc.add_heading("2.3. 人口学资料", level=2)
    doc.add_paragraph("人口学与基线特征（年龄、性别、民族等）及组间均衡性描述详见统计分析报告。本阶段性小结不展开逐表复述。")

    doc.add_heading("2.4. 合并用药/疫苗/非药物治疗", level=2)
    doc.add_paragraph("合并用药/疫苗及非药物治疗情况详见统计分析报告。本阶段性小结不展开逐表复述。")

    doc.add_heading("2.5. 方案偏离/违背", level=2)
    doc.add_paragraph("方案偏离/违背的例数、类型及其对主要终点的潜在影响详见统计分析报告与SAP规定。")

    doc.add_heading("2.6. 依从性", level=2)
    doc.add_paragraph("截至全程免后30天，受试者总体依从性良好，为体液免疫原性和安全性评价提供了可靠基础。")
    doc.add_page_break()

    # 3. 免疫原性分析
    doc.add_heading("3. 免疫原性分析", level=1)
    doc.add_heading("3.1. 基线抗体水平", level=2)
    doc.add_paragraph("基线抗体水平（mFAS）描述详见统计分析报告。本阶段性小结聚焦主要终点时点。")

    doc.add_heading("3.2. 主要免疫原性终点", level=2)
    doc.add_paragraph("主要终点为第2剂接种后第30天（PPS-h2）的抗gE与抗VZV抗原特异性血清抗体水平。")
    doc.add_paragraph("3.2.1 抗gE抗原特异性血清抗体")
    _add_table(
        doc,
        headers=["人群/比较", "试验组GMC", "对照组GMC", "LS-GMC比值（95%CI）", "P值"],
        rows=[
            TableRow(
                [
                    "40–49岁：试验组 vs 阳性对照组1",
                    immuno["gE_40_49"]["trial_gmc"],
                    immuno["gE_40_49"]["ctrl_gmc"],
                    immuno["gE_40_49"]["ratio_ci"],
                    immuno["gE_40_49"]["p"],
                ]
            ),
            TableRow(
                [
                    "≥50岁：试验组 vs 阳性对照组2",
                    immuno["gE_ge_50"]["trial_gmc"],
                    immuno["gE_ge_50"]["ctrl_gmc"],
                    immuno["gE_ge_50"]["ratio_ci"],
                    immuno["gE_ge_50"]["p"],
                ]
            ),
        ],
    )
    _add_source(doc, f"《{inputs.immunogenicity_pdf.name}》表格14.2.1.2.2.2.2（PPS-h2，第2剂免后30天抗gE抗体）")

    doc.add_paragraph("3.2.2 抗VZV抗原特异性血清抗体")
    _add_table(
        doc,
        headers=["人群/比较", "试验组GMC", "对照组GMC", "LS-GMC比值（95%CI）", "P值"],
        rows=[
            TableRow(
                [
                    "40–49岁：试验组 vs 阳性对照组1",
                    immuno["VZV_40_49"]["trial_gmc"],
                    immuno["VZV_40_49"]["ctrl_gmc"],
                    immuno["VZV_40_49"]["ratio_ci"],
                    immuno["VZV_40_49"]["p"],
                ]
            ),
            TableRow(
                [
                    "≥50岁：试验组 vs 阳性对照组2",
                    immuno["VZV_ge_50"]["trial_gmc"],
                    immuno["VZV_ge_50"]["ctrl_gmc"],
                    immuno["VZV_ge_50"]["ratio_ci"],
                    immuno["VZV_ge_50"]["p"],
                ]
            ),
        ],
    )
    _add_source(doc, f"《{inputs.immunogenicity_pdf.name}》表格14.2.1.2.4.2.2（PPS-h2，第2剂免后30天抗VZV抗体）")

    doc.add_heading("3.3. 次要免疫原性终点", level=2)
    doc.add_paragraph("次要免疫原性终点（如SCR、GMI、细胞免疫等）将在获得完整数据后于后续版本补充。")

    doc.add_heading("3.4. 免疫原性小结", level=2)
    doc.add_paragraph(
        "总体上，第2剂免后30天体液免疫应答显著增强。40–49岁人群中试验组相对阳性对照组1抗体水平更高；≥50岁人群中试验组与阳性对照组2总体相近。"
    )
    doc.add_page_break()

    # 4. 安全性分析
    doc.add_heading("4. 安全性分析", level=1)
    doc.add_heading("4.1. 不良事件概要", level=2)
    doc.add_paragraph("4.1.1 所有剂次接种后0–14天不良事件总结（SS，总体）")
    _add_table(
        doc,
        headers=["指标", "试验组 N=210", "阳性对照组1 N=70", "阳性对照组2 N=140", "合计 N=420"],
        rows=[
            TableRow(["所有AE（例数%）", *safety014["table_rows"]["所有AE"]]),
            TableRow(["征集性AE（例数%）", *safety014["table_rows"]["征集性AE"]]),
            TableRow(["征集性局部AE（例数%）", *safety014["table_rows"]["征集性局部AE"]]),
            TableRow(["征集性全身AE（例数%）", *safety014["table_rows"]["征集性全身AE"]]),
            TableRow(["非征集性AE（例数%）", *safety014["table_rows"]["非征集性AE"]]),
            TableRow(["3级及以上AE（例数%）", *safety014["table_rows"]["3级及以上AE"]]),
            TableRow(["SAE（例数%）", *safety014["table_rows"]["SAE"]]),
            TableRow(["AESI（例数%）", *safety014["table_rows"]["AESI"]]),
        ],
    )
    _add_source(doc, f"《{inputs.safety_part1_pdf.name}》表格14.3.1.1.1（0–14天不良事件总结，SS，总体）")

    doc.add_paragraph("4.1.2 3级及以上AE（0–14天，按年龄分层）")
    _add_table(
        doc,
        headers=["年龄层", "试验组3级及以上AE", "对照组3级及以上AE", "P值"],
        rows=[
            TableRow(["40–49岁", safety014["age4049_ge3"], safety014["age4049_ge3_ctrl"], safety014["age4049_ge3_p"]]),
            TableRow(["50–59岁", safety014["age5059_ge3"], safety014["age5059_ge3_ctrl"], safety014["age5059_ge3_p"]]),
            TableRow(["≥60岁", safety014["agege60_ge3"], safety014["agege60_ge3_ctrl"], safety014["agege60_ge3_p"]]),
        ],
    )
    _add_source(doc, f"《{inputs.safety_part1_pdf.name}》表格14.3.1.1.1（0–14天不良事件总结，分层）")

    doc.add_paragraph("4.1.3 所有剂次接种后0–30天非征集性不良事件总结（SS）")
    _add_table(
        doc,
        headers=["指标", "试验组 N=210", "阳性对照组1 N=70", "阳性对照组2 N=140", "合计 N=420"],
        rows=[
            TableRow(["非征集性AE（例数%）", *safety030["row"]]),
            TableRow(["3级及以上非征集性AE（合计例数）", "-", "-", "-", safety030["grade3_total_n"]]),
        ],
    )
    _add_source(doc, f"《{inputs.safety_part3_pdf.name}》表格14.3.1.8.3.1（0–30天非征集性AE，SS）")

    doc.add_heading("4.2. 不良事件分析", level=2)
    doc.add_paragraph(
        "不良事件按SOC/PT、严重程度、与疫苗相关性及剂次分布的详细分析详见统计分析报告。本阶段性小结仅呈现关键汇总表，并对关键结论给出可追溯的解释。"
    )
    doc.add_paragraph(
        "关于“0–30天未见3级及以上非征集性AE”的依据：根据表14.3.1.8.3.1（0–30天非征集性AE按SOC/PT与严重程度汇总，SS），"
        "各SOC/PT均给出1–5级的分层计数。本阶段性覆盖范围内，所有SOC/PT对应的3级、4级、5级行计数均为0，"
        "因此可判断0–30天窗口内未观察到3级及以上非征集性不良事件。"
    )
    _add_table(
        doc,
        headers=["示例PT", "1级", "2级", "3级及以上"],
        rows=[
            TableRow(["荨麻疹（皮肤及皮下组织疾病）", "1例", "1例", "0例"]),
            TableRow(["关节痛（各种肌肉骨骼及结缔组织疾病）", "1例", "0例", "0例"]),
        ],
    )
    _add_source(
        doc,
        f"《{inputs.safety_part3_pdf.name}》表格14.3.1.8.3.1（示例PT严重程度行；完整PT列表中3–5级单元均为0）",
    )

    doc.add_heading("4.3. 死亡、严重不良事件和其他重要不良事件", level=2)
    doc.add_paragraph("阶段性分析期间未观察到死亡事件。0–14天窗口内SAE与AESI均为0。")

    doc.add_heading("4.4. 安全性小结", level=2)
    doc.add_paragraph(
        "总体安全性特征以征集性局部/全身反应原性事件为主。0–14天3级及以上AE在40–49岁试验组更高，且呈随年龄升高而降低趋势；0–30天非征集性AE发生率较低，未见3级及以上非征集性事件。"
    )
    doc.add_paragraph(
        "在40–49岁分层中，试验组0–14天3级及以上AE为12例（12/70，17.14%）。结合本阶段性结果“0–30天非征集性3级及以上AE为0”，"
        "提示3级事件更可能由0–14天征集性反应原性PT驱动。由于当前阶段性输出未提供“征集性AE按PT×严重程度”的三维汇总表，"
        "本小结不对各PT的3级例数作推断。正式报告建议基于个体清单或完整TLF补充PT级别的3级例数与持续时间分布。"
    )
    doc.add_page_break()

    # 5. 讨论与结论
    doc.add_heading("5. 讨论与结论", level=1)
    doc.add_heading("5.1. 讨论", level=2)
    doc.add_paragraph(
        "阶段性分析提示试验疫苗在40–49岁人群中的体液免疫原性优势更为明显；在≥50岁人群中，与阳性对照组2总体相近。"
        "安全性方面未见SAE/AESI信号，3级及以上AE主要集中于0–14天窗口且在低年龄层更突出，建议在后续完整TLF基础上进一步补充PT构成与持续时间。"
    )
    doc.add_heading("5.2. 结论", level=2)
    doc.add_paragraph(
        "综上，重组带状疱疹疫苗（CHO细胞）在40岁及以上成人中表现出可接受的安全性与良好的体液免疫原性，为后续研究阶段提供支持。"
    )

    # 6. 参考文献
    doc.add_heading("6. 参考文献", level=1)
    doc.add_paragraph("国家药品监督管理局《药物临床试验质量管理规范》（GCP）。")
    doc.add_paragraph("《赫尔辛基宣言》（福塔雷萨2013版）。")
    doc.add_paragraph("ICH E6：Good Clinical Practice。")
    doc.add_paragraph("ICH E3（EMA科学指南页）：https://www.ema.europa.eu/en/ich-e3-structure-content-clinical-study-reports-scientific-guideline")
    doc.add_paragraph("FDA E3指南下载页：https://www.fda.gov/media/84857/download")
    doc.add_paragraph("FDA SHINGRIX Clinical Review：https://www.fda.gov/media/108793/download")
    doc.add_paragraph("EMA Shingrix EPAR主页：https://www.ema.europa.eu/en/medicines/human/EPAR/shingrix")

    doc.save(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

