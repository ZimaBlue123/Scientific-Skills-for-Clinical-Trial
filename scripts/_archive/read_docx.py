import sys

import docx


def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Also get text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return "\n".join(full_text)


if __name__ == "__main__":
    text = read_docx(sys.argv[1])
    with open("output.txt", "w", encoding="utf-8") as f:
        f.write(text)
