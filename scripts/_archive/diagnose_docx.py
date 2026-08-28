"""Diagnose DOCX structure to understand paragraph styles and content mapping."""

from __future__ import annotations

import logging
import sys

from docx import Document  # type: ignore
from docx.document import Document as _DocxDocument  # type: ignore

LOGGER = logging.getLogger("diagnose_docx")
if not LOGGER.handlers:
    _handler = logging.StreamHandler(stream=sys.stderr)
    _handler.setFormatter(logging.Formatter("[%(levelname)s] %(name)s: %(message)s"))
    LOGGER.addHandler(_handler)
    LOGGER.setLevel(logging.INFO)


def diagnose(docx_path: str) -> None:
    """Print structural summary of *docx_path*.

    The function is best-effort: malformed files are logged and re-raised
    so the caller may decide how to react.
    """
    if not docx_path:
        raise ValueError("docx_path must be a non-empty string")

    try:
        doc: _DocxDocument = Document(docx_path)
    except (OSError, ValueError) as exc:
        LOGGER.error("Cannot open %s: %s", docx_path, exc)
        raise

    print(f"\n=== {docx_path} ===")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    # Section info
    for i, section in enumerate(doc.sections):
        print(f"\nSection {i}: width={section.page_width}, height={section.page_height}")
        print(
            f"  Margins: top={section.top_margin}, bottom={section.bottom_margin},"
            f" left={section.left_margin}, right={section.right_margin}"
        )

    # First 30 paragraphs with style info
    print("\n--- First 30 paragraphs ---")
    for i, para in enumerate(doc.paragraphs[:30]):
        text = (para.text or "")[:120] or "(empty)"
        style = para.style.name if para.style else "None"
        fmt_info = ""
        if para.runs:
            run = para.runs[0]
            fmt_info = f"font={run.font.name}, size={run.font.size}, bold={run.font.bold}"
        print(f"  P{i}: style={style}, {fmt_info}")
        print(f"       text: {text}")

    # Tables summary
    print(f"\n--- Tables ({len(doc.tables)}) ---")
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        header = [cell.text[:40] for cell in table.rows[0].cells] if rows else []
        print(f"  Table {i}: {rows}x{cols}, header: {header}")


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: python diagnose_docx.py <path-to-docx>", file=sys.stderr)
        return 2
    try:
        diagnose(argv[1])
    except (OSError, ValueError) as exc:
        LOGGER.error("Diagnosis failed: %s", exc)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
