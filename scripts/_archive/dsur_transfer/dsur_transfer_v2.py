"""
DSUR Content Transfer Script v2 - Improved section mapping.
Distinguishes TOC from body, handles sub-sections properly.
"""

import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def identify_section_key(text):
    """Identify DSUR section key from heading text."""
    t = text.strip().lower()

    mappings = [
        ("development safety update report (dsur)", "title_dsur"),
        ("executive summary", "exec_summary"),
        ("table of contents", "toc_heading"),
        ("confidentiality statement", "confidentiality"),
        ("1. introduction", "sec1"),
        ("2. worldwide marketing approval status", "sec2"),
        ("3. actions taken in the reporting period for safety reasons", "sec3"),
        ("4. changes to reference safety information", "sec4"),
        ("5. inventory of clinical trials", "sec5"),
        ("6. estimated cumulative exposure", "sec6"),
        ("6.1 cumulative subject exposure in development program", "sec6_1"),
        ("6.2 patient exposure from marketing experience", "sec6_2"),
        ("7. data in line listings and summary tabulations", "sec7"),
        ("7.1 reference information", "sec7_1"),
        ("7.2 line listings of serious adverse reactions", "sec7_2"),
        ("7.3 cumulative summary tabulations of serious adverse events", "sec7_3"),
        ("8. significant findings from clinical trials", "sec8"),
        ("8.1 completed clinical trials", "sec8_1"),
        ("8.2 ongoing clinical trials", "sec8_2"),
        ("8.3 long-term follow-up", "sec8_3"),
        ("8.4 other therapeutic use", "sec8_4"),
        ("8.5 new safety data related to combination therapies", "sec8_5"),
        ("9. safety findings from non-interventional studies", "sec9"),
        ("10. other clinical trial/study safety information", "sec10"),
        ("11. safety findings from marketing experience", "sec11"),
        ("12. non-clinical data", "sec12"),
        ("13. literature", "sec13"),
        ("14. other dsurs", "sec14"),
        ("15. lack of efficacy", "sec15"),
        ("16. region-specific information", "sec16"),
        ("17. late-breaking information", "sec17"),
        ("18. overall safety assessment", "sec18"),
        ("18.1 risk assessment", "sec18_1"),
        ("18.1.1 known adverse reactions", "sec18_1_1"),
        ("18.1.2 potential risks", "sec18_1_2"),
        ("18.1.3 adverse events resulting in death", "sec18_1_3"),
        ("18.1.4 potential impact of concomitant use", "sec18_1_4"),
        ("18.1.5", "sec18_1_5"),
        ("18.2 benefit-risk considerations", "sec18_2"),
        ("18.2.1 benefit assessment", "sec18_2_1"),
        (
            "18.2.2 pharmacodynamics suggested by completed clinical studies",
            "sec18_2_2",
        ),
        ("18.2.3 impact of identified adverse reactions", "sec18_2_3"),
        ("18.2.4 do other potential risks have clinical significance", "sec18_2_4"),
        ("18.2.5 are there any events requiring close attention", "sec18_2_5"),
        ("19. summary of important risks", "sec19"),
        ("19.1 important risks in the previous cycle", "sec19_1"),
        ("19.2 important risks in the current cycle", "sec19_2"),
        ("20. conclusions", "sec20"),
        ("appendices", "appendices"),
        ("appendix 1 -", "appendix1"),
        ("appendix 2 -", "appendix2"),
        ("appendix 3 -", "appendix3"),
        ("appendix 4 -", "appendix4"),
        ("appendix 5 -", "appendix5"),
        ("appendix 6 -", "appendix6"),
        ("appendix 7 -", "appendix7"),
        ("regional appendices", "regional"),
        ("appendix r1 -", "appendix_r1"),
        ("appendix r2 -", "appendix_r2"),
        ("appendix r3 -", "appendix_r3"),
        ("appendix r4 -", "appendix_r4"),
        ("appendix r5 -", "appendix_r5"),
        ("r1:", "appendix_r1"),
        ("r2:", "appendix_r2"),
        ("r3:", "appendix_r3"),
        ("r4:", "appendix_r4"),
        ("r5:", "appendix_r5"),
    ]

    for pattern, key in mappings:
        if t.startswith(pattern):
            return key
    return None


def build_source_index(doc):
    """Build source content index: section_key -> list of content paragraph texts."""
    sections = {}
    current_key = "__preamble__"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        key = identify_section_key(text)
        if key:
            current_key = key
            continue  # headings are separate

        sections.setdefault(current_key, []).append(text)

    return sections


def replace_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run of new_text,
    preserving the paragraph-level formatting from the style."""
    # Remove all existing runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    if not new_text:
        return

    # Create new run
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Copy formatting from paragraph style
    style = para.style
    if style:
        try:
            if style.font.name:
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), style.font.name)
                rFonts.set(qn("w:hAnsi"), style.font.name)
                rFonts.set(qn("w:eastAsia"), style.font.name)
                rPr.append(rFonts)
        except:
            pass
        try:
            if style.font.size:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(style.font.size.pt * 2)))
                rPr.append(sz)
        except:
            pass
        try:
            if style.font.bold:
                b = OxmlElement("w:b")
                rPr.append(b)
        except:
            pass

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    new_run.append(t)
    para._element.append(new_run)


def find_toc_boundaries(doc):
    """Find the TOC start and end paragraph indices."""
    toc_start = None
    toc_end = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        style_name = (para.style.name or "").lower()

        if text == "table of contents":
            toc_start = i
            continue

        if toc_start is not None and toc_end is None:
            # Check if we've exited TOC - looking for the first main section heading
            key = identify_section_key(text)
            if key and key not in ("toc_heading",):
                if "toc" not in style_name:
                    toc_end = i
                    break

    return toc_start, toc_end


def get_title_value(source_doc, keyword):
    """Extract a title page value from source paragraphs."""
    for para in source_doc.paragraphs:
        text = para.text.strip()
        if text.lower().startswith(keyword.lower()):
            return text
    return None


def main(template_path, source_path, output_path):
    template = Document(template_path)
    source = Document(source_path)

    # Build source content index
    src_index = build_source_index(source)
    print("=== Source Content Index ===")
    for k, v in src_index.items():
        if v:
            print(f"  {k}: {len(v)} paragraphs")

    # Find TOC boundaries
    toc_start, toc_end = find_toc_boundaries(template)
    print(f"\nTOC range: {toc_start} to {toc_end}")

    # Step 1: Handle title page fields
    for i, para in enumerate(template.paragraphs):
        text = para.text.strip()

        if text == "Development Safety Update Report (DSUR) No. 2":
            replace_para_text(para, "Development Safety Update Report (DSUR) No. 1")
        elif text == "Recombinant Hexavalent Norovirus Vaccine (Hansenula polymorpha)":
            replace_para_text(para, "Recombinant Varicella Vaccine (CHO Cell)")
        elif "Reporting Period: 25 April 2025" in text:
            replace_para_text(para, "Reporting Period: 26-Jun-2025 to 25-Jun-2026")
        elif "Date of Report: 05 June 2026" in text:
            replace_para_text(para, "Date of Report: 09-Jul-2026")
        elif text.startswith(
            "All information contained in this document is the property of Grand Theravac Life Science"
        ):
            replace_para_text(
                para,
                "All information contained in this document is the exclusive property of Grand Theravac Life Sciences (Nanjing) Co., Ltd. and Grand Theravac Life Sciences (Hangzhou) Co., Ltd., and is strictly confidential. It may not be disclosed or reproduced, in whole or in part, without prior written consent from the Sponsor.",
            )

    # Step 1b: Handle sponsor info table
    if template.tables:
        t = template.tables[0]
        try:
            t.rows[0].cells[
                1
            ].text = "Grand Theravac Life Sciences (Nanjing) Co., Ltd. / Grand Theravac Life Sciences (Hangzhou) Co., Ltd."
        except:
            pass
        try:
            t.rows[1].cells[1].text = ""
        except:
            pass
        try:
            t.rows[2].cells[1].text = ""
        except:
            pass
        try:
            t.rows[3].cells[1].text = ""
        except:
            pass
        try:
            t.rows[4].cells[1].text = ""
        except:
            pass

    # Step 2: Clear TOC (between toc_start and toc_end)
    if toc_start is not None and toc_end is not None:
        for i in range(toc_start, toc_end):
            para = template.paragraphs[i]
            text = para.text.strip()
            style_name = (para.style.name or "").lower()
            # Only clear actual TOC entries, keep the heading
            if "toc" in style_name or ("\t" in text):
                replace_para_text(para, "")

    # Step 3: Build template body structure (paragraph index -> section key)
    # Start after TOC
    body_start = toc_end if toc_end else 0

    template_sections = (
        {}
    )  # section_key -> list of paragraph indices (content only, no headings)
    current_section = None
    content_buffer = []

    for i in range(body_start, len(template.paragraphs)):
        para = template.paragraphs[i]
        text = para.text.strip()
        style_name = (para.style.name or "").lower()

        # Skip TOC style paragraphs in body
        if "toc" in style_name:
            continue

        sec_key = identify_section_key(text) if text else None

        if sec_key:
            # Save previous section's content paragraphs
            if current_section and content_buffer:
                template_sections.setdefault(current_section, []).extend(content_buffer)
                content_buffer = []
            current_section = sec_key
        elif text and current_section:
            # This is a content paragraph
            content_buffer.append(i)

    # Save last section
    if current_section and content_buffer:
        template_sections.setdefault(current_section, []).extend(content_buffer)

    print("\n=== Template Body Sections ===")
    for k, v in template_sections.items():
        print(f"  {k}: {len(v)} paragraphs (indices {v[0] if v else 'N/A'}...)")

    # Step 4: Replace content paragraph by paragraph
    replaced = 0
    cleared = 0

    for sec_key, tpl_indices in template_sections.items():
        src_paras = src_index.get(sec_key, [])

        if not src_paras:
            # No source content for this section - clear template content
            for idx in tpl_indices:
                para = template.paragraphs[idx]
                if para.text.strip():
                    replace_para_text(para, "")
                    cleared += 1
            continue

        # Replace template content with source content
        for j in range(min(len(src_paras), len(tpl_indices))):
            idx = tpl_indices[j]
            para = template.paragraphs[idx]
            replace_para_text(para, src_paras[j])
            replaced += 1

        # If source has more paragraphs than template, we have extra content
        # If template has more paragraphs, clear the extras
        if len(tpl_indices) > len(src_paras):
            for j in range(len(src_paras), len(tpl_indices)):
                idx = tpl_indices[j]
                para = template.paragraphs[idx]
                if para.text.strip():
                    replace_para_text(para, "")
                    cleared += 1

    print(f"\nReplaced: {replaced} paragraphs")
    print(f"Cleared: {cleared} paragraphs")

    # Step 5: Clear data tables (since source DSUR #1 has no clinical data)
    # Tables to clear based on source content
    tables_to_clear = {
        # Table index -> description
        1: "Estimated Cumulative Subject Exposure",  # T1 -> blank
        8: "Line Listings of Serious Adverse Reactions",  # T9 -> blank
        9: "Cumulative Summary Tabulations of SAEs",  # T10 -> blank
        11: "Cumulative Summary Tabulation of SARs",  # T12 -> blank
        12: "List of Subjects who Died",  # T13 -> blank
        13: "List of Subjects who Dropped Out",  # T14 -> blank
    }

    # For key data tables, clear content and write "Not Applicable" or keep as empty
    for t_idx in range(1, len(template.tables)):
        table = template.tables[t_idx]
        # For tables 7 (demographics), clear since no subjects
        if t_idx == 7:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and not any(
                        h in cell.text.lower()
                        for h in [
                            "age",
                            "male",
                            "female",
                            "total",
                            "race",
                            "asian",
                            "black",
                            "white",
                            "other",
                            "unknown",
                        ]
                    ):
                        for p in cell.paragraphs:
                            replace_para_text(p, "")

    # Step 6: Handle Appendix 3 table specifically
    # Source has a different table structure for Appendix 3
    # Template Table 5 (index 5) is "Ongoing Clinical Trials" for a study that doesn't exist in source
    # Source Table 0 (index 0) is the equivalent
    if template.tables and len(template.tables) > 5:
        tpl_table = template.tables[5]  # Appendix 3 - Ongoing Clinical Trials table

        if source.tables:
            src_table = source.tables[0]
            # Clear template table and replace with source data
            # First clear all data rows
            for r_idx in range(1, len(tpl_table.rows)):
                for cell in tpl_table.rows[r_idx].cells:
                    for p in cell.paragraphs:
                        replace_para_text(p, "")

            # Now set data from source
            if len(src_table.rows) > 1:  # has data row
                src_row = src_table.rows[1]
                tpl_row = tpl_table.rows[1]
                for c_idx in range(min(len(src_row.cells), len(tpl_row.cells))):
                    src_text = src_row.cells[c_idx].text.strip()
                    tpl_cell = tpl_row.cells[c_idx]
                    if tpl_cell.paragraphs:
                        replace_para_text(tpl_cell.paragraphs[0], src_text)

    # Step 7: Handle special empty tables for "not applicable" scenarios
    # Table 6 (index 6) - Completed Clinical Trials - already has "Not Applicable" style
    # Keep as is

    # Step 8: Clear the old detailed AE data tables (Template tables 2, 3, 9, 10)
    # These are Phase I/II AE tables specific to Norovirus study
    tables_to_empty = [2, 3, 9, 10]  # Indices of big data tables
    for t_idx in tables_to_empty:
        if t_idx < len(template.tables):
            table = template.tables[t_idx]
            for row in table.rows[1:]:  # Keep header
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            replace_para_text(p, "")

    # Save
    template.save(output_path)
    print(f"\nOutput saved to: {output_path}")


if __name__ == "__main__":
    tp = sys.argv[1] if len(sys.argv) > 1 else "template.docx"
    sp = sys.argv[2] if len(sys.argv) > 2 else "source.docx"
    op = sys.argv[3] if len(sys.argv) > 3 else "output.docx"
    main(tp, sp, op)
