"""
DSUR Content Transfer Script v4 - FINAL.
Key fixes:
1. Process exec_summary section directly
2. Filter TOC paragraphs from body sections
3. Handle missing section mappings for sec18 subsections
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def identify_section_key(text):
    t = text.strip().lower()
    mappings = [
        ('development safety update report (dsur)', 'title_dsur'),
        ('executive summary', 'exec_summary'),
        ('table of contents', 'toc_heading'),
        ('confidentiality statement', 'confidentiality'),
        ('1. introduction', 'sec1'),
        ('2. worldwide marketing approval status', 'sec2'),
        ('3. actions taken', 'sec3'),
        ('4. changes to reference safety information', 'sec4'),
        ('5. inventory of clinical trials', 'sec5'),
        ('6.1 cumulative subject exposure in development program', 'sec6_1'),
        ('6.2 patient exposure from marketing experience', 'sec6_2'),
        ('6. estimated cumulative exposure', 'sec6'),
        ('7.1 reference information', 'sec7_1'),
        ('7.2 line listings of serious adverse reactions', 'sec7_2'),
        ('7.3 cumulative summary tabulations of serious adverse events', 'sec7_3'),
        ('7. data in line listings and summary tabulations', 'sec7'),
        ('8.1 completed clinical trials', 'sec8_1'),
        ('8.2 ongoing clinical trials', 'sec8_2'),
        ('8.3 long-term follow-up', 'sec8_3'),
        ('8.4 other therapeutic use', 'sec8_4'),
        ('8.5 new safety data related to combination therapies', 'sec8_5'),
        ('8. significant findings from clinical trials', 'sec8'),
        ('9. safety findings from non-interventional studies', 'sec9'),
        ('10. other clinical trial/study safety information', 'sec10'),
        ('11. safety findings from marketing experience', 'sec11'),
        ('12. non-clinical data', 'sec12'),
        ('13. literature', 'sec13'),
        ('14. other dsurs', 'sec14'),
        ('15. lack of efficacy', 'sec15'),
        ('16. region-specific information', 'sec16'),
        ('17. late-breaking information', 'sec17'),
        ('18.1.1 known adverse reactions', 'sec18_1_1'),
        ('18.1.2 potential risks', 'sec18_1_2'),
        ('18.1.3 adverse events resulting in death', 'sec18_1_3'),
        ('18.1.4 potential impact of concomitant use', 'sec18_1_4'),
        ('18.1.5', 'sec18_1_5'),
        ('18.1 risk assessment', 'sec18_1'),
        ('18.2.1 benefit assessment', 'sec18_2_1'),
        ('18.2.2 pharmacodynamics suggested by completed clinical studies', 'sec18_2_2'),
        ('18.2.3 impact of identified adverse reactions', 'sec18_2_3'),
        ('18.2.4 do other potential risks have clinical significance', 'sec18_2_4'),
        ('18.2.5 are there any events requiring close attention', 'sec18_2_5'),
        ('18.2 benefit-risk considerations', 'sec18_2'),
        ('18. overall safety assessment', 'sec18'),
        ('19.1 important risks in the previous cycle', 'sec19_1'),
        ('19.2 important risks in the current cycle', 'sec19_2'),
        ('19. summary of important risks', 'sec19'),
        ('20. conclusions', 'sec20'),
        ('appendices', 'appendices'),
        ('appendix 1 -', 'appendix1'),
        ('appendix 2 -', 'appendix2'),
        ('appendix 3 -', 'appendix3'),
        ('appendix 4 -', 'appendix4'),
        ('appendix 5 -', 'appendix5'),
        ('appendix 6 -', 'appendix6'),
        ('appendix 7 -', 'appendix7'),
        ('regional appendices', 'regional'),
        ('appendix r1 -', 'appendix_r1'),
        ('appendix r2 -', 'appendix_r2'),
        ('appendix r3 -', 'appendix_r3'),
        ('appendix r4 -', 'appendix_r4'),
        ('appendix r5 -', 'appendix_r5'),
    ]
    for pattern, key in mappings:
        if t.startswith(pattern):
            return key
    return None


def is_toc_paragraph(para):
    """Check if paragraph is a TOC entry."""
    style_name = (para.style.name or '').lower()
    text = para.text.strip()
    if 'toc' in style_name:
        return True
    if '\t' in text and identify_section_key(text):
        return True
    return False


def build_para_index(doc):
    """Build section -> list of paragraph indices, filtering out TOC entries."""
    index = {}
    current_key = '__preamble__'
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Skip TOC paragraphs
        if is_toc_paragraph(para):
            continue
        
        if not text:
            continue
        
        key = identify_section_key(text)
        if key:
            current_key = key
            index.setdefault(current_key, []).append(i)
            continue
        
        index.setdefault(current_key, []).append(i)
    
    return index


def replace_para_text(para, new_text):
    """Replace paragraph text, preserving paragraph style formatting."""
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    
    if not new_text:
        return
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    style = para.style
    if style:
        try:
            if style.font.name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), style.font.name)
                rFonts.set(qn('w:hAnsi'), style.font.name)
                rFonts.set(qn('w:eastAsia'), style.font.name)
                rPr.append(rFonts)
        except:
            pass
        try:
            if style.font.size:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(style.font.size.pt * 2)))
                rPr.append(sz)
        except:
            pass
    
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = new_text
    new_run.append(t)
    para._element.append(new_run)


def clear_cell_text(cell):
    """Clear all text in a table cell."""
    for para in cell.paragraphs:
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = ''
        new_run = OxmlElement('w:r')
        new_run.append(t)
        para._element.append(new_run)


def clear_table_data(table, keep_header=True):
    """Clear data rows."""
    start_row = 1 if keep_header else 0
    for row in table.rows[start_row:]:
        for cell in row.cells:
            clear_cell_text(cell)


def get_content_indices(para_indices, doc):
    """From a list of paragraph indices, return only content (non-heading) indices."""
    result = []
    for idx in para_indices:
        text = doc.paragraphs[idx].text.strip()
        if not text:
            continue
        if identify_section_key(text):
            # This is a heading (could be sub-section)
            # Include sub-section headings as they may need content replacement
            continue
        result.append(idx)
    return result


def get_src_content_texts(para_indices, doc):
    """Extract content texts from source paragraphs."""
    texts = []
    for idx in para_indices:
        text = doc.paragraphs[idx].text.strip()
        if not text:
            continue
        if identify_section_key(text):
            continue
        texts.append(text)
    return texts


def replace_section_content(template_doc, tpl_indices, src_texts, operation_name=""):
    """Replace content in template paragraphs with source content."""
    tpl_content = get_content_indices(tpl_indices, template_doc)
    
    replaced = 0
    cleared = 0
    
    for j in range(min(len(src_texts), len(tpl_content))):
        para = template_doc.paragraphs[tpl_content[j]]
        replace_para_text(para, src_texts[j])
        replaced += 1
    
    for j in range(len(src_texts), len(tpl_content)):
        para = template_doc.paragraphs[tpl_content[j]]
        if para.text.strip():
            replace_para_text(para, '')
            cleared += 1
    
    if operation_name:
        print(f"  {operation_name}: replaced {replaced}, cleared {cleared}")
    
    return replaced, cleared


def main(template_path, source_path, output_path):
    template = Document(template_path)
    source = Document(source_path)
    
    tpl_index = build_para_index(template)
    src_index = build_para_index(source)
    
    print("=== Template Sections (TOC filtered) ===")
    for k, v in sorted(tpl_index.items()):
        hdr = template.paragraphs[v[0]].text[:60] if v else ''
        print(f"  {k}: {len(v)} paras, '{hdr}'")
    
    print("\n=== Source Sections ===")
    for k, v in sorted(src_index.items()):
        hdr = source.paragraphs[v[0]].text[:60] if v else ''
        print(f"  {k}: {len(v)} paras, '{hdr}'")
    
    total_replaced = 0
    total_cleared = 0
    
    # ---- PHASE 1: Title Page ----
    print("\n--- Title Page ---")
    for i, para in enumerate(template.paragraphs):
        text = para.text.strip()
        if 'Development Safety Update Report (DSUR) No. 2' in text and 'Executive Summary' not in text:
            replace_para_text(para, 'Development Safety Update Report (DSUR) No. 1')
        elif text == 'Recombinant Hexavalent Norovirus Vaccine (Hansenula polymorpha)':
            replace_para_text(para, 'Recombinant Varicella Vaccine (CHO Cell)')
        elif 'Reporting Period: 25 April 2025 to 24 April 2026' in text:
            replace_para_text(para, 'Reporting Period: 26-Jun-2025 to 25-Jun-2026')
        elif 'Date of Report: 05 June 2026' in text:
            replace_para_text(para, 'Date of Report: 09-Jul-2026')
        elif text.startswith('All information contained in this document is the property of Grand Theravac Life Science'):
            replace_para_text(para, 'All information contained in this document is the exclusive property of Grand Theravac Life Sciences (Nanjing) Co., Ltd. and Grand Theravac Life Sciences (Hangzhou) Co., Ltd., and is strictly confidential. It may not be disclosed or reproduced, in whole or in part, without prior written consent from the Sponsor.')
    
    # Sponsor info table
    if template.tables:
        t = template.tables[0]
        cells_to_set = [
            (0, 'Grand Theravac Life Sciences (Nanjing) Co., Ltd. / Grand Theravac Life Sciences (Hangzhou) Co., Ltd.'),
            (1, ''), (2, ''), (3, ''), (4, ''),
        ]
        for row_idx, val in cells_to_set:
            try:
                clear_cell_text(t.rows[row_idx].cells[1])
                if val:
                    t.rows[row_idx].cells[1].paragraphs[0].add_run(val)
            except:
                pass
    
    # ---- PHASE 2: Executive Summary (pre-TOC) ----
    print("\n--- Executive Summary ---")
    if 'exec_summary' in tpl_index and 'exec_summary' in src_index:
        tpl_paras = tpl_index['exec_summary']
        src_paras = src_index['exec_summary']
        src_texts = get_src_content_texts(src_paras, source)
        r, c = replace_section_content(template, tpl_paras, src_texts, "Exec Summary")
        total_replaced += r
        total_cleared += c
    
    # ---- PHASE 3: Clear TOC ----
    print("\n--- TOC Clearing ---")
    # Find TOC boundary and clear all TOC paragraphs
    toc_found = False
    for i, para in enumerate(template.paragraphs):
        text = para.text.strip().lower()
        style_name = (para.style.name or '').lower()
        
        if text == 'table of contents':
            toc_found = True
            continue
        
        if toc_found and 'toc' in style_name:
            replace_para_text(para, '')
        
        if toc_found:
            # Check if we've passed the TOC
            key = identify_section_key(text)
            if key and key != 'toc_heading' and 'toc' not in style_name:
                # Also check this is actually a real heading, not a TOC entry
                if not is_toc_paragraph(para):
                    toc_found = False
    
    # ---- PHASE 4: Main Body Sections ----
    print("\n--- Main Body ---")
    
    # Define sections to skip (containers, title page, already processed, pure heading)
    skip_sections = {
        '__preamble__', 'toc_heading', 'exec_summary', 
        'title_dsur', 'confidentiality',
        'appendices', 'regional',
        'sec6', 'sec7', 'sec8', 'sec18', 'sec18_1', 'sec19',
        # Container headings with no direct content - only sub-sections
    }
    
    # Handle mapping for sec18_2: source has combined benefit-risk in sec18_2,
    # template has separate sec18_2_1, sec18_2_2, sec18_2_3, sec18_2_4, sec18_2_5
    # Source puts all benefit-risk content (3 paragraphs) under sec18_2
    if 'sec18_2' in src_index and 'sec18_2_1' in tpl_index:
        src_texts_all = get_src_content_texts(src_index['sec18_2'], source)
        # Assign to sec18_2_1 in template (the benefit assessment sub-section)
        # Template expects individual sub-sections; source combines them
        if src_texts_all:
            tpl_paras = tpl_index['sec18_2_1']
            r, c = replace_section_content(template, tpl_paras, src_texts_all[:min(3, len(src_texts_all))], "sec18_2_1 (combined)")
            total_replaced += r
            total_cleared += c
        # Clear remaining benefit-risk sub-sections
        for sub in ['sec18_2_2', 'sec18_2_3', 'sec18_2_4', 'sec18_2_5']:
            if sub in tpl_index:
                r, c = replace_section_content(template, tpl_index[sub], [], sub)
                total_replaced += r
                total_cleared += c
    
    for sec_key in sorted(tpl_index.keys()):
        if sec_key in skip_sections:
            continue
        # Skip already processed sec18_2 subsections
        if sec_key.startswith('sec18_2'):
            continue
        
        if sec_key not in src_index:
            # No source content - clear
            r, c = replace_section_content(template, tpl_index[sec_key], [], sec_key)
            total_replaced += r
            total_cleared += c
            continue
        
        src_texts = get_src_content_texts(src_index[sec_key], source)
        r, c = replace_section_content(template, tpl_index[sec_key], src_texts, sec_key)
        total_replaced += r
        total_cleared += c
    
    # Handle sec18_1 heading content (source has no direct content under sec18_1, only subsections)
    # Already handled via sub-sections
    
    print(f"\nTotal: replaced {total_replaced}, cleared {total_cleared}")
    
    # ---- PHASE 5: Tables ----
    print("\n--- Tables ---")
    
    # Clear data tables
    data_tables = [1, 2, 3, 7, 9]  # Exposure, AE I, AE II, Demographics, SAE cumulative
    for t_idx in data_tables:
        if t_idx < len(template.tables):
            clear_table_data(template.tables[t_idx], keep_header=True)
    
    # Clear Appendix R4 tables
    for t_idx in range(14, min(20, len(template.tables))):
        clear_table_data(template.tables[t_idx], keep_header=True)
    
    # Handle Appendix 3 table (Table 5): Replace with source's trial table
    if len(template.tables) > 5 and source.tables:
        tpl_ap3 = template.tables[5]
        src_ap3 = source.tables[0]
        
        # Copy source data
        if len(src_ap3.rows) > 1:
            src_row = src_ap3.rows[1]
            for r_idx in range(1, min(len(tpl_ap3.rows), 2)):
                tpl_row = tpl_ap3.rows[r_idx]
                for c_idx in range(min(len(src_row.cells), len(tpl_row.cells))):
                    src_text = src_row.cells[c_idx].text.strip()
                    clear_cell_text(tpl_row.cells[c_idx])
                    if src_text:
                        tpl_row.cells[c_idx].paragraphs[0].add_run(src_text)
            # Clear remaining rows
            for r_idx in range(2, len(tpl_ap3.rows)):
                for cell in tpl_ap3.rows[r_idx].cells:
                    clear_cell_text(cell)
    
    template.save(output_path)
    print(f"\nSaved: {output_path}")


if __name__ == "__main__":
    tp = sys.argv[1] if len(sys.argv) > 1 else "template.docx"
    sp = sys.argv[2] if len(sys.argv) > 2 else "source.docx"
    op = sys.argv[3] if len(sys.argv) > 3 else "output.docx"
    main(tp, sp, op)
