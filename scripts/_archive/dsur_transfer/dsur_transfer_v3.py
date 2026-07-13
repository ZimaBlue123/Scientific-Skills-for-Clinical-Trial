"""
DSUR Content Transfer Script v3.
Complete rewrite with proper section handling:
1. Pre-TOC: title page + executive summary
2. TOC: cleared
3. Main body: sections 1-20 + appendices
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def identify_section_key(text):
    """Identify DSUR section key from heading text."""
    t = text.strip().lower()
    
    mappings = [
        ('development safety update report (dsur)', 'title_dsur'),
        ('executive summary', 'exec_summary'),
        ('table of contents', 'toc_heading'),
        ('confidentiality statement', 'confidentiality'),
        ('1. introduction', 'sec1'),
        ('2. worldwide marketing approval status', 'sec2'),
        ('3. actions taken', 'sec3'),
        ('4. changes to reference safety information', 'sec4'),
        ('5. inventory of clinical trials', 'sec5'),
        ('6.1 cumulative subject exposure in development program', 'sec6_1'),
        ('6.2 patient exposure from marketing experience', 'sec6_2'),
        ('6. estimated cumulative exposure', 'sec6'),
        ('7.1 reference information', 'sec7_1'),
        ('7.2 line listings of serious adverse reactions', 'sec7_2'),
        ('7.3 cumulative summary tabulations of serious adverse events', 'sec7_3'),
        ('7. data in line listings and summary tabulations', 'sec7'),
        ('8.1 completed clinical trials', 'sec8_1'),
        ('8.2 ongoing clinical trials', 'sec8_2'),
        ('8.3 long-term follow-up', 'sec8_3'),
        ('8.4 other therapeutic use', 'sec8_4'),
        ('8.5 new safety data related to combination therapies', 'sec8_5'),
        ('8. significant findings from clinical trials', 'sec8'),
        ('9. safety findings from non-interventional studies', 'sec9'),
        ('10. other clinical trial/study safety information', 'sec10'),
        ('11. safety findings from marketing experience', 'sec11'),
        ('12. non-clinical data', 'sec12'),
        ('13. literature', 'sec13'),
        ('14. other dsurs', 'sec14'),
        ('15. lack of efficacy', 'sec15'),
        ('16. region-specific information', 'sec16'),
        ('17. late-breaking information', 'sec17'),
        ('18.1.1 known adverse reactions', 'sec18_1_1'),
        ('18.1.2 potential risks', 'sec18_1_2'),
        ('18.1.3 adverse events resulting in death', 'sec18_1_3'),
        ('18.1.4 potential impact of concomitant use', 'sec18_1_4'),
        ('18.1.5', 'sec18_1_5'),
        ('18.1 risk assessment', 'sec18_1'),
        ('18.2.1 benefit assessment', 'sec18_2_1'),
        ('18.2.2 pharmacodynamics suggested by completed clinical studies', 'sec18_2_2'),
        ('18.2.3 impact of identified adverse reactions', 'sec18_2_3'),
        ('18.2.4 do other potential risks have clinical significance', 'sec18_2_4'),
        ('18.2.5 are there any events requiring close attention', 'sec18_2_5'),
        ('18.2 benefit-risk considerations', 'sec18_2'),
        ('18. overall safety assessment', 'sec18'),
        ('19.1 important risks in the previous cycle', 'sec19_1'),
        ('19.2 important risks in the current cycle', 'sec19_2'),
        ('19. summary of important risks', 'sec19'),
        ('20. conclusions', 'sec20'),
        ('appendices', 'appendices'),
        ('appendix 1 -', 'appendix1'),
        ('appendix 2 -', 'appendix2'),
        ('appendix 3 -', 'appendix3'),
        ('appendix 4 -', 'appendix4'),
        ('appendix 5 -', 'appendix5'),
        ('appendix 6 -', 'appendix6'),
        ('appendix 7 -', 'appendix7'),
        ('regional appendices', 'regional'),
        ('appendix r1 -', 'appendix_r1'),
        ('r1:', 'appendix_r1'),
        ('appendix r2 -', 'appendix_r2'),
        ('r2:', 'appendix_r2'),
        ('appendix r3 -', 'appendix_r3'),
        ('r3:', 'appendix_r3'),
        ('appendix r4 -', 'appendix_r4'),
        ('r4:', 'appendix_r4'),
        ('appendix r5 -', 'appendix_r5'),
        ('r5:', 'appendix_r5'),
    ]
    
    for pattern, key in mappings:
        if t.startswith(pattern):
            return key
    return None


def build_para_index(doc):
    """Build a paragraph index mapping: para_index -> section_key or None.
    For heading paragraphs, returns the section key.
    For content paragraphs, returns None (they belong to the previous heading's section).
    """
    index = {}  # section_key -> list of paragraph indices
    current_key = '__preamble__'
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        key = identify_section_key(text)
        if key:
            current_key = key
            index.setdefault(current_key, []).append(i)  # heading paragraph
            continue
        
        index.setdefault(current_key, []).append(i)  # content paragraph
    
    return index


def replace_para_text(para, new_text):
    """Replace all content in a paragraph with new text, preserving style formatting."""
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    
    if not new_text:
        return
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    style = para.style
    if style:
        try:
            if style.font.name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), style.font.name)
                rFonts.set(qn('w:hAnsi'), style.font.name)
                rFonts.set(qn('w:eastAsia'), style.font.name)
                rPr.append(rFonts)
        except:
            pass
        try:
            if style.font.size:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(style.font.size.pt * 2)))
                rPr.append(sz)
        except:
            pass
        try:
            if style.font.bold:
                b = OxmlElement('w:b')
                rPr.append(b)
        except:
            pass
    
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = new_text
    new_run.append(t)
    para._element.append(new_run)


def clear_cell_text(cell):
    """Clear all text in a table cell."""
    for para in cell.paragraphs:
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = ''
        new_run = OxmlElement('w:r')
        new_run.append(t)
        para._element.append(new_run)


def clear_table_data(table, keep_header=True):
    """Clear all data rows in a table, keeping or clearing header."""
    start_row = 1 if keep_header else 0
    for row in table.rows[start_row:]:
        for cell in row.cells:
            clear_cell_text(cell)


def main(template_path, source_path, output_path):
    template = Document(template_path)
    source = Document(source_path)
    
    # Build paragraph indices for both documents
    tpl_index = build_para_index(template)
    src_index = build_para_index(source)
    
    print("=== Template Sections ===")
    for k, v in tpl_index.items():
        hdr = template.paragraphs[v[0]].text[:60] if v else ''
        print(f"  {k}: {len(v)} paras, heading='{hdr}'")
    
    print("\n=== Source Sections ===")
    for k, v in src_index.items():
        if len(v) <= 1:
            continue
        hdr = source.paragraphs[v[0]].text[:60] if v else ''
        print(f"  {k}: {len(v)} paras, heading='{hdr}'")
    
    # ---- PHASE 1: Title Page ----
    # Replace title fields
    for i, para in enumerate(template.paragraphs):
        text = para.text.strip()
        
        if 'Development Safety Update Report (DSUR) No. 2' in text:
            replace_para_text(para, 'Development Safety Update Report (DSUR) No. 1')
        elif text == 'Recombinant Hexavalent Norovirus Vaccine (Hansenula polymorpha)':
            replace_para_text(para, 'Recombinant Varicella Vaccine (CHO Cell)')
        elif 'Reporting Period: 25 April 2025 to 24 April 2026' in text:
            replace_para_text(para, 'Reporting Period: 26-Jun-2025 to 25-Jun-2026')
        elif 'Date of Report: 05 June 2026' in text:
            replace_para_text(para, 'Date of Report: 09-Jul-2026')
        elif text.startswith('All information contained in this document is the property of Grand Theravac Life Science'):
            replace_para_text(para, 'All information contained in this document is the exclusive property of Grand Theravac Life Sciences (Nanjing) Co., Ltd. and Grand Theravac Life Sciences (Hangzhou) Co., Ltd., and is strictly confidential. It may not be disclosed or reproduced, in whole or in part, without prior written consent from the Sponsor.')
    
    # Sponsor info table
    if template.tables:
        t = template.tables[0]
        cells_to_set = [
            (0, 'Grand Theravac Life Sciences (Nanjing) Co., Ltd. / Grand Theravac Life Sciences (Hangzhou) Co., Ltd.'),
            (1, ''),
            (2, ''),
            (3, ''),
            (4, ''),
        ]
        for row_idx, val in cells_to_set:
            try:
                clear_cell_text(t.rows[row_idx].cells[1])
                t.rows[row_idx].cells[1].paragraphs[0].add_run(val)
            except:
                pass
    
    # ---- PHASE 2: Pre-TOC Executive Summary ----
    # The exec summary content is between "Executive Summary" heading and "Table of Contents" heading
    # In the template, these are in __preamble__ section
    
    preamble_tpl = tpl_index.get('__preamble__', [])
    preamble_src = src_index.get('__preamble__', [])
    
    # Find exec_summary and confidentiality content paragraphs in template preamble
    exec_summary_src_paras = src_index.get('exec_summary', [])
    confidentiality_src_paras = src_index.get('confidentiality', [])
    
    # The template preamble contains: title stuff, empty paragraphs, confidentiality heading+content, exec summary heading+content
    # We need to replace exec summary content paragraphs only
    
    # Identify content-only (non-heading, non-title) paragraphs in preamble
    tpl_preamble_content = []
    in_exec = False
    for idx in preamble_tpl:
        para = template.paragraphs[idx]
        text = para.text.strip()
        
        # Skip if it's a heading
        if identify_section_key(text):
            # Check if we're entering executive summary
            if 'executive summary' in text.lower():
                in_exec = True
            elif 'table of contents' in text.lower() or 'confidentiality' in text.lower():
                in_exec = False
            continue
        
        if in_exec and text:
            tpl_preamble_content.append(idx)
    
    # Get exec summary content from source (excluding the heading)
    src_exec_content = []
    for idx in exec_summary_src_paras:
        text = source.paragraphs[idx].text.strip()
        if identify_section_key(text):
            continue
        if text:
            src_exec_content.append(text)
    
    print(f"\nExec Summary: template has {len(tpl_preamble_content)} content paras, source has {len(src_exec_content)}")
    
    # Replace exec summary content
    for j, src_text in enumerate(src_exec_content):
        if j < len(tpl_preamble_content):
            para = template.paragraphs[tpl_preamble_content[j]]
            replace_para_text(para, src_text)
    
    # Clear extra template exec summary paragraphs
    for j in range(len(src_exec_content), len(tpl_preamble_content)):
        para = template.paragraphs[tpl_preamble_content[j]]
        replace_para_text(para, '')
    
    # ---- PHASE 3: Clear Table of Contents ----
    toc_indices = tpl_index.get('toc_heading', [])
    if toc_indices:
        # Find TOC start (the TOC heading paragraph)
        toc_start = toc_indices[0]
        
        # Find where TOC ends - next major section heading
        toc_end = None
        for i in range(toc_start + 1, len(template.paragraphs)):
            text = template.paragraphs[i].text.strip()
            key = identify_section_key(text)
            # First real section heading after TOC
            if key and key not in ('toc_heading', 'title_dsur', 'exec_summary'):
                style_name = (template.paragraphs[i].style.name or '').lower()
                if 'toc' not in style_name:
                    toc_end = i
                    break
        
        print(f"TOC range: {toc_start} to {toc_end}")
        
        if toc_end:
            for i in range(toc_start + 1, toc_end):
                para = template.paragraphs[i]
                style_name = (para.style.name or '').lower()
                text = para.text.strip()
                if 'toc' in style_name or '\t' in text or text:
                    replace_para_text(para, '')
    
    # ---- PHASE 4: Main Body Sections ----
    # Map section keys between template and source
    replaced_count = 0
    cleared_count = 0
    
    for sec_key in sorted(tpl_index.keys()):
        # Skip preamble, TOC, and title-page sections
        if sec_key in ('__preamble__', 'toc_heading', 'exec_summary', 
                        'title_dsur', 'confidentiality', 'appendices', 'regional'):
            # 'appendices' and 'regional' are container headings, content is in sub-sections
            continue
        
        tpl_paras = tpl_index[sec_key]
        if len(tpl_paras) <= 1:
            continue  # Only heading, no content
        
        src_paras = src_index.get(sec_key, [])
        
        # Get content-only paragraphs (skip heading)
        tpl_content = [idx for idx in tpl_paras 
                       if not identify_section_key(template.paragraphs[idx].text.strip())]
        
        src_content_texts = []
        for idx in src_paras:
            text = source.paragraphs[idx].text.strip()
            if identify_section_key(text):
                continue
            if text:
                src_content_texts.append(text)
        
        if not src_content_texts:
            # No source content - clear template content
            for idx in tpl_content:
                para = template.paragraphs[idx]
                if para.text.strip():
                    replace_para_text(para, '')
                    cleared_count += 1
            continue
        
        # Replace content
        for j in range(min(len(src_content_texts), len(tpl_content))):
            para = template.paragraphs[tpl_content[j]]
            replace_para_text(para, src_content_texts[j])
            replaced_count += 1
        
        # Clear extra template paragraphs
        for j in range(len(src_content_texts), len(tpl_content)):
            para = template.paragraphs[tpl_content[j]]
            if para.text.strip():
                replace_para_text(para, '')
                cleared_count += 1
    
    print(f"\nBody: replaced {replaced_count}, cleared {cleared_count}")
    
    # ---- PHASE 5: Handle Tables ----
    # Table 0: Sponsor info (already done above)
    # Table 1 (index 1): Estimated Cumulative Subject Exposure table -> clear all data
    # Table 2 (index 2): Phase I AE summary table -> clear
    # Table 3 (index 3): Phase II AE summary table -> clear
    # Table 4 (index 4): Appendix 2 regulatory requirements -> keep as "Not Applicable"
    # Table 5 (index 5): Appendix 3 Ongoing Clinical Trials -> replace with source data
    # Table 6 (index 6): Appendix 3 Completed Clinical Trials -> keep "Not Applicable"
    # Table 7 (index 7): Appendix 4 Demographics -> clear
    # Table 8 (index 8): Appendix 5 Line Listing SARs -> keep blank
    # Table 9 (index 9): Appendix 6 Cumulative SAEs -> clear
    # Table 10 (index 10): Appendix 7 Literature -> keep or clear
    # Table 11 (index 11): Appendix R1 SARs -> keep blank
    # Table 12 (index 12): Appendix R2 Deaths -> keep current
    # Table 13 (index 13): Appendix R3 Dropouts -> keep current
    # Tables 14-19: Appendix R4 -> clear/keep as N/A
    
    # Clear data tables that contain Norovirus-specific data
    data_table_indices = [1, 2, 3, 7, 9]  # Exposure, AE tables, Demographics, SAE cumulative
    for t_idx in data_table_indices:
        if t_idx < len(template.tables):
            clear_table_data(template.tables[t_idx], keep_header=True)
    
    # Clear Appendix R4 tables (14-19)
    for t_idx in range(14, min(20, len(template.tables))):
        clear_table_data(template.tables[t_idx], keep_header=True)
    
    # Handle Appendix 3 table (index 5): Replace with source's ongoing clinical trial table
    if len(template.tables) > 5 and source.tables:
        tpl_ap3 = template.tables[5]
        src_ap3 = source.tables[0]
        
        # Keep header, clear data
        for r_idx in range(1, len(tpl_ap3.rows)):
            for cell in tpl_ap3.rows[r_idx].cells:
                clear_cell_text(cell)
        
        # Copy source data row
        if len(src_ap3.rows) > 1:
            src_row = src_ap3.rows[1]
            tpl_row = tpl_ap3.rows[1]
            for c_idx in range(min(len(src_row.cells), len(tpl_row.cells))):
                src_text = src_row.cells[c_idx].text.strip()
                if src_text:
                    clear_cell_text(tpl_row.cells[c_idx])
                    tpl_row.cells[c_idx].paragraphs[0].add_run(src_text)
    
    # Keep tables 4, 6, 8, 10, 11, 12, 13 as they already have "Not Applicable" or blank content
    
    # ---- PHASE 6: Handle Appendix 7 References ----
    # Source appendix7 has 1 paragraph, template has many literature references
    # Clear all reference paragraphs except the heading
    app7 = tpl_index.get('appendix7', [])
    src_app7 = src_index.get('appendix7', [])
    
    if app7 and src_app7:
        src_texts = []
        for idx in src_app7:
            text = source.paragraphs[idx].text.strip()
            if identify_section_key(text):
                continue
            if text:
                src_texts.append(text)
        
        # Get content-only paragraphs
        app7_content = [idx for idx in app7 
                        if not identify_section_key(template.paragraphs[idx].text.strip())]
        
        for j, src_text in enumerate(src_texts):
            if j < len(app7_content):
                replace_para_text(template.paragraphs[app7_content[j]], src_text)
        
        for j in range(len(src_texts), len(app7_content)):
            replace_para_text(template.paragraphs[app7_content[j]], '')
    
    # Save
    template.save(output_path)
    print(f"\nFinal output saved to: {output_path}")


if __name__ == "__main__":
    tp = sys.argv[1] if len(sys.argv) > 1 else "template.docx"
    sp = sys.argv[2] if len(sys.argv) > 2 else "source.docx"
    op = sys.argv[3] if len(sys.argv) > 3 else "output.docx"
    main(tp, sp, op)
