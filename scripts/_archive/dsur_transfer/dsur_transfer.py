"""
DSUR Content Transfer Script.
Transfers content from source DSUR to template DSUR while preserving template formatting.
Strategy: Identify sections by heading patterns, replace content paragraph-by-paragraph,
preserving the template's paragraph styles and run-level formatting.
"""

import sys

from docx import Document
from docx.oxml.ns import qn


def get_para_text(para):
    """Get full paragraph text."""
    return para.text.strip()


def get_all_runs_text(runs):
    """Join all run texts."""
    return "".join(r.text for r in runs)


def identify_section(text):
    """
    Identify which DSUR section a paragraph belongs to.
    Returns section key or None.
    """
    t = text.strip().lower()

    # Title page
    if t.startswith("development safety update report"):
        return "title_dsur"
    if t.startswith("recombinant") and ("vaccine" in t):
        return "title_product"
    if t.startswith("reporting period"):
        return "title_reporting_period"
    if t.startswith("date of report"):
        return "title_date"
    if t.startswith("sponsor name"):
        return "title_sponsor"
    if t.startswith("address:"):
        return "title_address"
    if t.startswith("telephone:"):
        return "title_telephone"
    if t.startswith("signature of approver"):
        return "title_signature"
    if t.startswith("date of approval"):
        return "title_approval"
    if t == "confidentiality statement":
        return "confidentiality_heading"

    # Executive Summary
    if t == "executive summary":
        return "exec_summary_heading"

    # Table of Contents
    if t == "table of contents":
        return "toc_heading"

    # Main body sections
    section_patterns = [
        ("1. introduction", "sec1_intro"),
        ("2. worldwide marketing approval status", "sec2_marketing"),
        ("3. actions taken in the reporting period for safety reasons", "sec3_actions"),
        ("4. changes to reference safety information", "sec4_rsi"),
        ("5. inventory of clinical trials", "sec5_inventory"),
        (
            "5. inventory of clinical trials ongoing and completed during the reporting period",
            "sec5_inventory",
        ),
        ("6. estimated cumulative exposure", "sec6_exposure"),
        ("6.1 cumulative subject exposure in development program", "sec6_1"),
        ("6.2 patient exposure from marketing experience", "sec6_2"),
        ("7. data in line listings and summary tabulations", "sec7_data"),
        ("7.1 reference information", "sec7_1"),
        ("7.2 line listings of serious adverse reactions", "sec7_2"),
        ("7.3 cumulative summary tabulations of serious adverse events", "sec7_3"),
        ("8. significant findings from clinical trials", "sec8_findings"),
        ("8.1 completed clinical trials", "sec8_1"),
        ("8.2 ongoing clinical trials", "sec8_2"),
        ("8.3 long-term follow-up", "sec8_3"),
        ("8.4 other therapeutic use", "sec8_4"),
        ("8.5 new safety data related to combination therapies", "sec8_5"),
        ("9. safety findings from non-interventional studies", "sec9_noninterv"),
        ("10. other clinical trial/study safety information", "sec10_other"),
        ("11. safety findings from marketing experience", "sec11_marketing"),
        ("12. non-clinical data", "sec12_nonclinical"),
        ("13. literature", "sec13_literature"),
        ("14. other dsurs", "sec14_other_dsur"),
        ("15. lack of efficacy", "sec15_efficacy"),
        ("16. region-specific information", "sec16_region"),
        ("17. late-breaking information", "sec17_late"),
        ("18. overall safety assessment", "sec18_overall"),
        ("18.1 risk assessment", "sec18_1"),
        ("18.1.1 known adverse reactions", "sec18_1_1"),
        ("18.1.2 potential risks", "sec18_1_2"),
        ("18.1.3 adverse events resulting in death", "sec18_1_3"),
        ("18.1.4 potential impact of concomitant use", "sec18_1_4"),
        ("18.1.5", "sec18_1_5"),  # other missing info
        ("18.2 benefit-risk considerations", "sec18_2"),
        ("18.2.1 benefit assessment", "sec18_2_1"),
        ("18.2.2 pharmacodynamics", "sec18_2_2"),
        ("18.2.3 impact of identified adverse reactions", "sec18_2_3"),
        ("18.2.4 do other potential risks have clinical significance", "sec18_2_4"),
        ("18.2.5 are there any events requiring close attention", "sec18_2_5"),
        ("19. summary of important risks", "sec19_summary"),
        ("19.1 important risks in the previous cycle", "sec19_1"),
        ("19.2 important risks in the current cycle", "sec19_2"),
        ("20. conclusions", "sec20_conclusions"),
    ]

    for pattern, key in section_patterns:
        if t.startswith(pattern):
            return key

    # Appendices
    if t == "appendices":
        return "appendices_heading"
    if t.startswith("appendix 1"):
        return "appendix1"
    if t.startswith("appendix 2"):
        return "appendix2"
    if t.startswith("appendix 3"):
        return "appendix3"
    if t.startswith("appendix 4"):
        return "appendix4"
    if t.startswith("appendix 5"):
        return "appendix5"
    if t.startswith("appendix 6"):
        return "appendix6"
    if t.startswith("appendix 7"):
        return "appendix7"
    if t == "regional appendices":
        return "regional_heading"
    if t.startswith("appendix r1"):
        return "appendix_r1"
    if t.startswith("appendix r2"):
        return "appendix_r2"
    if t.startswith("appendix r3"):
        return "appendix_r3"
    if t.startswith("appendix r4"):
        return "appendix_r4"
    if t.startswith("appendix r5"):
        return "appendix_r5"

    return None


def build_content_index(doc):
    """
    Build a content index from the source document.
    Returns: dict of section_key -> list of paragraph texts
    """
    sections = {}
    current_section = "__preamble__"
    sections[current_section] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        section_key = identify_section(text)
        if section_key:
            current_section = section_key
            if current_section not in sections:
                sections[current_section] = []
            # Don't add the heading itself to the content list
            # We'll handle headings separately
        else:
            sections.setdefault(current_section, []).append(text)

    return sections


def get_paragraph_style_info(para):
    """Extract style and formatting info from a paragraph."""
    info = {
        "style": para.style.name if para.style else None,
        "alignment": para.alignment,
        "runs": [],
    }
    for run in para.runs:
        run_info = {
            "text": run.text,
            "bold": run.font.bold,
            "italic": run.font.italic,
            "size": run.font.size,
            "name": run.font.name,
            "color": (run.font.color.rgb if run.font.color and run.font.color.rgb else None),
        }
        info["runs"].append(run_info)
    return info


def replace_paragraph_text(para, new_text):
    """
    Replace paragraph text while preserving the first run's formatting.
    Clears all runs and creates one new run with the original formatting.
    """
    if not para.runs:
        return

    # Preserve formatting from the first run
    first_run = para.runs[0]

    # Clear all existing runs
    for run in para.runs:
        run.text = ""

    # Set text in first run
    first_run.text = new_text

    # Remove any extra empty runs
    # We keep the first run and clear the rest
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)


def replace_paragraph_full(para, new_text):
    """
    Completely replace paragraph content: remove all runs, add one new run
    with the paragraph's default formatting, using the new text.
    """
    # Store formatting from first run if available
    para._element.find(qn("w:pPr"))

    # Remove all existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # If no text, just clear
    if not new_text:
        return

    # Add a single new run with the text
    from docx.oxml import OxmlElement

    new_run_elem = OxmlElement("w:r")

    # Copy run properties from the paragraph's default style
    # For template-only, we use the paragraph style's formatting
    rPr = OxmlElement("w:rPr")

    # Try to get font info from the paragraph style
    style = para.style
    if style and style.font:
        if style.font.name:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), style.font.name)
            rFonts.set(qn("w:hAnsi"), style.font.name)
            rFonts.set(qn("w:eastAsia"), style.font.name)
            rPr.append(rFonts)
        if style.font.size:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(style.font.size.pt * 2)))
            rPr.append(sz)
        if style.font.bold:
            b = OxmlElement("w:b")
            rPr.append(b)

    new_run_elem.append(rPr)

    # Add text element
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    new_run_elem.append(t)

    para._element.append(new_run_elem)


def transfer_content(template_path, source_path, output_path):
    """
    Main transfer function.
    1. Reads both documents
    2. Builds content index from source
    3. Walks through template and replaces content
    """
    template_doc = Document(template_path)
    source_doc = Document(source_path)

    # Build content index from source
    source_sections = build_content_index(source_doc)

    # Build template structure
    template_structure = []
    current_section = "__preamble__"

    for i, para in enumerate(template_doc.paragraphs):
        text = para.text.strip()
        section_key = identify_section(text) if text else None
        if section_key:
            current_section = section_key
        template_structure.append(
            {
                "index": i,
                "text": text,
                "section": current_section,
                "is_heading": bool(section_key) if text else False,
                "is_toc": "toc" in (para.style.name or "").lower(),
                "is_empty": not text,
            }
        )

    # Print section mapping for debugging
    print("=== Source Sections ===")
    for key, paras in source_sections.items():
        print(f"  {key}: {len(paras)} paragraphs")

    print("\n=== Template Structure ===")
    current_sec = None
    for item in template_structure:
        if item["section"] != current_sec or item["is_heading"]:
            current_sec = item["section"]
            if item["is_heading"]:
                print(f"  [{current_sec}] HEADING: {item['text'][:80]}")

    # Now do the actual content replacement
    # Strategy: For each section in the template, find corresponding source section
    # and replace paragraph texts

    # Build a paragraph index per section in the template
    template_section_paras = {}
    for item in template_structure:
        sec = item["section"]
        template_section_paras.setdefault(sec, []).append(item["index"])

    # Sections we should NOT replace (they stay as template)
    # - __preamble__: title page content gets special handling
    # - toc-related items

    # Create content replacement plan
    # For each section that exists in both, replace template paragraphs with source paragraphs

    replaced_count = 0

    for sec_key in source_sections:
        if sec_key == "__preamble__":
            continue

        source_paras = source_sections[sec_key]
        if not source_paras:
            continue

        template_indices = template_section_paras.get(sec_key, [])
        if not template_indices:
            print(f"  SKIP {sec_key}: no matching template section")
            continue

        # Skip the heading paragraph (first one in the section)
        # The heading itself we handle based on what section it is
        content_indices = []
        for idx in template_indices:
            item = template_structure[idx] if idx < len(template_structure) else None
            if not item:
                continue
            # Skip heading paragraphs, TOC paragraphs, and empty paragraphs
            if item["is_heading"]:
                continue
            if item["is_toc"]:
                continue
            content_indices.append(idx)

        # Replace content paragraphs with source paragraphs
        # We might have more template paragraphs than source, or vice versa
        for j, src_text in enumerate(source_paras):
            if j < len(content_indices):
                tpl_idx = content_indices[j]
                para = template_doc.paragraphs[tpl_idx]
                replace_paragraph_full(para, src_text)
                replaced_count += 1
            else:
                # Source has more paragraphs than template content slots
                # We'll need to insert new paragraphs
                if content_indices:
                    # Insert after the last content paragraph
                    template_doc.paragraphs[content_indices[-1]]
                    # We can't easily insert, so we'll skip for now
                    pass

        # If template has more paragraphs than source, clear the extras
        for j in range(len(source_paras), len(content_indices)):
            tpl_idx = content_indices[j]
            para = template_doc.paragraphs[tpl_idx]
            replace_paragraph_full(para, "")
            replaced_count += 1

    print(f"\nReplaced {replaced_count} paragraphs")

    # Now handle special sections:
    # 1. Title page - replace key fields
    # 2. Tables need section-aware handling
    # 3. TOC - skip for now (will need manual update or can be removed)

    # Handle title page paragraphs
    # Find and replace DSUR number, product name, dates
    for i, para in enumerate(template_doc.paragraphs):
        text = para.text.strip()

        if text == "Development Safety Update Report (DSUR) No. 2":
            replace_paragraph_full(para, "Development Safety Update Report (DSUR) No. 1")
        elif text == "Recombinant Hexavalent Norovirus Vaccine (Hansenula polymorpha)":
            replace_paragraph_full(para, "Recombinant Varicella Vaccine (CHO Cell)")
        elif text.startswith("Reporting Period: 25 April 2025 to 24 April 2026"):
            replace_paragraph_full(para, "Reporting Period: 26-Jun-2025 to 25-Jun-2026")
        elif text.startswith("Date of Report: 05 June 2026"):
            replace_paragraph_full(para, "Date of Report: 09-Jul-2026")

    # Handle sponsor info table
    if template_doc.tables:
        sponsor_table = template_doc.tables[0]
        # Update sponsor name
        sponsor_table.rows[0].cells[
            1
        ].text = "Grand Theravac Life Sciences (Nanjing) Co., Ltd. / Grand Theravac Life Sciences (Hangzhou) Co., Ltd."
        # Update address
        sponsor_table.rows[1].cells[1].text = ""
        # Update drug safety officer
        sponsor_table.rows[2].cells[1].text = ""
        # Update telephone
        sponsor_table.rows[3].cells[1].text = ""
        # Update email
        sponsor_table.rows[4].cells[1].text = ""

    # Handle Confidentiality paragraph
    for i, para in enumerate(template_doc.paragraphs):
        text = para.text.strip()
        if text.startswith(
            "All information contained in this document is the property of Grand Theravac Life Science"
        ):
            replace_paragraph_full(
                para,
                "All information contained in this document is the exclusive property of Grand Theravac Life Sciences (Nanjing) Co., Ltd. and Grand Theravac Life Sciences (Hangzhou) Co., Ltd., and is strictly confidential. It may not be disclosed or reproduced, in whole or in part, without prior written consent from the Sponsor.",
            )

    # Handle TOC - mark for manual update
    for i, para in enumerate(template_doc.paragraphs):
        if para.style and "toc" in (para.style.name or "").lower():
            # Clear old TOC entries
            replace_paragraph_full(para, "")

    # Handle data tables - they need to be cleared or marked as N/A since
    # the source DSUR #1 has no clinical data
    # Tables to keep or modify based on source content

    # Now handle tables based on the source's situation (no clinical trials initiated)
    # Most data tables should be cleared/emptied

    # Save
    template_doc.save(output_path)
    print(f"\nOutput saved to: {output_path}")


if __name__ == "__main__":
    template_path = sys.argv[1] if len(sys.argv) > 1 else "template.docx"
    source_path = sys.argv[2] if len(sys.argv) > 2 else "source.docx"
    output_path = sys.argv[3] if len(sys.argv) > 3 else "output.docx"
    transfer_content(template_path, source_path, output_path)
