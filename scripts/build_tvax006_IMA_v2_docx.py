# -*- coding: utf-8 -*-
"""Merge TVAX-006_海外桥接法规清单_IMA_V2.md + IMA.docx URLs → TVAX-006_海外桥接法规清单_IMA_V2.docx."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

HERE = Path(__file__).resolve().parent
OUT = HERE / "TVAX-006_海外桥接法规清单_IMA_V2.docx"


def _strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s)


def _set_cell_font(cell, name: str = "Microsoft YaHei", size_pt: float = 9.0) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = name
            r._element.rPr.rFonts.set(qn("w:eastAsia"), name)
            r.font.size = Pt(size_pt)


def _add_table(doc: Document, headers: list[str], data_rows: list[list[str]]) -> Table:
    tbl = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = h
        _set_cell_font(hdr[j], size_pt=9)
        for p in hdr[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(data_rows, start=1):
        for j, val in enumerate(row):
            if j < len(tbl.rows[i].cells):
                tbl.rows[i].cells[j].text = _strip_md_bold(val)
                _set_cell_font(tbl.rows[i].cells[j])
    return tbl


def main() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    t = doc.add_heading("TVAX-006 海外桥接 — 相关法规与指导原则清单", level=0)
    for r in t.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    intro = doc.add_paragraph()
    intro.add_run(
        "生成时间：2026-05-12（V2更新版）\n"
        "更新说明：基于上传法规附件（EMA 2005/2023疫苗临床评价指南、WHO 2025带状疱疹疫苗立场文件）"
        "及知识库补充，聚焦Top法规调整，新增WHO HZ立场文件、补充EMA版本迭代关键差异；"
        "下载栏整合自 IMA.docx（2026-05-11）中的可核实链接。"
    )
    for r in intro.runs:
        r.font.size = Pt(10)

    # --- URLs from legacy IMA.docx (same rows) + WHO HZ 2025 official ---
    U_WHO_TRS1004 = "https://www.who.int/publications/m/item/WHO-TRS-1004-web-annex-9"
    U_WHO_TRS978 = "https://www.who.int/publications/m/item/TRS-978-61st-report-annex-6"
    U_WHO_PQ_PTC = "https://extranet.who.int/prequal/sites/default/files/document_files/clinical_considerations_oct10.pdf"
    U_WHO_COVID_LIST = "https://www.who.int/publications/m/item/considerations-for-the-assessment-of-covid-19-vaccines-for-listing-by-who"
    U_WHO_PQ_VAR = "https://extranet.who.int/prequal/sites/default/files/document_files/PQ_VXA_Variations_V7.pdf"
    U_WHO_HZ_2025 = "https://www.who.int/publications/i/item/WER10027-28-265-284"

    U_CFR600 = "https://www.ecfr.gov/on/2018-01-26/title-21/chapter-I/subchapter-F/part-600"
    U_CFR601 = "https://www.ecfr.gov/on/2024-01-14/title-21/part-601"
    U_FDA_FLU = "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/clinical-data-needed-support-licensure-seasonal-inactivated-influenza-vaccines"
    U_FDA_PAN = "https://www.fda.gov/downloads/BiologicsBloodVaccines/GuidanceComplianceRegulatoryInformation/Guidances/Vaccines/ucm091985.pdf"
    U_FDA_GLOBAL = "https://www.hhs.gov/guidance/sites/default/files/hhs-guidance-documents/FDA/Guidance-for-Industry--General-Principles-for-the-Development-of-Vaccines-to-Protect-Against-Global-Infectious-Diseases.pdf"
    U_FDA_COVID = "https://www.fda.gov/media/139638/download"
    U_FDA_SUR = "https://www.fda.gov/drugs/development-resources/table-surrogate-endpoints-were-basis-drug-approval-or-licensure"

    U_EMA_VAC = "https://www.ema.europa.eu/en/clinical-evaluation-new-vaccines-scientific-guideline"
    U_TGA_VAC = "https://www.tga.gov.au/sites/default/files/2024-05/Clinical-Evaluation-New-Vaccines-Current.pdf"
    U_EMA_IC = "https://www.ema.europa.eu/en/documents/scientific-guideline/addendum-guideline-clinical-development-vaccines-address-clinical-trials-immunocompromised-individuals_en.pdf"
    U_EMA_VAR = "https://www.ema.europa.eu/en/documents/scientific-guideline/reflection-paper-regulatory-requirements-vaccines-intended-provide-protection-against-variant_en.pdf"
    U_EMA_ADAPT = "https://www.ema.europa.eu/en/news/adapting-covid-19-vaccines-sars-cov-2-variants-guidance-vaccine-manufacturers"

    U_PMDA_E5 = "https://www.pmda.go.jp/files/000156836.pdf"
    U_EMA_E5QA = "https://www.ema.europa.eu/en/documents/scientific-guideline/ich-e-5-r1-questions-answers-ethnic-factors-acceptability-foreign-clinical-data_en.pdf"
    U_ICH_E17 = "https://database.ich.org/sites/default/files/E17EWG_Step4_2017_1116.pdf"
    U_EMA_E17 = "https://www.ema.europa.eu/en/documents/scientific-guideline/ich-guideline-e17-general-principles-planning-and-design-multi-regional-clinical-trials-step-5-first-version_en.pdf"
    U_ICH_E6 = "https://www.ich.org/page/efficacy-guidelines"

    U_CDE_BRIDGE = "https://www.cde.org.cn/main/news/viewInfoCommon/df73ae8ee00c3d983d17c82025339226"
    U_CDE_BRIDGE_FULL = "https://123.207.8.69/technical-guiding-principles-for-clinical-trials-of-vaccine-immunogenicity-bridging-trial-implementation/"
    U_NMPA_OS = "https://www.nmpa.gov.cn/directory/web/nmpa///yaowen/ypjgyw/ypyw/20180710152101322.html"
    U_CDE_OS = "https://www.cde.org.cn/main/news/viewInfoCommon/5f2ac28f0d19e36fd5b54ee015dfee43"
    U_CDE_HOME = "https://www.cde.org.cn"
    U_NMPA_2026 = "https://www.nmpa.gov.cn"

    U_RU_LAW = "https://base.garant.ru/12180594/"
    U_RU_MIN = "https://minzdrav.gov.ru"
    U_EAEU = "https://eec.eaeunion.org"
    U_BR_RDC36 = "https://www.in.gov.br/web/dou/-/resolucao-rdc-n-36-de-27-de-junho-de-2013-500390"
    U_ANVISA = "https://www.gov.br/anvisa"
    U_BR_CONEP = "https://conselho.saude.gov.br/resolucoes/2012/Res466.pdf"
    U_BPOM = "https://www.bpom.go.id"
    U_EDA = "https://www.edaegypt.gov.eg"

    U_ICMRA = "https://www.gov.uk/government/publications/access-consortium-alignment-with-icmra-consensus-on-immunobridging-for-authorising-new-covid-19-vaccines"
    U_MDPI = "https://www.mdpi.com/2076-393X/13/1/19"

    def h1(text: str) -> None:
        doc.add_heading(text, level=1)

    # 一、WHO
    h1("一、WHO（世界卫生组织）")
    _add_table(
        doc,
        ["序号", "法规/指导原则", "发布年份", "核心内容", "对TVAX-006桥接的关联", "下载链接"],
        [
            ["1", "Guidelines on Clinical Evaluation of Vaccines: Regulatory Expectations (WHO TRS No. 1004, Annex 9)", "2017", "疫苗临床评价总框架，含免疫原性替代终点、桥接研究等监管期望，替代2001版", "通用框架，所有疫苗桥接的基础依据", f"WHO官网\n{U_WHO_TRS1004}"],
            ["2", "Procedure for Assessing the Acceptability, in Principle, of Vaccines for Purchase by UN Agencies (WHO TRS No. 978, Annex 6)", "2012", "WHO疫苗预认证(PQ)程序，含临床数据要求", "PQ路径必参考", f"WHO官网\n{U_WHO_TRS978}"],
            ["3", "Points to Consider for Manufacturers of Human Vaccines: Clinical Considerations for Evaluation of Vaccines for Prequalification", "2010", "PQ临床评价要点，桥接数据与免疫原性证据要求", "PQ桥接数据具体要求", f"PDF下载\n{U_WHO_PQ_PTC}"],
            ["4", "Considerations for the Assessment of COVID-19 Vaccines for Listing by WHO", "2021", "COVID-19疫苗WHO列名评估考量，含免疫桥接策略", "免疫桥接策略参考", f"WHO官网\n{U_WHO_COVID_LIST}"],
            ["5", "Guidance on Variations to a Prequalified Vaccine", "持续更新", "预认证疫苗变更指导，含桥接研究要求", "变更场景桥接", f"PDF下载\n{U_WHO_PQ_VAR}"],
            [
                "6",
                "WHO Position Paper on Herpes Zoster Vaccines 🆕",
                "2025",
                "替代2014版水痘/带状疱疹疫苗立场文件；首次WHO推荐常规HZ疫苗接种；明确≥50岁目标人群、≥18岁免疫低下人群接种建议、2剂间隔2-6月程序；提供RZV免疫原性与效力数据基准（≥50岁VE 97.2%，≥70岁VE 89.8%）；免疫桥接考量：无免疫干扰证据、共接种规则、非RZV疫苗互换性数据缺乏；PHN预防效力91.2%",
                "⭐直接适用：为TVAX-006桥接设计提供WHO认可的效力基准与目标人群定义，免疫低下人群桥接需参照此文件",
                f"WHO WER（立场文件）\n{U_WHO_HZ_2025}",
            ],
        ],
    )
    doc.add_paragraph()

    # 二、FDA
    h1("二、FDA（美国食品药品监督管理局）")
    _add_table(
        doc,
        ["序号", "法规/指导原则", "发布年份", "核心内容", "对TVAX-006桥接的关联", "下载链接"],
        [
            ["7", "21 CFR Part 600 — Biological Products: General", "持续更新", "生物制品通用法规，含疫苗许可基础要求", "美国上市路径通用合规框架", f"eCFR在线\n{U_CFR600}"],
            ["8", "21 CFR Part 601 — Licensing", "持续更新", "生物制品许可证申请(BLA)程序要求", "BLA资料组织与审评程序参照", f"eCFR在线\n{U_CFR601}"],
            ["9", "Guidance for Industry: Clinical Data Needed to Support the Licensure of Seasonal Inactivated Influenza Vaccines", "2007", "季节性流感疫苗许可临床数据要求，含免疫原性桥接标准（HI抗体GMC比值、血清转换率）", "免疫原性桥接终点设定与可比性分析思路参考（病种不同需个案映射至gE/VZV）", f"FDA页面\n{U_FDA_FLU}"],
            ["10", "Guidance for Industry: Clinical Data Needed to Support the Licensure of Pandemic Influenza Vaccines", "2007", "大流行流感疫苗许可临床数据要求，含加速审批下免疫原性替代终点", "加速路径下免疫学替代终点论证逻辑参考", f"PDF下载\n{U_FDA_PAN}"],
            ["11", "Guidance for Industry: General Principles for the Development of Vaccines to Protect Against Global Infectious Diseases", "2009", "全球传染病疫苗开发一般原则，含境外数据利用与桥接策略", "境外临床数据外推与桥接策略表述参考", f"PDF下载\n{U_FDA_GLOBAL}"],
            ["12", "Development and Licensure of Vaccines to Prevent COVID-19: Guidance for Industry", "2020", "COVID-19疫苗开发与许可，含免疫桥接和替代终点策略", "免疫桥接试验设计与监管沟通案例参考", f"PDF下载\n{U_FDA_COVID}"],
            ["13", "Table of Surrogate Endpoints", "持续更新", "FDA认可的替代终点清单，含疫苗相关免疫学指标", "若主张免疫学终点支持效力推断，需对照清单与个案科学论证", f"FDA页面\n{U_FDA_SUR}"],
        ],
    )
    doc.add_paragraph()

    # 三、EMA
    h1("三、EMA（欧洲药品管理局）")
    _add_table(
        doc,
        ["序号", "法规/指导原则", "发布年份", "核心内容", "对TVAX-006桥接的关联", "下载链接"],
        [
            [
                "14",
                "Guideline on Clinical Evaluation of Vaccines — Revision 1 (EMEA/CHMP/VWP/164653/05 Rev.1) 📌",
                "2023生效（2023.8.1）",
                "疫苗临床评价核心指南；替代2005版（CHMP/VWP/164653/2005）及SPC附件、佐剂指南",
                "⭐最系统桥接技术指南",
                f"EMA页面\n{U_EMA_VAC}\nTGA镜像PDF\n{U_TGA_VAC}",
            ],
            [
                "15",
                "Addendum to the Guideline on Clinical Development of Vaccines to Address Clinical Trials in Immunocompromised Individuals (EMA/52912/2025)",
                "2025",
                "免疫功能低下人群疫苗桥接研究补充指导",
                "与WHO 2025 HZ立场文件免疫低下人群建议形成双重支撑",
                f"PDF下载\n{U_EMA_IC}",
            ],
            [
                "16",
                "Reflection Paper on the Regulatory Requirements for Vaccines Intended to Provide Protection Against Variant Strain(s) of SARS-CoV-2 (EMA/CHMP/VWP/5880/21)",
                "2021",
                "变异株疫苗监管要求，含免疫桥接设计原则与中和抗体比较标准",
                "免疫桥接非劣效设计参考",
                f"PDF下载\n{U_EMA_VAR}",
            ],
            [
                "17",
                "Adapting COVID-19 Vaccines to SARS-CoV-2 Variants: Guidance for Vaccine Manufacturers",
                "2021",
                "COVID-19变异株疫苗适配指南，含桥接试验可接受设计",
                "桥接试验设计参考",
                f"EMA页面\n{U_EMA_ADAPT}",
            ],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("📎 EMA 2005→2023 Rev.1 关键变更要点（基于上传附件对比）", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "原版：CHMP/VWP/164653/2005（2005年5月发布征求意见，替代1999版CPMP/EWP/463/97）\n"
        "新版：EMEA/CHMP/VWP/164653/05 Rev.1（2018年CHMP采纳+公众咨询，2023年1月CHMP正式采纳，8月生效）"
    )
    doc.add_paragraph()

    _add_table(
        doc,
        ["变更维度", "2005版", "2023 Rev.1版", "对TVAX-006桥接的影响"],
        [
            ["ICP框架", "提及免疫保护相关性概念，但未系统化识别路径", "新增ICP识别路径：效力试验子集分析→人体攻毒→动物模型→血清流行病学→被动免疫数据；明确ICP不可跨机制/年龄/亚型/给药途径推广", "TVAX-006无既定ICP，需依赖免疫桥接路径；gE抗体GMC/SCR作为推定ICP需个案论证"],
            ["免疫桥接", '桥接章节为新增特殊考量之一；原则为"保护水平比例相似"', "大幅扩展：无ICP时可通过与已上市疫苗免疫应答非劣效推断效力；参考疫苗退市时可与同样通过桥接获批的疫苗比较；多亚型场景可通过交叉保护评估支持全亚型效力推断", "⭐TVAX-006与Shingrix®免疫桥接的核心法规依据"],
            ["特殊人群桥接", "提及宿主因素影响，部分可上市后探索", "新增专项章节：老年人按年龄亚组（65-74/75-84/85+）分层；免疫缺陷人群仅入组特定高影响亚组；孕妇需先完成育龄非妊娠女性研究", "TVAX-006需关注≥50岁年龄亚组分层设计，免疫低下人群桥接需界定亚组"],
            ["免疫程序桥接", "未覆盖", "新增：支持多疫苗完成同一程序、异源加强（需非劣效）、异源初免-加强（需优效+广覆盖）", "如TVAX-006拟支持与Shingrix®互换或加强，需参照此条款"],
            ["安全数据库规模", "未明确分级标准", "新增分级：含未上市新组分→需估计不常见AE(1/100-1/1000)；全组分已上市→可参考已上市安全性数据；需覆盖所有目标年龄亚组", "TVAX-006含新佐剂TVA01，安全数据库规模需满足新组分标准"],
            ["佐剂", "单独佐剂指南(EMEA/CHMP/VEG/134716/04)", "整合入本指南，佐剂不再单独出指南", "TVAX-006佐剂TVA01临床评价需参照整合后要求"],
            ["SPC附件", "单独SPC附件(EMEA/CHMP/VWP/382702/06)", "整合入本指南", "—"],
        ],
    )
    doc.add_paragraph()

    h1("四、ICH（国际人用药品注册技术协调会）")
    _add_table(
        doc,
        ["序号", "法规/指导原则", "发布年份", "核心内容", "下载链接"],
        [
            ["18", "ICH E5(R1): Ethnic Factors in the Acceptability of Foreign Clinical Data", "1998/2003", "桥接研究的理论基础，种族敏感性三步评估法", f"PDF下载(PMDA)\n{U_PMDA_E5}"],
            ["19", "ICH E5(R1) Q&A: Questions and Answers on Ethnic Factors", "2003/2006", "E5实施问答集，细化桥接实操要求", f"PDF下载(EMA)\n{U_EMA_E5QA}"],
            [
                "20",
                "ICH E17: General Principles for Planning and Design of Multi-Regional Clinical Trials",
                "2017",
                "MRCT设计框架，区域样本量分配",
                f"PDF下载(ICH)\n{U_ICH_E17}\nPDF下载(EMA)\n{U_EMA_E17}",
            ],
            ["21", "ICH E6(R2): Guideline for Good Clinical Practice", "2016", "临床试验质量标准（GCP）", f"ICH页面\n{U_ICH_E6}"],
        ],
    )
    doc.add_paragraph()

    h1("五、中国 NMPA/CDE（原产国/数据来源国）")
    _add_table(
        doc,
        ["序号", "法规/指导原则", "发布年份", "核心内容", "下载链接"],
        [
            [
                "22",
                "《疫苗免疫原性桥接临床试验技术指导原则(试行)》——2024年第41号",
                "2024",
                "最直接相关：疫苗免疫桥接适用范围、设计要点、评价标准",
                f"CDE页面\n{U_CDE_BRIDGE}\n全文阅览\n{U_CDE_BRIDGE_FULL}",
            ],
            ["23", "《接受药品境外临床试验数据的技术指导原则》", "2018", "境外数据接受条件、种族敏感性评估要求", f"NMPA页面\n{U_NMPA_OS}"],
            ["24", "《境外已上市境内未上市药品临床技术要求》", "2020", "境外已上市药品桥接临床要求", f"CDE页面\n{U_CDE_OS}"],
            ["25", "《预防用疫苗临床可比性研究技术指导原则》", "—", "疫苗临床可比研究设计", f"CDE页面\n{U_CDE_HOME}"],
            ["26", "《国家药监局关于进一步优化临床急需境外已上市药品审评审批有关事项的公告》(2026年第3号)", "2026", "最新优化政策，加速境外已上市药品审批", f"NMPA页面\n{U_NMPA_2026}"],
        ],
    )
    doc.add_paragraph()

    h1("六、目标国法规")
    _add_table(
        doc,
        ["序号", "国家", "法规/指导原则", "核心内容", "下载链接"],
        [
            ["27", "俄罗斯", '联邦法 №61-ФЗ "药品流通法"', "俄罗斯药品注册基本法", f"在线查阅\n{U_RU_LAW}"],
            ["28", "俄罗斯", "Minzdrav 第200н号令", "临床试验审批程序", f"在线查阅\n{U_RU_MIN}"],
            ["29", "俄罗斯", "欧亚经济联盟(EAEU) TR CU 005/2011", "EAEU路径注册统一标准", f"EAEU官网\n{U_EAEU}"],
            ["30", "巴西", "RDC 36/2013", "ANVISA临床试验审批规范", f"PDF下载\n{U_BR_RDC36}"],
            ["31", "巴西", "RDC 39/2013", "药品注册技术要求", f"ANVISA官网\n{U_ANVISA}"],
            ["32", "巴西", "IN 74/2024", "等效外国监管机构审评利用（reliance机制）", f"ANVISA官网\n{U_ANVISA}"],
            ["33", "巴西", "RDC 505/2021", "生物制品和疫苗注册要求", f"ANVISA官网\n{U_ANVISA}"],
            ["34", "巴西", "CONEP Resolution 466/2012", "伦理审查要求", f"PDF下载\n{U_BR_CONEP}"],
            ["35", "印尼", "BPOM Regulation No. 6/2020", "药品注册技术要求", f"BPOM官网\n{U_BPOM}"],
            ["36", "埃及", "EDA Clinical Trial Registry Guidelines (2023)", "临床试验注册和审批要求", f"EDA官网\n{U_EDA}"],
            ["37", "埃及", "Ministerial Decree 413/2021", "药品注册法规", f"EDA官网\n{U_EDA}"],
        ],
    )
    doc.add_paragraph()

    h1("七、国际监管联盟共识")
    _add_table(
        doc,
        ["序号", "法规/指导原则", "发布年份", "核心内容", "下载链接"],
        [
            [
                "38",
                "Access Consortium: Alignment with ICMRA Consensus on Immunobridging for Authorising New COVID-19 Vaccines",
                "2021",
                "多国监管机构免疫桥接共识：非劣效设计、安全数据库≥3000人、12个月随访承诺",
                f"GOV.UK页面\n{U_ICMRA}",
            ],
            [
                "39",
                "Immunobridging Trials: An Important Tool to Protect Vulnerable and Immunocompromised Patients Against Evolving Pathogens (Vaccines, MDPI)",
                "2025",
                "综述：免疫桥接设计原则、WHO/FDA/EMA框架对比、成功案例",
                f"PDF下载\n{U_MDPI}",
            ],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("TOP法规调整摘要（V2更新重点）", level=1)
    doc.add_heading("新增高影响法规", level=2)
    _add_table(
        doc,
        ["调整项", "内容", "对TVAX-006的影响"],
        [
            [
                "WHO HZ立场文件(2025)",
                "替代2014版；首次推荐常规HZ疫苗接种；≥50岁目标人群+≥18岁免疫低下人群；RZV效力基准（≥50岁97.2%，≥70岁89.8%）；免疫桥接无干扰证据",
                "⭐直接适用：为TVAX-006桥接效力目标提供WHO认可的基准值；免疫低下人群桥接需参照此文件；目标人群年龄分层有据可依",
            ],
            [
                "EMA 2005→2023 Rev.1变更",
                "ICP框架系统化、免疫桥接路径扩展、特殊人群桥接细化、安全数据库分级、佐剂指南整合",
                "原清单仅列2023版，缺少版本差异；TVAX-006新佐剂TVA01需满足新组分安全数据库规模",
            ],
        ],
    )
    doc.add_paragraph()
    doc.add_heading("已有条目强化", level=2)
    _add_table(
        doc,
        ["条目", "强化内容"],
        [
            ["EMA免疫低下补充(#15)", "WHO 2025 HZ立场文件提供免疫低下人群（HIV、HSCT、自身免疫病）疫苗效果数据，与EMA 2025 addendum形成双重支撑"],
            ["ICMRA共识(#38)", "WHO 2025 HZ立场文件中免疫低下人群建议与ICMRA免疫桥接共识一致，互为佐证"],
        ],
    )
    doc.add_paragraph()

    doc.add_heading("TOP 6 — 对TVAX-006桥接设计最具参考价值（V2更新）", level=1)
    _add_table(
        doc,
        ["优先级", "法规/指导原则", "理由", "下载链接"],
        [
            [
                "1",
                "CDE《疫苗免疫原性桥接临床试验技术指导原则(试行)》(2024年第41号)",
                "直接适用，中国原产国数据桥接出境的起始框架",
                f"{U_CDE_BRIDGE}\n{U_CDE_BRIDGE_FULL}",
            ],
            ["2", "WHO TRS 1004 Annex 9: Guidelines on Clinical Evaluation of Vaccines", "全球疫苗桥接最高权威，各国PQ和注册的通用引用", U_WHO_TRS1004],
            [
                "3",
                "EMA CHMP/VWP/164653/05 Rev.1: Guideline on Clinical Evaluation of Vaccines",
                "疫苗免疫桥接设计最系统的技术指南；2023版新增ICP路径+特殊人群桥接+程序桥接",
                f"{U_EMA_VAC}\n{U_TGA_VAC}",
            ],
            [
                "4",
                "WHO Position Paper on Herpes Zoster Vaccines (2025)",
                "直接适用TVAX-006：首次WHO推荐HZ疫苗常规接种；提供RZV效力基准（VE≥89.8%）；明确免疫低下人群桥接要求；非RZV疫苗互换性数据缺乏提示桥接设计需独立论证",
                U_WHO_HZ_2025,
            ],
            [
                "5",
                "ICH E5(R1): Ethnic Factors in the Acceptability of Foreign Clinical Data",
                "种族敏感性评估+桥接研究判定的基础框架",
                f"{U_PMDA_E5}\n{U_EMA_E5QA}",
            ],
            ["6", "ICMRA Consensus on Immunobridging for Authorising New Vaccines", "多国监管机构对免疫桥接的统一立场，对新兴市场国家有说服力", U_ICMRA],
        ],
    )
    foot = doc.add_paragraph()
    foot.add_run(
        "V1→V2 Top清单变化：原Top 5扩展为Top 6，新增WHO 2025 HZ立场文件（优先级4）。"
        "EMA提供桥接设计方法论，WHO HZ立场文件提供带状疱疹疫苗专属效力与安全标准。"
    )
    doc.add_paragraph()
    endp = doc.add_paragraph()
    endp.add_run("本清单基于知识库006及上传法规附件整理，更新时间2026年5月12日。建议结合项目进展持续跟踪各国法规动态。")
    endp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()
