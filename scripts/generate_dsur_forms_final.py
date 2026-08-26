import os

import openpyxl
from docx import Document

base_dir = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\018-3 DSUR"
r4_path = os.path.join(
    base_dir,
    "区域附录R4：报告周期内发生的药物临床试验方案变更或者临床方面的新发现、非临床或者药学的变化或者新发现总结表.docx",
)
r5_path = os.path.join(base_dir, "区域附录R5：下一报告周期内总体研究计划概要.docx")
xlsx_path = os.path.join(
    base_dir, "远大重组破伤风疫苗（大肠埃希菌）-DSUR#2撰写所需资料清单.xlsx"
)


def process_r4(path):
    doc = Document(path)
    for i in range(3):
        if len(doc.tables) > i:
            table = doc.tables[i]
            if len(table.rows) > 1:
                table.rows[1].cells[0].text = "无"
                table.rows[1].cells[1].text = "无"
                table.rows[1].cells[2].text = "无"
                table.rows[1].cells[3].text = "无"
    for i in range(3, 6):
        if len(doc.tables) > i:
            table = doc.tables[i]
            if len(table.rows) > 1:
                table.rows[1].cells[0].text = "无"
                table.rows[1].cells[1].text = "无重大变更"
                table.rows[1].cells[2].text = "无"
                table.rows[1].cells[3].text = "无"
                table.rows[1].cells[4].text = "无"
    doc.save(path)


def process_r5(path):
    doc = Document(path)
    r5_content = {
        "（一）立题依据；": "破伤风是致死性的感染性疾病。本品为重组破伤风疫苗，现国内外尚无相同技术路线的产品批准。鉴于既往重组疫苗可引发强大且持久免疫反应的优势并具有较高的纯度，本品相对于传统类毒素疫苗可能具有更强的免疫原性和更好的安全性。",
        "（二）拟研究的适应症；": "预防破伤风梭状芽孢杆菌（破伤风杆菌）感染。",
        "（三）评价药物时所遵循的总体路径；": "通过Ⅰ期临床试验评价不同剂量组的安全性并初步探索免疫原性，在安全性和耐受性良好的前提下开展Ⅱ期试验，在较大样本量下进一步评价免疫原性和安全性，筛选出最佳剂量，随后推进后续确证性临床研究。",
        "（四）下一个报告周期内拟开展的临床试验；": "下一报告周期内计划主要推进Ⅱ期临床试验。本研究采用随机、盲法、阳性对照的试验设计，评价试验疫苗在18岁及以上健康人群中的免疫原性和安全性。",
        "（五）预计受试者人数；": "Ⅱ期临床试验计划纳入480例18岁及以上健康受试者。",
        "（六）预计的风险：": "基于现有数据，预计可能发生与疫苗接种相关的常见不良反应（如注射部位疼痛、红晕以及发热等），具体以临床试验实际观察结果为准。",
        "二、下一报告周期内非临床研究总体计划概要": "初步定稿后，由相应模块的老师确认。",
        "三、下一报告周期内药学研究总体计划概要": "初步定稿后，由相应模块的老师确认。",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, insert_text in r5_content.items():
            if text.startswith(key) or key in text:
                para.add_run("\n" + insert_text)
                r5_content[key] = ""
    doc.save(path)


def process_xlsx(path):
    wb = openpyxl.load_workbook(path)
    sheet = wb.active
    sheet.cell(row=3, column=5).value = "已提供：临床批件-临床意见-20240826.png"
    sheet.cell(row=4, column=5).value = "已提供：V1.2/2026年01月04日版本"
    sheet.cell(row=5, column=5).value = (
        "已提供Ⅰ期（V2.2，2025.06.16）、Ⅱ期（V0.4，2024.08.07）、Ⅲ期（V0.3，2024.04.08）方案"
    )
    sheet.cell(row=11, column=5).value = "已提供：风险管理计划（V1.0，2024年08月08日）"
    sheet.cell(row=12, column=5).value = "本报告期间内没有相应的CSR或小结"
    sheet.cell(row=13, column=5).value = "初步定稿后，由相应模块的老师确认"
    sheet.cell(row=14, column=5).value = "初步定稿后，由相应模块的老师确认"
    sheet.cell(row=15, column=5).value = "已在R5附录中提供初稿"
    wb.save(path)


if __name__ == "__main__":
    process_r4(r4_path)
    process_r5(r5_path)
    process_xlsx(xlsx_path)
