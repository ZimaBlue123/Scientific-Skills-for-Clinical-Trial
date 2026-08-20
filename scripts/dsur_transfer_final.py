"""
DSUR Content Transfer v7 - FINAL.

Fix: Remove TOC paragraphs at XML level (handle hyperlinks + PAGEREF fields).

Robustness layer (Stage-1 hardening):
- Replace bare ``except:`` blocks with explicit exception types and log via ``logging``.
- Add type hints across public helpers and ``main()``.
- Validate CLI arguments, input file existence and index bounds before mutating documents.
- Convert top-level flow into ``main() -> int`` for clean exit semantics.
"""

from __future__ import annotations

import logging
import os
import sys
from typing import Any

from docx import Document  # type: ignore
from docx.document import Document as _DocxDocument  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.table import Table, _Cell  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

LOGGER = logging.getLogger("dsur_transfer_v7")
if not LOGGER.handlers:
    _handler = logging.StreamHandler(stream=sys.stderr)
    _handler.setFormatter(logging.Formatter("[%(levelname)s] %(name)s: %(message)s"))
    LOGGER.addHandler(_handler)
    LOGGER.setLevel(logging.INFO)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def identify_section_key(text: str) -> str | None:
    """Map a paragraph text to one of the canonical DSUR section keys."""
    t = text.strip().lower()
    mappings: list[tuple[str, str]] = [
        ("development safety update report (dsur)", "title_dsur"),
        ("executive summary", "exec_summary"),
        ("table of contents", "toc_heading"),
        ("confidentiality statement", "confidentiality"),
        ("1. introduction", "sec1"),
        ("2. worldwide marketing approval status", "sec2"),
        ("3. actions taken", "sec3"),
        ("4. changes to reference safety information", "sec4"),
        ("5. inventory of clinical trials", "sec5"),
        ("6.1 cumulative subject exposure in development program", "sec6_1"),
        ("6.2 patient exposure from marketing experience", "sec6_2"),
        ("6. estimated cumulative exposure", "sec6"),
        ("7.1 reference information", "sec7_1"),
        ("7.2 line listings of serious adverse reactions", "sec7_2"),
        ("7.3 cumulative summary tabulations of serious adverse events", "sec7_3"),
        ("7. data in line listings and summary tabulations", "sec7"),
        ("8.1 completed clinical trials", "sec8_1"),
        ("8.2 ongoing clinical trials", "sec8_2"),
        ("8.3 long-term follow-up", "sec8_3"),
        ("8.4 other therapeutic use", "sec8_4"),
        ("8.5 new safety data related to combination therapies", "sec8_5"),
        ("8. significant findings from clinical trials", "sec8"),
        ("9. safety findings from non-interventional studies", "sec9"),
        ("10. other clinical trial/study safety information", "sec10"),
        ("11. safety findings from marketing experience", "sec11"),
        ("12. non-clinical data", "sec12"),
        ("13. literature", "sec13"),
        ("14. other dsurs", "sec14"),
        ("15. lack of efficacy", "sec15"),
        ("16. region-specific information", "sec16"),
        ("17. late-breaking information", "sec17"),
        ("18.1.1 known adverse reactions", "sec18_1_1"),
        ("18.1.2 potential risks", "sec18_1_2"),
        ("18.1.3 adverse events resulting in death", "sec18_1_3"),
        ("18.1.4 potential impact of concomitant use", "sec18_1_4"),
        ("18.1.5", "sec18_1_5"),
        ("18.1 risk assessment", "sec18_1"),
        ("18.2.1 benefit assessment", "sec18_2_1"),
        (
            "18.2.2 pharmacodynamics suggested by completed clinical studies",
            "sec18_2_2",
        ),
        ("18.2.3 impact of identified adverse reactions", "sec18_2_3"),
        ("18.2.4 do other potential risks have clinical significance", "sec18_2_4"),
        ("18.2.5 are there any events requiring close attention", "sec18_2_5"),
        ("18.2 benefit-risk considerations", "sec18_2"),
        ("18. overall safety assessment", "sec18"),
        ("19.1 important risks in the previous cycle", "sec19_1"),
        ("19.2 important risks in the current cycle", "sec19_2"),
        ("19. summary of important risks", "sec19"),
        ("20. conclusions", "sec20"),
        ("appendices", "appendices"),
        ("appendix 1 -", "appendix1"),
        ("appendix 2 -", "appendix2"),
        ("appendix 3 -", "appendix3"),
        ("appendix 4 -", "appendix4"),
        ("appendix 5 -", "appendix5"),
        ("appendix 6 -", "appendix6"),
        ("appendix 7 -", "appendix7"),
        ("regional appendices", "regional"),
        ("appendix r1 -", "appendix_r1"),
        ("appendix r2 -", "appendix_r2"),
        ("appendix r3 -", "appendix_r3"),
        ("appendix r4 -", "appendix_r4"),
        ("appendix r5 -", "appendix_r5"),
    ]
    for pattern, key in mappings:
        if t.startswith(pattern):
            return key
    return None


def is_toc_style(style_name: str | None) -> bool:
    return "toc" in (style_name or "").lower()


def build_para_index(doc: _DocxDocument) -> dict[str, list[int]]:
    index: dict[str, list[int]] = {}
    current_key = "__preamble__"
    for i, para in enumerate(doc.paragraphs):
        if is_toc_style(para.style.name):
            continue
        text = para.text.strip()
        if not text:
            continue
        key = identify_section_key(text)
        if key:
            current_key = key
        index.setdefault(current_key, []).append(i)
    return index


def replace_para_text(para: Paragraph, new_text: str | None) -> None:
    """Replace paragraph text at XML level. Handles hyperlinks and fields."""
    p_elem = para._element

    # Remove ALL child elements except pPr (paragraph properties)
    to_remove: list[Any] = []
    for child in p_elem:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "pPr":
            to_remove.append(child)
    for child in to_remove:
        p_elem.remove(child)

    if new_text is None:
        return

    # Create a simple run with text
    new_run = OxmlElement("w:r")
    if new_text:
        rPr = OxmlElement("w:rPr")
        try:
            style = para.style
            if style and style.font and style.font.name:
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), style.font.name)
                rFonts.set(qn("w:hAnsi"), style.font.name)
                rPr.append(rFonts)
            if style and style.font and style.font.size:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(style.font.size.pt * 2)))
                rPr.append(sz)
        except (AttributeError, ValueError, TypeError) as exc:
            LOGGER.debug("Skipping font/style hints for paragraph: %s", exc)
        new_run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text if new_text else ""
    new_run.append(t)
    p_elem.append(new_run)


def set_cell_text(cell: _Cell, text: str) -> None:
    for para in cell.paragraphs:
        replace_para_text(para, text)
        break


def clear_table_data(table: Table, keep_header: bool = True) -> None:
    start_row = 1 if keep_header else 0
    for row in table.rows[start_row:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_para_text(para, "")


def get_content_indices(para_indices: list[int], doc: _DocxDocument) -> list[int]:
    result: list[int] = []
    for idx in para_indices:
        text = doc.paragraphs[idx].text.strip()
        if not text:
            continue
        if identify_section_key(text):
            continue
        result.append(idx)
    return result


def get_src_content_texts(para_indices: list[int], doc: _DocxDocument) -> list[str]:
    texts: list[str] = []
    for idx in para_indices:
        text = doc.paragraphs[idx].text.strip()
        if not text:
            continue
        if identify_section_key(text):
            continue
        texts.append(text)
    return texts


def replace_section_content(
    template_doc: _DocxDocument,
    tpl_indices: list[int],
    src_texts: list[str],
    op_name: str = "",
) -> tuple[int, int]:
    tpl_content = get_content_indices(tpl_indices, template_doc)
    replaced = 0
    cleared = 0
    for j in range(min(len(src_texts), len(tpl_content))):
        replace_para_text(template_doc.paragraphs[tpl_content[j]], src_texts[j])
        replaced += 1
    for j in range(len(src_texts), len(tpl_content)):
        para = template_doc.paragraphs[tpl_content[j]]
        if para.text.strip():
            replace_para_text(para, "")
            cleared += 1
    if op_name:
        print(f"  {op_name}: replaced {replaced}, cleared {cleared}")
    return replaced, cleared


# ---------------------------------------------------------------------------
# Phase handlers
# ---------------------------------------------------------------------------
def _phase_title(template: _DocxDocument) -> None:
    print("--- Title Page ---")
    title_map = {
        "Development Safety Update Report (DSUR) No. 2": "Development Safety Update Report (DSUR) No. 1",
        "Recombinant Hexavalent Norovirus Vaccine (Hansenula polymorpha)": "Recombinant Varicella Vaccine (CHO Cell)",
        "Reporting Period: 25 April 2025 to 24 April 2026": "Reporting Period: 26-Jun-2025 to 25-Jun-2026",
        "Date of Report: 05 June 2026": "Date of Report: 09-Jul-2026",
    }
    conf_new = (
        "All information contained in this document is the exclusive property of "
        "Grand Theravac Life Sciences (Nanjing) Co., Ltd. and Grand Theravac Life Sciences "
        "(Hangzhou) Co., Ltd., and is strictly confidential. It may not be disclosed or "
        "reproduced, in whole or in part, without prior written consent from the Sponsor."
    )

    for para in template.paragraphs:
        if is_toc_style(para.style.name):
            continue
        text = para.text.strip()
        matched = False
        for old, new in title_map.items():
            if old in text and "\t" not in text:
                replace_para_text(para, new)
                matched = True
                break
        if not matched and text.startswith(
            "All information contained in this document is the property of Grand Theravac"
        ):
            replace_para_text(para, conf_new)


def _phase_sponsor_table(template: _DocxDocument) -> None:
    if not template.tables:
        LOGGER.warning("Template has no tables; skipping sponsor block.")
        return
    try:
        t = template.tables[0]
        sponsor_data = [
            "Grand Theravac Life Sciences (Nanjing) Co., Ltd. / Grand Theravac Life Sciences (Hangzhou) Co., Ltd.",
            "",
            "",
            "",
            "",
        ]
        for row_idx, val in enumerate(sponsor_data):
            if row_idx >= len(t.rows):
                break
            try:
                set_cell_text(t.rows[row_idx].cells[1], val)
            except (IndexError, AttributeError) as exc:
                LOGGER.warning("Skipping sponsor row %d: %s", row_idx, exc)
    except (IndexError, AttributeError) as exc:
        LOGGER.warning("Sponsor table update failed: %s", exc)


def _phase_exec_summary(
    template: _DocxDocument,
    source: _DocxDocument,
    tpl_index: dict[str, list[int]],
    src_index: dict[str, list[int]],
) -> tuple[int, int]:
    print("--- Executive Summary ---")
    if "exec_summary" in tpl_index and "exec_summary" in src_index:
        src_texts = get_src_content_texts(src_index["exec_summary"], source)
        return replace_section_content(
            template, tpl_index["exec_summary"], src_texts, "Exec Summary"
        )
    return 0, 0


def _phase_toc(template: _DocxDocument) -> None:
    print("--- TOC Clearing ---")
    toc_cleared = 0
    for para in template.paragraphs:
        if is_toc_style(para.style.name) and para.text.strip():
            replace_para_text(para, "")
            toc_cleared += 1
    print(f"  Cleared {toc_cleared} TOC paragraphs (XML level)")


def _phase_main_body(
    template: _DocxDocument,
    source: _DocxDocument,
    tpl_index: dict[str, list[int]],
    src_index: dict[str, list[int]],
) -> tuple[int, int]:
    print("--- Main Body ---")
    skip_sections = {
        "__preamble__",
        "toc_heading",
        "exec_summary",
        "title_dsur",
        "confidentiality",
        "appendices",
        "regional",
        "sec6",
        "sec7",
        "sec8",
        "sec18",
        "sec18_1",
        "sec19",
    }
    total_replaced = 0
    total_cleared = 0

    if "sec18_2" in src_index and "sec18_2_1" in tpl_index:
        src_texts_all = get_src_content_texts(src_index["sec18_2"], source)
        if src_texts_all:
            r, c = replace_section_content(
                template,
                tpl_index["sec18_2_1"],
                src_texts_all[: min(3, len(src_texts_all))],
                "sec18_2_1 (combined)",
            )
            total_replaced += r
            total_cleared += c
        for sub in ["sec18_2_2", "sec18_2_3", "sec18_2_4", "sec18_2_5"]:
            if sub in tpl_index:
                r, c = replace_section_content(template, tpl_index[sub], [], sub)
                total_replaced += r
                total_cleared += c

    for sec_key in sorted(tpl_index.keys()):
        if sec_key in skip_sections or sec_key.startswith("sec18_2"):
            continue
        if sec_key not in src_index:
            r, c = replace_section_content(template, tpl_index[sec_key], [], sec_key)
            total_replaced += r
            total_cleared += c
            continue
        src_texts = get_src_content_texts(src_index[sec_key], source)
        r, c = replace_section_content(template, tpl_index[sec_key], src_texts, sec_key)
        total_replaced += r
        total_cleared += c

    print(f"  Total: replaced {total_replaced}, cleared {total_cleared}")
    return total_replaced, total_cleared


def _phase_tables(
    template: _DocxDocument,
    source: _DocxDocument,
) -> None:
    print("--- Tables ---")
    for t_idx in (1, 2, 3, 7, 9):
        if t_idx < len(template.tables):
            clear_table_data(template.tables[t_idx], keep_header=True)

    for t_idx in range(14, min(20, len(template.tables))):
        clear_table_data(template.tables[t_idx], keep_header=True)

    if len(template.tables) > 5 and source.tables:
        tpl_table = template.tables[5]
        src_table = source.tables[0]

        for r_idx in range(1, len(tpl_table.rows)):
            for cell in tpl_table.rows[r_idx].cells:
                for para in cell.paragraphs:
                    replace_para_text(para, "")

        if len(src_table.rows) > 1:
            src_row = src_table.rows[1]
            tpl_row = tpl_table.rows[1]
            cell_count = min(len(src_row.cells), len(tpl_row.cells))
            for c_idx in range(cell_count):
                src_text = src_row.cells[c_idx].text.strip()
                set_cell_text(tpl_row.cells[c_idx], src_text)
            set_cell_text(tpl_row.cells[len(tpl_row.cells) - 1], "0")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------
def main(
    template_path: str,
    source_path: str,
    output_path: str,
) -> int:
    if not template_path or not os.path.exists(template_path):
        LOGGER.error("Template not found: %s", template_path)
        return 2
    if not source_path or not os.path.exists(source_path):
        LOGGER.error("Source not found: %s", source_path)
        return 2
    if not output_path:
        LOGGER.error("Output path must be provided")
        return 2

    try:
        template = Document(template_path)
    except (OSError, ValueError) as exc:
        LOGGER.error("Cannot open template %s: %s", template_path, exc)
        return 1
    try:
        source = Document(source_path)
    except (OSError, ValueError) as exc:
        LOGGER.error("Cannot open source %s: %s", source_path, exc)
        return 1

    try:
        tpl_index = build_para_index(template)
        src_index = build_para_index(source)

        _phase_title(template)
        _phase_sponsor_table(template)
        _phase_exec_summary(template, source, tpl_index, src_index)
        _phase_toc(template)
        _phase_main_body(template, source, tpl_index, src_index)
        _phase_tables(template, source)

        # Ensure output directory exists
        out_dir = os.path.dirname(os.path.abspath(output_path))
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)
        template.save(output_path)
    except (OSError, ValueError, IndexError, KeyError) as exc:
        LOGGER.exception("Transfer failed: %s", exc)
        return 1

    print(f"\nSaved: {output_path}")
    return 0


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] in ("-h", "--help"):
        print(
            "Usage: python dsur_transfer_v7.py <template.docx> <source.docx> <output.docx>",
            file=sys.stderr,
        )
        sys.exit(0)
    tp = sys.argv[1] if len(sys.argv) > 1 else "template.docx"
    sp = sys.argv[2] if len(sys.argv) > 2 else "source.docx"
    op = sys.argv[3] if len(sys.argv) > 3 else "output.docx"
    sys.exit(main(tp, sp, op))
