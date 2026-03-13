import os, sys
from pathlib import Path
sys.stdout.reconfigure(encoding='utf-8')
INPUT = Path('review_materials')
OUT = INPUT / 'converted'
OUT.mkdir(exist_ok=True)
created = []
for f in sorted(INPUT.iterdir()):
    if not f.is_file() or f.suffix.lower() not in ('.docx','.pdf'):
        continue
    out_name = f.stem
    for c in '\\/:*?\"<>|':
        out_name = out_name.replace(c,'_')
    out_name = out_name + '.md'
    out_path = OUT / out_name
    content = None
    if f.suffix.lower() == '.docx':
        from docx import Document
        doc = Document(str(f))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts.append(' | '.join(c.text for c in row.cells))
        content = '\n'.join(parts)
    else:
        import pypdf
        r = pypdf.PdfReader(str(f))
        content = '\n'.join((p.extract_text() or '') for p in r.pages)
    if content and content.strip():
        out_path.write_text(content, encoding='utf-8')
        created.append(str(out_path))
        print('OK:', f.name)
    else:
        print('SKIP:', f.name)
print('---')
for p in created:
    print(p)
