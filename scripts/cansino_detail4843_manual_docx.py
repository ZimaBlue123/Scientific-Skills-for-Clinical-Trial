# -*- coding: utf-8 -*-
"""
Download 康希诺官网 detail-4843 公示中的「说明书」两页（JPG），生成横版 Word。

要点：
- **不嵌入整页图片**：可编辑正文 + Word 表格由通用模块 **`scripts/extract_tables_to_docx.py`** 完成（img2table + 词级 OCR，含多尺度放大择优）。
- 本脚本只负责从康希诺官网 **detail-4843** 下载两页说明书 JPG 并调用该模块。
- 可选 `--fulltext-appendix`：在文末追加整页 Tesseract 纯文本（默认关闭）。

依赖：python-docx、Pillow、pytesseract、img2table、pandas、beautifulsoup4、
      opencv-contrib-python-headless（img2table 需要 cv2.ximgproc）。

用法：
  py -3.10 scripts/cansino_detail4843_manual_docx.py
  （默认输出：`仓库根/output/吸附无细胞百白破联合疫苗_说明书_官网4843.docx`）
  py -3.10 scripts/cansino_detail4843_manual_docx.py -o "D:\\\\out\\\\说明书.docx"
  py -3.10 scripts/cansino_detail4843_manual_docx.py --fulltext-appendix
  py -3.10 scripts/cansino_detail4843_manual_docx.py --table-psm 11 --prose-psm 6 --no-ocr
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import urllib.error
import urllib.request
from pathlib import Path
from typing import List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

try:
    from PIL import Image
except ImportError as e:  # pragma: no cover
    raise SystemExit("需要安装 Pillow：pip install pillow") from e

HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parent
DEFAULT_OUTPUT_DIR = (REPO_ROOT / "output").resolve()

UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)

CHI_SIM_URL = "https://github.com/tesseract-ocr/tessdata/raw/main/chi_sim.traineddata"
ENG_URL = "https://github.com/tesseract-ocr/tessdata/raw/main/eng.traineddata"

MANUAL_PAGES: Sequence[Tuple[str, str]] = (
    (
        "说明书 第1页",
        "https://www.cansinotech.com.cn/uploadfile/2026/04/13/10-08-06/%E8%AF%B4%E6%98%8E%E4%B9%A61.jpg",
    ),
    (
        "说明书 第2页",
        "https://www.cansinotech.com.cn/uploadfile/2026/04/13/10-08-08/%E8%AF%B4%E6%98%8E%E4%B9%A62.jpg",
    ),
)

SOURCE_PAGE = "https://www.cansinotech.com.cn/detail-4843"

DEFAULT_TESSERACT = Path(r"D:\tesseract-ocr\tesseract.exe")


def _download(url: str) -> bytes:
    req = urllib.request.Request(url, headers={"User-Agent": UA})
    with urllib.request.urlopen(req, timeout=300) as r:
        data = r.read()
    if len(data) < 2048:
        raise RuntimeError(f"下载内容异常偏小（{len(data)} bytes），请检查网络或 URL。")
    if not data.startswith(b"\xff\xd8") and not data.startswith(b"\x89PNG"):
        raise RuntimeError("下载内容不是 JPEG/PNG 头，可能被拦截或返回错误页。")
    return data


def _download_file(url: str, dest: Path, min_bytes: int = 10_000) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    req = urllib.request.Request(url, headers={"User-Agent": UA})
    tmp = dest.with_suffix(dest.suffix + ".download")
    with urllib.request.urlopen(req, timeout=600) as r:
        tmp.write_bytes(r.read())
    if tmp.stat().st_size < min_bytes:
        tmp.unlink(missing_ok=True)
        raise RuntimeError(f"下载异常：{url} → 体积过小")
    tmp.replace(dest)


def _resolve_tesseract_exe(cli_path: Path | None) -> Path:
    if cli_path is not None:
        return cli_path
    env = os.environ.get("TESSERACT_CMD", "").strip().strip('"')
    if env:
        return Path(env)
    return DEFAULT_TESSERACT


def _tessdata_directory(tesseract_exe: Path) -> Path:
    """返回存放 *.traineddata 的 tessdata 目录。"""
    prefix = os.environ.get("TESSDATA_PREFIX", "").strip().strip('"')
    if prefix:
        root = Path(prefix)
        cand_a = root / "tessdata"
        if cand_a.is_dir():
            return cand_a
        if root.name.lower() == "tessdata" and root.is_dir():
            return root
    return tesseract_exe.parent / "tessdata"


def ensure_traineddata_files(
    tesseract_exe: Path,
    *,
    allow_download: bool,
    names: Tuple[str, ...] = ("chi_sim", "eng"),
) -> Path:
    """确保 tessdata 目录存在且包含所需 *.traineddata。"""
    td = _tessdata_directory(tesseract_exe)
    td.mkdir(parents=True, exist_ok=True)

    need: List[Tuple[str, str]] = []
    for name in names:
        f = td / f"{name}.traineddata"
        if f.is_file() and f.stat().st_size > 50_000:
            continue
        need.append((name, f))

    if not need:
        os.environ["TESSDATA_PREFIX"] = str(td)
        return td

    if not allow_download:
        missing = ", ".join(n for n, _ in need)
        raise FileNotFoundError(
            f"tessdata 缺少语言文件：{missing}（目录 {td}）。"
            "请手动放入对应 .traineddata，或去掉 --no-download-lang。"
        )

    url_map = {"chi_sim": CHI_SIM_URL, "eng": ENG_URL}
    for name, dest in need:
        url = url_map.get(name)
        if not url:
            raise RuntimeError(f"未知语言包：{name}")
        print(f"正在下载 {name}.traineddata → {dest}", file=sys.stderr)
        _download_file(url, dest, min_bytes=50_000)

    # Windows 版 Tesseract：TESSDATA_PREFIX 指向 tessdata 目录本身（内含 *.traineddata）
    os.environ["TESSDATA_PREFIX"] = str(td)
    return td


def _jpeg_bytes_for_docx(raw: bytes) -> bytes:
    buf = io.BytesIO()
    with Image.open(io.BytesIO(raw)) as im:
        if im.mode in ("RGBA", "P"):
            im = im.convert("RGBA")
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=im.split()[-1])
            im = bg
        else:
            im = im.convert("RGB")
        im.save(buf, format="JPEG", quality=93, optimize=True, progressive=False)
    return buf.getvalue()


def _assert_jpeg_magic(jpeg_bytes: bytes) -> None:
    if not jpeg_bytes.startswith(b"\xff\xd8\xff"):
        raise RuntimeError("内部 JPEG 转换失败（缺少 SOI 标记）。")


def _set_cn_body_font(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(10.5)


def _configure_landscape_a4(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    m = Mm(12)
    section.left_margin = m
    section.right_margin = m
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)


def _content_box_mm(section) -> Tuple[float, float]:
    w_mm = float(section.page_width.mm - section.left_margin.mm - section.right_margin.mm)
    h_mm = float(section.page_height.mm - section.top_margin.mm - section.bottom_margin.mm)
    if w_mm <= 1 or h_mm <= 1:
        raise RuntimeError(f"版心尺寸异常：w={w_mm}mm h={h_mm}mm")
    return w_mm, h_mm


def _configure_tesseract(tesseract_exe: Path) -> None:
    try:
        import pytesseract  # type: ignore  # noqa: PLC0415
    except ImportError as e:  # pragma: no cover
        raise SystemExit("需要安装 pytesseract：pip install pytesseract") from e

    exe = tesseract_exe.expanduser()
    if not exe.is_file():
        raise FileNotFoundError(
            f"未找到 Tesseract：{exe}\n请安装或设置 TESSERACT_CMD / --tesseract。"
        )
    import pytesseract as pt  # type: ignore

    pt.pytesseract.tesseract_cmd = str(exe)


def _tesseract_plain_text(im: Image.Image, tesseract_exe: Path, *, psm: int) -> str:
    import pytesseract  # type: ignore  # noqa: PLC0415

    _configure_tesseract(tesseract_exe)
    cfg = f"--psm {int(psm)}"
    txt = pytesseract.image_to_string(im.convert("RGB"), lang="chi_sim+eng", config=cfg)
    return (txt or "").strip()


def build_docx(
    out_path: Path,
    *,
    ocr_appendix: bool,
    tesseract_exe: Path | None,
    download_lang: bool,
    fulltext_appendix: bool,
    table_psm: int,
    prose_psm: int,
    ocr_psm: int,
    dual_borderless_extract: bool = True,
    mask_pad_rel: float = 0.038,
    auto_upscale_pages: bool = True,
) -> None:
    doc = Document()
    _set_cn_body_font(doc)
    for sec in doc.sections:
        _configure_landscape_a4(sec)

    section = doc.sections[0]
    cw_mm, _ch_mm = _content_box_mm(section)
    width_in = max(4.0, (cw_mm / 25.4) * 0.94)

    tess_exe = _resolve_tesseract_exe(tesseract_exe)
    print(f"[cansino] 使用 Tesseract：{tess_exe}", file=sys.stderr, flush=True)
    if ocr_appendix or fulltext_appendix:
        ensure_traineddata_files(tess_exe, allow_download=download_lang)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("吸附无细胞百（三组分）白破联合疫苗 — 说明书（公示页）")
    tr.bold = True
    tr.font.size = Pt(15)
    tr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(3)
    sr = sub.add_run(f"来源：{SOURCE_PAGE}")
    sr.font.size = Pt(8.5)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note0 = doc.add_paragraph()
    n0r = note0.add_run(
        "以下为从公示页图像自动提取的可编辑正文与表格（img2table + Tesseract），"
        "未嵌入整页图。复杂版式、弱线框表或扫描质量差时可能有漏识或错格，请对照官网核对。"
    )
    n0r.font.size = Pt(9)
    n0r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    collected: List[Tuple[str, bytes]] = []
    for idx, (label, url) in enumerate(MANUAL_PAGES, start=1):
        raw = _download(url)
        jpeg_bytes = _jpeg_bytes_for_docx(raw)
        _assert_jpeg_magic(jpeg_bytes)
        if ocr_appendix or fulltext_appendix:
            collected.append((label, jpeg_bytes))

        if not ocr_appendix:
            h = doc.add_heading(f"{idx}. {label}", level=2)
            h.paragraph_format.keep_with_next = True
            for r in h.runs:
                r.font.name = "Microsoft YaHei"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if fulltext_appendix:
                doc.add_paragraph().add_run(
                    "（已使用 --no-ocr，主文未做表格与词级识别；若启用 --fulltext-appendix，见文末附录。）"
                ).font.size = Pt(9)
            else:
                doc.add_paragraph().add_run("（已使用 --no-ocr，未生成可编辑内容。）").font.size = Pt(9)
            if idx < len(MANUAL_PAGES):
                doc.add_page_break()

    if ocr_appendix and collected:
        try:
            import extract_tables_to_docx as etd  # type: ignore  # noqa: PLC0415
        except ImportError:
            doc.add_paragraph().add_run(
                "未加载 extract_tables_to_docx。请安装：pip install img2table opencv-contrib-python-headless "
                "pandas beautifulsoup4 pytesseract pillow"
            )
        else:
            try:
                etd.editable_docx_append_pages_with_headings(
                    doc,
                    collected,
                    tess_exe,
                    width_in=width_in,
                    table_psm=int(table_psm),
                    prose_psm=int(prose_psm),
                    dual_borderless=bool(dual_borderless_extract),
                    mask_pad_rel=float(mask_pad_rel),
                    auto_upscale=bool(auto_upscale_pages),
                )
            except Exception as e:
                doc.add_paragraph().add_run(f"可编辑排版失败：{e}")

    appendix_blocks: List[str] = []
    if fulltext_appendix and collected:
        doc.add_page_break()
        ah = doc.add_heading("附录：整页全文 OCR（与上文结构化排版无关）", level=2)
        for r in ah.runs:
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        note = doc.add_paragraph()
        note.add_run("以下为整页 Tesseract 输出，便于检索；与正文中的表格/段落划分无关。").font.size = Pt(9)
        for label, jb in collected:
            with Image.open(io.BytesIO(jb)) as im2:
                plain = _tesseract_plain_text(im2.convert("RGB"), tess_exe, psm=int(ocr_psm))
            appendix_blocks.append(f"【{label}】\n{plain}")
        body = doc.add_paragraph()
        br = body.add_run("\n\n".join(appendix_blocks))
        br.font.size = Pt(8)
        br.font.name = "Microsoft YaHei"
        br._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        br.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    out_final = Path(out_path).expanduser().resolve()
    out_final.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_final))
    if not out_final.is_file():
        raise RuntimeError(f"Word 保存失败：路径不存在 {out_final}")
    sz = out_final.stat().st_size
    if sz < 4000:
        raise RuntimeError(
            f"输出文件异常偏小（{sz} bytes），请检查是否误用 --no-ocr 或 OCR 失败。目标：{out_final}"
        )


def main(argv: List[str] | None = None) -> int:
    default_out = DEFAULT_OUTPUT_DIR / "吸附无细胞百白破联合疫苗_说明书_官网4843.docx"

    p = argparse.ArgumentParser(description="康希诺 detail-4843 说明书两页 → 横版 Word")
    p.add_argument(
        "-o",
        "--output",
        type=Path,
        default=default_out,
        help="输出 .docx（默认：仓库根目录下 output 文件夹内）",
    )
    p.add_argument("--no-ocr", action="store_true", help="不进行词级/表格识别；主文仅保留标题与说明（可与 --fulltext-appendix 同用）")
    p.add_argument(
        "--no-download-lang",
        action="store_true",
        help="禁止自动下载 chi_sim.traineddata（缺文件则报错）",
    )
    p.add_argument(
        "--tesseract",
        type=Path,
        default=None,
        help=r"指定 tesseract.exe，默认 TESSERACT_CMD 或 D:\tesseract-ocr\tesseract.exe",
    )
    p.add_argument(
        "--table-psm",
        type=int,
        default=11,
        help="img2table 内 Tesseract PSM，默认 11（稀疏表格）；可试 6",
    )
    p.add_argument(
        "--prose-psm",
        type=int,
        default=6,
        help="正文词级 OCR 的 Tesseract PSM，默认 6（单栏块）；可试 11",
    )
    p.add_argument(
        "--fulltext-appendix",
        action="store_true",
        help="文末追加整页全文 OCR（默认不追加）；可与 --no-ocr 同用以仅生成附录",
    )
    p.add_argument(
        "--ocr-psm",
        type=int,
        default=6,
        help="仅用于 --fulltext-appendix 的整页 OCR PSM，默认 6",
    )
    p.add_argument(
        "--no-dual-borderless",
        action="store_true",
        help="表格检测不比较 borderless 开/关两种模式（略快，复杂表可能略差）",
    )
    p.add_argument(
        "--mask-pad-rel",
        type=float,
        default=0.038,
        help="表格检测框相对膨胀（按表自身短边比例），默认 0.038",
    )
    p.add_argument(
        "--no-auto-upscale",
        action="store_true",
        help="关闭整页 1.5×/2× 放大重试（更快，难表可能更差）",
    )
    args = p.parse_args(argv)

    ocr_appendix = not bool(args.no_ocr)
    out = args.output.expanduser().resolve()

    try:
        build_docx(
            out,
            ocr_appendix=ocr_appendix,
            tesseract_exe=args.tesseract,
            download_lang=not bool(args.no_download_lang),
            fulltext_appendix=bool(args.fulltext_appendix),
            table_psm=int(args.table_psm),
            prose_psm=int(args.prose_psm),
            ocr_psm=int(args.ocr_psm),
            dual_borderless_extract=not bool(args.no_dual_borderless),
            mask_pad_rel=float(args.mask_pad_rel),
            auto_upscale_pages=not bool(args.no_auto_upscale),
        )
    except urllib.error.HTTPError as e:  # pragma: no cover
        print(f"下载失败：{e}", file=sys.stderr)
        return 1
    except Exception as e:  # pragma: no cover
        print(f"生成失败：{e}", file=sys.stderr)
        return 1

    if not out.is_file():
        print(f"生成失败：未找到输出文件 {out}", file=sys.stderr)
        return 1
    print(out, flush=True)
    print(f"[cansino] 已写入 Word：{out}", file=sys.stderr, flush=True)
    print(f"[cansino] 文件大小：{out.stat().st_size} 字节", file=sys.stderr, flush=True)
    try:
        marker = (REPO_ROOT / "output" / "LAST_IMAGE_OCR_DOCX_PATH.txt").resolve()
        marker.parent.mkdir(parents=True, exist_ok=True)
        marker.write_text(str(out.resolve()) + "\n", encoding="utf-8")
    except OSError:
        pass
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
