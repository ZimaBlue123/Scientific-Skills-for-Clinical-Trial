#!/usr/bin/env python3
"""
Extract text from .docx / .doc files.

Supports a single file, a list of files, or every supported file in a
folder. Modern .docx is handled via python-docx; legacy .doc is handled
via the Windows-only Word COM interface.

Output formats
--------------
* --format text  (default) plain paragraphs + tables-as-pipes
* --format md    Markdown-flavoured rendering. Replaces the now-archived
                 scripts/extract_docx_to_md.py.

Usage
-----
    py -3 scripts/extract_docx_full.py document.docx [output.txt]
    py -3 scripts/extract_docx_full.py document.docx --format md output.md
    py -3 scripts/extract_docx_full.py folder/ [output.txt]
"""
from __future__ import annotations

import argparse
import logging
import os
import sys
from collections.abc import Sequence
from pathlib import Path

LOG_FORMAT = "%(asctime)s [%(levelname)s] extract_docx_full: %(message)s"
logger = logging.getLogger("extract_docx_full")

ENCODING = "utf-8"


def extract_docx(filepath: Path) -> str:
    """Extract all text from a .docx using python-docx."""
    from docx import Document  # local import keeps module importable when missing

    doc = Document(str(filepath))
    full_text: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n\n".join(full_text)



def extract_docx_to_markdown(filepath: Path) -> str:
    """Render a .docx to Markdown.

    Paragraph styles named "Heading 1..3" / "Title" map to ``#``-style
    headings; tables render as GFM pipe tables. Replaces the archived
    ``extract_docx_to_md.py``.
    """
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(filepath))
    out: list[str] = []

    def _heading_level(para: Paragraph) -> int | None:
        name = (para.style.name or "").strip().lower()
        if name == "title":
            return 0
        if name.startswith("heading "):
            try:
                return int(name.split()[1])
            except (IndexError, ValueError):
                return None
        return None

    def _esc(s: str) -> str:
        return s.replace("|", "\\|").replace("\n", " ")

    def _render_table(tbl: Table) -> list[str]:
        rows = tbl.rows
        if not rows:
            return []
        md: list[str] = []
        header = [_esc(c.text.strip()) for c in rows[0].cells]
        md.append("| " + " | ".join(header) + " |")
        md.append("|" + "|".join(["---"] * len(header)) + "|")
        for row in rows[1:]:
            cells = [_esc(c.text.strip()) for c in row.cells]
            md.append("| " + " | ".join(cells) + " |")
        return md

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            text = para.text.rstrip()
            if not text:
                continue
            level = _heading_level(para)
            if level == 0:
                out.append(f"# {text}")
            elif level is not None and 1 <= level <= 6:
                out.append("#" * (level + 1) + " " + text)
            else:
                out.append(text)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            out.extend(_render_table(tbl))
    return "\n\n".join(out) + "\n"


def extract_doc_legacy(filepath: Path) -> str | None:
    """Extract text from a legacy .doc file using Word COM (Windows only)."""
    try:
        import pythoncom  # type: ignore[import-not-found]
        import win32com.client  # type: ignore[import-not-found]
    except ImportError:
        logger.error("pywin32 not available - cannot extract .doc files")
        return None

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        try:
            word = win32com.client.Dispatch("Word.Application")
        except Exception as exc:  # noqa: BLE001
            logger.error("failed to launch Word.Application: %s", exc)
            return None
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(filepath))
            return str(doc.Content.Text or "")
        except Exception as exc:  # noqa: BLE001
            logger.error("COM error while reading %s: %s", filepath, exc)
            return None
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:  # noqa: BLE001
                    logger.debug("doc.Close raised", exc_info=True)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:  # noqa: BLE001
                logger.debug("word.Quit raised", exc_info=True)
        try:
            pythoncom.CoUninitialize()
        except Exception:  # noqa: BLE001
            logger.debug("pythoncom.CoUninitialize raised", exc_info=True)


def extract_file(filepath: Path) -> str | None:
    """Dispatch to the right backend based on the file extension."""
    suffix = filepath.suffix.lower()
    if suffix == ".docx":
        return extract_docx(filepath)
    if suffix == ".doc":
        return extract_doc_legacy(filepath)
    logger.error("unsupported format: %s", suffix)
    return None


def extract_folder(
    folder_path: Path,
    output_path: Path | None = None,
) -> str:
    """Extract text from all docx/doc files in ``folder_path``."""
    files = [
        f for f in os.listdir(folder_path)
        if f.lower().endswith((".docx", ".doc"))
    ]

    combined: list[str] = [f"Found {len(files)} files", ""]

    for fname in files:
        path = Path(folder_path) / fname
        combined.append(f"\n========== {fname} ==========\n")
        try:
            text = extract_file(path)
            combined.append(text if text else "[FAILED to extract content]")
        except Exception as exc:  # noqa: BLE001
            combined.append(f"Error: {exc}")
        combined.append("")

    result = "\n".join(combined)

    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(result, encoding=ENCODING)
        print(f"Extracted {len(files)} files to: {output_path}")

    return result


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="extract_docx_full",
        description="Extract text from .docx and legacy .doc files.",
    )
    parser.add_argument("input", help="Input .docx/.doc file or a folder.")
    parser.add_argument("output", nargs="?", default=None, help="Optional output .txt path.")
    parser.add_argument(
        "--log-level", default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging verbosity (default: %(default)s).",
    )
    parser.add_argument(
        "--format", default="text",
        choices=("text", "md"),
        dest="fmt",
        help="Output format (default: %(default)s).",
    )
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    logging.basicConfig(level=getattr(logging, args.log_level), format=LOG_FORMAT)

    input_p = Path(args.input)
    output_p = Path(args.output) if args.output else None

    if not input_p.exists():
        logger.error("input does not exist: %s", input_p)
        return 2

    if input_p.is_dir():
        extract_folder(input_p, output_p)
        return 0

    if args.fmt == "md":
        try:
            text = extract_docx_to_markdown(input_p)
        except Exception as exc:  # noqa: BLE001
            logger.exception("failed to extract md from %s: %s", input_p, exc)
            return 1
    else:
        try:
            text = extract_file(input_p)
        except Exception as exc:  # noqa: BLE001
            logger.exception("failed to extract %s: %s", input_p, exc)
            return 1

    if not text:
        logger.error("failed to extract content: %s", input_p)
        return 1

    if output_p is not None:
        output_p.parent.mkdir(parents=True, exist_ok=True)
        output_p.write_text(text, encoding=ENCODING)
        print(f"Extracted to: {output_p}")
    else:
        print(text)
    return 0


if __name__ == "__main__":
    sys.exit(main())
