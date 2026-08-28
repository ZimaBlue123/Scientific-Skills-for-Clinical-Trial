"""
Audit DSUR document: extract full text + tables from three documents for cross-reference review.

Outputs to ``.workbuddy/audit/`` folder. The script tolerates missing input files,
records failures via ``logging``, and uses explicit type hints throughout.
"""

from __future__ import annotations

import json
import logging
import os
import sys
from collections.abc import Iterator

from docx import Document  # type: ignore
from docx.document import Document as _DocxDocument  # type: ignore
from docx.table import _Cell  # type: ignore

# ---------------------------------------------------------------------------
# Logging configuration
# ---------------------------------------------------------------------------
LOGGER = logging.getLogger("audit_dsur")
if not LOGGER.handlers:
    _handler = logging.StreamHandler(stream=sys.stderr)
    _handler.setFormatter(logging.Formatter("[%(levelname)s] %(name)s: %(message)s"))
    LOGGER.addHandler(_handler)
    LOGGER.setLevel(logging.INFO)

OUT_DIR = ".workbuddy/audit"

DOCS: dict[str, str] = {
    "DSUR1": "review_materials/Recombinant Varicella Vaccine (CHO Cell) DSUR1_V1.0_09Jul2026.docx",
    "PROTOCOL": "review_materials/Recombinant Varicella Vaccine (CHO Cell) Phase I Study_Protocol_V0.7_20250620-clean.docx",
    "IB": "review_materials/Recombinant Varicella Vaccine (CHO Cell)_Investigator's Brochure_V0.2_20250621-clean-updated.docx",
    "CLINICAL_OVERVIEW": "review_materials/Recombinant Varicella Vaccine (CHO Cell) _2.5 Clinical Overview_V0.5_20250514-clean.docx",
}


def get_block_items(doc: _DocxDocument) -> Iterator[tuple[str, object]]:
    """Yield ordered (kind, payload) items in body order: paragraphs and tables."""
    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    tbl_map = {t._tbl: t for t in doc.tables}
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            p = para_map.get(child)
            if p is None:
                continue
            yield ("p", p)
        elif tag.endswith("}tbl"):
            t = tbl_map.get(child)
            if t is None:
                continue
            yield ("t", t)


def cell_text(cell: _Cell) -> str:
    """Extract all paragraph text from a table cell."""
    parts = [p.text for p in cell.paragraphs]
    return "\n".join([x for x in parts if x is not None]).strip()


def extract_document(path: str) -> str:
    """Extract paragraphs and tables from *path* into a single human-readable string."""
    try:
        doc = Document(path)
    except (OSError, ValueError) as exc:
        LOGGER.error("Cannot open DOCX %s: %s", path, exc)
        raise

    out_lines: list[str] = []
    out_lines.append(f"===== FILE: {path} =====")
    out_lines.append(f"Total paragraphs: {len(doc.paragraphs)}")
    out_lines.append(f"Total tables: {len(doc.tables)}")
    out_lines.append("")

    table_idx = 0
    for kind, obj in get_block_items(doc):
        if kind == "p":
            para = obj  # type: Paragraph
            style = para.style.name if para.style else "?"
            text = para.text
            if text.strip():
                out_lines.append(f"[P style={style}] {text}")
            else:
                # keep blank line marker to preserve structure
                out_lines.append(f"[P style={style}] ")
        else:
            tbl = obj  # type: Table
            out_lines.append(f"[[TABLE {table_idx} rows={len(tbl.rows)} cols={len(tbl.columns)}]]")
            for r_i, row in enumerate(tbl.rows):
                cells = [cell_text(c).replace("\n", " | ") for c in row.cells]
                out_lines.append(f"  R{r_i}: " + " || ".join(cells))
            out_lines.append("[[/TABLE]]")
            table_idx += 1
    return "\n".join(out_lines)


def extract_keyword_hits(path: str, keywords: list[str]) -> list[tuple[int, str, str]]:
    """Return list of (line_no, keyword, line_text) hits (case-insensitive)."""
    try:
        full = extract_document(path).splitlines()
    except (OSError, ValueError) as exc:
        LOGGER.warning("Skipping keyword scan for %s: %s", path, exc)
        return []
    hits: list[tuple[int, str, str]] = []
    for i, line in enumerate(full, 1):
        low = line.lower()
        for kw in keywords:
            if kw.lower() in low:
                hits.append((i, kw, line))
                break
    return hits


def main(argv: list[str] | None = None) -> int:
    argv = argv if argv is not None else sys.argv[1:]
    os.makedirs(OUT_DIR, exist_ok=True)

    summary: dict[str, dict[str, object]] = {}
    for tag, path in DOCS.items():
        if not os.path.exists(path):
            LOGGER.error("Missing input file: %s", path)
            continue
        try:
            text = extract_document(path)
        except (OSError, ValueError) as exc:
            LOGGER.error("Failed to extract %s: %s", path, exc)
            continue

        out_path = os.path.join(OUT_DIR, f"{tag}_full.txt")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
        except OSError as exc:
            LOGGER.error("Cannot write %s: %s", out_path, exc)
            continue

        summary[tag] = {
            "src": path,
            "out": out_path,
            "lines": len(text.splitlines()),
        }
        print(f"[OK] {tag} -> {out_path} ({summary[tag]['lines']} lines)")

    try:
        with open(os.path.join(OUT_DIR, "_summary.json"), "w", encoding="utf-8") as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
    except OSError as exc:
        LOGGER.error("Cannot write summary JSON: %s", exc)
        return 1

    return 0 if summary else 2


if __name__ == "__main__":
    sys.exit(main())
