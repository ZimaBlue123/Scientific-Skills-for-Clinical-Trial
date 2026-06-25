#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extract text from Chinese and English IB (Investigator's Brochure) docx files
for side-by-side comparison.

This script:
- Loads two .docx files (Chinese IB and English IB).
- Extracts paragraphs and tables from each.
- Writes a numbered plain-text transcript and a separate tables dump
  for downstream comparison/diff workflows.

Usage
-----
    # Default: process the two IB files under <repo>/review_materials/
    py -3 scripts/extract_ib_texts.py

    # Custom folder / filenames / encoding
    py -3 scripts/extract_ib_texts.py \
        --input-dir review_materials \
        --cn-name "1.远大赛威信六价诺如疫苗 研究者手册-V2.0-20251225-清洁版.docx" \
        --en-name "Recombinant Hexavalent Norovirus Vaccine_IB_V2.0_25Dec2025.docx"

    # Programmatic API
    from extract_ib_texts import extract_ib
    extract_ib(input_dir=Path("review_materials"), output_dir=Path("review_materials"))
"""
from __future__ import annotations

import argparse
import logging
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Sequence

try:
    from docx import Document  # type: ignore
    from docx.document import Document as _DocumentT  # type: ignore
    from docx.table import Table as _TableT  # type: ignore
except ModuleNotFoundError as exc:  # pragma: no cover - import guard
    # Keep the module importable for tooling (compileall, IDE inspection).
    # The CLI entrypoint will surface a clear, actionable error to the user.
    Document = None  # type: ignore[assignment]
    _DOCX_IMPORT_ERROR: Optional[BaseException] = exc
else:
    _DOCX_IMPORT_ERROR = None


# ---------------------------------------------------------------------------
# Constants & defaults
# ---------------------------------------------------------------------------

DEFAULT_INPUT_DIR = "review_materials"
DEFAULT_CN_NAME = "1.远大赛威信六价诺如疫苗 研究者手册-V2.0-20251225-清洁版.docx"
DEFAULT_EN_NAME = "Recombinant Hexavalent Norovirus Vaccine_IB_V2.0_25Dec2025.docx"

CN_TEXT_OUT = "cn_ib_text.txt"
EN_TEXT_OUT = "en_ib_text.txt"
CN_TABLES_OUT = "cn_ib_tables.txt"
EN_TABLES_OUT = "en_ib_tables.txt"

ENCODING = "utf-8"
LOG_FORMAT = "%(asctime)s [%(levelname)s] %(name)s: %(message)s"

logger = logging.getLogger("extract_ib_texts")


# ---------------------------------------------------------------------------
# Data containers
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class IBExtractResult:
    """Aggregate of all artifacts produced for a single IB document."""

    source: Path
    paragraphs: int
    tables: int
    text_out: Path
    tables_out: Path


# ---------------------------------------------------------------------------
# Core extraction logic
# ---------------------------------------------------------------------------


def _paragraph_lines(paragraphs: Iterable[object]) -> List[str]:
    """Render paragraphs as `[P{n}] text` lines, skipping fully empty runs."""
    lines: List[str] = []
    for idx, para in enumerate(paragraphs, start=1):
        text = getattr(para, "text", "") or ""
        # Keep empty paragraphs as blank lines so the marker stays unique;
        # the caller can post-process if needed.
        lines.append(f"[P{idx}] {text}")
    return lines


def _table_to_text(table: _TableT) -> str:  # type: ignore[name-defined]
    """Render a single docx Table as a `[Table n]` header plus pipe-joined rows."""
    rendered_rows: List[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rendered_rows.append(" | ".join(cells))
    return "\n".join(rendered_rows)


def _table_blocks(tables: Iterable[_TableT]) -> List[str]:  # type: ignore[name-defined]
    """Render a list of tables into `[Table n]` blocks separated by blank lines."""
    blocks: List[str] = []
    for idx, table in enumerate(tables, start=1):
        body = _table_to_text(table)
        blocks.append(f"[Table {idx}]\n{body}" if body else f"[Table {idx}]")
    return blocks


def _write_text(path: Path, lines: Sequence[str], *, terminator: str = "\n") -> None:
    """Write `lines` to `path` atomically; creates parent dirs on demand."""
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = terminator.join(lines)
    if not payload.endswith(terminator):
        payload += terminator
    path.write_text(payload, encoding=ENCODING)
    logger.debug("wrote %d bytes to %s", len(payload.encode(ENCODING)), path)


def _load_doc(path: Path) -> "_DocumentT":  # type: ignore[name-defined]
    """Wrap `Document(path)` with a defensive import-error path."""
    if _DOCX_IMPORT_ERROR is not None or Document is None:  # pragma: no cover
        raise ModuleNotFoundError(
            "python-docx is required for IB extraction. "
            "Install via: python -m pip install python-docx"
        ) from _DOCX_IMPORT_ERROR
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"IB docx not found: {path}")
    return Document(str(path))  # type: ignore[misc]


def extract_one(
    source: Path,
    text_out: Path,
    tables_out: Path,
) -> IBExtractResult:
    """Extract paragraphs + tables from a single IB docx and persist as text."""
    logger.info("loading %s", source)
    doc = _load_doc(source)

    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)

    _write_text(text_out, _paragraph_lines(paragraphs))
    _write_text(tables_out, _table_blocks(tables), terminator="\n\n")

    result = IBExtractResult(
        source=source,
        paragraphs=len(paragraphs),
        tables=len(tables),
        text_out=text_out,
        tables_out=tables_out,
    )
    logger.info(
        "%s: %d paragraphs, %d tables extracted",
        source.name,
        result.paragraphs,
        result.tables,
    )
    return result


def extract_ib(
    input_dir: Path,
    output_dir: Optional[Path] = None,
    cn_name: str = DEFAULT_CN_NAME,
    en_name: str = DEFAULT_EN_NAME,
) -> List[IBExtractResult]:
    """Run the full extract pipeline for both Chinese and English IBs.

    Parameters
    ----------
    input_dir:
        Directory that contains the two source .docx files.
    output_dir:
        Directory to write text/tables artifacts to. Defaults to ``input_dir``.
    cn_name, en_name:
        Filenames of the Chinese and English IBs within ``input_dir``.
    """
    if output_dir is None:
        output_dir = input_dir

    if not input_dir.exists() or not input_dir.is_dir():
        raise FileNotFoundError(f"input directory not found: {input_dir}")

    cn_src = input_dir / cn_name
    en_src = input_dir / en_name

    jobs = [
        (cn_src, output_dir / CN_TEXT_OUT, output_dir / CN_TABLES_OUT),
        (en_src, output_dir / EN_TEXT_OUT, output_dir / EN_TABLES_OUT),
    ]

    results: List[IBExtractResult] = []
    for src, text_out, tables_out in jobs:
        try:
            results.append(extract_one(src, text_out, tables_out))
        except FileNotFoundError as e:
            logger.error("skip %s: %s", src.name, e)
        except Exception:  # noqa: BLE001 - top-level guard, log & rethrow below
            logger.exception("failed to extract %s", src)
            raise
    return results


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="extract_ib_texts",
        description=(
            "Extract paragraphs and tables from the Chinese and English IB "
            "docx files for comparison."
        ),
    )
    parser.add_argument(
        "--input-dir",
        default=DEFAULT_INPUT_DIR,
        help="Directory containing the source IB docx files (default: %(default)s).",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Directory to write extracted text/tables (default: same as --input-dir).",
    )
    parser.add_argument(
        "--cn-name",
        default=DEFAULT_CN_NAME,
        help="Filename of the Chinese IB inside --input-dir.",
    )
    parser.add_argument(
        "--en-name",
        default=DEFAULT_EN_NAME,
        help="Filename of the English IB inside --input-dir.",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        help="Logging verbosity (default: %(default)s).",
    )
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)

    logging.basicConfig(level=getattr(logging, args.log_level), format=LOG_FORMAT)

    if _DOCX_IMPORT_ERROR is not None:
        logger.error(
            "missing dependency 'python-docx'. Install via: "
            "python -m pip install python-docx"
        )
        return 2

    input_dir = Path(args.input_dir).resolve()
    output_dir = Path(args.output_dir).resolve() if args.output_dir else input_dir

    try:
        results = extract_ib(
            input_dir=input_dir,
            output_dir=output_dir,
            cn_name=args.cn_name,
            en_name=args.en_name,
        )
    except FileNotFoundError as e:
        logger.error("%s", e)
        return 2
    except Exception:  # noqa: BLE001
        logger.exception("IB extraction failed")
        return 1

    if not results:
        logger.error("No IB files were processed; check --input-dir and filenames.")
        return 2

    logger.info("done. processed %d IB(s).", len(results))
    return 0


if __name__ == "__main__":
    sys.exit(main())
