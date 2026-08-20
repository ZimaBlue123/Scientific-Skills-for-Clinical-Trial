# -*- coding: utf-8 -*-
"""Generate comprehensive Word document (V10) with corrected clinical trial conclusion wording."""

import sys
import re
try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_hyperlink(paragraph, url, text, size=7.5):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFont = OxmlElement('w:rFonts')
    rFont.set(qn('w:ascii'), 'Arial')
    rFont.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFont)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_cell_text(cell, text, bold=False, size=8, is_ref=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.1

    if is_ref:
        pattern = re.compile(r'(PMID:\s*)(\d+)|(DOI:\s*)(10\.\S+)|(NCT\d{8})|(CTR\d{8,10})')
        last_idx = 0
        for match in pattern.finditer(text):
            if match.start() > last_idx:
                run = p.add_run(text[last_idx:match.start()])
                run.font.size = Pt(size)
                run.font.name = 'Arial'
                
            if match.group(1):
                run = p.add_run(match.group(1))
                run.font.size = Pt(size)
                run.font.name = 'Arial'
                pmid = match.group(2)
                url = f'https://pubmed.ncbi.nlm.nih.gov/{pmid}/'
                add_hyperlink(p, url, pmid, size)
            elif match.group(3):
                run = p.add_run(match.group(3))
                run.font.size = Pt(size)
                run.font.name = 'Arial'
                doi = match.group(4)
                url = f'https://doi.org/{doi}'
                add_hyperlink(p, url, doi, size)
            elif match.group(5):
                nct = match.group(5)
                url = f'https://clinicaltrials.gov/study/{nct}'
                add_hyperlink(p, url, nct, size)
            else:
                pass 
            last_idx = match.end()
            
        if last_idx < len(text):
            run = p.add_run(text[last_idx:])
            run.font.size = Pt(size)
            run.font.name = 'Arial'
    else:
        parts = text.split('**')
        for idx, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            run.font.size = Pt(size)
            run.font.name = 'Arial'
            if len(parts) > 1 and idx % 2 != 0:
                run.font.bold = True
            else:
                run.font.bold = bold

doc = Document()
section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

title = doc.add_heading('CpG浣愬墏棰勯槻鎬х柅鑻楋細鏍稿績涓村簥璇曢獙涓庡畨鍏ㄦ€ф暟鎹眹鎬?(V10)', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('鏈枃浠舵眹鎬讳簡鍦?FDA 鍜?NMPA 绛夌櫥璁扮殑鍚?CpG 浣愬墏棰勯槻鎬х柅鑻楁牳蹇冧复搴婃暟鎹€傝瀺鍏?FDA 瀹樻柟瀹¤瘎(Clinical Review)绾ф牳蹇冨畨鍏ㄦ€ф暟鎹€?)
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

headers = [
    '鐤嫍鍚嶇О\n& 閫傚簲鐥?(鐮斿彂鐘舵€?',
    '娉ㄥ唽骞冲彴\n& 缂栧彿',
    '涓村簥璇曢獙鍩烘湰淇℃伅',
    '瀹夊叏鎬ф暟鎹眹鎬?,
    '鏍稿績鍙傝€冩枃鐚?
]

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(4.5), Cm(2.8), Cm(5.2), Cm(11.0), Cm(3.8)]
for i, width in enumerate(col_widths):
    table.columns[i].width = width

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    add_cell_text(hdr_cells[i], h, bold=True, size=9)
    set_cell_shading(hdr_cells[i], '2E75B6')
    for run in hdr_cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data = [
    {
        'vaccine': '**銆愨渽 宸蹭笂甯傘€?*\n(缇庡浗 FDA锛?017)\n\n**HEPLISAV-B**\n(HepB-CpG 1018)\n\n**閫傚簲鐥?*锛氶闃睭BV鎰熸煋 (鈮?8宀?\n**鐢冲姙鑰?*: Dynavax',
        'registry': 'FDA\nNCT01282762\n(HBV-23绛?',
        'clinical': (
            '**鍒嗘湡**: Phase 3\n'
            '**璁捐**: 闅忔満銆佽瀵熻€呯洸銆佹椿鎬у鐓n'
            '**鏍锋湰閲?*: ~6,665浜篭n'
            '**閫傚簲浜虹兢**: 18-70宀佹垚浜篭n'
            '**璇曢獙鍒嗙粍**: HEPLISAV-B缁?vs Engerix-B缁刓n'
            '**鍏嶇柅绋嬪簭**: 鍏?鍓?(0, 1涓湀锛岃倢鑲夋敞灏?\n'
            '**鐮旂┒缁堢偣**: 涔欒倽琛ㄩ潰鎶椾綋琛€娓呬繚鎶ょ巼(SPR鈮?0 mIU/mL)鍙婂畨鍏ㄦ€?
        ),
        'safety': (
            '銆怓DA Clinical Review 瀹樻柟瀹夊叏鏁版嵁銆慭n'
            '鈻?灞€閮ㄤ笉鑹弽搴?鏈€甯歌): 娉ㄥ皠閮ㄤ綅鐤肩棝(23%鈥?9%)銆佺孩鏂?2%鈥?%)銆佽偪鑳€(1%鈥?%) (ADR)\n'
            '鈻?鍏ㄨ韩涓嶈壇鍙嶅簲(鏈€甯歌): 鐤插姵(11%鈥?7%)銆佸ご鐥?8%鈥?7%)銆佽倢鐥?3%鈥?%) (ADR)\n'
            '鈻?涓ラ噸涓嶈壇浜嬩欢(SAE): HEPLISAV-B 缁勬€诲彂鐢熺巼 4.8%锛屾椿鎬у鐓х粍 4.8% (AE)\n'
            '鈻?鐗规畩鍏虫敞浜嬩欢(AESI - 鑷韩鍏嶇柅): 鍙戠敓鐜?0.1% (HEPLISAV-B) vs 0.1% (瀵圭収缁? (AE)\n'
            '鈻?鐗规畩鍏虫敞浜嬩欢(鎬ユ€у績鑲屾姝?AMI): 0.2% (HEPLISAV-B) vs 0.1% (瀵圭収缁?銆侳DA 鐙珛涓撳缁勮瘎浼板悗璁ゅ畾鈥滅己涔忕敓鐗╁鍚堢悊鎬р€濓紝鍒ゅ畾闈炵柅鑻楃浉鍏抽闄?(AE)銆?
        ),
        'ref': (
            'FDA Summary Basis for Regulatory Action (HEPLISAV-B)\n'
            'PMID: 37085451\n'
            'DOI: 10.1016/j.vaccine.2023.04.028'
        )
    },
    {
        'vaccine': '**銆愨渽 宸蹭笂甯傘€?*\n(涓浗 EUA锛?022)\n\n**SCB-2019**\n(閲嶇粍SARS-CoV-2涓夎仛浣揝铔嬬櫧鐤嫍 + CpG 1018)\n\n**閫傚簲鐥?*锛氶闃睠OVID-19\n**鐢冲姙鑰?*: 涓夊彾鑽夌敓鐗?,
        'registry': 'FDA\nNCT04672395\n(SPECTRA鍏ㄧ悆)',
        'clinical': (
            '**鍒嗘湡**: Phase 2/3 (SPECTRA)\n'
            '**璁捐**: 澶氫腑蹇冦€侀殢鏈恒€佸弻鐩层€佸畨鎱板墏瀵圭収\n'
            '**鏍锋湰閲?*: 30,137浜篭n'
            '**閫傚簲浜虹兢**: 鈮?8宀佹垚浜?鍚潚灏戝勾)\n'
            '**璇曢獙鍒嗙粍**: SCB-2019缁?vs 瀹夋叞鍓傜粍\n'
            '**鐮旂┒缁堢偣**: 淇濇姢鏁堝姏(VE)鍙婂畨鍏ㄦ€?
        ),
        'safety': (
            '銆愰潪寰侀泦鎬т笉鑹簨浠?(鑷矰ay43)銆慭n'
            '鈻?鐤嫍鐩稿叧鍏ㄨ韩涓嶈壇鍙嶅簲鎬诲彂鐢熺巼: 4.6% vs 瀹夋叞鍓?.0% (ADR)\n'
            '鈻?娉ㄥ皠閮ㄤ綅鐤肩棝: 2.0% (ADR) (鎬讳綋鏄捐憲杞诲井)\n'
            '鈻?涓ラ噸(Grade 3)灞€閮ㄥ弽搴? 浠?渚?0.02%) (AE)\n\n'
            '銆愰暱鏈熶弗閲?鐗规畩鍏虫敞浜嬩欢 (6涓湀闅忚)銆慭n'
            '鈻?SAE: 0.6% (90/15,070) (AE)\n'
            '鈻?SADR (鐤嫍鐩稿叧SAE): 浠?渚?0.027%) (ADR)\n'
            '鈻?AESI (鍙兘鑷厤绛?: 2.1% (323/15,070) (AE)'
        ),
        'ref': (
            'PMID: 36868877\n'
            'DOI: 10.1016/j.vaccine.2023.02.018'
        )
    },
    {
        'vaccine': '**銆愨渽 宸蹭笂甯傘€?*\n(鍙版咕鍦板尯 EUA)\n\n**MVC-COV1901**\n(閲嶇粍SARS-CoV-2 S铔嬬櫧鐤嫍 + CpG 1018)\n\n**閫傚簲鐥?*锛氶闃睠OVID-19\n**鐢冲姙鑰?*: 楂樼鐤嫍',
        'registry': 'FDA\nNCT04695652',
        'clinical': (
            '**鍒嗘湡**: Phase 2 (澶ц妯?\n'
            '**璁捐**: 澶氫腑蹇冦€侀殢鏈恒€佸弻鐩层€佸畨鎱板墏瀵圭収\n'
            '**鏍锋湰閲?*: 3,844浜篭n'
            '**璇曢獙鍒嗙粍**: MVC-COV1901缁?vs 瀹夋叞鍓傜粍\n'
            '**鐮旂┒缁堢偣**: 瀹夊叏鎬?鑰愬彈鎬э紝鍙婁腑鍜屾姉浣揋MT'
        ),
        'safety': (
            '銆愬緛闆嗘€т笉鑹簨浠?(鎺ョ鍚?澶?銆慭n'
            '鈻?娉ㄥ皠閮ㄤ綅鐤肩棝(Pain): 71.2% (2346/3295) (AE)\n'
            '鈻?涔忓姏/涓嶉€?Malaise/Fatigue): 36.0% (AE)\n'
            '鈻?鍙戠儹(Fever, 鈮?8掳C): 0.7% (23/3295) (AE)\n\n'
            '銆怱AE / SADR / AESI銆慭n'
            '鈻?SADR: 鏃犵柅鑻楃浉鍏充弗閲嶄笉鑹簨浠?(0 SADR)\n'
            '鈻?AESI: 1渚嬫殏鏃舵€ч潰绁炵粡楹荤椆(<0.1%)琚瘎浼板彲鑳戒笌鐤嫍鐩稿叧 (ADR)'
        ),
        'ref': (
            'PMID: 34655522\n'
            'DOI: 10.1016/S2213-2600(21)00402-1'
        )
    },
    {
        'vaccine': '**銆愨渽 宸蹭笂甯傘€?*\n(鍗板凹 BPOM EUA)\n\n**IndoVac**\n(閲嶇粍SARS-CoV-2铔嬬櫧鐤嫍 + CpG 1018)\n\n**閫傚簲鐥?*锛氶闃睠OVID-19\n**鐢冲姙鑰?*: Bio Farma / 璐濆嫆鍖诲闄?,
        'registry': 'FDA\nNCT05433285',
        'clinical': (
            '**鍒嗘湡**: Phase 3\n'
            '**璁捐**: 闅忔満銆佽瀵熻€呯洸銆佹椿鎬у鐓n'
            '**鏍锋湰閲?*: 4,050浜篭n'
            '**璇曢獙鍒嗙粍**: IndoVac缁?vs 瀹夋叞鍓傜粍\n'
            '**鐮旂┒缁堢偣**: 涓拰鎶椾綋浼樻晥/闈炲姡鏁堝強瀹夊叏鎬?
        ),
        'safety': (
            '銆愬緛闆嗘€т笉鑹簨浠?(Solicited AEs)銆慭n'
            '鈻?灞€閮?- 娉ㄥ皠閮ㄤ綅鐤肩棝(Pain): 14.69% (AE)\n'
            '鈻?鍏ㄨ韩 - 鎬讳綋鍙戠敓鐜? 27.95% (AE)\n'
            '鈻?鍏ㄨ韩 - 鑲岀棝(Myalgia): 7.48% (AE)\n'
            '鈻?鍏ㄨ韩 - 鐤插姵(Fatigue): 6.77% (AE)\n\n'
            '銆怱AE / SADR銆慭n'
            '鈻?SAE: 鏈彂鐜版瀬鍙兘涓庣柅鑻楃浉鍏崇殑SAE (0 SADR)銆?
        ),
        'ref': (
            'PMID: 38575433\n'
            'DOI: 10.1016/j.vaccine.2024.03.077'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺銆?*\n(Phase 1/2)\n\n**ZR202-CoV**\n(閲嶇粍鏂板啝S铔嬬櫧涓夎仛浣撶柅鑻?+ CpG 7909)\n\n**閫傚簲鐥?*锛氶闃睠OVID-19\n**鐢冲姙鑰?*: 娉芥鼎鐢熺墿',
        'registry': 'NMPA\nChiCTR2200057758\nNCT04990544',
        'clinical': (
            '**鍒嗘湡**: Phase 1/2\n'
            '**璁捐**: 闅忔満銆佸弻鐩层€佸畨鎱板墏瀵圭収\n'
            '**鏍锋湰閲?*: 72浜?P1) / 1,056浜?P2)\n'
            '**璇曢獙鍒嗙粍**: 鐤嫍缁?vs 瀹夋叞鍓傜粍\n'
            '**鐮旂┒缁堢偣**: 瀹夊叏鎬у強鍋囩梾姣掍腑鍜屾姉浣撴淮搴?
        ),
        'safety': (
            '銆愬緛闆嗘€т笉鑹簨浠躲€慭n'
            '鈻?娉ㄥ皠閮ㄤ綅鐤肩棝(Pain): 杞昏嚦涓害 (AE)\n'
            '鈻?鍙戠儹(Fever): 缃曡 (AE)\n'
            '鈻?鈮?绾?Severe)灞€閮ㄦ垨鍏ㄨ韩涓嶈壇浜嬩欢: 0% (AE)\n\n'
            '銆怱AE / SADR / AESI銆慭n'
            '鈻?SADR: 鏈彂鐢熶换浣曠柅鑻楃浉鍏充弗閲嶄笉鑹簨浠?(0渚? (ADR)\n'
            '鈻?AESI: 鏈瀵熷埌鐗规畩鍏虫敞浜嬩欢 (AE)'
        ),
        'ref': (
            'PMID: 37881130\n'
            'DOI: 10.1080/21645515.2023.2262635'
        )
    },
    {
        'vaccine': '**銆愨渽 宸蹭笂甯傘€?*\n(缇庡浗 FDA锛?023)\n\n**AV7909 / CYFENDUS庐**\n(BioThrax + CPG 7909浣愬墏)\n\n**閫傚簲鐥?*锛氱偔鐤芥毚闇插悗棰勯槻\n**鐢冲姙鑰?*: Emergent BioSolutions',
        'registry': 'FDA\nNCT03877926\n(Phase 3)',
        'clinical': (
            '**鍒嗘湡**: Phase 3 (鍏抽敭娉ㄥ唽璇曢獙)\n'
            '**璁捐**: 闅忔満銆佸弻鐩层€佹椿鎬у鐓n'
            '**鏍锋湰閲?*: 3,689浜篭n'
            '**閫傚簲浜虹兢**: 18-65宀佸仴搴锋垚浜篭n'
            '**璇曢獙鍒嗙粍**: AV7909缁?vs BioThrax瀵圭収缁刓n'
            '**鐮旂┒缁堢偣**: 鍏嶇柅鍘熸€?TNA锛屾浛浠ｇ粓鐐?鍙婂畨鍏ㄦ€?
        ),
        'safety': (
            '銆怓DA Package Insert 瀹樻柟瀹夊叏鏁版嵁銆慭n'
            '鈻?鏈€甯歌灞€閮ㄥ弽搴?鍙戠敓鐜?10%): 瑙︾棝(Tenderness, 74%)銆佺柤鐥?Pain, 51%)銆佸彂绾?Redness, 42%)銆佹墜鑷傛椿鍔ㄥ彈闄?29%)銆佽偪鑳€(Swelling, 22%) (ADR)\n'
            '鈻?鏈€甯歌鍏ㄨ韩鍙嶅簲: 鑲岃倝閰哥棝(Muscle Aches, 40%)銆佺柌鍔?Tiredness, 32%)銆佸ご鐥?Headache, 24%) (ADR)\n'
            '鈻?鐗瑰埆浣撳緛(涓€杩囨€?: 鍦ㄦ帴鍙楀惈 CpG 7909 浣愬墏鐨勫彈璇曡€呬腑锛岃瀵熷埌涓€杩囨€х殑缁濆娣嬪反缁嗚優璁℃暟涓嬮檷锛孎DA璇勪及瑙嗕负浣愬墏鐨勪竴杩囨€х敓鐞嗗綊宸㈡晥搴旓紝鏃犱复搴婄梾鐞嗘剰涔夈€俓n'
            '鈻?SADR: 鎬讳綋涓村簥姹犳湭鎶ュ憡涓庣柅鑻楀洜鏋滃叧绯绘槑纭殑 SAE (0 SADR)銆?
        ),
        'ref': (
            'FDA Package Insert (CYFENDUS)\n'
            'PMID: 41401704\n'
            'DOI: 10.1016/j.vaccine.2025.128068'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺銆?*\n(Phase 1b)\n\n**BK-SE36/CpG**\n(閲嶇粍鐤熷師铏玈E36鎶楀師 + CpG-ODN K3)\n\n**閫傚簲鐥?*锛氶闃茬枱鐤綷n**鐢冲姙鑰?*: BIKEN / 璐靛窞鐧剧伒',
        'registry': 'PACTR\nPACTR201701001921166',
        'clinical': (
            '**鍒嗘湡**: Phase 1b\n'
            '**璁捐**: 闅忔満銆佸弻鐩层€佸勾榫勯檷绾n'
            '**鏍锋湰閲?*: 135浜篭n'
            '**璇曢獙鍒嗙粍**: BK-SE36/CpG缁?vs 鍗曢摑瀵圭収缁刓n'
            '**鐮旂┒缁堢偣**: 瀹夊叏鎬у強鎶桽E36 IgG婊村害'
        ),
        'safety': (
            '銆愬緛闆嗘€у弽搴斿師鎬?(Day 1-7)銆慭n'
            '鈻?鎬讳綋鐤嫍鐩稿叧浜嬩欢鍙戠敓鐜? 38% vs 瀵圭収缁?4% (ADR)\n'
            '鈻?灞€閮ㄧ柤鐥?娲诲姩鍙楅檺: 鎴愪汉 17%鈥?3%; 鍎跨 40%鈥?7%; 骞煎効 6%鈥?9% (ADR)\n'
            '鈻?鍙戠儹(Fever): 5-10宀佺粍0-13%; 12-24鏈堥緞缁?3-29% (ADR)\n\n'
            '銆怱AE / SADR銆慭n'
            '鈻?SAE: 5渚?(閲嶇棁鐤熺柧)锛屽潎璇勪及涓庣柅鑻楁棤鍏?(AE)\n'
            '鈻?SADR / SUSAR: 闆舵姤鍛?(0渚? (ADR)'
        ),
        'ref': (
            'PMID: 37908361\n'
            'DOI: 10.3389/fimmu.2023.1267372'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺銆?*\n(Phase 2)\n\n**Na-GST-1/Al + CpG 10104**\n(閽╄櫕鐥呴噸缁勭柅鑻?\n\n**閫傚簲鐥?*锛氶闃查挬铏劅鏌揬n**鐢冲姙鑰?*: Sabin Vaccine Institute',
        'registry': 'FDA\nNCT03172975',
        'clinical': (
            '**鍒嗘湡**: Phase 2 (鍚獵HHI鍙楁帶鎰熸煋)\n'
            '**璁捐**: 闅忔満銆佸弻鐩层€佸畨鎱板墏瀵圭収\n'
            '**鏍锋湰閲?*: 39浜篭n'
            '**璇曢獙鍒嗙粍**: CpG鑱斿悎浣愬墏缁?vs 鍗曢摑缁刓n'
            '**鐮旂┒缁堢偣**: CHHI妯″瀷涓挬铏劅鏌撳己搴﹀強瀹夊叏鎬?
        ),
        'safety': (
            '銆愬弽搴斿師鎬т笌甯歌涓嶈壇浜嬩欢銆慭n'
            '鈻?灞€閮ㄥ強鍏ㄨ韩鍙嶅簲: 澶у鏁拌〃鐜颁负杞诲害(Mild) (AE)\n'
            '鈻?鍏嶇柅/琛€娑叉寚鏍? CpG浣愬墏缁勬樉钁楁姂鍒朵簡閽╄櫕寮曞彂鐨勫鍛ㄨ鍡滈吀鎬х矑缁嗚優澧炲 (涓綅鍊?0.6脳10鲁/渭L vs 瀹夋叞鍓傜粍 3.1脳10鲁/渭L, p=0.027) (ADR)\n\n'
            '銆怱AE / SADR銆慭n'
            '鈻?SADR: 鍏ㄧ▼鏈瀵熷埌鐤嫍鐩稿叧鐨勪弗閲嶄笉鑹簨浠?(0 SADR)銆?
        ),
        'ref': (
            'PMID: 41861834\n'
            'DOI: 10.1016/S1473-3099(26)00018-6'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺銆?*\n(Phase 1)\n\n**閲嶇粍涔欒倽鐤嫍** (姹夐€婇叺姣?\n(浣愬墏: CpG ODN 250渭g)\n\n**閫傚簲鐥?*锛氶闃睭BV\n**鐢冲姙鑰?*: 鍗庢櫘鐢熺墿/鍖椾含鐢熺墿鍒跺搧鐮旂┒鎵€',
        'registry': '鍥藉唴鍗曚腑蹇僜n(鏃燙TR鐧昏锛?016骞?',
        'clinical': (
            '**鍒嗘湡**: Phase 1\n'
            '**璁捐**: 闅忔満銆佸弻鐩层€佸鐓n'
            '**鏍锋湰閲?*: 48浜篭n'
            '**璇曢獙鍒嗙粍**: CpG璇曢獙缁?vs 閾濆鐓х粍\n'
            '**鐮旂┒缁堢偣**: 瀹夊叏鎬с€佽€愬彈鎬у強鍒濇鍏嶇柅鍘熸€?
        ),
        'safety': (
            '銆愭€讳綋涓嶈壇浜嬩欢 (AE)銆慭n'
            '鈻?CpG缁?66.67% (16/24) vs 閾濅綈鍓傜粍 54.17% (13/24) (AE)\n'
            ' (P=0.556, 宸紓鏃犵粺璁″鎰忎箟)\n'
            '鈻?涓ラ噸绋嬪害: 鍏ㄩ儴涓?Grade 1-2 杞讳腑搴﹀弽搴?(AE)\n'
            '鈻?鏃?鈮?绾?涓嶈壇浜嬩欢鎶ュ憡 (AE)\n\n'
            '銆怱AE / AESI銆慭n'
            '鈻?SADR / AESI: 鏈姤鍛婁换浣曚弗閲嶄笉鑹簨浠舵垨鑷厤浜嬩欢 (0渚? (AE)'
        ),
        'ref': (
            'PMID: 32842315\n'
            'DOI: 10.3760/cma.j.cn112150-20200401-00490'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺銆?*\n(Phase 1/2)\n\n**Z-1018**\n(甯︾姸鐤辩柟鐤嫍 + CpG 1018)\n\n**閫傚簲鐥?*锛氶闃插甫鐘剁柋鐤筡n**鐢冲姙鑰?*: Dynavax',
        'registry': 'FDA\nNCT06569823',
        'clinical': (
            '**鍒嗘湡**: Phase 1/2\n'
            '**璁捐**: 闅忔満銆佽瀵熺洸銆佸鐓?Shingrix)\n'
            '**鏍锋湰閲?*: 441浜?(Part 1)\n'
            '**璇曢獙鍒嗙粍**: Z-1018澶氬墏閲忕粍 vs Shingrix\n'
            '**鐮旂┒缁堢偣**: 鑰愬彈鎬у強鎶梘E IgG闃宠浆鐜?
        ),
        'safety': (
            '銆愭牳蹇冨畨鍏ㄦ暟鎹?(Part 1鏈熶腑鍒嗘瀽)銆慭n'
            '鈻?灞€閮ㄥ弽搴斿彂鐢熺巼(PIR): 涓噸搴︾柤鐥?绾㈣偪绛?Z-1018 (7.7%鈥?5.0%) 鏄捐憲浣庝簬 Shingrix (52.6%) (AE)\n'
            '鈻?鍏ㄨ韩鍙嶅簲鍙戠敓鐜? 涓噸搴﹁倢鐥?鐤插姵绛?Z-1018 (17.5%鈥?6.2%) 鏄捐憲浣庝簬 Shingrix (63.2%) (AE)\n'
            ' (鏍稿績缁撹: Z-1018 鍦ㄦ彁渚涘彲姣旀姉浣撳簲绛旂殑鍚屾椂锛岀郴缁熶笌灞€閮ㄥ弽搴斿師鎬уぇ骞呬笅闄?\n'
        ),
        'ref': (
            'DOI: 10.1093/ofid/ofaf695.018\n'
            '(OFID 2026浼氳鎽樿)'
        )
    },
    {
        'vaccine': '**銆愨渽 宸蹭笂甯傘€?*\n(鍗板害 EUA锛?021)\n\n**CORBEVAX**\n(閲嶇粍RBD铔嬬櫧鐤嫍 + CpG 1018)\n\n**閫傚簲鐥?*锛氶闃睠OVID-19\n**鐢冲姙鑰?*: Biological E / 璐濆嫆鍖诲闄?,
        'registry': 'CTRI\nCTRI/2021/08/036074',
        'clinical': (
            '**鍒嗘湡**: Phase 3\n'
            '**璁捐**: 鍗曠洸銆侀殢鏈恒€佹椿鎬у鐓?COVISHIELD)\n'
            '**鏍锋湰閲?*: 2,139浜篭n'
            '**璇曢獙鍒嗙粍**: CORBEVAX缁?vs COVISHIELD\n'
            '**鐮旂┒缁堢偣**: 涓拰鎶椾綋GMT浼樻晥鎬у強瀹夊叏鎬?
        ),
        'safety': (
            '銆愬緛闆嗘€т笉鑹簨浠?(鍚堝苟鏁版嵁)銆慭n'
            '鈻?娉ㄥ皠閮ㄤ綅鐤肩棝: 16.49% vs COVISHIELD 15.00% (AE)\n'
            '鈻?鍙戠儹: 11.00% vs 15.63% (AE)\n'
            '鈻?澶寸棝: 7.09% vs 6.56% (AE)\n'
            '鈻?鐤插姵: 6.05% vs 2.50% (AE)\n\n'
            '銆怱AE / SADR銆慭n'
            '鈻?SAE: 2渚?(鐧婚潻鐑瓑)锛屽潎璇勪及涓轰笌鐤嫍鏃犲叧 (AE)\n'
            '鈻?鐤嫍鐩稿叧SAE: 0渚?(0 SADR)'
        ),
        'ref': (
            'PMID: 37113012\n'
            'DOI: 10.1080/21645515.2023.2203632'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺銆?*\n(Phase 2)\n\n**VN-0200**\n(RSV F绯栬泲鐧?+ 尾-钁¤仛绯?CpG)\n\n**閫傚簲鐥?*锛氶闃睷SV鎰熸煋\n**鐢冲姙鑰?*: 绗竴涓夊叡',
        'registry': '鏃ユ湰jRCT\njRCT2071220051',
        'clinical': (
            '**鍒嗘湡**: Phase 2\n'
            '**璁捐**: 闅忔満銆佸弻鐩层€佸畨鎱板墏瀵圭収\n'
            '**鏍锋湰閲?*: 342浜篭n'
            '**璇曢獙鍒嗙粍**: 鐤嫍缁?vs 瀹夋叞鍓傜粍\n'
            '**鐮旂┒缁堢偣**: RSV涓拰鎶椾綋鍙婂畨鍏ㄦ€?
        ),
        'safety': (
            '銆愭牳蹇冨畨鍏ㄦ€ф暟鎹?(Phase 2)銆慭n'
            '鈻?寰侀泦鎬E鍙戠敓鐜? 楂樺墏閲忕粍 78.0% (32/41) (AE)\n'
            '鈻?涓ラ噸TEAE: 4渚?(1.2%)锛屽潎鍒ゅ畾涓庣柅鑻楁棤鍏?(AE)\n'
            '鈻?鐤嫍鐩稿叧涓ラ噸TEAE: 0渚?(0 SADR)\n'
            '鈻?鍥燭EAE鍋滆嵂: 4渚?(1.2%)锛屽叾涓?渚嬭偄浣撲笉閫傚垽涓虹浉鍏?(ADR)'
        ),
        'ref': (
            'PMID: 40257186\n'
            'DOI: 10.1080/21645515.2025.2489900'
        )
    },
    {
        'vaccine': '**銆愷煣?鍦ㄧ爺椤圭洰姹囨€汇€?*\n\n**涓村簥鐮斿彂绠＄嚎涓殑鏂板瀷 CpG 鐤嫍**\n(棰勯槻鎬х柅鑻?\n\n**閫傚簲鐥?*锛氬甫鐘剁柋鐤广€佷箼鑲濄€佹祦鎰熴€佺媯鐘梾绛塡n**鐢冲姙鑰?*: 鍚勫ぇ鍒涙柊鐤嫍浼佷笟',
        'registry': 'NMPA / FDA\n(涓浗鍙婃捣澶栧涓績)',
        'clinical': (
            '鈻?**甯︾姸鐤辩柟**: 绠€杈剧敓鐗?Phase 2); 鍚夎鍗敓鐗?Phase 2); 鎬￠亾/涓収鍏冮€?Phase 3瀹屾垚,宸叉姤NDA); 鏄庣憺浣矼RJ103(IND); 鍗庢櫘鐢熺墿HP2001(Phase 1); 杩滃ぇ鐢熺墿TVAX-006(Phase 2)銆俓n'
            '鈻?**涔欏瀷鑲濈値**: 杩滃ぇ鐢熺墿TVAX-008(Phase 3), TVAX-009(EOP2), 杩滃ぇ涔欒倽(Phase 2); 鍗庢櫘鐢熺墿HP2002(Phase 1)銆俓n'
            '鈻?**娴佹劅**: 鍗庢櫘鐢熺墿HP-3001瑁傝В鐤嫍(NMPA IND); 绠€杈剧敓鐗╅噸缁勮泲鐧芥祦鎰熺柅鑻?FDA IND)銆俓n'
            '鈻?**鍏朵粬鐤剧梾**: 闀挎槬鍗撹皧鐙傜姮(IND); Dynavax閲嶇粍榧犵柅(Phase 2); 杈夌憺RSV+CpG(Phase 1/2); Uvax Bio HIV棰勯槻鐤嫍(Phase 1)銆?
        ),
        'safety': (
            '銆愮绾垮畨鍏ㄦ€ц拷韪姸鎬併€慭n'
            '鈻?缁濆ぇ澶氭暟鍦ㄧ爺绠＄嚎澶勪簬鍙岀洸杩涘睍鏈燂紝鏈姭闇茬粓鐐瑰畨鍏ㄦ暟鎹€俓n'
            '鈻?鐗瑰埆娉ㄦ剰锛氫互涓婂垪琛ㄥ凡涓ユ牸鎺掗櫎鈥滄不鐤楁€ц偪鐦ょ柅鑻椻€?濡侶PV娌荤枟鎬х柅鑻?銆俓n'
            '鈻?琛屼笟鍔ㄦ€佹彁绀猴細鐩墠鍥藉唴鐢虫姤鐨勫ぇ閲?CpG 鐤嫍锛堢壒鍒槸甯︾姸鐤辩柟鍜屼箼鑲濓級澶у閲囩敤鈥滈噸缁勮泲鐧芥姉鍘?+ CpG 1018 绛夋晥鐗?+ 閾濅綈鍓傗€濈殑鑱斿悎鎶€鏈矾绾裤€傞壌浜庡厛鍙戜笂甯備骇鍝佺疮绉殑鏁颁竾渚嬬‘鍑夸复搴婅瘯楠屾暟鎹敮鎾戯紝杩欐壒绠＄嚎鐤嫍鍦ㄧ郴缁熸€у畨鍏ㄦ€у拰鑷厤椋庨櫓鎺у埗涓婂叿鏈夎緝楂樼殑鍓嶇疆纭畾鎬с€?
        ),
        'ref': (
            '鏁版嵁鏉ユ簮: \n'
            'ClinicalTrials.gov 鍙奬nNMPA/CDE 鏈€鏂板叕绀烘暟鎹?
        )
    }
]

for row_data in data:
    row_cells = table.add_row().cells
    fields = ['vaccine', 'registry', 'clinical', 'safety', 'ref']
    for i, field in enumerate(fields):
        is_ref = (field == 'ref')
        add_cell_text(row_cells[i], row_data[field], bold=False, size=7.5, is_ref=is_ref)

for idx, row in enumerate(table.rows[1:], start=1):
    if idx % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, 'F2F6FB')

doc.add_paragraph('')
footnote = doc.add_paragraph()
run = footnote.add_run(
    '銆愭爣娉ㄨ鏄庛€戯細\n'
    '1. 鏍囩 (AE) 琛ㄧず璇ユ暟鍊间负鈥滀笉鑹簨浠垛€濓紙Adverse Event锛夛紝涓嶄竴瀹氬叿鏈夊洜鏋滃叧绯汇€俓n'
    '2. 鏍囩 (ADR) 琛ㄧず璇ユ暟鍊间负鈥滀笉鑹弽搴斺€濓紙Adverse Drug Reaction锛夛紝鎸囨湁鐞嗙敱璁や负涓庣柅鑻楁帴绉嶅瓨鍦ㄥ洜鏋滃叧鑱旂殑浜嬩欢銆俓n'
    '3. 绗竴鍒楁爣璇嗕簡鍚勯」浜у搧鐩墠鐨勬渶楂樼爺鍙戜笌瀹℃壒鐘舵€侊細鉁?宸蹭笂甯?/ 馃И 鍦ㄧ爺銆?
)
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Add Summary Section
doc.add_paragraph('')
summary_title = doc.add_heading('銆愭€讳綋瀹夊叏鎬ф€荤粨銆戝熀浜庡惈 CpG 浣愬墏棰勯槻鎬х柅鑻椾复搴婃暟鎹殑缁煎悎璇勪环', level=2)
summary_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

summary_text = (
    "閫氳繃姹囨€诲叏鐞冭寖鍥村唴宸蹭笂甯傚強澶勪簬涓村簥涓悗鏈熺殑 CpG 浣愬墏棰勯槻鎬х柅鑻楁暟鎹紝鎴戜滑鍙互寰楀嚭浠ヤ笅绯荤粺鎬у畨鍏ㄦ€х粨璁猴細\n\n"
    "1. 浼樿秺鐨勫眬閮ㄤ笌鍏ㄨ韩鑰愬彈鎬?(Reactogenicity)锛歕n"
    "   浠ユ敞灏勯儴浣嶇柤鐥涳紙閫氬父鍦?0%-40%宸﹀彸锛夈€佺孩鑲夸互鍙婅交搴﹁嚦涓害鐨勭郴缁熸€х柌鍔炽€佸ご鐥涘拰鑲岀棝涓烘渶甯歌鐨勯潪涓ラ噸涓嶈壇鍙嶅簲锛圓DR锛夈€?
    "鍊煎緱娉ㄦ剰鐨勬槸锛岀浉杈冧簬濡?Shingrix (AS01B 鑴傝川浣撲綈鍓傜郴缁? 绛夊叾浠栧己鏁堟柊鍨嬩綈鍓傦紝CpG 1018 鍙婂叾鑱斿悎閾濅綈鍓傜殑鏂规鍦ㄨ瀵奸珮姘村钩浣撴恫鍙婄粏鑳炲厤鐤殑鍚屾椂锛?
    "灞曠幇鍑烘洿浣庣殑 3 绾э紙涓ラ噸锛夊眬閮ㄥ強鍏ㄨ韩鍙嶅簲鍘熸€у彂鐢熺巼锛堜互 Z-1018 璇曢獙涓樉钁楅檷浣庣殑涓噸搴﹀弽搴斿彂鐢熺巼涓哄吀鍨嬩唬琛級銆俓n\n"
    "2. 鏃犳槑鏄剧殑鑷韩鍏嶇柅鎬х柧鐥呭強涓ラ噸椋庨櫓淇″彿锛歕n"
    "   鍖呭惈 HEPLISAV-B 鍦ㄨ秴 6000 浜虹殑鏍稿績涓夋湡璇曢獙鍙婂ぇ瑙勬ā瀹夊叏鎬х洃娴嬶紝浠ュ強 SCB-2019 绛夋柊鍐犵柅鑻?3 涓囦汉鐨勫ぇ瑙勬ā鏁版嵁鍧囪〃鏄庯紝"
    "閲囩敤 CpG 浣愬墏鐨勯闃叉€х柅鑻椾笉浼氭彁鍗囨綔鍦ㄧ殑鍏嶇柅浠嬪鎬х柧鐥咃紙AESI锛屽 Guillain-Barr茅 缁煎悎寰併€佺被椋庢箍鎬у叧鑺傜値绛夛級鍙婂績琛€绠′簨浠讹紙濡傛€ユ€у績鑲屾姝伙級鐨勫彂鐢熷熀绾跨巼锛?
    "鍚勭粍涓庡畨鎱板墏鎴栦紶缁熼摑浣愬墏瀵圭収缁勪繚鎸佷簡楂樺害鐩镐技涓斿浜庢瀬浣庢按骞筹紙澶氭暟<0.1%锛夈€俓n\n"
    "3. 鐗瑰緛鎬х敓鐞嗗弽搴旂殑瀹夊叏鍙帶鎬э細\n"
    "   閮ㄥ垎璇曢獙锛堝 CYFENDUS銆丯a-GST-1锛夎瀵熷埌鐭殏涓斿彲閫嗙殑琛€娑插鏀瑰彉锛堝涓€杩囨€ф穻宸寸粏鑳炶鏁颁笅闄嶆垨澶栧懆琛€鍡滈吀鎬х矑缁嗚優璋冩帶锛夛紝"
    "FDA 涓村簥瀹¤瘎璁や负姝や负瀵℃牳鑻烽吀浣愬墏鐗规湁鐨勬€ユ€ч潪鐥呯悊鎬у厤鐤綊宸㈡晥搴旓紙Homing Effect锛夛紝鏃犻暱杩滅梾鐞嗗鎰忎箟锛屽弽鑰屼綈璇佷簡浣愬墏杩呴€熷惎鍔ㄦ穻宸寸粨鍐呭厤鐤簲绛旂殑鏈哄埗銆俓n\n"
    "缁撹锛欳pG 1018 鍙婂悓绫诲鏍歌嫹閰镐綈鍓傞€氳繃閽堝 TLR-9 鐨勯珮鐗瑰紓鎬ф縺娲伙紝鍦ㄤ繚璇佹瀬楂樻姉鍘熻娓呬繚鎶ょ巼锛圫PR锛夌殑鍚屾椂锛屾垚鍔熷墺绂讳簡浼犵粺寮烘晥浣愬墏甯镐即闅忕殑楂樺弽搴斿師鎬с€傚熀浜庣疮璁℃暟涓囦緥涓ヨ皑鐨勫ぇ瑙勬ā涓村簥璇曢獙瀹夊叏鎬ф暟鎹殑纭瘉锛屼互鍙?FDA 瀹樻柟瀹¤瘎鐨勮儗涔︼紝璇佸疄浜嗗叾浣滀负涓嬩竴浠ｂ€滈珮鏁堜綆姣掆€濋€氱敤鍨嬩汉鐢ㄧ柅鑻椾綈鍓傚钩鍙扮殑鍧氬疄瀹夊叏鎬ф綔鍔涖€?
)
p_summary = doc.add_paragraph(summary_text)
p_summary.paragraph_format.line_spacing = 1.3
p_summary.runs[0].font.size = Pt(9.5)
p_summary.runs[0].font.name = 'Microsoft YaHei'

output_path = r'E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\CpG_Vaccine_Safety_Summary-V10-20260820.docx'
doc.save(output_path)
import logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logging.info(f'Document successfully generated and saved to: {output_path}')

