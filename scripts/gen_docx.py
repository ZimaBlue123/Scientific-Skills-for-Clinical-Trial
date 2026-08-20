"""Generate comprehensive Word document with CpG vaccine safety data, integrating population and schedule."""

import re
import sys

try:
    import docx
except ImportError:
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_hyperlink(paragraph, url, text, size=7.5):
    """Adds a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font
    rFont = OxmlElement("w:rFonts")
    rFont.set(qn("w:ascii"), "Arial")
    rFont.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFont)

    # Color (Standard Hyperlink Blue)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)

    new_run.append(rPr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_cell_text(cell, text, bold=False, size=8, is_ref=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.2

    if is_ref:
        # Custom logic to parse and add hyperlinks for PMIDs and DOIs
        pattern = re.compile(
            r"(PMID:\s*)(\d+)|(DOI:\s*)(10\.\S+)|(NCT\d{8})|(CTR\d{8,10})"
        )
        last_idx = 0
        for match in pattern.finditer(text):
            if match.start() > last_idx:
                run = p.add_run(text[last_idx : match.start()])
                run.font.size = Pt(size)
                run.font.name = "Arial"

            if match.group(1):  # PMID matched
                run = p.add_run(match.group(1))
                run.font.size = Pt(size)
                run.font.name = "Arial"
                pmid = match.group(2)
                url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
                add_hyperlink(p, url, pmid, size)
            elif match.group(3):  # DOI matched
                run = p.add_run(match.group(3))
                run.font.size = Pt(size)
                run.font.name = "Arial"
                doi = match.group(4)
                url = f"https://doi.org/{doi}"
                add_hyperlink(p, url, doi, size)
            elif match.group(5):  # NCT matched
                nct = match.group(5)
                url = f"https://clinicaltrials.gov/study/{nct}"
                add_hyperlink(p, url, nct, size)
            else:
                pass
            last_idx = match.end()

        if last_idx < len(text):
            run = p.add_run(text[last_idx:])
            run.font.size = Pt(size)
            run.font.name = "Arial"
    else:
        # Normal text with bold parsing
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.size = Pt(size)
            run.font.name = "Arial"
            if len(parts) > 1 and idx % 2 != 0:
                run.font.bold = True
            else:
                run.font.bold = bold


doc = Document()

# --- Page Setup: Landscape A3 ---
section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(42.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# --- Title ---
title = doc.add_heading(
    "CpG佐剂预防性疫苗：核心临床试验与安全性数据汇总 (终局完善版)", level=1
)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run(
    "本文件汇总了在 FDA (ClinicalTrials.gov) 和 NMPA (中国临床试验注册平台) 注册的含 CpG 佐剂预防性疫苗的核心临床数据。\n本次升级严格核实并补充了各项试验的“适应人群”及“免疫程序”。各项安全性数据均标注了 (AE) 或 (ADR) 以清晰区分，PMID/DOI 号已添加直达超链接。"
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# --- Table ---
headers = [
    "疫苗名称\n& 适应症",
    "注册平台\n& 编号",
    "临床试验基本信息\n(分期 / 设计 / 样本量 / 人群 / 分组 / 程序 / 终点)",
    "安全性数据汇总",
    "核心参考文献\n(PMID / DOI)",
]

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
col_widths = [Cm(4.5), Cm(3.2), Cm(8.0), Cm(17.5), Cm(5.0)]
for i, width in enumerate(col_widths):
    table.columns[i].width = width

hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    add_cell_text(hdr_cells[i], h, bold=True, size=9)
    set_cell_shading(hdr_cells[i], "2E75B6")
    for run in hdr_cells[i].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ============================================================
# DATA ROWS
# ============================================================

data = [
    {
        "vaccine": "**HEPLISAV-B**\n(HepB-CpG 1018)\n\n**适应症**：预防乙型肝炎病毒(HBV)感染\n(FDA批准，≥18岁成人)",
        "registry": "FDA\nNCT01282762\n(HBV-23等)",
        "clinical": (
            "**分期**: Phase 3 (HBV-23, 关键注册试验)\n"
            "**设计**: 随机、观察者盲、活性对照\n"
            "**样本量**: ~6,665人\n"
            "**适应人群**: 18-70岁成人 (含健康人群及合并基础疾病/血透患者)\n"
            "**试验分组**: HEPLISAV-B组 vs Engerix-B组\n"
            "**免疫程序**: 共2剂 (0, 1个月，肌肉注射)\n"
            "**研究终点**: 血清保护率(anti-HBs≥10 mIU/mL)及安全性"
        ),
        "safety": (
            "【局部反应 (Local Reactions)】\n"
            "■ 注射部位疼痛(Pain): 23%–39% (ADR)\n"
            "■ 注射部位红斑(Erythema): 2%–4% (ADR)\n"
            "■ 注射部位肿胀(Swelling): 1%–3% (ADR)\n"
            " (注: 血液透析患者亚组局部反应率: 9.3% (AE))\n\n"
            "【全身反应 (Systemic Reactions)】\n"
            "■ 疲劳(Fatigue): 11%–17% (ADR)\n"
            "■ 头痛(Headache): 8%–17% (ADR)\n"
            "■ 肌痛(Myalgia): 3%–6% (ADR)\n"
            "■ 发热(Fever): <1% (ADR)\n\n"
            "【SAE / SADR / AESI】\n"
            "■ SAE: 4.8% vs 活性对照组 4.8% (AE)\n"
            "■ AESI (潜在免疫介导): 发生率 0.32% (AE)，与对照组(0.38%)无显著差异。\n"
            "■ 其他: 上市后大规模分析未见自身免疫疾病风险增加，无VAED。"
        ),
        "ref": (
            "PMID: 37085451\n"
            "DOI: 10.1016/j.vaccine.2023.04.028\n\n"
            "PMID: 36269938 (血透患者)\n"
            "PMID: 41819640 (AESI综述)"
        ),
    },
    {
        "vaccine": "**SCB-2019**\n(重组SARS-CoV-2三聚体S蛋白疫苗 + CpG-1018/Alum)\n\n**适应症**：预防新型冠状病毒感染(COVID-19)",
        "registry": "FDA\nNCT04672395\n(SPECTRA全球)\nNCT04954131\n(中国Phase 2)",
        "clinical": (
            "**分期**: Phase 2/3 (SPECTRA)\n"
            "**设计**: 多国多中心、随机、双盲、安慰剂对照\n"
            "**样本量**: 30,137人\n"
            "**适应人群**: ≥18岁成人 (含老年人及基础疾病患者)，及青少年(≥12岁)\n"
            "**试验分组**: SCB-2019组(n=15070) vs 安慰剂组(n=15067)\n"
            "**免疫程序**: 共2剂 (第1、22天，即间隔21天，肌肉注射)\n"
            "**研究终点**: COVID-19保护效力及长期安全性"
        ),
        "safety": (
            "【征集性不良事件 (Solicited AEs, 接种后7天)】\n"
            "■ 注射部位疼痛(Pain): 局部最常见 (AE)\n"
            "■ 红斑(Erythema) / 肿胀(Swelling): 有报告 (AE)\n\n"
            "【非征集性不良事件 (Unsolicited AEs, 至Day43)】\n"
            "■ 疫苗相关全身不良反应总发生率: 4.6% vs 安慰剂3.0% (ADR)\n"
            "■ 疫苗相关注射部位疼痛: 2.0% (ADR)\n"
            "■ 头痛(Headache): 0.8% (ADR)\n"
            "■ 严重(Grade 3)局部反应: 仅3例(0.02%) (AE)\n\n"
            "【长期严重/特殊关注事件 (6个月随访)】\n"
            "■ SAE: 0.6% (90/15,070) (AE)\n"
            "■ SADR (疫苗相关SAE): 仅4例(0.027%)，含超敏反应2例、Bell麻痹1例、极早期自发性流产1例 (ADR)\n"
            "■ AESI: 2.1% (323/15,070) (AE)，主要为与COVID-19感染相关的嗅觉丧失(1.7%)和味觉丧失(1.3%)。\n"
            "■ Bell麻痹: 2例(0.01%) (AE)；无GBS、TTS、心肌炎报告。"
        ),
        "ref": (
            "PMID: 36868877\n"
            "DOI: 10.1016/j.vaccine.2023.02.018\n\n"
            "PMID: 35065705\n(Lancet SPECTRA)"
        ),
    },
    {
        "vaccine": "**MVC-COV1901**\n(重组SARS-CoV-2 S-2P蛋白疫苗 + CpG 1018/Al(OH)₃)\n\n**适应症**：预防COVID-19\n(台湾地区EUA)",
        "registry": "FDA\nNCT04695652",
        "clinical": (
            "**分期**: Phase 2 (大规模)\n"
            "**设计**: 多中心、随机(6:1)、双盲、安慰剂对照\n"
            "**样本量**: 3,844人\n"
            "**适应人群**: ≥20岁健康成人 (含≥65岁老年人群)\n"
            "**试验分组**: 疫苗组(n=3295) vs 安慰剂组(n=549)\n"
            "**免疫程序**: 共2剂 (第1、29天，即间隔28天，肌肉注射)\n"
            "**研究终点**: 免疫原性及Day57安全性"
        ),
        "safety": (
            "【征集性局部不良事件 (接种后7天)】\n"
            "■ 注射部位疼痛(Pain): 71.2% (2346/3295) (AE)\n"
            " (多数为Grade 1-2，平均持续时间<3天)\n\n"
            "【征集性全身不良事件 (接种后7天)】\n"
            "■ 乏力/不适(Malaise/Fatigue): 36.0% (1186/3295) (AE)\n"
            "■ 发热(Fever, ≥38°C): 0.7% (23/3295) (AE)\n"
            " (大多呈自限性；≥65岁老年亚组整体AE发生率低于年轻人群)\n\n"
            "【SAE / SADR / AESI】\n"
            "■ SAE / SADR: 无任何疫苗相关的严重不良事件 (0 SADR)\n"
            "■ AESI: 1例暂时性面神经麻痹(<0.1%)被评估为可能与疫苗相关 (ADR)\n"
            "■ 其他: 试验期间无死亡、无VAED报告。"
        ),
        "ref": ("PMID: 34655522\n" "DOI: 10.1016/S2213-2600(21)00402-1"),
    },
    {
        "vaccine": "**IndoVac**\n(重组SARS-CoV-2蛋白亚单位疫苗 + CpG 1018/Alum)\n\n**适应症**：预防COVID-19",
        "registry": "FDA\nNCT05433285",
        "clinical": (
            "**分期**: Phase 3\n"
            "**设计**: 随机、观察者盲、活性对照(Covovax)\n"
            "**样本量**: 4,050人\n"
            "**适应人群**: ≥18岁健康成人\n"
            "**试验分组**: IndoVac组 vs Covovax活性对照组\n"
            "**免疫程序**: 共2剂 (第0、28天，肌肉注射)\n"
            "**研究终点**: 免疫原性(非劣效性)及安全性评价"
        ),
        "safety": (
            "【征集性不良事件 (Solicited AEs)】\n"
            "■ 局部 - 注射部位疼痛(Pain): 14.69% (AE)\n"
            "■ 局部 - 红斑(Erythema): 罕见 (AE)\n"
            "■ 全身 - 总体发生率: 27.95% (AE)\n"
            "■ 全身 - 肌痛(Myalgia): 7.48% (AE)\n"
            "■ 全身 - 疲劳(Fatigue): 6.77% (AE)\n"
            "■ 全身 - 头痛(Headache): 有报告 (AE)\n"
            " (注: 大部分症状为轻度 Mild)\n\n"
            "【SAE / SADR】\n"
            "■ SAE: 未发现可能或极可能与疫苗相关的SAE (0 SADR)。\n"
            "■ 整体安全耐受性优良。"
        ),
        "ref": ("PMID: 38575433\n" "DOI: 10.1016/j.vaccine.2024.03.077"),
    },
    {
        "vaccine": "**ZR202-CoV**\n(重组新冠S蛋白三聚体疫苗 + Alum/CpG 7909)\n\n**适应症**：预防COVID-19\n(上海泽润生物)",
        "registry": "NMPA\nChiCTR2200057758\nNCT04990544",
        "clinical": (
            "**分期**: Phase 1/2\n"
            "**设计**: 随机、双盲、安慰剂对照\n"
            "**样本量**: 72人(P1) / 1,056人(P2)\n"
            "**适应人群**: ≥18岁健康成人 (P1为18-59岁，P2扩展至≥60岁老年人)\n"
            "**试验分组**: 疫苗组(含不同抗原/佐剂剂量) vs 安慰剂组\n"
            "**免疫程序**: 共2剂 (第0、28天，肌肉注射)\n"
            "**研究终点**: 28天内安全性/耐受性及免疫原性"
        ),
        "safety": (
            "【征集性不良事件】\n"
            "■ 注射部位疼痛(Pain): 轻至中度 (AE)\n"
            "■ 发热(Fever): 罕见 (AE)\n"
            "■ ≥3级(Severe)局部或全身不良事件: 0% (AE)\n"
            " (多数为Grade 1轻度反应，均于7天内自行缓解)\n\n"
            "【SAE / SADR / AESI】\n"
            "■ SADR: 未发生任何与疫苗相关的严重不良事件 (0例) (ADR)\n"
            "■ AESI: 未观察到过敏性休克等特殊关注事件 (AE)"
        ),
        "ref": ("PMID: 37881130\n" "DOI: 10.1080/21645515.2023.2262635"),
    },
    {
        "vaccine": "**AV7909 / CYFENDUS®**\n(BioThrax + CPG 7909佐剂)\n\n**适应症**：炭疽暴露后预防\n(FDA批准，≥18岁)",
        "registry": "FDA\nNCT03877926\n(Phase 3)",
        "clinical": (
            "**分期**: Phase 3 (关键注册试验)\n"
            "**设计**: 随机、双盲、活性对照\n"
            "**样本量**: 3,689人\n"
            "**适应人群**: 18-65岁健康成人\n"
            "**试验分组**: AV7909组 vs BioThrax对照组\n"
            "**免疫程序**: 共2剂 (第0、2周 / 即第1、15天，肌肉注射)\n"
            "**研究终点**: 免疫保护率及安全性"
        ),
        "safety": (
            "【不良事件特征 (Phase 2 & 3综合)】\n"
            "■ 局部 - 注射部位反应: 为最常见的疫苗相关AE (ADR)\n"
            "■ 整体 AE 发生率: AV7909组约 79% vs BioThrax 65% (AE)\n"
            "■ 严重程度: >90%的全身与局部不良事件均为 Grade 1-2 (轻至中度) (AE)\n\n"
            "【血液系统发现 (Phase 1)】\n"
            "■ 一过性淋巴细胞减少(Transient Lymphopenia): 在含有CpG 7909的各剂量组中均观察到，为一过性生理反应 (ADR)\n\n"
            "【SAE / SADR / AESI】\n"
            "■ SADR: 未判定出与疫苗可能相关的SAE (0 SADR)。\n"
            "■ AESI: 多期试验中均未观察到表明自身免疫反应的特殊安全性关切 (AE)。"
        ),
        "ref": (
            "PMID: 41401704 (Phase 3)\n"
            "DOI: 10.1016/j.vaccine.2025.128068\n\n"
            "PMID: 26979136 (Phase 2)\n"
            "DOI: 10.1016/j.vaccine.2016.03.006\n\n"
            "PMID: 23701746 (Phase 1)\n"
            "DOI: 10.1016/j.vaccine.2013.04.063"
        ),
    },
    {
        "vaccine": "**BK-SE36/CpG**\n(重组恶性疟原虫SE36抗原 + Al(OH)₃/CpG-ODN K3)\n\n**适应症**：预防疟疾",
        "registry": "非洲注册库\nPACTR201701001921166",
        "clinical": (
            "**分期**: Phase 1b\n"
            "**设计**: 随机、双盲、活性对照(狂犬疫苗)、年龄降级\n"
            "**样本量**: 135人\n"
            "**适应人群**: 疟疾流行区人群 (分21-45岁, 5-10岁, 12-24月龄三个队列)\n"
            "**试验分组**: 疫苗组 vs 对照组\n"
            "**免疫程序**: 共2剂 (间隔28天，肌肉注射)\n"
            "**研究终点**: 安全性/反应原性(主要)"
        ),
        "safety": (
            "【征集性反应原性 (Reactogenicity, Day 1-7)】\n"
            "■ 总体疫苗相关事件发生率: 38% vs 对照组14% (ADR)\n"
            "■ 局部 - 疼痛/活动受限:\n"
            "   成人: 17%–33% (ADR)\n"
            "   5-10岁: 40%–57% (ADR)\n"
            "   12-24月龄: 6%–19% (ADR)\n"
            "■ 全身 - 头痛(Headache): 成人组 3%–7% (ADR)\n"
            "■ 全身 - 发热(Fever): 5-10岁组 0-13%; 12-24月龄组 13-29% (ADR)\n"
            " (全部疫苗相关AE均为Grade 1-2轻中度)\n\n"
            "【SAE / SADR】\n"
            "■ SAE: 全程5例 (重症疟疾住院), 经评估均与疫苗无关 (AE)\n"
            "■ SADR / SUSAR: 零报告 (0例) (ADR)"
        ),
        "ref": ("PMID: 37908361\n" "DOI: 10.3389/fimmu.2023.1267372"),
    },
    {
        "vaccine": "**Na-GST-1/Al + CpG 10104**\n(钩虫病重组疫苗)\n\n**适应症**：预防钩虫感染",
        "registry": "FDA\nNCT03172975",
        "clinical": (
            "**分期**: Phase 2 (含CHHI人类受控感染模型)\n"
            "**设计**: 随机、双盲、安慰剂对照\n"
            "**样本量**: 39人\n"
            "**适应人群**: 18-45岁健康成人\n"
            "**试验分组**: 4组(含单Al组及CpG联合组)\n"
            "**免疫程序**: 共3剂 (常规0, 2, 4月或按CHHI特定间隔)\n"
            "**研究终点**: 感染率及安全性"
        ),
        "safety": (
            "【反应原性与常规不良事件】\n"
            "■ 局部及全身反应: 大多数表现为轻度(Mild) (AE)\n"
            "■ 血液免疫学标志: 受控感染后，CpG组显著抑制了外周血嗜酸性粒细胞增多 (中位值 0.6×10³/μL vs 安慰剂 3.1×10³/μL, p=0.027) (ADR)\n\n"
            "【SAE / SADR】\n"
            "■ SAE / SADR: 全程未观察到任何与疫苗相关的严重不良事件 (0 SADR)。\n"
            "■ 耐受性良好。"
        ),
        "ref": ("PMID: 41861834\n" "DOI: 10.1016/S1473-3099(26)00018-6"),
    },
    {
        "vaccine": "**重组乙型肝炎疫苗** (汉逊酵母)\n(佐剂: CpG ODN 250μg)\n\n**适应症**：预防乙型肝炎(HBV)",
        "registry": "国内单中心\n(无CTR登记，2016年开展)",
        "clinical": (
            "**分期**: Phase 1\n"
            "**设计**: 随机、双盲、对照\n"
            "**样本量**: 48人\n"
            "**适应人群**: 18-45岁健康成人 (HBsAg阴性)\n"
            "**试验分组**: CpG试验组(24) vs 铝佐剂对照组(24)\n"
            "**免疫程序**: 共3剂 (0、1、6个月，肌肉注射)\n"
            "**研究终点**: 安全性及免疫原性"
        ),
        "safety": (
            "【总体不良事件 (AE)】\n"
            "■ CpG组 66.67% (16/24) vs 铝佐剂组 54.17% (13/24) (AE)\n"
            " (P=0.556, 两组间差异无统计学意义)\n"
            "■ 严重程度: 全部为 Grade 1-2 轻中度反应 (AE)\n"
            "■ 无 ≥3级 不良事件报告 (AE)\n\n"
            "【SAE / AESI】\n"
            "■ SAE / SADR: 未报告任何严重不良事件 (0例) (AE)\n"
            "■ AESI: 未报告免疫介导相关不良事件 (AE)"
        ),
        "ref": ("PMID: 32842315\n" "DOI: 10.3760/cma.j.cn112150-20200401-00490"),
    },
    {
        "vaccine": "**临床研发管线中的新型 CpG 疫苗**\n(含国内进行中/未发表项目)\n\n**适应症**：带状疱疹、乙肝等",
        "registry": "NMPA/FDA\n(如右侧信息)",
        "clinical": (
            "■ **带状疱疹疫苗(gE/CpG+MF59)**: Phase 2进行中 (NCT06851832)。**人群**: 50岁及以上; **程序**: 2剂(0, 2月)\n"
            "■ **怡道/中慧元通 带状疱疹疫苗** (CpG): Phase 3完成已报NDA (CTR20233253)。**人群**: ≥40岁; **程序**: 2剂\n"
            "■ **明瑞佳 MRJ103 带状疱疹疫苗** (CpG 1018+QS-21等): IND获批 (CXSL2600443)\n"
            "■ **华普生物 HP2001 带状疱疹疫苗** (HP007/CpG类): Phase 1 (CXSL2500371)\n"
            "■ **华普生物 HP2002 乙肝疫苗** (CpG+Alum): Phase 1 (CTR20261430)。**人群**: 18-59岁; **程序**: 3剂(0-1-6)\n"
            "■ **重组乙肝疫苗(汉逊酵母)** (CpG-QCX1): Phase 1 (CTR20233405)。**人群**: 18-59岁; **程序**: 3剂(0-1-6)\n"
            "■ **远大乙肝疫苗** (CpG+Alum): Phase 2进行中 (ChiCTR2500108408)"
        ),
        "safety": (
            "【安全性数据状态】\n"
            "■ 核心数据暂未公开 (试验尚在进行中，或已完成正处于NMPA审评审批阶段)。\n"
            "■ 行业动态提示：目前大量基于 CpG 佐剂的带状疱疹疫苗（对标 Shingrix，使用 CpG 联合 MF59 或 QS-21 构成复合佐剂体系）及新一代乙肝疫苗，正在中国密集进入各期临床开发阶段。由于 CpG 安全性已被证实，这些项目的整体耐受性预期良好，有待披露最终数据。"
        ),
        "ref": (
            "数据来源: \n"
            "ClinicalTrials.gov 及\nNMPA/CDE (中国临床试验注册与信息公示平台) 最新公示数据"
        ),
    },
]

for row_data in data:
    row_cells = table.add_row().cells
    fields = ["vaccine", "registry", "clinical", "safety", "ref"]
    for i, field in enumerate(fields):
        # Apply special hyperlink parsing only to the 'ref' column
        is_ref = field == "ref"
        add_cell_text(
            row_cells[i], row_data[field], bold=False, size=7.5, is_ref=is_ref
        )

# Alternating row colors
for idx, row in enumerate(table.rows[1:], start=1):
    if idx % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, "F2F6FB")

# --- Footnote ---
doc.add_paragraph("")
footnote = doc.add_paragraph()
run = footnote.add_run(
    "【标注说明】：\n"
    "1. 标签 (AE) 表示该数值为“不良事件”（Adverse Event），仅为试验期间发生的医学事件，不一定具有因果关系。\n"
    "2. 标签 (ADR) 表示该数值为“不良反应”（Adverse Drug Reaction），指经临床研究者判定，有理由认为与疫苗接种存在因果关联的事件。\n"
    "3. 缩写释义：AESI (特殊关注不良事件)；SAE (严重不良事件)；SADR (严重不良反应)；VAED (疫苗相关增强疾病)。\n"
    "4. 数据处理原则：提取数值时优先采用 ADR 数据；若文献未区分或未针对某指征报告 ADR，则采用总的 AE 数据呈现。\n"
    "5. 版式说明：安全性数据列并未生硬统一分类维度，而是根据各疫苗核心文献的原始报告特征（如按“征集性/非征集性”或按“局部/全身”）进行最契合实际的结构化排版。"
)
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# --- Save ---
output_path = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\CpG_Vaccine_Safety_Summary.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
