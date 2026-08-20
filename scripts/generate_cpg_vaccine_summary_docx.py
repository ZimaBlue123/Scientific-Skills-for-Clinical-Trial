"""
国内外 CpG 佐剂预防性疫苗安全性数据汇总 - v3.0 单总表生成脚本
输出: review_materials/CpG_Vaccine_Safety_Summary.docx
版式: 无封面/无目录/无章节编号; 横向 A4; 单张总表 5 列; 表头浅灰底纹; 地域+状态分块
v3.0 要点:
  - 5 列: 疫苗名称&适应症 / 注册平台&编号 / 临床试验基本信息 / 安全性数据汇总 / 核心参考文献
  - 安全性列按原文分类逻辑动态分层, 每项数据末尾标注 (AE)/(ADR), 含 SAE/SADR/AESI
  - 禁止定性描述, 全部数值来自原文提取
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_PATH = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\CpG_Vaccine_Safety_Summary.docx"

# ================================================================ data ====
# 每行: (列1 疫苗名称&适应症, 列2 注册平台&编号, 列3 基本信息, 列4 安全性行列表, 列5 参考文献)
# 列4 safety lines: ('h', 分类标题-加粗) / ('i', 数据条目-缩进) / ('b', 加粗说明段)
ROWS = [
    # ================= 国外(已上市/已发表) =================
    (
        "HEPLISAV-B（HBsAg-1018）\n佐剂: CpG 1018（TLR9激动剂）\n预防适应症: 乙型肝炎（HBV）",
        "FDA（ClinicalTrials.gov）\nNCT00435812\nNCT01005407\nNCT02117934",
        "试验分期: 上市（FDA 2017）; Ⅲ期×3\n整体设计: 随机、观察者盲、活性对照（Engerix-B）、多中心\n样本量: 14,238（10,038疫苗/4,200对照）\n试验分组: 2剂(0,4周)+24周安慰剂 vs Engerix-B 3剂(0,4,24周)\n研究终点: 主-28周血清保护率(anti-HBs≥10mIU/mL); 次-安全性",
        [
            ("h", "【未征集AE(28天)】(AE)"),
            ("i", "T1 42.0% vs 41.3%; T2 35.4% vs 36.2%; T3 20.1% vs 20.1%"),
            ("h", "【轻中度AE/心血管事件】(AE)"),
            ("i", "轻中度AE 45.6% vs 45.7%; 心血管事件 0.27% vs 0.14%"),
            ("h", "【SAE】(AE)"),
            ("i", "T1 1.5% vs 2.1%; T2 3.9% vs 4.8%; T3 6.2% vs 5.3%"),
            ("h", "【AESI-免疫介导AE(IMAE)】(AE)"),
            (
                "i",
                "T1 0.2% vs 0.7%; T2 0.3% vs 0.0%; T3 0.1% vs 0%; 三试验汇总 0.32% vs 0.38%（综述）",
            ),
        ],
        "Hyer R, et al. Vaccine. 2018;36(19):2604-2611\nPMID: 29628151 | DOI: 10.1016/j.vaccine.2018.03.067\nhttps://pubmed.ncbi.nlm.nih.gov/29628151/\n另: Janssen RS, Coffman RL. Vaccine. 2026;79:128437 (IMAE汇总)\nPMID: 41819640 | DOI: 10.1016/j.vaccine.2026.128437",
    ),
    (
        "AV7909（AVA+CpG 7909, NuThrax）\n佐剂: CpG 7909（TLR9激动剂）\n预防适应症: 炭疽（暴露后预防）",
        "FDA（ClinicalTrials.gov）\nNCT01770743",
        "试验分期: Ⅱ期（上市 FDA 2023）\n整体设计: 随机、双盲、BioThrax活性对照、多中心\n样本量: 168\n试验分组: AV7909 0/14d、0/28d、0/14/28d、半剂量0/14/28d vs BioThrax 0/14/28d\n研究终点: 主-D63 TNA应答率; 次-AE/SAE/反应原性发生率",
        [
            ("h", "【总体AE】(AE)"),
            ("i", "79%（AV7909）vs 65%（BioThrax）; 92% vs 87% 为1-2级"),
            ("h", "【反应原性】(AE)"),
            ("i", "原文未逐项报道局部/全身各反应发生率数值"),
            ("h", "【SAE】(AE)"),
            (
                "i",
                "2/44（4.55%, 0/14d组）; 0/34、0/23、0/44; BioThrax 1/23（4.35%）; 无疫苗相关SAE（注册结果页）",
            ),
            ("h", "【AESI】(AE)"),
            ("i", "无自身免疫病因AE报告"),
        ],
        "Hopkins RJ, et al. Vaccine. 2016;34(18):2096-2105\nPMID: 26979136 | DOI: 10.1016/j.vaccine.2016.03.006\n注册结果页: https://clinicaltrials.gov/study/NCT01770743",
    ),
    (
        "SCB-2019（S-Trimer）\n佐剂: CpG 1018 + Alum\n预防适应症: COVID-19（SARS-CoV-2）",
        "FDA（ClinicalTrials.gov）\nNCT04672395",
        "试验分期: Ⅱ/Ⅲ期\n整体设计: 多中心、双盲、随机、安慰剂对照（5国）\n样本量: 30,174入组; 30,128接种（15,064/15,064）\n试验分组: SCB-2019 2剂(0,21天) vs 安慰剂\n研究终点: 主-任何严重程度COVID-19保护效力; 次-中重度/重症保护效力、安全性",
        [
            ("h", "【征集性局部反应】(AE)"),
            (
                "i",
                "第1剂后 36.1%（290/803）vs 11.3%（89/786）; 注射部位疼痛 35.7% vs 10.3%",
            ),
            ("i", "第2剂后 28.2% vs 8.2%; 疼痛 26.9%（189/702）vs 7.4%（52/699）"),
            ("h", "【征集性全身反应】(AE)"),
            ("i", "第1剂后 36% vs 34%; 第2剂后 23% vs 21%"),
            ("h", "【非征集性AE】(AE)"),
            ("i", "12.3% vs 12.4%（两组均衡）"),
            ("h", "【SAE】(AE)"),
            (
                "i",
                "0.3%（49/15,064）vs 0.4%（59/15,064）; 治疗相关4例（中度过敏、轻度贝尔氏麻痹、自发流产、过敏反应）vs 安慰剂1例",
            ),
            ("h", "【死亡】(AE)"),
            ("i", "疫苗组3例 vs 安慰剂组13例（3例COVID-19相关, 均在安慰剂组）"),
        ],
        "Bravo L, et al. Lancet. 2022;399:461-472\nPMID: 35065705 | DOI: 10.1016/S0140-6736(22)00055-1\nhttps://pmc.ncbi.nlm.nih.gov/articles/PMC8776284/",
    ),
    (
        "BK-SE36/CpG\n佐剂: CpG ODN K3 + Alum\n预防适应症: 疟疾（恶性疟原虫）",
        "国际（非FDA注册）\n（ClinicalTrials.gov可查）",
        "试验分期: Ⅰb期\n整体设计: 单中心、双盲、随机、对照、年龄降阶梯\n样本量: 135（91疫苗/44对照）\n试验分组: BK-SE36/CpG 3剂(0,28,112天) vs 狂犬疫苗对照; 队列21-45岁/5-10岁/12-24月龄\n研究终点: 主-安全性与反应原性; 次-免疫原性",
        [
            ("h", "【局部AE】(AE)"),
            ("i", "46% vs 18%; 疼痛/活动受限（成人17-33%、儿童40-57%、婴幼儿0-19%）"),
            ("h", "【全身AE】(AE)"),
            ("i", "发热（婴幼儿第1剂29%）; 头痛（成人第3剂20%）; AEFI 56% vs 19%"),
            ("h", "【SAE】(AE)"),
            ("i", "5例（全部重症疟疾, 均判定与疫苗无关）; SUSAR 0例; 死亡0例"),
            ("h", "【≥3级事件/AESI】(AE)"),
            (
                "i",
                "6例（4人）均判定不相关; 自身免疫标志物（ANA/anti-dsDNA/ANCA）无临床异常",
            ),
        ],
        "Safety and immunogenicity of BK-SE36/CpG… Front Immunol. 2023;14:1267372\nPMID: 37638048 | DOI: 10.3389/fimmu.2023.1267372\nhttps://www.frontiersin.org/journals/immunology/articles/10.3389/fimmu.2023.1267372/full",
    ),
    (
        "Na-GST-1/Alhydrogel ± CpG 10104\n佐剂: CpG 10104（TLR9激动剂）\n预防适应症: 钩虫病（美洲钩虫）",
        "FDA（ClinicalTrials.gov）\nNCT03172975（CHHI）",
        "试验分期: Ⅰ期\n整体设计: 随机、观察者盲、对照、剂量递增、单中心\n样本量: 24（CpG组16/无CpG组8）\n试验分组: 30μg+500μg CpG / 100μg / 100μg+500μg CpG 各8例; 3剂(0,56,112天)\n研究终点: 主-安全性/耐受性; 次-抗原特异性IgG",
        [
            ("h", "【局部】(AE)"),
            (
                "i",
                "注射部位疼痛: 30μg+CpG 7/8、100μg无CpG 7/8、100μg+CpG 9例次; 压痛最常见; 重度仅2例（压痛1例CpG组、肿胀1例无CpG组）",
            ),
            ("h", "【全身】(AE)"),
            (
                "i",
                "头痛最常见（25-50%）; 肌痛 62.5%（100μg+CpG组, 其中中度4例）; 发热 0-25%",
            ),
            ("h", "【SAE】(AE)"),
            ("i", "0例; 无疫苗相关SAE"),
            ("h", "【AESI】(AE)"),
            ("i", "0例; 无自身免疫症状; 1例ANA升高停药（揭盲后为无CpG组）"),
        ],
        "Diemert DJ, et al. PLoS Negl Trop Dis. 2024;18(12):e0012788\nPMID: 39775205 | DOI: 10.1371/journal.pntd.0012788\nhttps://pmc.ncbi.nlm.nih.gov/articles/PMC11717351",
    ),
    (
        "rF1V + CpG 1018\n佐剂: CpG 1018（TLR9激动剂）\n预防适应症: 鼠疫（肺鼠疫）",
        "FDA（ClinicalTrials.gov）\nNCT05506969",
        "试验分期: Ⅱ期（已完成）\n整体设计: 随机、活性对照、观察者盲、多中心\n样本量: 200\n试验分组: rF1V-1018 2剂(0,4周) vs rF1V单疫苗 3剂\n研究终点: 主-反应原性与安全性(D1-W56); 结果未发布",
        [("b", "未公开（试验已完成, 结果未发布）")],
        "Dynavax/DoD rF1V+CpG 1018 Phase 2 注册信息\nhttps://clinicaltrials.gov/study/NCT05506969",
    ),
    (
        "H5N1流感疫苗 + CpG 1018\n佐剂: CpG 1018（TLR9激动剂）\n预防适应症: 流感（H5N1大流行）",
        "FDA（ClinicalTrials.gov）\n注册号未公开",
        "试验分期: Ⅰ/Ⅱ期（Part 1完成）\n整体设计: 随机、活性对照\n样本量: 101（18-49岁）\n试验分组: 单剂/双剂多种CpG 1018配方\n研究终点: 主-安全性与免疫原性; Part 2数据预计2026",
        [("b", "未公开（Part 1完成; Part 2数据预计2026）")],
        "Dynavax Pandemic Influenza Adjuvant Program（公司公告2025）\nhttps://www.nasdaq.com/press-release/dynavax-reports-third-quarter-2025-financial-results-and-announces-new-100-million",
    ),
    (
        "Z-1018 带状疱疹疫苗\n佐剂: CpG 1018（±Alum, 6000μg/剂）\n预防适应症: 带状疱疹（VZV）",
        "FDA（ClinicalTrials.gov）\n注册号未公开（IDWeek摘要已发表）",
        "试验分期: Ⅰ/Ⅱ期（中期数据已发表）\n整体设计: 随机、观察者盲、Shingrix对照（10:1）\n样本量: 441（中位年龄58.0岁, 68.3%女性）\n试验分组: Z-1018（gE 50/100/200μg + CpG 1018 6000μg ± Alum 750μg, 8或12周间隔）vs Shingrix\n研究终点: 主-安全性（PIR 7天内）与免疫原性（anti-gE IgG VRR）",
        [
            ("h", "【中重度接种后反应PIR(7天内)】(AE)"),
            ("i", "局部: 7.7-35.0%（Z-1018各剂量组）vs 52.6%（Shingrix）"),
            ("i", "全身: 17.5-46.2% vs 63.2%（Shingrix）"),
            ("h", "【总体安全性】(AE)"),
            ("i", "无安全性顾虑报告; 中期数据支持继续开发"),
            ("h", "【SAE】(AE)"),
            ("i", "摘要未披露具体例数"),
        ],
        "Medzihradsky O, et al. Open Forum Infect Dis. 2026;13(Suppl 1):ofaf695.018\nDOI: 10.1093/ofid/ofaf695.018 | PMCID: PMC12792159\nhttps://pmc.ncbi.nlm.nih.gov/articles/PMC12792159/",
    ),
    # ================= 国内(已发表) =================
    (
        "重组乙型肝炎疫苗（汉逊酵母）\n佐剂: CpG ODN（250μg）\n预防适应症: 乙型肝炎（HBV）",
        "无NCT/CTR（2016年开展, 平台建立前）\n溯源: PMID 32842315",
        "试验分期: Ⅰ期（2020年发表）\n整体设计: 随机、双盲、对照、单中心（浙江开化）\n样本量: 48（24/24）\n试验分组: CpG ODN 250μg试验组 vs 市售铝佐剂对照（大连汉信）; 0-1-6月\n研究终点: 主-安全性; 次-anti-HBs阳转率/GMC",
        [
            ("h", "【总体AE】(AE)"),
            (
                "i",
                "66.67%（16/24）vs 54.17%（13/24）, P=0.556; 全部1-2级, 无≥3级（原文未细分局部/全身）",
            ),
            ("h", "【SAE】(AE)"),
            ("i", "未报告"),
            ("h", "【AESI】(AE)"),
            ("i", "无免疫介导AE报道"),
        ],
        "梁贞贞, 邵彦, 王慎玉, 等. 中华预防医学杂志. 2020;54(8):854-860\nPMID: 32842315 | DOI: 10.3760/cma.j.cn112150-20200401-00490\nhttps://pubmed.ncbi.nlm.nih.gov/32842315/",
    ),
    (
        "SCB-2019（S-Trimer）中国Ⅱ期\n佐剂: CpG 1018 + Alum\n预防适应症: COVID-19（SARS-CoV-2）",
        "FDA（ClinicalTrials.gov）\nNCT04954131",
        "试验分期: Ⅱ期\n整体设计: 双盲、随机、安慰剂对照（山西太原/四川成都CDC）\n样本量: 766\n试验分组: SCB-2019（30μg抗原+1.5mg CpG 1018+0.75mg Alum）2剂 vs 安慰剂\n研究终点: 主-免疫原性; 次-安全性",
        [
            (
                "b",
                "中国Ⅱ期(n=766)安全性数值未单独发表; 下表引自同产品核心试验SPECTRA全球Ⅲ期(含中国受试者)",
            ),
            ("h", "【征集性局部反应】(AE)"),
            ("i", "第1剂后 36.1% vs 11.3%; 注射部位疼痛 35.7% vs 10.3%"),
            ("h", "【征集性全身反应】(AE)"),
            ("i", "第1剂后 36% vs 34%"),
            ("h", "【SAE】(AE)"),
            ("i", "0.3% vs 0.4%（与安慰剂无差异）"),
        ],
        "Bravo L, et al. Lancet. 2022;399:461-472（SPECTRA, 含中国受试者）\nPMID: 35065705 | DOI: 10.1016/S0140-6736(22)00055-1\nhttps://pmc.ncbi.nlm.nih.gov/articles/PMC8776284/",
    ),
    (
        "ZR202-CoV\n佐剂: CpG 7909 + Alum\n预防适应症: COVID-19（SARS-CoV-2）",
        "FDA（ClinicalTrials.gov）\nNCT04982068",
        "试验分期: Ⅰ期（2023年发表）\n整体设计: 随机、双盲、安慰剂对照、剂量递增（河南许昌）\n样本量: 72（18/18/18/18）\n试验分组: A 50μg+250μg CpG / B 25μg+500μg CpG / C 50μg+500μg CpG / 安慰剂; 2剂(0,28天)\n研究终点: 主-安全性/耐受性; 次-免疫原性",
        [
            ("h", "【征集性局部症状】(AE, 均为1-2级)"),
            (
                "i",
                "任一局部症状: A 5.6%(1/18)、B 22.2%(4/18)、C 38.9%(7/18)、安慰剂5.6%(1/18)",
            ),
            ("i", "疼痛 5.6%/22.2%/27.8%/5.6%; 肿胀 C组16.7%(3/18)"),
            ("h", "【征集性全身症状】(AE)"),
            (
                "i",
                "任一全身症状: A 11.1%、B 11.1%、C 5.6%、安慰剂5.6%; 发热(腋温≥37.3℃)同比例; 恶心 A组5.6%(1/18)",
            ),
            ("h", "【SAE】(AE)"),
            ("i", "1例（A组, 输尿管结石, 判定与疫苗无关, 5.6%）"),
            ("h", "【非征集性AE≥3级/AESI】(AE)"),
            (
                "i",
                "C组1例(尿潜血, 月经相关)、安慰剂2例; 无AESI; 纤维蛋白原升高 22.2%/22.2%/17.3%/5.6%",
            ),
        ],
        "Feng GW, et al. Hum Vaccin Immunother. 2023;19(3):2262635\nPMID: 37881130 | DOI: 10.1080/21645515.2023.2262635\nhttps://pmc.ncbi.nlm.nih.gov/articles/PMC10644802/",
    ),
    # ================= 国内(进行中, 数据未公开) =================
    (
        "重组带状疱疹疫苗（gE）\n佐剂: CpG 1018 + MF59\n预防适应症: 带状疱疹（VZV）",
        "FDA（ClinicalTrials.gov）\nNCT06851832",
        "试验分期: Ⅰ/Ⅱ期（进行中, 2026-09主要完成）\n整体设计: 随机、双盲、活性/安慰剂对照、多中心\n样本量: 960\n试验分组: 低剂量gE50μg+MF59+CpG 50μg / 标准gE50μg+MF59+CpG 100μg / 减毒活疫苗 / Shingrix配方 / 佐剂对照 / 安慰剂\n研究终点: 主-各剂0-14天征集性AE、0-30天非征集性AE、SAE/AESI",
        [("b", "未公开（进行中）")],
        "重组带状疱疹疫苗(gE+CpG1018+MF59) Ⅰ/Ⅱ期注册信息\nhttps://clinicaltrials.gov/study/NCT06851832",
    ),
    (
        "华普 HP2001 带状疱疹疫苗\n佐剂: HP007（CpG类）\n预防适应症: 带状疱疹（VZV）",
        "NMPA（CDE受理号）\nCXSL2500371",
        "试验分期: Ⅰ期（2026-03启动）\n整体设计: 双盲、2剂、40岁+\n样本量: 待披露\n试验分组: HP007佐剂带状疱疹疫苗 2剂\n研究终点: 安全性/免疫原性",
        [("b", "未公开（进行中）")],
        "华普生物 HP2001 带状疱疹疫苗Ⅰ期受理信息(CXSL2500371)\nhttps://parrbio.com/newsinfo/11232956.html",
    ),
    (
        "华普 HP2002 乙肝疫苗\n佐剂: CpG + Alum\n预防适应症: 乙型肝炎（HBV）",
        "NMPA（Chinadrugtrials.org.cn）\nCTR20261430",
        "试验分期: Ⅰ期（2026-04启动）\n整体设计: 随机、双盲、平行对照、单中心\n样本量: 待披露\n试验分组: 不同CpG配比浓度, 0-1-6月3剂, 18岁+\n研究终点: 主-安全性/耐受性; 探索-免疫原性",
        [("b", "未公开（进行中）")],
        "华普 HP2002 乙肝疫苗Ⅰ期登记信息(CTR20261430)\nhttps://www.chinadrugtrials.org.cn",
    ),
    (
        "重组乙型肝炎疫苗（汉逊酵母）\n佐剂: CpG + Alum（CpG-QCX1）\n预防适应症: 乙型肝炎（HBV）",
        "NMPA（Chinadrugtrials.org.cn）\nCTR20233405",
        "试验分期: Ⅰ期（2023-11公示, 招募完成）\n整体设计: 单中心、随机、双盲、对照、平行分组\n样本量: 120（目标）\n试验分组: 不同CpG-QCX1佐剂配比浓度, 0-1-6月3剂, 18岁+\n研究终点: 主-安全性和耐受性; 次-免疫原性",
        [("b", "未公开（招募完成, 数据未发表）")],
        "重组乙型肝炎疫苗(汉逊酵母,CpG和铝佐剂)Ⅰ期登记信息(CTR20233405)\nhttps://www.chinadrugtrials.org.cn",
    ),
    (
        "远大乙肝疫苗\n佐剂: CpG + Alum\n预防适应症: 乙型肝炎（HBV）",
        "中国临床试验注册中心（ChiCTR）\nChiCTR2500108408",
        "试验分期: Ⅱ期（进行中）\n整体设计: 随机、双盲、对照\n样本量: 待披露\n试验分组: CpG+铝佐剂乙肝疫苗\n研究终点: 安全性/免疫原性",
        [("b", "未公开（进行中）")],
        "远大乙肝疫苗Ⅱ期注册信息(ChiCTR2500108408)\nhttps://www.chictr.org.cn",
    ),
    (
        "怡道/中慧元通 带状疱疹疫苗\n佐剂: CpG类\n预防适应症: 带状疱疹（VZV）",
        "NMPA（Chinadrugtrials.org.cn）\nCTR20233253",
        "试验分期: Ⅲ期（2023-11启动, 已报NDA）\n整体设计: 随机、双盲、对照（40岁+）\n样本量: 待披露\n试验分组: CpG佐剂带状疱疹疫苗 vs 对照\n研究终点: 主-保护效力; 次-安全性",
        [("b", "未公开（已报NDA, 数据未发表）")],
        "怡道/中慧元通带状疱疹疫苗Ⅲ期登记信息(CTR20233253)\nhttps://www.chinadrugtrials.org.cn",
    ),
    (
        "明瑞佳 MRJ103 带状疱疹疫苗\n佐剂: CpG 1018 + QS-21 + 脂质体\n预防适应症: 带状疱疹（VZV）",
        "NMPA（CDE受理号）\nCXSL2600443",
        "试验分期: 临床批件（IND获批2026）\n整体设计: 待启动\n样本量: 待披露\n试验分组: MRJ103复合佐剂带状疱疹疫苗\n研究终点: 安全性/免疫原性",
        [("b", "未公开（IND获批, 试验未启动）")],
        "明瑞佳带状疱疹疫苗临床批件(CXSL2600443)\nhttps://www.chinadrugtrials.org.cn",
    ),
]

# 数据缺口(不列入表格, 文末带过)
GAPS = (
    "HPV疫苗: 仅见CpG佐剂治疗性(肿瘤)HPV疫苗临床前/临床研究, 无CpG佐剂预防性HPV疫苗人体试验, 按规则不列入; "
    "吉诺卫、昆明简达(带状疱疹/流感)、远大带状疱疹(TVA01)、长春卓谊(CpG狂犬疫苗, 已获临床批件、Ⅰ期待启动): "
    "临床试验信息未公开或试验未启动, 按规则不列入; "
    "昆明简达CpG复合佐剂流感疫苗已于2026-07获FDA临床许可(临床试验信息尚未公示)。"
)

CONCLUSIONS = [
    "1. 已上市产品(Heplisav-B、AV7909)及大型Ⅲ期试验(SPECTRA, n=30,128)显示CpG佐剂疫苗总体安全性可控: SPECTRA疫苗组SAE 0.3% vs 安慰剂0.4%, 治疗相关SAE 4例 vs 1例;",
    "2. 国内已发表Ⅰ期数据(乙肝48例、ZR202-CoV 72例): 不良事件以1-2级为主, 无疫苗相关≥3级AE或SAE;",
    "3. AESI无异常信号: Heplisav-B三试验汇总免疫介导AE 0.32% vs 铝佐剂0.38%; 疟疾/钩虫疫苗均无自身免疫事件;",
    "4. 局部注射部位疼痛为最常见反应: SPECTRA 35.7% vs 10.3%; Z-1018中重度局部PIR 7.7-35% vs Shingrix 52.6%, 多轻中度、一过性;",
    "5. 国内带状疱疹/流感/狂犬等CpG佐剂疫苗管线尚在推进中, 安全性数据未公开, 待后续补充。",
]


# ============================================================ helpers ====
def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def set_font(run, size=9, bold=False, name_ea="宋体", name_ascii="Times New Roman"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_ascii
    r = run._element.rPr
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_ea)


def add_para(
    doc,
    text,
    size=10.5,
    bold=False,
    align=None,
    space_after=6,
    ea="宋体",
    ascii_f="Times New Roman",
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, name_ea=ea, name_ascii=ascii_f)
    return p


def fill_cell(cell, text, size=8.5, bold_first=False, width=None):
    """多行文本填入单元格, 按\n分段"""
    if width is not None:
        cell.width = Cm(width)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        set_font(run, size=size, bold=(bold_first and i == 0))


def fill_safety_cell(cell, lines, size=8):
    """安全性数据列: 层次化排版. ('h')分类标题加粗, ('i')条目缩进, ('b')加粗说明段"""
    for i, (kind, text) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        if kind == "h":
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(text)
            set_font(run, size=size, bold=True, name_ea="黑体", name_ascii="Arial")
        elif kind == "i":
            p.paragraph_format.left_indent = Pt(7)
            run = p.add_run("· " + text)
            set_font(run, size=size, bold=False)
        else:  # 'b'
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.left_indent = Pt(3)
            run = p.add_run(text)
            set_font(run, size=size, bold=True, name_ea="黑体", name_ascii="Arial")


def add_block_row(table, text, ncols, fill="BDD7EE"):
    row = table.add_row()
    row.cells[0].merge(row.cells[ncols - 1])
    c = row.cells[0]
    p = c.paragraphs[0]
    run = p.add_run(text)
    set_font(run, size=9.5, bold=True, name_ea="黑体", name_ascii="Arial")
    set_cell_shading(c, fill)


# =============================================================== build ===
doc = Document()
sec = doc.sections[0]
sec.orientation = 1  # 横向
sec.page_width = Cm(29.7)
sec.page_height = Cm(21.0)
sec.left_margin = Cm(1.5)
sec.right_margin = Cm(1.5)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)

# 标题
add_para(
    doc,
    "国内外CpG佐剂预防性疫苗安全性数据汇总",
    size=16,
    bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
add_para(
    doc,
    "检索日期: 2026-08-20 | 数据截止: 2026-08-20 | 数据源: ClinicalTrials.gov / 中国药物临床试验登记与信息公示平台(Chinadrugtrials.org.cn) / ChiCTR / PubMed / 期刊原文",
    size=9,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=8,
    ea="宋体",
)

# 说明段
add_para(
    doc,
    "本表汇总国内外使用CpG类佐剂(含CpG 1018、CpG 7909、CpG ODN K3、CpG 10104、HP007等)的预防性疫苗临床试验安全性数据，"
    "覆盖乙型肝炎、炭疽、COVID-19、疟疾、钩虫、鼠疫、流感、带状疱疹等适应症；治疗性肿瘤疫苗不在范围内。"
    "安全性数据均来自文献原文的仔细阅读与提取，逐项标注 (AE)（不良事件）或 (ADR)（不良反应）：各试验原文以AE报告为主，"
    "未见以ADR单独汇总报道的数据，故本表数据性质均为AE。"
    "征集性=Solicited（预设收集）、非征集性=Unsolicited、PIR=接种后反应(Post-injection Reactions)、AESI=特殊关注不良事件。"
    "数据未公开的试验（进行中）在表中明确标注。",
    size=10.5,
    space_after=10,
)

# 总表
HEADERS = [
    "疫苗名称 & 适应症",
    "注册平台 & 编号",
    "临床试验基本信息",
    "安全性数据汇总（逐项标注 AE/ADR）",
    "核心参考文献",
]
WIDTHS = [3.7, 2.9, 6.3, 9.6, 4.2]  # 合计 26.7cm

t = doc.add_table(rows=1, cols=5)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False

hdr = t.rows[0]
for j, h in enumerate(HEADERS):
    c = hdr.cells[j]
    c.width = Cm(WIDTHS[j])
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, size=9, bold=True, name_ea="黑体", name_ascii="Arial")
    set_cell_shading(c, "D9E2F3")

BLOCKS = [
    ("◆ 国外部分（已上市 / 已发表）", 0, 8),
    ("◆ 国内部分（已发表安全性数据）", 8, 3),
    ("◆ 国内部分（临床试验信息可查, 安全性数据未公开）", 11, 7),
]
for block_text, start, n in BLOCKS:
    add_block_row(t, block_text, 5)
    for row in ROWS[start : start + n]:
        cells = t.add_row().cells
        fill_cell(cells[0], row[0], size=8.5, bold_first=True, width=WIDTHS[0])
        fill_cell(cells[1], row[1], size=8, width=WIDTHS[1])
        fill_cell(cells[2], row[2], size=8, width=WIDTHS[2])
        fill_safety_cell(cells[3], row[3], size=8)
        fill_cell(cells[4], row[4], size=7.5, width=WIDTHS[4])

doc.add_paragraph()

# 数据缺口
add_para(
    doc,
    "数据缺口与排除说明",
    size=13,
    bold=True,
    space_after=4,
    ea="黑体",
    ascii_f="Arial",
)
add_para(doc, GAPS, size=9.5, space_after=6)

# 关键安全性结论
add_para(
    doc, "关键安全性结论", size=13, bold=True, space_after=4, ea="黑体", ascii_f="Arial"
)
for s in CONCLUSIONS:
    add_para(doc, s, size=10, space_after=2)

doc.save(OUT_PATH)
print("SAVED:", OUT_PATH)
