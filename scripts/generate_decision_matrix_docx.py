import os
import subprocess
import sys


def install_and_import(package):
    try:
        import docx
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
    finally:
        globals()["docx"] = __import__("docx")


install_and_import("python-docx")

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_docx():
    import docx

    doc = docx.Document()

    # Title
    title = doc.add_heading("防火墙团队 — 交叉验证决策表 (Decision Matrix)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Core principles
    doc.add_heading("核心原则", level=2)
    p = doc.add_paragraph()
    p.add_run("当 EAC 判定结果与 qPCR 结果不一致时，").bold = True
    p.add_run("以 qPCR 结果为病例判定的依据。").bold = True
    p.add_run("\n仅当无皮损样本或 qPCR 结果无法判定时，才采用 EAC 的临床诊断结论。")

    # Phase 1
    doc.add_heading("第一阶段：盲评与实验室结果交叉验证", level=2)
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = "Table Grid"
    hdr_cells1 = table1.rows[0].cells
    hdr_cells1[0].text = "qPCR 检测结果"
    hdr_cells1[1].text = "EAC 第一阶段盲评结果"
    hdr_cells1[2].text = "防火墙团队最终判定结论"
    hdr_cells1[3].text = "判定依据/说明"

    for cell in hdr_cells1:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    data1 = [
        (
            "阳性 (≥1个样本阳性)",
            "是 (符合带状疱疹)",
            "✅ 实验室确诊病例",
            "qPCR 阳性为金标准，直接确诊",
        ),
        ("阳性", "否 (不符合)", "✅ 实验室确诊病例", "qPCR 结果优先于 EAC 判定"),
        (
            "真阴性 (内参+，靶基因-)",
            "是 (符合带状疱疹)",
            "❌ 非带状疱疹排除",
            "qPCR 结果优先于 EAC 判定",
        ),
        ("真阴性", "否 (不符合)", "❌ 非带状疱疹排除", "实验室与临床意见一致排除"),
        (
            "无法判定 / 无样本 (内参-且靶基因-)",
            "否 (不符合)",
            "❌ 非带状疱疹排除",
            "缺乏实验室依据且临床不支持",
        ),
        (
            "无法判定 / 无样本",
            "是 (符合带状疱疹)",
            "🔄 进入第二阶段复判",
            "实验室无结论但临床支持，防火墙团队汇总该情况交由 EAC 重新审视",
        ),
    ]

    for item in data1:
        row_cells = table1.add_row().cells
        for i in range(4):
            row_cells[i].text = item[i]

    doc.add_paragraph()

    # Phase 2
    doc.add_heading("第二阶段：EAC 复判（仅针对第一阶段最后一种情况）", level=2)
    doc.add_paragraph(
        "防火墙团队将“qPCR无法判定/无结果”的信息汇总给 EAC，由 EAC 结合该客观信息对病例进行第二阶段复判："
    )

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = "EAC 第二阶段复判结果"
    hdr_cells2[1].text = "最终判定结论"
    hdr_cells2[2].text = "归类说明"

    for cell in hdr_cells2:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    data2 = [
        (
            "仍符合带状疱疹",
            "🟣 临床诊断确认病例",
            "最终确诊（但属于临床确诊，非实验室确诊）",
        ),
        ("改判为 不符合", "❌ 非带状疱疹排除", "被排除的疑似病例"),
        ("无法判定", "⬜ 不能确诊", "证据均不足，列为不能确诊病例"),
    ]

    for item in data2:
        row_cells = table2.add_row().cells
        for i in range(3):
            row_cells[i].text = item[i]

    doc.add_paragraph()

    # Final stats
    doc.add_heading("💡 最终确诊病例统计口径", level=2)
    p2 = doc.add_paragraph()
    run = p2.add_run("带状疱疹确诊病例总数 = 实验室确诊病例 + 临床诊断确认病例")
    run.bold = True
    run.font.size = Pt(12)

    # Save document
    out_dir = r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "hz_decision_matrix.docx")

    doc.save(out_path)
    print(f"Word document successfully saved to: {out_path}")


if __name__ == "__main__":
    create_docx()
