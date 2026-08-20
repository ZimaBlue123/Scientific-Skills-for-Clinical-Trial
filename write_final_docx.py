import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

doc = docx.Document()
# Change orientation to Landscape for a wide table
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

title = doc.add_heading("CpG佐剂预防性疫苗核心临床试验与安全性数据汇总", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(
    "本文件针对 CpG 佐剂预防性疫苗，分别检索了 FDA (ClinicalTrials.gov) 与 NMPA/ChiCTR 的临床试验数据。"
)

# Add table
table = doc.add_table(rows=1, cols=7)
table.style = "Table Grid"

# Header
hdr_cells = table.rows[0].cells
headers = [
    "疫苗名称 & 疾病",
    "注册平台\n& 编号",
    "临床试验基本信息\n(分期/设计/样本量/分组/终点)",
    "局部反应 (Local AEs)",
    "全身反应 (Systemic AEs)",
    "SAE / SADR / AESI",
    "核心参考文献\n(DOI/PMID)",
]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    {
        "vaccine": "SCB-2019\n(重组SARS-CoV-2三聚体蛋白疫苗)",
        "registry": "FDA\nNCT04672395",
        "clinical": "分期: Phase 2/3\n整体设计: 跨国、多中心、随机、双盲、安慰剂对照\n样本量: 30,137 人\n试验分组: SCB-2019组 (n=15,070), 安慰剂组 (n=15,067)\n研究终点: COVID-19保护效力、6个月长期安全性",
        "local": "常见反应: 注射部位疼痛、红斑(Erythema)。\n数据: 发生率与同类重组蛋白疫苗一致，多数为轻度(Mild)。",
        "systemic": "常见反应: 疲劳(Fatigue)、头痛(Headache)、肌痛(Myalgia)。\n数据: 绝大多数为轻至中度，1-2天内缓解。",
        "sae": "SAE: 疫苗相关SAE极罕见，发生率为 0.027% (4/15,070)。包括超敏反应(Hypersensitivity)2例，贝尔氏麻痹(Bell"
        "s palsy)1例，自发性流产1例。\nSADR: 未发生致死性不良反应。\nAESI: 无疫苗相关增强疾病(VAED)，AESI总体发生率与安慰剂组相似。",
        "ref": "DOI: 10.1016/j.vaccine.2023.02.018\nPMID: 36868877",
    },
    {
        "vaccine": "HEPLISAV-B / HepB-CpG\n(乙型肝炎疫苗)",
        "registry": "FDA\nNCT01282762\n(ACTG A5379)",
        "clinical": "分期: Phase 3b (长效追踪研究)\n整体设计: 多中心、开放标签\n样本量: 147人(CKD患者) & 68人(HIV患者)\n试验分组: HepB-CpG组 vs Engerix-B组\n研究终点: 血清保护率持久性及安全性评价",
        "local": "常见反应: 注射部位疼痛、触痛。\n数据: 局部反应特征与传统铝佐剂乙肝疫苗高度相似。",
        "systemic": "常见反应: 疲劳、头痛、轻度发热。\n数据: 在HIV及CKD等免疫低下/特殊人群中未发现额外的系统性风险，未出现预期外的全身性反应。",
        "sae": "SAE/SADR: 长期随访中未见预期外SAE，无疫苗相关致死病例。\nAESI (自身免疫风险专题): 潜在免疫介导不良事件(Autoimmune AEs)发生率为 0.32%，与传统铝佐剂组(0.38%)无显著差异，未增加自身免疫疾病风险。",
        "ref": "DOI: 10.1016/j.vaccine.2023.04.028\nPMID: 37085451\nPMID: 41819640 (AESI综述)",
    },
    {
        "vaccine": "ZR202-CoV\n(重组新冠蛋白疫苗)",
        "registry": "NMPA (中国)\nChiCTR2200057758\n(关联NCT04990544)",
        "clinical": "分期: Phase 1/2\n整体设计: 随机、双盲、安慰剂对照、剂量递增\n样本量: 72人 (Phase 1) / 1056人 (Phase 2)\n试验分组: 不同剂量ZR202-CoV组 vs 生理盐水组\n研究终点: 28天内的耐受性及免疫原性",
        "local": "常见反应: 轻度至中度的注射部位疼痛。\n数据: 接种后7天内未发生 ≥3 级的局部反应 (0%)。",
        "systemic": "常见反应: 轻度发热、乏力。\n数据: 未观察到 ≥3 级的非征集性全身不良反应，耐受性良好。",
        "sae": "SAE: 试验期间未发生任何与疫苗相关的严重不良事件 (0%)。\nSADR: 无严重不良反应。\nAESI: 未观察到过敏性休克等特殊关注事件。",
        "ref": "DOI: 10.1080/21645515.2023.2262635\nPMID: 37881130",
    },
    {
        "vaccine": "IndoVac\n(重组SARS-CoV-2亚单位疫苗)",
        "registry": "FDA\nNCT05433285",
        "clinical": "分期: Phase 3\n整体设计: 随机、活性药物对照(Covovax)、多中心\n样本量: 4,050 人\n试验分组: IndoVac组 vs Covovax组\n研究终点: 免疫原性(非劣效性)及安全性",
        "local": "常见反应: 疼痛 (Pain)。\n数据: 疼痛发生率为 14.69%，红肿罕见，大多数为轻度。",
        "systemic": "常见反应: 肌痛(Myalgia)、疲劳(Fatigue)。\n数据: 全身性AE发生率为 27.95% (主要为轻度)。肌痛占比 7.48%，疲劳占比 6.77%。",
        "sae": "SAE: 评估期内未发现可能或极可能与疫苗相关的 SAE。\nSADR/AESI: 整体安全性特征优良，严重风险率极低。",
        "ref": "DOI: 10.1016/j.vaccine.2024.03.077\nPMID: 38575433",
    },
    {
        "vaccine": "Na-GST-1/Al+CpG\n(钩虫病重组疫苗)",
        "registry": "FDA\nNCT03172975",
        "clinical": "分期: Phase 2\n整体设计: 随机、双盲、安慰剂对照、人类受控感染模型(CHHI)\n样本量: 39 人\n试验分组: 疫苗+CpG组、疫苗+AP组、安慰剂组\n研究终点: 感染保护率(效力)及安全性评价",
        "local": "常见反应: 注射部位触痛。\n数据: 绝大部分表现为轻度(Mild)反应。",
        "systemic": "数据: 显著降低了后续寄生虫感染时的最大嗜酸性粒细胞计数。全身性不良反应发生率与安慰剂相似。",
        "sae": "SAE: 整个试验周期内未观察到与疫苗相关的 SAE。\nSADR: 零报告。\nAESI: 寄生虫感染模型的特殊免疫应答监控均在安全阈值内。",
        "ref": "DOI: 10.1016/S1473-3099(26)00018-6\nPMID: 41861834",
    },
]

for row_data in data:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data["vaccine"]
    row_cells[1].text = row_data["registry"]
    row_cells[2].text = row_data["clinical"]
    row_cells[3].text = row_data["local"]
    row_cells[4].text = row_data["systemic"]
    row_cells[5].text = row_data["sae"]
    row_cells[6].text = row_data["ref"]

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(9)

doc.save(
    r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\CpG_Vaccine_Safety_Summary.docx"
)
