import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

doc = docx.Document()
# Change orientation to Landscape for a wide table
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

title = doc.add_heading("CpG佐剂预防性疫苗临床试验与安全性数据汇总", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph(
    "本表格汇总了在 FDA (ClinicalTrials.gov) 与 NMPA/ChiCTR (中国药物临床试验登记平台) 注册的 CpG 佐剂（TLR9 激动剂）预防性疫苗的核心临床试验信息及详细安全性数据。"
)

# Add table
table = doc.add_table(rows=1, cols=7)
table.style = "Table Grid"

# Header
hdr_cells = table.rows[0].cells
headers = [
    "疫苗名称\n& 疾病",
    "注册平台\n& 编号",
    "临床试验基本信息\n(分期/设计/样本量/分组/终点)",
    "局部反应\n(接种部位)",
    "全身反应\n(Systemic AEs)",
    "严重不良事件与特殊关注\n(SAE / SADR / AESI)",
    "核心参考文献\n(PMID / DOI / URL)",
]
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.font.bold = True

data = [
    {
        "vaccine": "SCB-2019\n(重组新型冠状病毒三聚体疫苗)",
        "registry": "FDA\nNCT04672395",
        "clinical": "分期: Phase 2/3\n设计: 跨国、多中心、随机、双盲、安慰剂对照\n样本量: 30,137 人\n分组: SCB-2019组 (n=15,070), 安慰剂组 (n=15,067)\n终点: 预防COVID-19的保护效力及6个月长期安全性",
        "local": "常见，主要为轻至中度的注射部位疼痛、红斑。发生率与同类疫苗一致。",
        "systemic": "常见不良反应 (AE) 包括疲劳、头痛和肌痛。",
        "sae": "在15,070名接种者中，非征集性AE和SAE发生率与安慰剂组相似。\n仅4例判定为疫苗相关SAE（超敏反应2例，贝尔氏麻痹1例，自发性流产1例）。\n无疫苗相关增强疾病(VAED)，AESI发生率与安慰剂无显著差异。",
        "ref": "PMID: 36868877\nDOI: 10.1016/j.vaccine.2023.02.018",
    },
    {
        "vaccine": "IndoVac\n(重组SARS-CoV-2蛋白亚单位疫苗)",
        "registry": "FDA\nNCT05433285",
        "clinical": "分期: Phase 3\n设计: 随机、阳性药物对照、多中心\n样本量: 4,050 人\n分组: IndoVac组 vs Covovax组\n终点: 免疫原性(非劣效性)及安全性评价",
        "local": "注射部位疼痛发生率为 14.69%（多数为轻度）。",
        "systemic": "总体不良事件发生率为 27.95% (主要为轻度)。\n全身反应常见肌痛 (7.48%) 和疲劳 (6.77%)。",
        "sae": "未发现可能或极可能与疫苗相关的严重不良事件 (SAE) 或 SADR。\n整体安全性特征良好。",
        "ref": "PMID: 38575433\nDOI: 10.1016/j.vaccine.2024.03.077",
    },
    {
        "vaccine": "ZR202-CoV\n(重组新冠蛋白疫苗)",
        "registry": "NMPA (ChiCTR)\nChiCTR2200057758\n(早期部分数据)",
        "clinical": "分期: Phase 1/2\n设计: 随机、双盲、安慰剂对照、剂量递增试验\n样本量: 72人(Phase1) / 后续扩展\n分组: 不同剂型组 vs 生理盐水组\n终点: 接种后28天的安全性及免疫原性",
        "local": "耐受性良好。轻中度注射部位疼痛。接种后7天内未发生 ≥3 级的局部反应。",
        "systemic": "未报告严重的全身性反应。",
        "sae": "未发生任何疫苗相关的严重不良事件 (SAE)。\n无 ≥3 级的非征集性不良事件。无严重不良反应 (SADR)。",
        "ref": "PMID: 37881130\nDOI: 10.1080/21645515.2023.2262635",
    },
    {
        "vaccine": "HEPLISAV-B / HepB-CpG\n(乙型肝炎疫苗)",
        "registry": "FDA\nNCT01282762",
        "clinical": "分期: Phase 3b (长效追踪)\n设计: 多中心、开放标签\n样本量: 147人 (针对慢性肾病患者等脆弱人群)\n分组: HepB-CpG组 vs Engerix-B组\n终点: 血清保护率的持久性及安全性",
        "local": "与传统铝佐剂乙肝疫苗一致，主要为轻微疼痛。",
        "systemic": "与传统疫苗相似。未在HIV人群(A5379试验)中观察到额外系统性副作用。",
        "sae": "专题研究 (PMID: 41819640) 表明：\n潜在免疫介导不良事件(AESI)发生率为 0.32% (与铝佐剂的 0.38% 相当)。\n未增加自身免疫疾病风险，无超预期的 SAE/SADR。",
        "ref": "PMID: 37085451\nDOI: 10.1016/j.vaccine.2023.04.028\nPMID: 41819640",
    },
    {
        "vaccine": "Na-GST-1/Al+CpG\n(钩虫病重组疫苗)",
        "registry": "FDA\nNCT03172975",
        "clinical": "分期: Phase 2\n设计: 随机、双盲、安慰剂对照、人类受控感染模型(CHHI)\n样本量: 39人\n分组: Na-GST-1/CpG组、Na-GST-1/AP组、无佐剂组、安慰剂组\n终点: 感染发生率、感染强度及安全性",
        "local": "绝大多数为轻度(Mild)反应。",
        "systemic": "无严重全身性反应。显著降低了受控感染后的最大嗜酸性粒细胞计数。",
        "sae": "试验全程未观察到与疫苗相关的严重不良事件 (SAE) 或 SADR。",
        "ref": "PMID: 41861834\nDOI: 10.1016/S1473-3099(26)00018-6",
    },
]

for row_data in data:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data["vaccine"]
    row_cells[1].text = row_data["registry"]
    row_cells[2].text = row_data["clinical"]
    row_cells[3].text = row_data["local"]
    row_cells[4].text = row_data["systemic"]
    row_cells[5].text = row_data["sae"]
    row_cells[6].text = row_data["ref"]

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(9)

doc.save(
    r"E:\Cursor Project\2-Scientific-Skills-for-Clinical_Trial\review_materials\CpG_Vaccine_Safety_Summary.docx"
)
